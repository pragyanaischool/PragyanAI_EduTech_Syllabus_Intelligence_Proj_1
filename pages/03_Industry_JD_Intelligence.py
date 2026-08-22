# ============================================================
# CHUNK 1/10
# 03_💼_Industry_JD_Intelligence.py
# IMPORTS + CONFIGURATION + SESSION STATE
# ============================================================

import json
import re
from datetime import datetime

import pandas as pd
import streamlit as st


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="Industry & JD Intelligence",
    page_icon="💼",
    layout="wide",
)


# ============================================================
# PAGE TITLE
# ============================================================

st.title(
    "💼 Industry & JD Intelligence"
)

st.caption(
    "Analyse Job Descriptions and map industry skills "
    "against the curriculum."
)


# ============================================================
# SESSION STATE INITIALIZATION
# ============================================================

SESSION_DEFAULTS = {

    "jd_documents": [],

    "jd_text": "",

    "jd_analysis": None,

    "jd_skills": [],

    "jd_keywords": [],

    "jd_roles": [],

    "jd_tools": [],

    "jd_projects": [],

    "curriculum_skills": [],

    "skill_match_results": [],

    "covered_skills": [],

    "partial_skills": [],

    "missing_skills": [],

    "industry_gaps": [],

    "industry_intelligence": None,

    "jd_report": None,

    "jd_master_package": None,

    "industry_jd_complete": False,

}


for key, default_value in SESSION_DEFAULTS.items():

    if key not in st.session_state:

        st.session_state[key] = default_value


# ============================================================
# HELPER — SAFE TEXT
# ============================================================

def safe_text(
    value,
    default="",
):
    """
    Convert a value safely into readable text.
    """

    if value is None:
        return default

    if isinstance(
        value,
        (dict, list),
    ):
        return str(value)

    value = str(value).strip()

    return value if value else default


# ============================================================
# HELPER — JSON SERIALIZATION
# ============================================================

def serialize_json(
    data,
):
    """
    Safely serialize Python objects for download.
    """

    return json.dumps(
        data,
        indent=2,
        ensure_ascii=False,
        default=str,
    )


# ============================================================
# HELPER — NORMALIZE LIST
# ============================================================

def normalize_list(
    value,
):
    """
    Convert arbitrary values into a clean list.
    """

    if value is None:
        return []

    if isinstance(
        value,
        list,
    ):

        values = value

    else:

        values = [value]


    cleaned = []

    for item in values:

        text = safe_text(item)

        if text:
            cleaned.append(text)


    return cleaned


# ============================================================
# HELPER — NORMALIZE SKILL
# ============================================================

def normalize_skill(
    skill,
):
    """
    Normalize a skill name for comparison.
    """

    skill = safe_text(
        skill
    ).lower()

    skill = re.sub(
        r"[^a-z0-9+#.\-/ ]",
        " ",
        skill,
    )

    skill = re.sub(
        r"\s+",
        " ",
        skill,
    ).strip()

    return skill


# ============================================================
# HELPER — UNIQUE VALUES
# ============================================================

def unique_values(
    values,
):
    """
    Return unique values while preserving readable form.
    """

    result = []

    seen = set()

    for value in values:

        text = safe_text(
            value
        )

        if not text:
            continue

        normalized = normalize_skill(
            text
        )

        if normalized in seen:
            continue

        seen.add(
            normalized
        )

        result.append(
            text
        )

    return result


# ============================================================
# HELPER — EXTRACT LIST FROM DICT
# ============================================================

def get_dict_list(
    data,
    *keys,
):
    """
    Return the first available list from a dictionary.
    """

    if not isinstance(
        data,
        dict,
    ):
        return []


    for key in keys:

        value = data.get(
            key
        )

        if value:

            return normalize_list(
                value
            )


    return []


# ============================================================
# HELPER — FLATTEN CURRICULUM TOPICS
# ============================================================

def flatten_curriculum_topics(
    curriculum,
):
    """
    Extract all topic names from the structured curriculum.
    """

    if not isinstance(
        curriculum,
        dict,
    ):
        return []


    topics = []


    modules = curriculum.get(
        "modules",
        []
    )


    if not isinstance(
        modules,
        list,
    ):
        return []


    for module in modules:

        if not isinstance(
            module,
            dict,
        ):
            continue


        module_topics = module.get(
            "topics",
            []
        )


        if isinstance(
            module_topics,
            list,
        ):

            for topic in module_topics:

                if isinstance(
                    topic,
                    dict,
                ):

                    topic_name = (
                        topic.get(
                            "name"
                        )
                        or
                        topic.get(
                            "topic"
                        )
                        or
                        topic.get(
                            "title"
                        )
                    )

                else:

                    topic_name = topic


                if topic_name:

                    topics.append(
                        safe_text(
                            topic_name
                        )
                    )


        # --------------------------------------------
        # Some extraction schemas may use concepts
        # --------------------------------------------

        concepts = module.get(
            "concepts",
            []
        )


        if isinstance(
            concepts,
            list,
        ):

            topics.extend(
                normalize_list(
                    concepts
                )
            )


    return unique_values(
        topics
    )


# ============================================================
# HELPER — EXTRACT CURRICULUM SKILLS
# ============================================================

def extract_curriculum_skills(
    curriculum,
):
    """
    Collect skills, tools, technologies and concepts from
    the structured curriculum.
    """

    if not isinstance(
        curriculum,
        dict,
    ):
        return []


    skills = []


    # --------------------------------------------
    # Top-level skills
    # --------------------------------------------

    skills.extend(
        get_dict_list(
            curriculum,
            "skills",
            "skill",
            "technical_skills",
            "technologies",
        )
    )


    # --------------------------------------------
    # Tools
    # --------------------------------------------

    skills.extend(
        get_dict_list(
            curriculum,
            "tools",
            "technologies",
            "frameworks",
            "libraries",
        )
    )


    # --------------------------------------------
    # Topics
    # --------------------------------------------

    skills.extend(
        flatten_curriculum_topics(
            curriculum
        )
    )


    # --------------------------------------------
    # Module-level extraction
    # --------------------------------------------

    modules = curriculum.get(
        "modules",
        []
    )


    if isinstance(
        modules,
        list,
    ):

        for module in modules:

            if not isinstance(
                module,
                dict,
            ):
                continue


            skills.extend(
                get_dict_list(
                    module,
                    "skills",
                    "tools",
                    "technologies",
                    "concepts",
                )
            )


    return unique_values(
        skills
    )


# ============================================================
# LOAD PAGE-02 OUTPUT
# ============================================================

curriculum_package = st.session_state.get(
    "curriculum_master_package"
)


primary_curriculum = st.session_state.get(
    "primary_syllabus"
)


comparison_data = st.session_state.get(
    "curriculum_comparison"
)


# ============================================================
# CURRICULUM AVAILABILITY
# ============================================================

curriculum_available = bool(
    primary_curriculum
)


# ============================================================
# PAGE-02 STATUS
# ============================================================

if curriculum_available:

    st.success(
        """
        ✅ Curriculum Intelligence data detected from Page 02.
        """
    )

else:

    st.warning(
        """
        ⚠️ Curriculum Intelligence data is not available.

        You can still upload a JD, but curriculum-to-JD matching
        will require the syllabus from Page 02.
        """
    )


# ============================================================
# PAGE ARCHITECTURE
# ============================================================

with st.expander(
    "ℹ️ Industry & JD Intelligence Workflow",
    expanded=False,
):

    st.markdown(
        """
        ### Workflow

        **1. Upload Job Description**

        PDF / DOCX / Image

        ↓

        **2. Extract JD**

        Role, company, experience, responsibilities,
        qualifications, skills and tools.

        ↓

        **3. Skill Intelligence**

        Normalize and classify JD skills.

        ↓

        **4. Curriculum Skill Extraction**

        Extract skills, concepts, tools and technologies
        from the curriculum.

        ↓

        **5. Skill Matching**

        Each JD skill is classified as:

        🟢 Covered

        🟡 Partially Covered

        🔴 Missing

        ↓

        **6. Industry Gap Intelligence**

        Identify missing concepts, tools, projects and
        technologies.

        ↓

        **7. Enhancement Input**

        Pass structured recommendations to:

        `04_🔍_Gap_Enhancement.py`
        """
    )


# ============================================================
# END OF CHUNK 1
# ============================================================
# ============================================================
# CHUNK 2/10
# JD UPLOAD + DOCUMENT EXTRACTION
# ============================================================

"""
Responsibilities
----------------
1. Upload Job Descriptions
2. Support PDF
3. Support DOCX
4. Support TXT
5. Support Images
6. Extract text
7. OCR image-based documents
8. Combine multiple JD documents
9. Display extracted text
10. Store JD text in session state
"""


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.header(
    "1️⃣ 📄 Upload Job Description"
)

st.markdown(
    """
Upload one or more Job Descriptions.

Supported formats:

- PDF
- DOCX
- TXT
- PNG
- JPG / JPEG
- WEBP

You can upload:

- A single JD
- Multiple JDs for the same role
- JDs from different companies
- Campus placement JDs
- Industry hiring JDs
"""
)


# ============================================================
# FILE UPLOADER
# ============================================================

uploaded_jds = st.file_uploader(

    "Upload Job Description(s)",

    type=[
        "pdf",
        "docx",
        "txt",
        "png",
        "jpg",
        "jpeg",
        "webp",
    ],

    accept_multiple_files=True,

    key="industry_jd_uploader",

)


# ============================================================
# DOCUMENT EXTRACTION HELPERS
# ============================================================

def extract_pdf_text(
    uploaded_file,
):
    """
    Extract text from PDF using the project's PDF loader
    when available.
    """

    try:

        from rag.pdf_loader import (
            load_pdf,
        )

        return load_pdf(
            uploaded_file
        )

    except ImportError:

        # ----------------------------------------------------
        # Fallback using PyMuPDF
        # ----------------------------------------------------

        import fitz

        pdf_bytes = uploaded_file.getvalue()

        document = fitz.open(
            stream=pdf_bytes,
            filetype="pdf",
        )

        pages = []

        for page in document:

            pages.append(
                page.get_text(
                    "text"
                )
            )

        document.close()

        return "\n".join(
            pages
        )


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(
    uploaded_file,
):
    """
    Extract text from DOCX.
    """

    try:

        from rag.docx_loader import (
            load_docx,
        )

        return load_docx(
            uploaded_file
        )

    except ImportError:

        from docx import Document

        document = Document(
            uploaded_file
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                paragraphs.append(
                    text
                )


        # ----------------------------------------------------
        # Extract tables
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = (
                        cell.text.strip()
                    )

                    if cell_text:

                        row_text.append(
                            cell_text
                        )


                if row_text:

                    paragraphs.append(
                        " | ".join(
                            row_text
                        )
                    )


        return "\n".join(
            paragraphs
        )


# ============================================================
# TXT EXTRACTION
# ============================================================

def extract_txt_text(
    uploaded_file,
):
    """
    Extract text from TXT.
    """

    raw_bytes = uploaded_file.getvalue()


    # --------------------------------------------------------
    # Try UTF-8
    # --------------------------------------------------------

    try:

        return raw_bytes.decode(
            "utf-8"
        )

    except UnicodeDecodeError:

        pass


    # --------------------------------------------------------
    # Fallback encodings
    # --------------------------------------------------------

    for encoding in [
        "utf-16",
        "latin-1",
        "cp1252",
    ]:

        try:

            return raw_bytes.decode(
                encoding
            )

        except UnicodeDecodeError:

            continue


    return raw_bytes.decode(
        "utf-8",
        errors="ignore",
    )


# ============================================================
# IMAGE / OCR EXTRACTION
# ============================================================

def extract_image_text(
    uploaded_file,
):
    """
    Extract text from image using project OCR loader when
    available, otherwise fall back to pytesseract.
    """

    try:

        from rag.ocr_loader import (
            load_image,
        )

        return load_image(
            uploaded_file
        )

    except ImportError:

        pass


    # --------------------------------------------------------
    # Fallback OCR
    # --------------------------------------------------------

    try:

        import pytesseract

        from PIL import Image

        image = Image.open(
            uploaded_file
        )

        return pytesseract.image_to_string(
            image
        )

    except ImportError:

        raise RuntimeError(
            """
            OCR dependencies are not installed.

            Install:

            pip install pytesseract pillow

            Also install the Tesseract OCR engine.
            """
        )


# ============================================================
# UNIVERSAL DOCUMENT EXTRACTION
# ============================================================

def extract_jd_document(
    uploaded_file,
):
    """
    Detect file type and extract text.
    """

    filename = (
        uploaded_file.name
        or ""
    )

    extension = (
        filename
        .lower()
        .split(".")
        [-1]
    )


    if extension == "pdf":

        return extract_pdf_text(
            uploaded_file
        )


    if extension == "docx":

        return extract_docx_text(
            uploaded_file
        )


    if extension == "txt":

        return extract_txt_text(
            uploaded_file
        )


    if extension in [
        "png",
        "jpg",
        "jpeg",
        "webp",
    ]:

        return extract_image_text(
            uploaded_file
        )


    raise ValueError(
        f"Unsupported file type: {extension}"
    )


# ============================================================
# NORMALIZE EXTRACTED TEXT
# ============================================================

def clean_extracted_text(
    text,
):
    """
    Clean extracted document text while preserving meaningful
    line structure.
    """

    if text is None:

        return ""


    text = str(
        text
    )


    # --------------------------------------------------------
    # Normalize line endings
    # --------------------------------------------------------

    text = text.replace(
        "\r\n",
        "\n",
    )

    text = text.replace(
        "\r",
        "\n",
    )


    # --------------------------------------------------------
    # Remove null characters
    # --------------------------------------------------------

    text = text.replace(
        "\x00",
        " ",
    )


    # --------------------------------------------------------
    # Normalize whitespace
    # --------------------------------------------------------

    lines = []

    for line in text.split(
        "\n"
    ):

        line = re.sub(
            r"[ \t]+",
            " ",
            line,
        )

        line = line.strip()

        if line:

            lines.append(
                line
            )


    return "\n".join(
        lines
    )


# ============================================================
# EXTRACT UPLOADED DOCUMENTS
# ============================================================

if uploaded_jds:

    st.divider()

    st.subheader(
        "📥 Uploaded Job Descriptions"
    )


    uploaded_summary = []


    for uploaded_file in uploaded_jds:

        uploaded_summary.append({

            "File":
                uploaded_file.name,

            "Type":
                uploaded_file.type,

            "Size":
                f"{uploaded_file.size / 1024:.1f} KB",

        })


    uploaded_df = pd.DataFrame(
        uploaded_summary
    )


    st.dataframe(

        uploaded_df,

        use_container_width=True,

        hide_index=True,

    )


    # ========================================================
    # EXTRACTION BUTTON
    # ========================================================

    extract_button = st.button(

        "📄 Extract JD Text",

        type="primary",

        use_container_width=True,

        key="extract_jd_button",

    )


    if extract_button:

        extracted_documents = []

        extraction_errors = []


        progress = st.progress(
            0
        )


        status = st.empty()


        total_files = len(
            uploaded_jds
        )


        # ====================================================
        # PROCESS EACH FILE
        # ====================================================

        for index, uploaded_file in enumerate(
            uploaded_jds,
            start=1,
        ):

            status.info(

                f"""
                Extracting:

                **{uploaded_file.name}**

                File {index} of {total_files}
                """
            )


            try:

                # --------------------------------------------
                # Extract
                # --------------------------------------------

                extracted_text = (
                    extract_jd_document(
                        uploaded_file
                    )
                )


                # --------------------------------------------
                # Clean
                # --------------------------------------------

                extracted_text = (
                    clean_extracted_text(
                        extracted_text
                    )
                )


                # --------------------------------------------
                # Validate
                # --------------------------------------------

                if not extracted_text:

                    raise ValueError(
                        "No text could be extracted."
                    )


                # --------------------------------------------
                # Store document
                # --------------------------------------------

                extracted_documents.append({

                    "filename":
                        uploaded_file.name,

                    "file_type":
                        uploaded_file.type,

                    "text":
                        extracted_text,

                    "character_count":
                        len(
                            extracted_text
                        ),

                    "word_count":
                        len(
                            extracted_text.split()
                        ),

                })


            except Exception as exc:

                extraction_errors.append({

                    "filename":
                        uploaded_file.name,

                    "error":
                        str(
                            exc
                        ),

                })


            progress.progress(
                index / total_files
            )


        # ====================================================
        # SAVE SUCCESSFUL EXTRACTIONS
        # ====================================================

        st.session_state[
            "jd_documents"
        ] = extracted_documents


        # ====================================================
        # COMBINE DOCUMENTS
        # ====================================================

        combined_sections = []


        for document in extracted_documents:

            combined_sections.append(

                f"""
============================================================
JOB DESCRIPTION: {document["filename"]}
============================================================

{document["text"]}
"""

            )


        combined_jd_text = "\n".join(
            combined_sections
        )


        st.session_state[
            "jd_text"
        ] = combined_jd_text.strip()


        # ====================================================
        # COMPLETION MESSAGE
        # ====================================================

        status.success(

            f"""
            ✅ Extraction completed.

            Successfully processed:
            **{len(extracted_documents)} / {total_files}**
            document(s).
            """
        )


        # ====================================================
        # ERROR DISPLAY
        # ====================================================

        if extraction_errors:

            st.warning(
                f"""
                {len(extraction_errors)}
                document(s) could not be processed.
                """
            )


            with st.expander(
                "⚠️ View Extraction Errors"
            ):

                st.dataframe(

                    pd.DataFrame(
                        extraction_errors
                    ),

                    use_container_width=True,

                    hide_index=True,

                )


# ============================================================
# LOAD EXTRACTED TEXT
# ============================================================

jd_text = st.session_state.get(
    "jd_text",
    "",
)


# ============================================================
# DISPLAY EXTRACTION RESULT
# ============================================================

if jd_text:

    st.divider()

    st.subheader(
        "📋 Extracted Job Description"
    )


    # ========================================================
    # TEXT METRICS
    # ========================================================

    words = jd_text.split()

    characters = len(
        jd_text
    )

    lines = len(
        jd_text.splitlines()
    )


    metric_col1, metric_col2, metric_col3 = (
        st.columns(3)
    )


    with metric_col1:

        st.metric(
            "Characters",
            f"{characters:,}",
        )


    with metric_col2:

        st.metric(
            "Words",
            f"{len(words):,}",
        )


    with metric_col3:

        st.metric(
            "Lines",
            f"{lines:,}",
        )


    # ========================================================
    # TEXT PREVIEW
    # ========================================================

    with st.expander(
        "👁️ View Extracted JD Text",
        expanded=True,
    ):

        st.text_area(

            "Extracted Text",

            value=jd_text,

            height=500,

            disabled=True,

            label_visibility="collapsed",

            key="extracted_jd_preview",

        )


    # ========================================================
    # DOWNLOAD RAW TEXT
    # ========================================================

    st.download_button(

        label="⬇️ Download Extracted JD Text",

        data=jd_text,

        file_name="extracted_job_description.txt",

        mime="text/plain",

        key="download_extracted_jd",

    )


# ============================================================
# NO JD AVAILABLE
# ============================================================

else:

    st.info(
        """
        Upload one or more Job Descriptions and click
        **Extract JD Text** to continue.
        """
    )


# ============================================================
# EXTRACTION STATUS
# ============================================================

st.session_state[
    "jd_extraction_complete"
] = bool(
    jd_text
)


# ============================================================
# END OF CHUNK 2
# ============================================================
# ============================================================
# CHUNK 3/10
# JD PARSING + STRUCTURED INTELLIGENCE
# ============================================================

"""
Responsibilities
----------------
1. Read extracted JD text
2. Send JD to LLM
3. Extract structured job information
4. Parse JSON safely
5. Normalize the response
6. Support multiple JDs
7. Save structured JD intelligence
8. Display structured JD information

Expected output:

{
    company,
    job_title,
    job_family,
    department,
    location,
    employment_type,
    experience,
    education,
    responsibilities,
    required_skills,
    preferred_skills,
    tools,
    technologies,
    frameworks,
    certifications,
    domain_knowledge,
    soft_skills,
    project_expectations,
    keywords
}
"""


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.header(
    "2️⃣ 🧠 JD Intelligence & Structured Analysis"
)

st.markdown(
    """
The extracted Job Description is now converted into a
structured representation.

The AI identifies:

- Company
- Job title
- Role / job family
- Department
- Location
- Experience
- Education
- Responsibilities
- Required skills
- Preferred skills
- Tools
- Technologies
- Frameworks
- Certifications
- Domain knowledge
- Soft skills
- Project expectations
- Important keywords
"""
)


# ============================================================
# LOAD JD TEXT
# ============================================================

jd_text = st.session_state.get(
    "jd_text",
    "",
)


jd_documents = st.session_state.get(
    "jd_documents",
    [],
)


# ============================================================
# JSON EXTRACTION HELPER
# ============================================================

def extract_json_from_response(
    response,
):
    """
    Extract JSON from an LLM response.

    Handles:

    1. Pure JSON
    2. Markdown JSON code blocks
    3. JSON embedded in explanatory text
    """

    if response is None:

        raise ValueError(
            "LLM returned an empty response."
        )


    # --------------------------------------------------------
    # Convert response to text
    # --------------------------------------------------------

    if isinstance(
        response,
        dict,
    ):

        return response


    text = str(
        response
    ).strip()


    if not text:

        raise ValueError(
            "LLM response is empty."
        )


    # --------------------------------------------------------
    # Remove markdown fences
    # --------------------------------------------------------

    text = re.sub(
        r"^```json\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )


    text = re.sub(
        r"^```\s*",
        "",
        text,
    )


    text = re.sub(
        r"\s*```$",
        "",
        text,
    )


    text = text.strip()


    # --------------------------------------------------------
    # Direct JSON
    # --------------------------------------------------------

    try:

        return json.loads(
            text
        )

    except json.JSONDecodeError:

        pass


    # --------------------------------------------------------
    # Find first JSON object
    # --------------------------------------------------------

    start = text.find(
        "{"
    )

    end = text.rfind(
        "}"
    )


    if start != -1 and end != -1:

        candidate = text[
            start:end + 1
        ]


        try:

            return json.loads(
                candidate
            )

        except json.JSONDecodeError:

            pass


    # --------------------------------------------------------
    # Find JSON array
    # --------------------------------------------------------

    start = text.find(
        "["
    )

    end = text.rfind(
        "]"
    )


    if start != -1 and end != -1:

        candidate = text[
            start:end + 1
        ]


        try:

            return json.loads(
                candidate
            )

        except json.JSONDecodeError:

            pass


    raise ValueError(
        "Could not parse valid JSON from LLM response."
    )


# ============================================================
# NORMALIZE JD ANALYSIS
# ============================================================

def normalize_jd_analysis(
    data,
):
    """
    Normalize the LLM response into a predictable schema.
    """

    if not isinstance(
        data,
        dict,
    ):

        data = {}


    normalized = {

        "company":
            safe_text(
                data.get(
                    "company"
                )
            ),

        "job_title":
            safe_text(
                data.get(
                    "job_title"
                )
                or
                data.get(
                    "role"
                )
            ),

        "job_family":
            safe_text(
                data.get(
                    "job_family"
                )
            ),

        "department":
            safe_text(
                data.get(
                    "department"
                )
            ),

        "location":
            safe_text(
                data.get(
                    "location"
                )
            ),

        "employment_type":
            safe_text(
                data.get(
                    "employment_type"
                )
            ),

        "experience":
            safe_text(
                data.get(
                    "experience"
                )
                or
                data.get(
                    "experience_required"
                )
            ),

        "education":
            normalize_list(
                data.get(
                    "education"
                )
            ),

        "responsibilities":
            normalize_list(
                data.get(
                    "responsibilities"
                )
                or
                data.get(
                    "job_responsibilities"
                )
            ),

        "required_skills":
            unique_values(
                normalize_list(
                    data.get(
                        "required_skills"
                    )
                )
            ),

        "preferred_skills":
            unique_values(
                normalize_list(
                    data.get(
                        "preferred_skills"
                    )
                )
            ),

        "tools":
            unique_values(
                normalize_list(
                    data.get(
                        "tools"
                    )
                )
            ),

        "technologies":
            unique_values(
                normalize_list(
                    data.get(
                        "technologies"
                    )
                )
            ),

        "frameworks":
            unique_values(
                normalize_list(
                    data.get(
                        "frameworks"
                    )
                )
            ),

        "libraries":
            unique_values(
                normalize_list(
                    data.get(
                        "libraries"
                    )
                )
            ),

        "cloud_platforms":
            unique_values(
                normalize_list(
                    data.get(
                        "cloud_platforms"
                    )
            )
            ),

        "databases":
            unique_values(
                normalize_list(
                    data.get(
                        "databases"
                    )
                )
            ),

        "programming_languages":
            unique_values(
                normalize_list(
                    data.get(
                        "programming_languages"
                    )
                )
            ),

        "certifications":
            normalize_list(
                data.get(
                    "certifications"
                )
            ),

        "domain_knowledge":
            unique_values(
                normalize_list(
                    data.get(
                        "domain_knowledge"
                    )
                )
            ),

        "soft_skills":
            unique_values(
                normalize_list(
                    data.get(
                        "soft_skills"
                    )
                )
            ),

        "project_expectations":
            normalize_list(
                data.get(
                    "project_expectations"
                )
            ),

        "keywords":
            unique_values(
                normalize_list(
                    data.get(
                        "keywords"
                    )
                )
            ),

        "job_summary":
            safe_text(
                data.get(
                    "job_summary"
                )
            ),

    }


    # ========================================================
    # BUILD MASTER SKILL LIST
    # ========================================================

    normalized["all_skills"] = unique_values(

        normalized["required_skills"]

        +

        normalized["preferred_skills"]

        +

        normalized["technologies"]

        +

        normalized["frameworks"]

        +

        normalized["libraries"]

        +

        normalized["programming_languages"]

        +

        normalized["cloud_platforms"]

        +

        normalized["databases"]

    )


    return normalized


# ============================================================
# FALLBACK JD PARSER
# ============================================================

def fallback_parse_jd(
    text,
):
    """
    Basic rule-based fallback.

    This is intentionally lightweight.

    The primary analysis should use the LLM.
    """

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]


    joined_text = " ".join(
        lines
    )


    result = {

        "company": "",

        "job_title": "",

        "job_family": "",

        "department": "",

        "location": "",

        "employment_type": "",

        "experience": "",

        "education": [],

        "responsibilities": [],

        "required_skills": [],

        "preferred_skills": [],

        "tools": [],

        "technologies": [],

        "frameworks": [],

        "libraries": [],

        "cloud_platforms": [],

        "databases": [],

        "programming_languages": [],

        "certifications": [],

        "domain_knowledge": [],

        "soft_skills": [],

        "project_expectations": [],

        "keywords": [],

        "job_summary": "",

    }


    # ========================================================
    # COMMON SECTION PATTERNS
    # ========================================================

    section_patterns = {

        "responsibilities": [
            "responsibilities",
            "what you will do",
            "what you'll do",
            "job responsibilities",
            "roles and responsibilities",
        ],

        "required_skills": [
            "required skills",
            "required qualifications",
            "must have",
            "requirements",
            "technical skills",
        ],

        "preferred_skills": [
            "preferred skills",
            "good to have",
            "nice to have",
            "preferred qualifications",
        ],

        "education": [
            "education",
            "educational qualification",
            "qualification",
        ],

        "certifications": [
            "certifications",
            "certification",
        ],

    }


    current_section = None


    for line in lines:

        normalized_line = (
            line.lower()
            .strip(":")
            .strip()
        )


        matched_section = None


        for section, patterns in (
            section_patterns.items()
        ):

            if any(
                pattern in normalized_line
                for pattern in patterns
            ):

                matched_section = section

                break


        if matched_section:

            current_section = (
                matched_section
            )

            continue


        # ----------------------------------------------------
        # Add bullet to active section
        # ----------------------------------------------------

        cleaned_line = re.sub(
            r"^[\-\*\•\▪\d\.\)\s]+",
            "",
            line,
        ).strip()


        if not cleaned_line:

            continue


        if current_section:

            if len(cleaned_line) > 2:

                result[
                    current_section
                ].append(
                    cleaned_line
                )


    # ========================================================
    # EXPERIENCE
    # ========================================================

    experience_patterns = [

        r"\b\d+\s*[-to]+\s*\d+\s*years?\b",

        r"\b\d+\+?\s*years?\b",

        r"\bminimum\s+\d+\s*years?\b",

    ]


    for pattern in experience_patterns:

        match = re.search(
            pattern,
            joined_text,
            flags=re.IGNORECASE,
        )


        if match:

            result[
                "experience"
            ] = match.group(
                0
            )

            break


    # ========================================================
    # PROGRAMMING LANGUAGES
    # ========================================================

    language_candidates = [

        "Python",

        "Java",

        "JavaScript",

        "TypeScript",

        "C",

        "C++",

        "C#",

        "Go",

        "Golang",

        "Rust",

        "R",

        "Scala",

        "Kotlin",

        "Swift",

        "SQL",

    ]


    for language in language_candidates:

        if re.search(

            rf"\b{re.escape(language)}\b",

            joined_text,

            flags=re.IGNORECASE,

        ):

            result[
                "programming_languages"
            ].append(
                language
            )


    # ========================================================
    # COMMON AI / DATA TECHNOLOGIES
    # ========================================================

    technology_candidates = [

        "Machine Learning",

        "Deep Learning",

        "Generative AI",

        "Artificial Intelligence",

        "NLP",

        "Natural Language Processing",

        "Computer Vision",

        "LLM",

        "Large Language Models",

        "RAG",

        "Retrieval Augmented Generation",

        "Agentic AI",

        "Data Science",

        "Data Engineering",

        "MLOps",

        "LLMOps",

        "DevOps",

    ]


    for technology in technology_candidates:

        if re.search(

            rf"\b{re.escape(technology)}\b",

            joined_text,

            flags=re.IGNORECASE,

        ):

            result[
                "technologies"
            ].append(
                technology
            )


    # ========================================================
    # FRAMEWORKS
    # ========================================================

    framework_candidates = [

        "LangChain",

        "LangGraph",

        "LlamaIndex",

        "TensorFlow",

        "PyTorch",

        "Keras",

        "Scikit-learn",

        "Hugging Face",

        "Transformers",

        "FastAPI",

        "Django",

        "Flask",

        "Streamlit",

        "React",

        "Angular",

        "Vue",

        "Spring",

        "Spring Boot",

    ]


    for framework in framework_candidates:

        if re.search(

            rf"\b{re.escape(framework)}\b",

            joined_text,

            flags=re.IGNORECASE,

        ):

            result[
                "frameworks"
            ].append(
                framework
            )


    # ========================================================
    # CLOUD
    # ========================================================

    cloud_candidates = [

        "AWS",

        "Amazon Web Services",

        "Azure",

        "Microsoft Azure",

        "GCP",

        "Google Cloud",

        "Google Cloud Platform",

    ]


    for cloud in cloud_candidates:

        if re.search(

            rf"\b{re.escape(cloud)}\b",

            joined_text,

            flags=re.IGNORECASE,

        ):

            result[
                "cloud_platforms"
            ].append(
                cloud
            )


    # ========================================================
    # DATABASES
    # ========================================================

    database_candidates = [

        "PostgreSQL",

        "MySQL",

        "MongoDB",

        "Oracle",

        "SQL Server",

        "Redis",

        "Elasticsearch",

        "Pinecone",

        "Chroma",

        "FAISS",

        "Weaviate",

    ]


    for database in database_candidates:

        if re.search(

            rf"\b{re.escape(database)}\b",

            joined_text,

            flags=re.IGNORECASE,

        ):

            result[
                "databases"
            ].append(
                database
            )


    # ========================================================
    # JOB TITLE HEURISTIC
    # ========================================================

    title_patterns = [

        r"(?:job title|position|role)\s*[:\-]\s*(.+)",

        r"(?:designation)\s*[:\-]\s*(.+)",

    ]


    for pattern in title_patterns:

        match = re.search(

            pattern,

            joined_text,

            flags=re.IGNORECASE,

        )


        if match:

            result[
                "job_title"
            ] = match.group(
                1
            ).strip()

            break


    # ========================================================
    # SUMMARY
    # ========================================================

    if lines:

        result[
            "job_summary"
        ] = " ".join(
            lines[:5]
        )


    return normalize_jd_analysis(
        result
    )


# ============================================================
# GROQ / LLM CALL
# ============================================================

def analyze_jd_with_llm(
    text,
):
    """
    Use the project's LLM implementation.

    Preferred architecture:

        Groq
          ↓
        Llama
          ↓
        Structured JSON
    """

    # --------------------------------------------------------
    # Try project-level LLM helper first
    # --------------------------------------------------------

    try:

        from llm import get_llm

        llm = get_llm()

    except ImportError:

        llm = None


    # --------------------------------------------------------
    # Direct LangChain Groq fallback
    # --------------------------------------------------------

    if llm is None:

        try:

            from langchain_groq import (
                ChatGroq,
            )

            import os


            api_key = os.getenv(
                "GROQ_API_KEY"
            )


            if not api_key:

                raise RuntimeError(
                    "GROQ_API_KEY is not configured."
                )


            llm = ChatGroq(

                model=(
                    os.getenv(
                        "GROQ_MODEL"
                    )
                    or
                    "llama-3.3-70b-versatile"
                ),

                temperature=0,

                api_key=api_key,

            )

        except ImportError:

            raise RuntimeError(
                """
                Groq integration is not installed.

                Install:

                pip install langchain-groq
                """
            )


    # ========================================================
    # SYSTEM PROMPT
    # ========================================================

    system_prompt = """
You are an expert Job Description Intelligence Analyst.

Analyze the supplied Job Description and return ONLY valid JSON.

Do not return markdown.

Do not return explanations outside the JSON.

Extract information accurately.

Do not invent information that is not present.

If a field is unavailable, return an empty string or empty list.

Required JSON structure:

{
  "company": "",
  "job_title": "",
  "job_family": "",
  "department": "",
  "location": "",
  "employment_type": "",
  "experience": "",
  "education": [],
  "responsibilities": [],
  "required_skills": [],
  "preferred_skills": [],
  "tools": [],
  "technologies": [],
  "frameworks": [],
  "libraries": [],
  "cloud_platforms": [],
  "databases": [],
  "programming_languages": [],
  "certifications": [],
  "domain_knowledge": [],
  "soft_skills": [],
  "project_expectations": [],
  "keywords": [],
  "job_summary": ""
}

Important:

- required_skills = explicitly required technical or professional skills
- preferred_skills = optional / preferred skills
- technologies = technologies and technical concepts
- frameworks = software frameworks
- tools = named tools/platforms
- libraries = named libraries
- cloud_platforms = AWS/Azure/GCP etc.
- databases = SQL/NoSQL/vector databases
- programming_languages = Python/Java/etc.
- domain_knowledge = domain-specific knowledge
- project_expectations = projects, systems or practical work expected
- keywords = important terms useful for later skill matching
"""


    user_prompt = f"""
Analyze this Job Description.

JOB DESCRIPTION
---------------
{text}

Return only the requested JSON object.
"""


    # ========================================================
    # INVOKE LLM
    # ========================================================

    try:

        # ----------------------------------------------------
        # LangChain message interface
        # ----------------------------------------------------

        from langchain_core.messages import (
            SystemMessage,
            HumanMessage,
        )


        response = llm.invoke([

            SystemMessage(
                content=system_prompt
            ),

            HumanMessage(
                content=user_prompt
            ),

        ])


        # ----------------------------------------------------
        # Extract response content
        # ----------------------------------------------------

        if hasattr(
            response,
            "content",
        ):

            response_content = (
                response.content
            )

        else:

            response_content = (
                str(response)
            )


        # ----------------------------------------------------
        # Parse JSON
        # ----------------------------------------------------

        parsed = extract_json_from_response(
            response_content
        )


        return normalize_jd_analysis(
            parsed
        )


    except Exception as exc:

        raise RuntimeError(
            f"LLM JD analysis failed: {exc}"
        )


# ============================================================
# ANALYZE JD BUTTON
# ============================================================

if jd_text:

    st.divider()

    analyze_col1, analyze_col2 = (
        st.columns([3, 1])
    )


    with analyze_col1:

        st.markdown(
            "### 🧠 Generate Structured JD Intelligence"
        )

        st.caption(
            """
            The JD will be analysed using Groq + Llama when
            configured. A rule-based fallback is available if
            the LLM cannot be used.
            """
        )


    with analyze_col2:

        analyze_jd_button = st.button(

            "🚀 Analyse JD",

            type="primary",

            use_container_width=True,

            key="analyze_jd_button",

        )


    # ========================================================
    # EXECUTE ANALYSIS
    # ========================================================

    if analyze_jd_button:

        analysis_results = []


        analysis_errors = []


        progress = st.progress(
            0
        )


        status = st.empty()


        total_documents = max(
            len(jd_documents),
            1,
        )


        # ====================================================
        # MULTIPLE DOCUMENTS
        # ====================================================

        if jd_documents:

            documents_to_analyze = (
                jd_documents
            )

        else:

            documents_to_analyze = [{

                "filename":
                    "Uploaded JD",

                "text":
                    jd_text,

            }]


        total_documents = len(
            documents_to_analyze
        )


        for index, document in enumerate(

            documents_to_analyze,

            start=1,

        ):

            filename = safe_text(

                document.get(
                    "filename"
                ),

                f"JD {index}",

            )


            document_text = safe_text(

                document.get(
                    "text"
                )

            )


            status.info(

                f"""
                🧠 Analysing **{filename}**

                Document {index} of {total_documents}
                """
            )


            try:

                # ==========================================
                # TRY LLM
                # ==========================================

                try:

                    analysis = (
                        analyze_jd_with_llm(
                            document_text
                        )
                    )


                    analysis_source = (
                        "Groq + Llama"
                    )


                except Exception as llm_error:

                    # ======================================
                    # FALLBACK
                    # ======================================

                    analysis = (
                        fallback_parse_jd(
                            document_text
                        )
                    )


                    analysis_source = (
                        "Rule-based fallback"
                    )


                    analysis[
                        "_llm_error"
                    ] = str(
                        llm_error
                    )


                # ==========================================
                # METADATA
                # ==========================================

                analysis[
                    "_source_file"
                ] = filename


                analysis[
                    "_analysis_source"
                ] = analysis_source


                analysis[
                    "_analyzed_at"
                ] = datetime.now().isoformat()


                analysis_results.append(
                    analysis
                )


            except Exception as exc:

                analysis_errors.append({

                    "filename":
                        filename,

                    "error":
                        str(
                            exc
                        ),

                })


            progress.progress(
                index / total_documents
            )


        # ====================================================
        # SAVE RESULTS
        # ====================================================

        st.session_state[
            "jd_analysis"
        ] = analysis_results


        # ====================================================
        # COMBINE SKILLS
        # ====================================================

        all_required_skills = []

        all_preferred_skills = []

        all_tools = []

        all_technologies = []

        all_frameworks = []

        all_libraries = []

        all_cloud = []

        all_databases = []

        all_languages = []

        all_keywords = []


        for analysis in analysis_results:

            all_required_skills.extend(

                analysis.get(
                    "required_skills",
                    []
                )

            )


            all_preferred_skills.extend(

                analysis.get(
                    "preferred_skills",
                    []
                )

            )


            all_tools.extend(

                analysis.get(
                    "tools",
                    []
                )

            )


            all_technologies.extend(

                analysis.get(
                    "technologies",
                    []
                )

            )


            all_frameworks.extend(

                analysis.get(
                    "frameworks",
                    []
                )

            )


            all_libraries.extend(

                analysis.get(
                    "libraries",
                    []
                )

            )


            all_cloud.extend(

                analysis.get(
                    "cloud_platforms",
                    []
                )

            )


            all_databases.extend(

                analysis.get(
                    "databases",
                    []
                )

            )


            all_languages.extend(

                analysis.get(
                    "programming_languages",
                    []
                )

            )


            all_keywords.extend(

                analysis.get(
                    "keywords",
                    []
                )

            )


        # ====================================================
        # SAVE NORMALIZED JD SKILLS
        # ====================================================

        st.session_state[
            "jd_skills"
        ] = unique_values(

            all_required_skills

            +

            all_preferred_skills

            +

            all_technologies

            +

            all_frameworks

            +

            all_libraries

            +

            all_cloud

            +

            all_databases

            +

            all_languages

        )


        st.session_state[
            "jd_keywords"
        ] = unique_values(
            all_keywords
        )


        st.session_state[
            "jd_tools"
        ] = unique_values(

            all_tools

            +

            all_frameworks

            +

            all_libraries

        )


        # ====================================================
        # SAVE ANALYSIS METADATA
        # ====================================================

        st.session_state[
            "jd_analysis_metadata"
        ] = {

            "documents_analyzed":
                len(
                    analysis_results
                ),

            "errors":
                analysis_errors,

            "generated_at":
                datetime.now().isoformat(),

        }


        status.success(
            f"""
            ✅ JD analysis completed.

            Analysed:
            **{len(analysis_results)}**
            document(s).
            """
        )


        # ====================================================
        # ANALYSIS ERRORS
        # ====================================================

        if analysis_errors:

            with st.expander(
                "⚠️ Analysis Errors"
            ):

                st.dataframe(

                    pd.DataFrame(
                        analysis_errors
                    ),

                    use_container_width=True,

                    hide_index=True,

                )


# ============================================================
# DISPLAY STRUCTURED JD RESULTS
# ============================================================

jd_analysis = st.session_state.get(
    "jd_analysis",
    []
)


if jd_analysis:

    st.divider()

    st.subheader(
        "📊 Structured JD Intelligence"
    )


    # ========================================================
    # MULTI-JD TABS
    # ========================================================

    if len(jd_analysis) > 1:

        jd_tabs = st.tabs([

            safe_text(

                item.get(
                    "_source_file"
                ),

                f"JD {index + 1}",

            )

            for index, item in enumerate(
                jd_analysis
            )

        ])

    else:

        jd_tabs = [
            st.container()
        ]


    # ========================================================
    # DISPLAY EACH JD
    # ========================================================

    for index, analysis in enumerate(
        jd_analysis
    ):

        with jd_tabs[index]:

            # ================================================
            # BASIC INFORMATION
            # ================================================

            info_col1, info_col2, info_col3 = (
                st.columns(3)
            )


            with info_col1:

                st.markdown(
                    f"""
                    **Company**

                    {safe_text(
                        analysis.get("company"),
                        "Not specified"
                    )}
                    """
                )


            with info_col2:

                st.markdown(
                    f"""
                    **Job Title**

                    {safe_text(
                        analysis.get("job_title"),
                        "Not specified"
                    )}
                    """
                )


            with info_col3:

                st.markdown(
                    f"""
                    **Experience**

                    {safe_text(
                        analysis.get("experience"),
                        "Not specified"
                    )}
                    """
                )


            # ================================================
            # ROLE DETAILS
            # ================================================

            st.markdown(
                "### 🏢 Role Details"
            )


            role_df = pd.DataFrame([{

                "Job Family":
                    analysis.get(
                        "job_family"
                    ),

                "Department":
                    analysis.get(
                        "department"
                    ),

                "Location":
                    analysis.get(
                        "location"
                    ),

                "Employment Type":
                    analysis.get(
                        "employment_type"
                    ),

                "Experience":
                    analysis.get(
                        "experience"
                    ),

            }])


            st.dataframe(

                role_df,

                use_container_width=True,

                hide_index=True,

            )


            # ================================================
            # RESPONSIBILITIES
            # ================================================

            st.markdown(
                "### 📌 Responsibilities"
            )


            responsibilities = analysis.get(
                "responsibilities",
                []
            )


            if responsibilities:

                for item in responsibilities:

                    st.write(
                        f"• {item}"
                    )

            else:

                st.info(
                    "No responsibilities extracted."
                )


            # ================================================
            # REQUIRED SKILLS
            # ================================================

            st.markdown(
                "### 🎯 Required Skills"
            )


            required_skills = analysis.get(
                "required_skills",
                []
            )


            if required_skills:

                skill_columns = st.columns(
                    min(
                        3,
                        len(
                            required_skills
                        )
                    )
                )


                for skill_index, skill in enumerate(
                    required_skills
                ):

                    skill_columns[
                        skill_index
                        % len(skill_columns)
                    ].success(
                        skill
                    )

            else:

                st.info(
                    "No required skills extracted."
                )


            # ================================================
            # PREFERRED SKILLS
            # ================================================

            st.markdown(
                "### ⭐ Preferred Skills"
            )


            preferred_skills = analysis.get(
                "preferred_skills",
                []
            )


            if preferred_skills:

                for skill in preferred_skills:

                    st.write(
                        f"• {skill}"
                    )

            else:

                st.info(
                    "No preferred skills extracted."
                )


            # ================================================
            # TOOLS / TECHNOLOGIES
            # ================================================

            st.markdown(
                "### 🛠 Tools & Technologies"
            )


            technology_data = {

                "Programming Languages":
                    analysis.get(
                        "programming_languages",
                        []
                    ),

                "Technologies":
                    analysis.get(
                        "technologies",
                        []
                    ),

                "Frameworks":
                    analysis.get(
                        "frameworks",
                        []
                    ),

                "Libraries":
                    analysis.get(
                        "libraries",
                        []
                    ),

                "Cloud":
                    analysis.get(
                        "cloud_platforms",
                        []
                    ),

                "Databases":
                    analysis.get(
                        "databases",
                        []
                    ),

                "Tools":
                    analysis.get(
                        "tools",
                        []
                    ),

            }


            for category, values in (
                technology_data.items()
            ):

                values = normalize_list(
                    values
                )


                if values:

                    st.markdown(
                        f"**{category}**"
                    )


                    st.write(
                        ", ".join(
                            values
                        )
                    )


            # ================================================
            # DOMAIN KNOWLEDGE
            # ================================================

            domain_knowledge = analysis.get(
                "domain_knowledge",
                []
            )


            if domain_knowledge:

                st.markdown(
                    "### 🏭 Domain Knowledge"
                )


                for item in domain_knowledge:

                    st.write(
                        f"• {item}"
                    )


            # ================================================
            # PROJECT EXPECTATIONS
            # ================================================

            project_expectations = analysis.get(
                "project_expectations",
                []
            )


            if project_expectations:

                st.markdown(
                    "### 🚀 Project Expectations"
                )


                for item in project_expectations:

                    st.write(
                        f"• {item}"
                    )


            # ================================================
            # EDUCATION / CERTIFICATIONS
            # ================================================

            education_col, certification_col = (
                st.columns(2)
            )


            with education_col:

                st.markdown(
                    "### 🎓 Education"
                )


                for item in normalize_list(
                    analysis.get(
                        "education"
                    )
                ):

                    st.write(
                        f"• {item}"
                    )


            with certification_col:

                st.markdown(
                    "### 📜 Certifications"
                )


                for item in normalize_list(
                    analysis.get(
                        "certifications"
                    )
                ):

                    st.write(
                        f"• {item}"
                    )


            # ================================================
            # JOB SUMMARY
            # ================================================

            summary = safe_text(
                analysis.get(
                    "job_summary"
                )
            )


            if summary:

                st.markdown(
                    "### 📝 Job Summary"
                )


                st.write(
                    summary
                )


            # ================================================
            # RAW STRUCTURED DATA
            # ================================================

            with st.expander(
                "🔧 View Structured JSON"
            ):

                st.json(
                    analysis
                )


# ============================================================
# END OF CHUNK 3
# ============================================================
# ============================================================
# CHUNK 4/10
# JD SKILL INTELLIGENCE
# TAXONOMY + NORMALIZATION + KEYWORD ANALYSIS
# ============================================================

"""
Responsibilities
----------------

1. Collect skills from structured JD analysis
2. Normalize skill names
3. Categorize skills
4. Separate mandatory / preferred skills
5. Identify technologies
6. Identify tools
7. Identify frameworks
8. Identify programming languages
9. Identify cloud platforms
10. Identify databases
11. Identify domain skills
12. Identify soft skills
13. Identify project-related skills
14. Build a unified JD skill taxonomy
15. Prepare data for curriculum matching

Output:

st.session_state["jd_skill_intelligence"]

Structure:

{
    "all_skills": [],
    "required_skills": [],
    "preferred_skills": [],
    "technical_skills": [],
    "tools": [],
    "technologies": [],
    "frameworks": [],
    "languages": [],
    "cloud": [],
    "databases": [],
    "domain_skills": [],
    "soft_skills": [],
    "project_skills": [],
    "keywords": [],
    "skill_records": []
}
"""


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.header(
    "3️⃣ 🎯 JD Skill Intelligence"
)

st.markdown(
    """
The structured Job Description is now converted into a
normalized **Industry Skill Profile**.

The platform separates:

- Required skills
- Preferred skills
- Technical skills
- Programming languages
- Frameworks
- Tools
- Cloud platforms
- Databases
- Domain knowledge
- Soft skills
- Project skills
- Industry keywords
"""
)


# ============================================================
# LOAD JD ANALYSIS
# ============================================================

jd_analysis = st.session_state.get(
    "jd_analysis",
    []
)


if not jd_analysis:

    st.warning(
        """
        ⚠️ Structured JD analysis is not available.

        Complete **Chunk 3 — JD Intelligence** first.
        """
    )

else:

    st.success(
        f"""
        ✅ {len(jd_analysis)} structured JD document(s)
        available for skill intelligence.
        """
    )


# ============================================================
# SKILL TAXONOMY
# ============================================================

SKILL_TAXONOMY = {

    # ========================================================
    # AI / ML
    # ========================================================

    "Artificial Intelligence": [
        "artificial intelligence",
        "ai",
    ],

    "Machine Learning": [
        "machine learning",
        "ml",
    ],

    "Deep Learning": [
        "deep learning",
        "dl",
    ],

    "Generative AI": [
        "generative ai",
        "genai",
        "gen ai",
    ],

    "Large Language Models": [
        "large language model",
        "large language models",
        "llm",
        "llms",
    ],

    "Natural Language Processing": [
        "natural language processing",
        "nlp",
    ],

    "Computer Vision": [
        "computer vision",
        "cv",
    ],

    "Retrieval Augmented Generation": [
        "retrieval augmented generation",
        "retrieval augmented generation",
        "rag",
    ],

    "Agentic AI": [
        "agentic ai",
        "agentic",
        "ai agents",
        "intelligent agents",
    ],

    "Prompt Engineering": [
        "prompt engineering",
        "prompt design",
    ],

    "Fine Tuning": [
        "fine tuning",
        "fine-tuning",
        "model fine tuning",
    ],


    # ========================================================
    # DATA
    # ========================================================

    "Data Science": [
        "data science",
    ],

    "Data Analytics": [
        "data analytics",
        "data analysis",
    ],

    "Data Engineering": [
        "data engineering",
    ],

    "Data Visualization": [
        "data visualization",
        "data visualisation",
    ],

    "Statistics": [
        "statistics",
        "statistical analysis",
    ],


    # ========================================================
    # PROGRAMMING
    # ========================================================

    "Python": [
        "python",
    ],

    "Java": [
        "java",
    ],

    "JavaScript": [
        "javascript",
        "js",
    ],

    "TypeScript": [
        "typescript",
        "ts",
    ],

    "C": [
        " c ",
        "c programming",
    ],

    "C++": [
        "c++",
    ],

    "C#": [
        "c#",
        "c sharp",
    ],

    "Go": [
        "golang",
        "go programming",
    ],

    "R": [
        " r ",
        "r programming",
    ],

    "SQL": [
        "sql",
    ],


    # ========================================================
    # ML FRAMEWORKS
    # ========================================================

    "TensorFlow": [
        "tensorflow",
    ],

    "PyTorch": [
        "pytorch",
    ],

    "Keras": [
        "keras",
    ],

    "Scikit-learn": [
        "scikit-learn",
        "sklearn",
        "scikit learn",
    ],

    "Hugging Face": [
        "hugging face",
        "huggingface",
    ],

    "Transformers": [
        "transformers",
        "hugging face transformers",
    ],


    # ========================================================
    # LLM FRAMEWORKS
    # ========================================================

    "LangChain": [
        "langchain",
    ],

    "LangGraph": [
        "langgraph",
    ],

    "LlamaIndex": [
        "llamaindex",
        "llama index",
    ],

    "CrewAI": [
        "crewai",
        "crew ai",
    ],


    # ========================================================
    # WEB / API
    # ========================================================

    "FastAPI": [
        "fastapi",
        "fast api",
    ],

    "Flask": [
        "flask",
    ],

    "Django": [
        "django",
    ],

    "REST API": [
        "rest api",
        "restful api",
        "rest services",
    ],

    "GraphQL": [
        "graphql",
    ],


    # ========================================================
    # CLOUD
    # ========================================================

    "AWS": [
        "aws",
        "amazon web services",
    ],

    "Microsoft Azure": [
        "azure",
        "microsoft azure",
    ],

    "Google Cloud": [
        "gcp",
        "google cloud",
        "google cloud platform",
    ],


    # ========================================================
    # DEVOPS / MLOPS
    # ========================================================

    "Docker": [
        "docker",
        "containerization",
        "containerisation",
    ],

    "Kubernetes": [
        "kubernetes",
        "k8s",
    ],

    "Git": [
        "git",
    ],

    "GitHub": [
        "github",
    ],

    "Jenkins": [
        "jenkins",
    ],

    "CI/CD": [
        "ci/cd",
        "cicd",
        "continuous integration",
        "continuous deployment",
    ],

    "MLOps": [
        "mlops",
        "machine learning operations",
    ],

    "LLMOps": [
        "llmops",
        "llm operations",
    ],


    # ========================================================
    # DATABASES
    # ========================================================

    "PostgreSQL": [
        "postgresql",
        "postgres",
    ],

    "MySQL": [
        "mysql",
    ],

    "MongoDB": [
        "mongodb",
        "mongo db",
    ],

    "Oracle": [
        "oracle database",
        "oracle db",
    ],

    "Redis": [
        "redis",
    ],

    "Elasticsearch": [
        "elasticsearch",
        "elastic search",
    ],


    # ========================================================
    # VECTOR DATABASES
    # ========================================================

    "FAISS": [
        "faiss",
    ],

    "Chroma": [
        "chroma",
        "chromadb",
    ],

    "Pinecone": [
        "pinecone",
    ],

    "Weaviate": [
        "weaviate",
    ],


    # ========================================================
    # BI / ANALYTICS
    # ========================================================

    "Power BI": [
        "power bi",
        "powerbi",
    ],

    "Tableau": [
        "tableau",
    ],

    "Excel": [
        "excel",
        "microsoft excel",
    ],

    "Power Query": [
        "power query",
    ],

    "DAX": [
        "dax",
    ],


    # ========================================================
    # VERSION CONTROL
    # ========================================================

    "GitLab": [
        "gitlab",
    ],

    "Bitbucket": [
        "bitbucket",
    ],


    # ========================================================
    # PROJECT / ENGINEERING
    # ========================================================

    "Agile": [
        "agile",
        "agile methodology",
    ],

    "Scrum": [
        "scrum",
    ],

    "Software Development": [
        "software development",
        "software engineering",
    ],

    "System Design": [
        "system design",
        "software architecture",
    ],

}


# ============================================================
# CATEGORY MAPPING
# ============================================================

SKILL_CATEGORIES = {

    "programming_languages": [

        "Python",
        "Java",
        "JavaScript",
        "TypeScript",
        "C",
        "C++",
        "C#",
        "Go",
        "R",
        "SQL",

    ],

    "ai_ml": [

        "Artificial Intelligence",
        "Machine Learning",
        "Deep Learning",
        "Generative AI",
        "Large Language Models",
        "Natural Language Processing",
        "Computer Vision",
        "Retrieval Augmented Generation",
        "Agentic AI",
        "Prompt Engineering",
        "Fine Tuning",

    ],

    "data": [

        "Data Science",
        "Data Analytics",
        "Data Engineering",
        "Data Visualization",
        "Statistics",

    ],

    "frameworks": [

        "TensorFlow",
        "PyTorch",
        "Keras",
        "Scikit-learn",
        "Hugging Face",
        "Transformers",
        "LangChain",
        "LangGraph",
        "LlamaIndex",
        "CrewAI",
        "FastAPI",
        "Flask",
        "Django",

    ],

    "cloud": [

        "AWS",
        "Microsoft Azure",
        "Google Cloud",

    ],

    "devops": [

        "Docker",
        "Kubernetes",
        "Git",
        "GitHub",
        "Jenkins",
        "CI/CD",
        "MLOps",
        "LLMOps",

    ],

    "databases": [

        "PostgreSQL",
        "MySQL",
        "MongoDB",
        "Oracle",
        "Redis",
        "Elasticsearch",
        "FAISS",
        "Chroma",
        "Pinecone",
        "Weaviate",

    ],

    "business_intelligence": [

        "Power BI",
        "Tableau",
        "Excel",
        "Power Query",
        "DAX",

    ],

    "engineering": [

        "REST API",
        "GraphQL",
        "Agile",
        "Scrum",
        "Software Development",
        "System Design",

    ],

}


# ============================================================
# TAXONOMY NORMALIZATION
# ============================================================

def normalize_taxonomy_skill(
    skill,
):
    """
    Map a raw skill to a canonical taxonomy name.
    """

    raw = normalize_skill(
        skill
    )


    if not raw:

        return ""


    # --------------------------------------------------------
    # Exact / phrase matching
    # --------------------------------------------------------

    for canonical, aliases in (
        SKILL_TAXONOMY.items()
    ):

        canonical_normalized = (
            normalize_skill(
                canonical
            )
        )


        if raw == canonical_normalized:

            return canonical


        for alias in aliases:

            alias_normalized = (
                normalize_skill(
                    alias
                )
            )


            if not alias_normalized:

                continue


            if (
                raw == alias_normalized
                or
                alias_normalized in raw
                or
                raw in alias_normalized
            ):

                return canonical


    # --------------------------------------------------------
    # Unknown skill
    # --------------------------------------------------------

    return safe_text(
        skill
    )


# ============================================================
# CLASSIFY SKILL
# ============================================================

def classify_skill(
    skill,
):
    """
    Return the taxonomy category for a canonical skill.
    """

    canonical = normalize_taxonomy_skill(
        skill
    )


    if not canonical:

        return "other"


    for category, skills in (
        SKILL_CATEGORIES.items()
    ):

        if canonical in skills:

            return category


    return "other"


# ============================================================
# BUILD SKILL RECORD
# ============================================================

def build_skill_record(
    skill,
    source="JD",
    priority="required",
):
    """
    Create one normalized skill record.
    """

    canonical = normalize_taxonomy_skill(
        skill
    )


    return {

        "raw_skill":
            safe_text(
                skill
            ),

        "canonical_skill":
            canonical,

        "category":
            classify_skill(
                canonical
            ),

        "priority":
            priority,

        "source":
            source,

        "normalized":
            normalize_skill(
                canonical
            ),

    }


# ============================================================
# EXTRACT RAW SKILLS FROM ONE JD
# ============================================================

def extract_raw_jd_skills(
    analysis,
):
    """
    Collect all relevant skill fields from a structured JD.
    """

    if not isinstance(
        analysis,
        dict,
    ):

        return []


    records = []


    # ========================================================
    # REQUIRED SKILLS
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "required_skills"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # PREFERRED SKILLS
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "preferred_skills"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="preferred",
            )
        )


    # ========================================================
    # TECHNOLOGIES
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "technologies"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # TOOLS
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "tools"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # FRAMEWORKS
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "frameworks"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # LIBRARIES
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "libraries"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # LANGUAGES
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "programming_languages"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # CLOUD
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "cloud_platforms"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # DATABASES
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "databases"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="required",
            )
        )


    # ========================================================
    # DOMAIN KNOWLEDGE
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "domain_knowledge"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="domain",
            )
        )


    # ========================================================
    # SOFT SKILLS
    # ========================================================

    for skill in normalize_list(
        analysis.get(
            "soft_skills"
        )
    ):

        records.append(
            build_skill_record(
                skill,
                source="JD",
                priority="soft",
            )
        )


    return records


# ============================================================
# DEDUPLICATE SKILL RECORDS
# ============================================================

def deduplicate_skill_records(
    records,
):
    """
    Deduplicate skill records while preserving the strongest
    priority.
    """

    priority_rank = {

        "required": 1,

        "preferred": 2,

        "domain": 3,

        "soft": 4,

        "other": 5,

    }


    merged = {}


    for record in records:

        canonical = safe_text(
            record.get(
                "canonical_skill"
            )
        )


        if not canonical:

            continue


        key = normalize_skill(
            canonical
        )


        if key not in merged:

            merged[key] = record

            continue


        existing_priority = (
            merged[key].get(
                "priority",
                "other"
            )
        )


        current_priority = (
            record.get(
                "priority",
                "other"
            )
        )


        if (
            priority_rank.get(
                current_priority,
                99
            )
            <
            priority_rank.get(
                existing_priority,
                99
            )
        ):

            merged[key] = record


    return list(
        merged.values()
    )


# ============================================================
# BUILD COMPLETE JD SKILL INTELLIGENCE
# ============================================================

def build_jd_skill_intelligence(
    analyses,
):
    """
    Build complete normalized JD skill intelligence.
    """

    all_records = []


    for analysis in analyses:

        all_records.extend(
            extract_raw_jd_skills(
                analysis
            )
        )


    all_records = (
        deduplicate_skill_records(
            all_records
        )
    )


    # ========================================================
    # CATEGORY COLLECTION
    # ========================================================

    categorized = {

        "required_skills": [],

        "preferred_skills": [],

        "technical_skills": [],

        "tools": [],

        "technologies": [],

        "frameworks": [],

        "languages": [],

        "cloud": [],

        "databases": [],

        "domain_skills": [],

        "soft_skills": [],

        "project_skills": [],

        "keywords": [],

    }


    # ========================================================
    # PROCESS RECORDS
    # ========================================================

    for record in all_records:

        skill = safe_text(
            record.get(
                "canonical_skill"
            )
        )


        category = safe_text(
            record.get(
                "category"
            ),
            "other",
        )


        priority = safe_text(
            record.get(
                "priority"
            ),
            "required",
        )


        if not skill:

            continue


        # ----------------------------------------------------
        # Required / preferred
        # ----------------------------------------------------

        if priority == "required":

            categorized[
                "required_skills"
            ].append(
                skill
            )

        elif priority == "preferred":

            categorized[
                "preferred_skills"
            ].append(
                skill
            )


        # ----------------------------------------------------
        # Category
        # ----------------------------------------------------

        if category in [
            "ai_ml",
            "data",
            "programming_languages",
            "frameworks",
            "cloud",
            "devops",
            "databases",
            "business_intelligence",
            "engineering",
        ]:

            categorized[
                "technical_skills"
            ].append(
                skill
            )


        if category == "programming_languages":

            categorized[
                "languages"
            ].append(
                skill
            )


        elif category == "frameworks":

            categorized[
                "frameworks"
            ].append(
                skill
            )


        elif category == "cloud":

            categorized[
                "cloud"
            ].append(
                skill
            )


        elif category == "databases":

            categorized[
                "databases"
            ].append(
                skill
            )


        elif category == "devops":

            categorized[
                "tools"
            ].append(
                skill
            )


        elif category == "ai_ml":

            categorized[
                "technologies"
            ].append(
                skill
            )


        elif category == "data":

            categorized[
                "technologies"
            ].append(
                skill
            )


        elif category == "business_intelligence":

            categorized[
                "technologies"
            ].append(
                skill
            )


        elif category == "engineering":

            categorized[
                "technologies"
            ].append(
                skill
            )


        # ----------------------------------------------------
        # Domain
        # ----------------------------------------------------

        if priority == "domain":

            categorized[
                "domain_skills"
            ].append(
                skill
            )


        # ----------------------------------------------------
        # Soft skill
        # ----------------------------------------------------

        if priority == "soft":

            categorized[
                "soft_skills"
            ].append(
                skill
            )


    # ========================================================
    # PROJECT EXPECTATIONS
    # ========================================================

    project_terms = []


    for analysis in analyses:

        project_terms.extend(

            normalize_list(
                analysis.get(
                    "project_expectations"
                )
            )

        )


    categorized[
        "project_skills"
    ] = unique_values(
        project_terms
    )


    # ========================================================
    # KEYWORDS
    # ========================================================

    for analysis in analyses:

        categorized[
            "keywords"
        ].extend(

            normalize_list(
                analysis.get(
                    "keywords"
                )
            )

        )


    # ========================================================
    # CLEAN ALL CATEGORIES
    # ========================================================

    for key in categorized:

        categorized[key] = unique_values(
            categorized[key]
        )


    # ========================================================
    # ALL SKILLS
    # ========================================================

    all_skills = unique_values(

        categorized[
            "required_skills"
        ]

        +

        categorized[
            "preferred_skills"
        ]

        +

        categorized[
            "technical_skills"
        ]

        +

        categorized[
            "domain_skills"
        ]

        +

        categorized[
            "soft_skills"
        ]

    )


    # ========================================================
    # RETURN FINAL OBJECT
    # ========================================================

    return {

        "all_skills":
            all_skills,

        "required_skills":
            categorized[
                "required_skills"
            ],

        "preferred_skills":
            categorized[
                "preferred_skills"
            ],

        "technical_skills":
            categorized[
                "technical_skills"
            ],

        "tools":
            categorized[
                "tools"
            ],

        "technologies":
            categorized[
                "technologies"
            ],

        "frameworks":
            categorized[
                "frameworks"
            ],

        "languages":
            categorized[
                "languages"
            ],

        "cloud":
            categorized[
                "cloud"
            ],

        "databases":
            categorized[
                "databases"
            ],

        "domain_skills":
            categorized[
                "domain_skills"
            ],

        "soft_skills":
            categorized[
                "soft_skills"
            ],

        "project_skills":
            categorized[
                "project_skills"
            ],

        "keywords":
            categorized[
                "keywords"
            ],

        "skill_records":
            all_records,

        "generated_at":
            datetime.now().isoformat(),

    }


# ============================================================
# GENERATE SKILL INTELLIGENCE
# ============================================================

if jd_analysis:

    generate_skill_intelligence = st.button(

        "🎯 Generate JD Skill Intelligence",

        type="primary",

        use_container_width=True,

        key="generate_jd_skill_intelligence",

    )


    if generate_skill_intelligence:

        with st.spinner(
            "Normalizing and classifying JD skills..."
        ):

            jd_skill_intelligence = (
                build_jd_skill_intelligence(
                    jd_analysis
                )
            )


        st.session_state[
            "jd_skill_intelligence"
        ] = jd_skill_intelligence


        # ----------------------------------------------------
        # Backward-compatible session state
        # ----------------------------------------------------

        st.session_state[
            "jd_skills"
        ] = jd_skill_intelligence[
            "all_skills"
        ]


        st.session_state[
            "jd_tools"
        ] = jd_skill_intelligence[
            "tools"
        ]


        st.session_state[
            "jd_keywords"
        ] = jd_skill_intelligence[
            "keywords"
        ]


        st.success(
            """
            ✅ JD Skill Intelligence generated successfully.
            """
        )


# ============================================================
# LOAD SKILL INTELLIGENCE
# ============================================================

jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence"
)


# ============================================================
# DISPLAY SKILL INTELLIGENCE
# ============================================================

if jd_skill_intelligence:

    st.divider()

    st.subheader(
        "📊 Industry Skill Profile"
    )


    # ========================================================
    # TOP METRICS
    # ========================================================

    col1, col2, col3, col4 = (
        st.columns(4)
    )


    with col1:

        st.metric(
            "Total Skills",
            len(
                jd_skill_intelligence.get(
                    "all_skills",
                    []
                )
            ),
        )


    with col2:

        st.metric(
            "Required",
            len(
                jd_skill_intelligence.get(
                    "required_skills",
                    []
                )
            ),
        )


    with col3:

        st.metric(
            "Preferred",
            len(
                jd_skill_intelligence.get(
                    "preferred_skills",
                    []
                )
            ),
        )


    with col4:

        st.metric(
            "Technical",
            len(
                jd_skill_intelligence.get(
                    "technical_skills",
                    []
                )
            ),
        )


    # ========================================================
    # CATEGORY TABS
    # ========================================================

    tabs = st.tabs([

        "🎯 Required",

        "⭐ Preferred",

        "🤖 AI / ML",

        "💻 Programming",

        "🧰 Tools & Frameworks",

        "☁️ Cloud",

        "🗄️ Databases",

        "🏭 Domain",

        "🚀 Projects",

        "🔑 Keywords",

    ])


    # ========================================================
    # REQUIRED
    # ========================================================

    with tabs[0]:

        values = jd_skill_intelligence.get(
            "required_skills",
            []
        )


        if values:

            for skill in values:

                st.success(
                    skill
                )

        else:

            st.info(
                "No required skills identified."
            )


    # ========================================================
    # PREFERRED
    # ========================================================

    with tabs[1]:

        values = jd_skill_intelligence.get(
            "preferred_skills",
            []
        )


        if values:

            for skill in values:

                st.write(
                    f"⭐ {skill}"
                )

        else:

            st.info(
                "No preferred skills identified."
            )


    # ========================================================
    # AI / ML
    # ========================================================

    with tabs[2]:

        ai_skills = [

            skill

            for skill in jd_skill_intelligence.get(
                "technical_skills",
                []
            )

            if classify_skill(
                skill
            ) == "ai_ml"

        ]


        if ai_skills:

            for skill in unique_values(
                ai_skills
            ):

                st.write(
                    f"🤖 {skill}"
                )

        else:

            st.info(
                "No AI/ML-specific skills identified."
            )


    # ========================================================
    # PROGRAMMING
    # ========================================================

    with tabs[3]:

        values = jd_skill_intelligence.get(
            "languages",
            []
        )


        if values:

            for skill in values:

                st.code(
                    skill
                )

        else:

            st.info(
                "No programming languages identified."
            )


    # ========================================================
    # TOOLS / FRAMEWORKS
    # ========================================================

    with tabs[4]:

        tools = jd_skill_intelligence.get(
            "tools",
            []
        )


        frameworks = jd_skill_intelligence.get(
            "frameworks",
            []
        )


        if tools:

            st.markdown(
                "**Tools**"
            )

            st.write(
                ", ".join(
                    tools
                )
            )


        if frameworks:

            st.markdown(
                "**Frameworks**"
            )

            st.write(
                ", ".join(
                    frameworks
                )
            )


        if not tools and not frameworks:

            st.info(
                "No tools or frameworks identified."
            )


    # ========================================================
    # CLOUD
    # ========================================================

    with tabs[5]:

        values = jd_skill_intelligence.get(
            "cloud",
            []
        )


        if values:

            for skill in values:

                st.write(
                    f"☁️ {skill}"
                )

        else:

            st.info(
                "No cloud platforms identified."
            )


    # ========================================================
    # DATABASES
    # ========================================================

    with tabs[6]:

        values = jd_skill_intelligence.get(
            "databases",
            []
        )


        if values:

            for skill in values:

                st.write(
                    f"🗄️ {skill}"
                )

        else:

            st.info(
                "No databases identified."
            )


    # ========================================================
    # DOMAIN
    # ========================================================

    with tabs[7]:

        values = jd_skill_intelligence.get(
            "domain_skills",
            []
        )


        if values:

            for skill in values:

                st.write(
                    f"🏭 {skill}"
                )

        else:

            st.info(
                "No domain-specific skills identified."
            )


    # ========================================================
    # PROJECTS
    # ========================================================

    with tabs[8]:

        values = jd_skill_intelligence.get(
            "project_skills",
            []
        )


        if values:

            for project in values:

                st.write(
                    f"🚀 {project}"
                )

        else:

            st.info(
                "No explicit project expectations identified."
            )


    # ========================================================
    # KEYWORDS
    # ========================================================

    with tabs[9]:

        values = jd_skill_intelligence.get(
            "keywords",
            []
        )


        if values:

            st.write(
                ", ".join(
                    values
                )
            )

        else:

            st.info(
                "No explicit keywords identified."
            )


# ============================================================
# SKILL RECORD TABLE
# ============================================================

if jd_skill_intelligence:

    with st.expander(
        "🔍 View Normalized Skill Records"
    ):

        records = jd_skill_intelligence.get(
            "skill_records",
            []
        )


        if records:

            records_df = pd.DataFrame(
                records
            )


            display_columns = [

                "raw_skill",

                "canonical_skill",

                "category",

                "priority",

                "source",

            ]


            available_columns = [

                column

                for column in display_columns

                if column in records_df.columns

            ]


            st.dataframe(

                records_df[
                    available_columns
                ],

                use_container_width=True,

                hide_index=True,

            )

        else:

            st.info(
                "No skill records available."
            )


# ============================================================
# SAVE COMPLETION STATE
# ============================================================

st.session_state[
    "jd_skill_intelligence_complete"
] = bool(
    jd_skill_intelligence
)


# ============================================================
# END OF CHUNK 4
# ============================================================
# ============================================================
# CHUNK 5/10
# CURRICULUM SKILL INTELLIGENCE
# ============================================================

"""
Purpose
-------
Load the curriculum generated by Page 02 and convert it into
a normalized skill intelligence structure that can be compared
with Job Description skills in Chunk 6.

Output
------
st.session_state["curriculum_skill_intelligence"]

Main fields
-----------
all_skills
technical_skills
concepts
tools
technologies
frameworks
languages
cloud
databases
projects
topics
modules
skill_records
"""


# ============================================================
# 1. LOAD CURRICULUM FROM PAGE 02
# ============================================================

primary_curriculum = st.session_state.get(
    "primary_syllabus"
)

curriculum_package = st.session_state.get(
    "curriculum_master_package"
)


# ------------------------------------------------------------
# Fallback: curriculum may be inside master package
# ------------------------------------------------------------

if not primary_curriculum and isinstance(
    curriculum_package,
    dict,
):

    primary_curriculum = (
        curriculum_package.get(
            "primary_syllabus"
        )
        or
        curriculum_package.get(
            "curriculum"
        )
    )


# ============================================================
# 2. CURRICULUM VALIDATION
# ============================================================

if not primary_curriculum:

    st.warning(
        """
        ⚠️ Curriculum not available.

        Please complete Page 02:
        **📚 Curriculum Intelligence**
        """
    )

else:

    st.success(
        "✅ Curriculum loaded successfully."
    )


# ============================================================
# 3. GENERIC CURRICULUM VALUE EXTRACTOR
# ============================================================

def curriculum_value(
    value,
):
    """
    Convert a curriculum item into a clean string.
    """

    if value is None:
        return ""

    if isinstance(
        value,
        str,
    ):
        return value.strip()

    if isinstance(
        value,
        dict,
    ):

        for key in (
            "name",
            "title",
            "topic",
            "concept",
            "skill",
            "technology",
            "tool",
            "framework",
            "subject",
        ):

            if value.get(key):

                return safe_text(
                    value.get(key)
                )

    return safe_text(
        value
    )


# ============================================================
# 4. LIST NORMALIZATION
# ============================================================

def curriculum_list(
    value,
):
    """
    Convert curriculum field into a clean unique list.
    """

    if value is None:
        return []

    if not isinstance(
        value,
        list,
    ):

        value = [value]


    values = []

    for item in value:

        text = curriculum_value(
            item
        )

        if text:
            values.append(
                text
            )


    return unique_values(
        values
    )


# ============================================================
# 5. EXTRACT FIELD
# ============================================================

def extract_field(
    curriculum,
    field_names,
):
    """
    Extract one or multiple possible fields from curriculum.
    """

    if not isinstance(
        curriculum,
        dict,
    ):
        return []


    values = []


    for field in field_names:

        if field not in curriculum:
            continue


        values.extend(
            curriculum_list(
                curriculum.get(field)
            )
        )


    return unique_values(
        values
    )


# ============================================================
# 6. EXTRACT MODULES
# ============================================================

def extract_modules(
    curriculum,
):
    """
    Extract module information.
    """

    if not isinstance(
        curriculum,
        dict,
    ):
        return []


    modules = curriculum.get(
        "modules",
        []
    )


    if not isinstance(
        modules,
        list,
    ):
        return []


    results = []


    for index, module in enumerate(
        modules,
        start=1,
    ):

        if isinstance(
            module,
            dict,
        ):

            name = (
                module.get("name")
                or
                module.get("title")
                or
                module.get("module")
                or
                f"Module {index}"
            )


            description = safe_text(
                module.get(
                    "description"
                )
            )


            topics = curriculum_list(
                module.get(
                    "topics",
                    []
                )
            )


            results.append({

                "module_id":
                    index,

                "module_name":
                    safe_text(
                        name
                    ),

                "description":
                    description,

                "topics":
                    topics,

            })


        else:

            results.append({

                "module_id":
                    index,

                "module_name":
                    curriculum_value(
                        module
                    ),

                "description":
                    "",

                "topics":
                    [],

            })


    return results


# ============================================================
# 7. EXTRACT TOPICS
# ============================================================

def extract_topics(
    modules,
):
    """
    Create module-topic records.
    """

    records = []


    for module in modules:

        module_name = module[
            "module_name"
        ]


        for topic in module[
            "topics"
        ]:

            records.append({

                "module":
                    module_name,

                "topic":
                    topic,

            })


    return records


# ============================================================
# 8. EXTRACT MODULE-LEVEL FIELDS
# ============================================================

def extract_module_fields(
    curriculum,
    fields,
):
    """
    Extract skills/tools/etc. from individual modules.
    """

    if not isinstance(
        curriculum,
        dict,
    ):
        return []


    modules = curriculum.get(
        "modules",
        []
    )


    if not isinstance(
        modules,
        list,
    ):
        return []


    values = []


    for module in modules:

        if not isinstance(
            module,
            dict,
        ):
            continue


        for field in fields:

            value = module.get(
                field
            )


            if value is None:
                continue


            values.extend(
                curriculum_list(
                    value
                )
            )


    return unique_values(
        values
    )


# ============================================================
# 9. EXTRACT CURRICULUM DATA
# ============================================================

def build_curriculum_profile(
    curriculum,
):
    """
    Build normalized curriculum profile.
    """

    modules = extract_modules(
        curriculum
    )


    topic_records = extract_topics(
        modules
    )


    topics = unique_values([

        item["topic"]

        for item in topic_records

    ])


    # --------------------------------------------------------
    # Concepts
    # --------------------------------------------------------

    concepts = extract_field(

        curriculum,

        [
            "concepts",
            "key_concepts",
            "knowledge_areas",
        ]

    )


    concepts.extend(

        extract_module_fields(

            curriculum,

            [
                "concepts",
                "key_concepts",
                "knowledge_areas",
            ],

        )

    )


    concepts = unique_values(
        concepts
    )


    # --------------------------------------------------------
    # Skills
    # --------------------------------------------------------

    explicit_skills = extract_field(

        curriculum,

        [
            "skills",
            "skill",
            "technical_skills",
            "competencies",
        ]

    )


    explicit_skills.extend(

        extract_module_fields(

            curriculum,

            [
                "skills",
                "technical_skills",
                "competencies",
            ],

        )

    )


    explicit_skills = unique_values(
        explicit_skills
    )


    # --------------------------------------------------------
    # Tools
    # --------------------------------------------------------

    tools = extract_field(

        curriculum,

        [
            "tools",
            "software",
            "platforms",
        ]

    )


    tools.extend(

        extract_module_fields(

            curriculum,

            [
                "tools",
                "software",
                "platforms",
            ],

        )

    )


    tools = unique_values(
        tools
    )


    # --------------------------------------------------------
    # Technologies
    # --------------------------------------------------------

    technologies = extract_field(

        curriculum,

        [
            "technologies",
            "technology",
        ]

    )


    technologies.extend(

        extract_module_fields(

            curriculum,

            [
                "technologies",
                "technology",
            ],

        )

    )


    technologies = unique_values(
        technologies
    )


    # --------------------------------------------------------
    # Frameworks
    # --------------------------------------------------------

    frameworks = extract_field(

        curriculum,

        [
            "frameworks",
            "libraries",
        ]

    )


    frameworks.extend(

        extract_module_fields(

            curriculum,

            [
                "frameworks",
                "libraries",
            ],

        )

    )


    frameworks = unique_values(
        frameworks
    )


    # --------------------------------------------------------
    # Programming Languages
    # --------------------------------------------------------

    languages = extract_field(

        curriculum,

        [
            "programming_languages",
            "languages",
            "coding_languages",
        ]

    )


    languages.extend(

        extract_module_fields(

            curriculum,

            [
                "programming_languages",
                "languages",
                "coding_languages",
            ],

        )

    )


    languages = unique_values(
        languages
    )


    # --------------------------------------------------------
    # Cloud
    # --------------------------------------------------------

    cloud = extract_field(

        curriculum,

        [
            "cloud",
            "cloud_platforms",
            "cloud_services",
        ]

    )


    cloud.extend(

        extract_module_fields(

            curriculum,

            [
                "cloud",
                "cloud_platforms",
                "cloud_services",
            ],

        )

    )


    cloud = unique_values(
        cloud
    )


    # --------------------------------------------------------
    # Databases
    # --------------------------------------------------------

    databases = extract_field(

        curriculum,

        [
            "databases",
            "database",
            "vector_databases",
        ]

    )


    databases.extend(

        extract_module_fields(

            curriculum,

            [
                "databases",
                "database",
                "vector_databases",
            ],

        )

    )


    databases = unique_values(
        databases
    )


    # --------------------------------------------------------
    # Projects
    # --------------------------------------------------------

    projects = extract_field(

        curriculum,

        [
            "projects",
            "project_areas",
            "project_work",
            "capstone_projects",
        ]

    )


    projects.extend(

        extract_module_fields(

            curriculum,

            [
                "projects",
                "project_work",
                "project_areas",
            ],

        )

    )


    projects = unique_values(
        projects
    )


    return {

        "modules":
            modules,

        "topics":
            topics,

        "topic_records":
            topic_records,

        "concepts":
            concepts,

        "skills":
            explicit_skills,

        "tools":
            tools,

        "technologies":
            technologies,

        "frameworks":
            frameworks,

        "languages":
            languages,

        "cloud":
            cloud,

        "databases":
            databases,

        "projects":
            projects,

    }


# ============================================================
# 10. BUILD SKILL RECORDS
# ============================================================

def create_curriculum_skill_records(
    profile,
):
    """
    Convert curriculum elements into normalized records.
    """

    records = []


    # --------------------------------------------------------
    # Explicit skills
    # --------------------------------------------------------

    for skill in profile[
        "skills"
    ]:

        canonical = normalize_taxonomy_skill(
            skill
        )


        records.append({

            "raw_skill":
                skill,

            "canonical_skill":
                canonical,

            "category":
                classify_skill(
                    canonical
                ),

            "source":
                "curriculum_skill",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Concepts
    # --------------------------------------------------------

    for concept in profile[
        "concepts"
    ]:

        canonical = normalize_taxonomy_skill(
            concept
        )


        records.append({

            "raw_skill":
                concept,

            "canonical_skill":
                canonical,

            "category":
                classify_skill(
                    canonical
                ),

            "source":
                "curriculum_concept",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Tools
    # --------------------------------------------------------

    for tool in profile[
        "tools"
    ]:

        canonical = normalize_taxonomy_skill(
            tool
        )


        records.append({

            "raw_skill":
                tool,

            "canonical_skill":
                canonical,

            "category":
                "tools",

            "source":
                "curriculum_tool",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Technologies
    # --------------------------------------------------------

    for technology in profile[
        "technologies"
    ]:

        canonical = normalize_taxonomy_skill(
            technology
        )


        records.append({

            "raw_skill":
                technology,

            "canonical_skill":
                canonical,

            "category":
                classify_skill(
                    canonical
                ),

            "source":
                "curriculum_technology",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Frameworks
    # --------------------------------------------------------

    for framework in profile[
        "frameworks"
    ]:

        canonical = normalize_taxonomy_skill(
            framework
        )


        records.append({

            "raw_skill":
                framework,

            "canonical_skill":
                canonical,

            "category":
                "frameworks",

            "source":
                "curriculum_framework",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Programming Languages
    # --------------------------------------------------------

    for language in profile[
        "languages"
    ]:

        canonical = normalize_taxonomy_skill(
            language
        )


        records.append({

            "raw_skill":
                language,

            "canonical_skill":
                canonical,

            "category":
                "programming_languages",

            "source":
                "curriculum_language",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Cloud
    # --------------------------------------------------------

    for cloud in profile[
        "cloud"
    ]:

        canonical = normalize_taxonomy_skill(
            cloud
        )


        records.append({

            "raw_skill":
                cloud,

            "canonical_skill":
                canonical,

            "category":
                "cloud",

            "source":
                "curriculum_cloud",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Databases
    # --------------------------------------------------------

    for database in profile[
        "databases"
    ]:

        canonical = normalize_taxonomy_skill(
            database
        )


        records.append({

            "raw_skill":
                database,

            "canonical_skill":
                canonical,

            "category":
                "databases",

            "source":
                "curriculum_database",

            "module":
                "",

            "topic":
                "",

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # --------------------------------------------------------
    # Topic-level records
    # --------------------------------------------------------

    for item in profile[
        "topic_records"
    ]:

        topic = safe_text(
            item.get(
                "topic"
            )
        )


        module = safe_text(
            item.get(
                "module"
            )
        )


        if not topic:
            continue


        canonical = normalize_taxonomy_skill(
            topic
        )


        records.append({

            "raw_skill":
                topic,

            "canonical_skill":
                canonical,

            "category":
                classify_skill(
                    canonical
                ),

            "source":
                "curriculum_topic",

            "module":
                module,

            "topic":
                topic,

            "normalized":
                normalize_skill(
                    canonical
                ),

        })


    # ========================================================
    # DEDUPLICATE
    # ========================================================

    unique_records = {}


    for record in records:

        key = normalize_skill(
            record.get(
                "canonical_skill"
            )
        )


        if not key:
            continue


        if key not in unique_records:

            unique_records[
                key
            ] = record

            continue


        existing = unique_records[
            key
        ]


        # Preserve module context
        if (
            not existing.get("module")
            and
            record.get("module")
        ):

            existing[
                "module"
            ] = record[
                "module"
            ]


        # Preserve topic context
        if (
            not existing.get("topic")
            and
            record.get("topic")
        ):

            existing[
                "topic"
            ] = record[
                "topic"
            ]


    return list(
        unique_records.values()
    )


# ============================================================
# 11. BUILD COMPLETE CURRICULUM INTELLIGENCE
# ============================================================

def build_curriculum_skill_intelligence(
    curriculum,
):
    """
    Create final normalized curriculum intelligence object.
    """

    profile = build_curriculum_profile(
        curriculum
    )


    records = create_curriculum_skill_records(
        profile
    )


    all_skills = unique_values([

        record[
            "canonical_skill"
        ]

        for record in records

    ])


    technical_categories = {

        "ai_ml",

        "data",

        "programming_languages",

        "frameworks",

        "cloud",

        "devops",

        "databases",

        "business_intelligence",

        "engineering",

        "tools",

    }


    technical_skills = unique_values([

        record[
            "canonical_skill"
        ]

        for record in records

        if record.get(
            "category"
        ) in technical_categories

    ])


    return {

        "all_skills":
            all_skills,

        "technical_skills":
            technical_skills,

        "concepts":
            profile[
                "concepts"
            ],

        "skills":
            profile[
                "skills"
            ],

        "tools":
            profile[
                "tools"
            ],

        "technologies":
            profile[
                "technologies"
            ],

        "frameworks":
            profile[
                "frameworks"
            ],

        "languages":
            profile[
                "languages"
            ],

        "cloud":
            profile[
                "cloud"
            ],

        "databases":
            profile[
                "databases"
            ],

        "projects":
            profile[
                "projects"
            ],

        "topics":
            profile[
                "topics"
            ],

        "topic_records":
            profile[
                "topic_records"
            ],

        "modules":
            profile[
                "modules"
            ],

        "skill_records":
            records,

        "generated_at":
            datetime.now().isoformat(),

    }


# ============================================================
# 12. GENERATE BUTTON
# ============================================================

if primary_curriculum:

    if st.button(

        "📚 Build Curriculum Skill Intelligence",

        type="primary",

        use_container_width=True,

        key="build_curriculum_skill_intelligence",

    ):

        with st.spinner(
            "Analysing curriculum skills, topics and technologies..."
        ):

            curriculum_skill_intelligence = (
                build_curriculum_skill_intelligence(
                    primary_curriculum
                )
            )


        st.session_state[
            "curriculum_skill_intelligence"
        ] = curriculum_skill_intelligence


        st.success(
            "✅ Curriculum Skill Intelligence generated."
        )


# ============================================================
# 13. LOAD RESULTS
# ============================================================

curriculum_skill_intelligence = (
    st.session_state.get(
        "curriculum_skill_intelligence"
    )
)


# ============================================================
# 14. DISPLAY RESULTS
# ============================================================

if curriculum_skill_intelligence:

    st.divider()

    st.subheader(
        "📊 Curriculum Skill Profile"
    )


    # ========================================================
    # METRICS
    # ========================================================

    col1, col2, col3, col4, col5 = (
        st.columns(5)
    )


    with col1:

        st.metric(
            "Modules",
            len(
                curriculum_skill_intelligence.get(
                    "modules",
                    []
                )
            ),
        )


    with col2:

        st.metric(
            "Topics",
            len(
                curriculum_skill_intelligence.get(
                    "topics",
                    []
                )
            ),
        )


    with col3:

        st.metric(
            "Skills",
            len(
                curriculum_skill_intelligence.get(
                    "all_skills",
                    []
                )
            ),
        )


    with col4:

        st.metric(
            "Tools",
            len(
                curriculum_skill_intelligence.get(
                    "tools",
                    []
                )
            ),
        )


    with col5:

        st.metric(
            "Projects",
            len(
                curriculum_skill_intelligence.get(
                    "projects",
                    []
                )
            ),
        )


    # ========================================================
    # TABS
    # ========================================================

    tabs = st.tabs([

        "📚 Topics",

        "🧠 Concepts",

        "🎯 Skills",

        "🛠 Tools",

        "⚙️ Technologies",

        "🧩 Frameworks",

        "💻 Languages",

        "☁️ Cloud",

        "🗄️ Databases",

        "🚀 Projects",

    ])


    # ========================================================
    # TOPICS
    # ========================================================

    with tabs[0]:

        values = curriculum_skill_intelligence.get(
            "topics",
            []
        )


        if values:

            st.dataframe(

                pd.DataFrame({
                    "Topic": values
                }),

                use_container_width=True,

                hide_index=True,

            )

        else:

            st.info(
                "No topics found."
            )


    # ========================================================
    # CONCEPTS
    # ========================================================

    with tabs[1]:

        values = curriculum_skill_intelligence.get(
            "concepts",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"🧠 {value}"
                )

        else:

            st.info(
                "No concepts found."
            )


    # ========================================================
    # SKILLS
    # ========================================================

    with tabs[2]:

        values = curriculum_skill_intelligence.get(
            "all_skills",
            []
        )


        if values:

            for value in values:

                st.success(
                    value
                )

        else:

            st.info(
                "No skills found."
            )


    # ========================================================
    # TOOLS
    # ========================================================

    with tabs[3]:

        values = curriculum_skill_intelligence.get(
            "tools",
            []
        )


        if values:

            st.write(
                ", ".join(
                    values
                )
            )

        else:

            st.info(
                "No tools found."
            )


    # ========================================================
    # TECHNOLOGIES
    # ========================================================

    with tabs[4]:

        values = curriculum_skill_intelligence.get(
            "technologies",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"⚙️ {value}"
                )

        else:

            st.info(
                "No technologies found."
            )


    # ========================================================
    # FRAMEWORKS
    # ========================================================

    with tabs[5]:

        values = curriculum_skill_intelligence.get(
            "frameworks",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"🧩 {value}"
                )

        else:

            st.info(
                "No frameworks found."
            )


    # ========================================================
    # LANGUAGES
    # ========================================================

    with tabs[6]:

        values = curriculum_skill_intelligence.get(
            "languages",
            []
        )


        if values:

            for value in values:

                st.code(
                    value
                )

        else:

            st.info(
                "No languages found."
            )


    # ========================================================
    # CLOUD
    # ========================================================

    with tabs[7]:

        values = curriculum_skill_intelligence.get(
            "cloud",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"☁️ {value}"
                )

        else:

            st.info(
                "No cloud platforms found."
            )


    # ========================================================
    # DATABASES
    # ========================================================

    with tabs[8]:

        values = curriculum_skill_intelligence.get(
            "databases",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"🗄️ {value}"
                )

        else:

            st.info(
                "No databases found."
            )


    # ========================================================
    # PROJECTS
    # ========================================================

    with tabs[9]:

        values = curriculum_skill_intelligence.get(
            "projects",
            []
        )


        if values:

            for value in values:

                st.write(
                    f"🚀 {value}"
                )

        else:

            st.info(
                "No projects found."
            )


# ============================================================
# 15. MODULE → TOPIC VIEW
# ============================================================

if curriculum_skill_intelligence:

    with st.expander(
        "📖 Module → Topic Mapping",
        expanded=False,
    ):

        topic_records = (
            curriculum_skill_intelligence.get(
                "topic_records",
                []
            )
        )


        if topic_records:

            st.dataframe(

                pd.DataFrame(
                    topic_records
                ),

                use_container_width=True,

                hide_index=True,

            )

        else:

            st.info(
                "No module-topic mapping available."
            )


# ============================================================
# 16. NORMALIZED SKILL RECORDS
# ============================================================

if curriculum_skill_intelligence:

    with st.expander(
        "🔍 Normalized Curriculum Skill Records",
        expanded=False,
    ):

        records = (
            curriculum_skill_intelligence.get(
                "skill_records",
                []
            )
        )


        if records:

            st.dataframe(

                pd.DataFrame(
                    records
                ),

                use_container_width=True,

                hide_index=True,

            )

        else:

            st.info(
                "No skill records available."
            )


# ============================================================
# 17. DOWNLOAD JSON
# ============================================================

if curriculum_skill_intelligence:

    st.download_button(

        "⬇️ Download Curriculum Skill Intelligence",

        data=serialize_json(
            curriculum_skill_intelligence
        ),

        file_name=(
            "curriculum_skill_intelligence.json"
        ),

        mime="application/json",

        key=(
            "download_curriculum_skill_intelligence"
        ),

    )


# ============================================================
# 18. COMPLETION FLAG
# ============================================================

st.session_state[
    "curriculum_skill_intelligence_complete"
] = bool(
    curriculum_skill_intelligence
)


# ============================================================
# END OF CHUNK 5
# ============================================================
# ============================================================
# CHUNK 6/10
# JD ↔ CURRICULUM SKILL MATCHER
# ============================================================

"""
Purpose
-------
Compare Industry / JD skills against Curriculum skills.

Matching levels
---------------

1. EXACT
2. ALIAS
3. TOKEN / TEXT SIMILARITY
4. SEMANTIC SIMILARITY (optional)
5. PARTIAL
6. MISSING

Output
------

st.session_state["skill_match_results"]

st.session_state["skill_match_summary"]

st.session_state["covered_skills"]

st.session_state["partial_skills"]

st.session_state["missing_skills"]

Coverage metrics
----------------

Required Skill Coverage
Preferred Skill Coverage
Overall Skill Coverage
Technical Skill Coverage
"""


# ============================================================
# 1. SECTION HEADER
# ============================================================

st.divider()

st.header(
    "5️⃣ 🔗 JD ↔ Curriculum Skill Matching"
)

st.markdown(
    """
Compare the skills demanded by industry with the skills
covered by the current curriculum.

Each JD requirement is classified as:

🟢 **Covered**  
The curriculum clearly contains the required skill.

🟡 **Partial**  
A related concept exists, but the exact industry skill,
tool or depth may be missing.

🔴 **Missing**  
The curriculum does not currently contain sufficient
evidence of the required skill.
"""
)


# ============================================================
# 2. LOAD INPUT DATA
# ============================================================

jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence"
)

curriculum_skill_intelligence = (
    st.session_state.get(
        "curriculum_skill_intelligence"
    )
)


# ============================================================
# 3. VALIDATION
# ============================================================

jd_available = bool(
    jd_skill_intelligence
    and
    jd_skill_intelligence.get(
        "skill_records"
    )
)


curriculum_available = bool(
    curriculum_skill_intelligence
    and
    curriculum_skill_intelligence.get(
        "skill_records"
    )
)


if not jd_available:

    st.warning(
        """
        ⚠️ JD Skill Intelligence is not available.

        Complete Chunk 4 first.
        """
    )


if not curriculum_available:

    st.warning(
        """
        ⚠️ Curriculum Skill Intelligence is not available.

        Complete Chunk 5 first.
        """
    )


# ============================================================
# 4. MATCHING CONFIGURATION
# ============================================================

EXACT_THRESHOLD = 1.0

ALIAS_THRESHOLD = 0.90

TOKEN_THRESHOLD = 0.72

PARTIAL_THRESHOLD = 0.45

SEMANTIC_THRESHOLD = 0.78


# ============================================================
# 5. COMMON ALIAS MAP
# ============================================================

SKILL_ALIASES = {

    "python": {
        "python programming",
        "python language",
        "python development",
    },

    "machine learning": {
        "machine learning",
        "ml",
        "machine learning algorithms",
    },

    "deep learning": {
        "deep learning",
        "dl",
        "deep neural networks",
    },

    "artificial intelligence": {
        "artificial intelligence",
        "ai",
    },

    "generative ai": {
        "generative ai",
        "genai",
        "gen ai",
    },

    "large language models": {
        "large language models",
        "large language model",
        "llm",
        "llms",
    },

    "retrieval augmented generation": {
        "retrieval augmented generation",
        "retrieval augmented generation",
        "rag",
    },

    "natural language processing": {
        "natural language processing",
        "nlp",
    },

    "computer vision": {
        "computer vision",
        "cv",
    },

    "scikit-learn": {
        "scikit-learn",
        "scikit learn",
        "sklearn",
    },

    "hugging face": {
        "hugging face",
        "huggingface",
    },

    "langchain": {
        "langchain",
    },

    "langgraph": {
        "langgraph",
    },

    "llamaindex": {
        "llamaindex",
        "llama index",
    },

    "tensorflow": {
        "tensorflow",
    },

    "pytorch": {
        "pytorch",
    },

    "postgresql": {
        "postgresql",
        "postgres",
    },

    "mongodb": {
        "mongodb",
        "mongo db",
        "mongo",
    },

    "amazon web services": {
        "aws",
        "amazon web services",
    },

    "microsoft azure": {
        "azure",
        "microsoft azure",
    },

    "google cloud": {
        "gcp",
        "google cloud",
        "google cloud platform",
    },

    "kubernetes": {
        "kubernetes",
        "k8s",
    },

    "ci/cd": {
        "ci/cd",
        "cicd",
        "continuous integration",
        "continuous deployment",
    },

    "machine learning operations": {
        "mlops",
        "ml ops",
        "machine learning operations",
    },

    "llmops": {
        "llmops",
        "llm ops",
        "llm operations",
    },

}


# ============================================================
# 6. NORMALIZE MATCH TEXT
# ============================================================

def match_normalize(
    text,
):
    """
    Normalize text specifically for matching.
    """

    text = safe_text(
        text
    ).lower()


    text = text.replace(
        "&",
        " and ",
    )


    text = text.replace(
        "/",
        " ",
    )


    text = text.replace(
        "-",
        " ",
    )


    text = re.sub(
        r"[^a-z0-9+#.\s]",
        " ",
        text,
    )


    text = re.sub(
        r"\s+",
        " ",
        text,
    )


    return text.strip()


# ============================================================
# 7. CANONICAL MATCH NAME
# ============================================================

def canonical_match_name(
    skill,
):
    """
    Convert a skill to its canonical matching representation.
    """

    normalized = match_normalize(
        skill
    )


    if not normalized:

        return ""


    # --------------------------------------------------------
    # Existing taxonomy
    # --------------------------------------------------------

    taxonomy_name = normalize_taxonomy_skill(
        skill
    )


    if taxonomy_name:

        normalized_taxonomy = match_normalize(
            taxonomy_name
        )


        if normalized_taxonomy:

            return normalized_taxonomy


    # --------------------------------------------------------
    # Alias map
    # --------------------------------------------------------

    for canonical, aliases in (
        SKILL_ALIASES.items()
    ):

        canonical_normalized = (
            match_normalize(
                canonical
            )
        )


        alias_values = {

            match_normalize(
                alias
            )

            for alias in aliases

        }


        if (
            normalized
            ==
            canonical_normalized
        ):

            return canonical_normalized


        if normalized in alias_values:

            return canonical_normalized


    return normalized


# ============================================================
# 8. ALIAS MATCH
# ============================================================

def alias_match(
    jd_skill,
    curriculum_skill,
):
    """
    Determine whether two skills are aliases.
    """

    jd_normalized = canonical_match_name(
        jd_skill
    )


    curriculum_normalized = canonical_match_name(
        curriculum_skill
    )


    if not jd_normalized or not curriculum_normalized:

        return False


    if jd_normalized == curriculum_normalized:

        return True


    return False


# ============================================================
# 9. TOKEN SIMILARITY
# ============================================================

def token_similarity(
    text_a,
    text_b,
):
    """
    Lightweight token-based similarity.

    Does not require external libraries.
    """

    a = set(
        match_normalize(
            text_a
        ).split()
    )


    b = set(
        match_normalize(
            text_b
        ).split()
    )


    if not a or not b:

        return 0.0


    intersection = len(
        a.intersection(
            b
        )
    )


    union = len(
        a.union(
            b
        )
    )


    if union == 0:

        return 0.0


    jaccard = (
        intersection
        /
        union
    )


    # --------------------------------------------------------
    # Containment bonus
    # --------------------------------------------------------

    normalized_a = match_normalize(
        text_a
    )

    normalized_b = match_normalize(
        text_b
    )


    containment = 0.0


    if (
        normalized_a in normalized_b
        or
        normalized_b in normalized_a
    ):

        containment = 0.20


    return min(
        1.0,
        jaccard + containment
    )


# ============================================================
# 10. OPTIONAL SEMANTIC MODEL
# ============================================================

@st.cache_resource(
    show_spinner=False
)
def load_semantic_model():
    """
    Load SentenceTransformer only when available.

    Semantic matching is optional.

    If unavailable, the matcher continues using exact,
    alias and token matching.
    """

    try:

        from sentence_transformers import (
            SentenceTransformer,
        )


        model = SentenceTransformer(
            "all-MiniLM-L6-v2"
        )


        return model


    except Exception:

        return None


# ============================================================
# 11. SEMANTIC SIMILARITY
# ============================================================

def semantic_similarity(
    text_a,
    text_b,
):
    """
    Calculate semantic similarity.

    Returns:

        0.0 when semantic model unavailable.
    """

    try:

        model = load_semantic_model()


        if model is None:

            return 0.0


        embeddings = model.encode(

            [
                text_a,
                text_b,
            ],

            normalize_embeddings=True,

        )


        similarity = float(

            embeddings[0]
            @
            embeddings[1]

        )


        return max(
            0.0,
            min(
                1.0,
                similarity
            )
        )


    except Exception:

        return 0.0


# ============================================================
# 12. MATCH ONE JD SKILL
# ============================================================

def match_one_skill(
    jd_record,
    curriculum_records,
    use_semantic=True,
):
    """
    Find the strongest curriculum match for one JD skill.
    """

    jd_skill = safe_text(
        jd_record.get(
            "canonical_skill"
        )
        or
        jd_record.get(
            "raw_skill"
        )
    )


    if not jd_skill:

        return None


    jd_canonical = canonical_match_name(
        jd_skill
    )


    best_match = None


    for curriculum_record in curriculum_records:

        curriculum_skill = safe_text(

            curriculum_record.get(
                "canonical_skill"
            )
            or
            curriculum_record.get(
                "raw_skill"
            )

        )


        if not curriculum_skill:

            continue


        curriculum_canonical = (
            canonical_match_name(
                curriculum_skill
            )
        )


        # ====================================================
        # EXACT MATCH
        # ====================================================

        if (
            jd_canonical
            ==
            curriculum_canonical
        ):

            score = 1.0

            method = "exact"


            semantic_score = 1.0

            token_score = 1.0


        else:

            # ================================================
            # ALIAS MATCH
            # ================================================

            if alias_match(
                jd_skill,
                curriculum_skill,
            ):

                score = 0.95

                method = "alias"

                semantic_score = 0.0

                token_score = 1.0


            else:

                # ============================================
                # TOKEN MATCH
                # ============================================

                token_score = token_similarity(

                    jd_skill,

                    curriculum_skill,

                )


                semantic_score = 0.0


                # ============================================
                # SEMANTIC MATCH
                # ============================================

                if use_semantic:

                    if token_score < TOKEN_THRESHOLD:

                        semantic_score = (
                            semantic_similarity(
                                jd_skill,
                                curriculum_skill,
                            )
                        )


                # ============================================
                # COMBINE SCORES
                # ============================================

                score = max(

                    token_score,

                    semantic_score,

                )


                if semantic_score >= SEMANTIC_THRESHOLD:

                    method = "semantic"

                elif token_score >= TOKEN_THRESHOLD:

                    method = "token"

                elif score >= PARTIAL_THRESHOLD:

                    method = "partial"

                else:

                    method = "weak"


        # ====================================================
        # KEEP BEST MATCH
        # ====================================================

        if (
            best_match is None
            or
            score
            >
            best_match[
                "score"
            ]
        ):

            best_match = {

                "jd_skill":
                    jd_skill,

                "jd_canonical":
                    jd_canonical,

                "curriculum_skill":
                    curriculum_skill,

                "curriculum_canonical":
                    curriculum_canonical,

                "score":
                    round(
                        score,
                        4
                    ),

                "token_score":
                    round(
                        token_score,
                        4
                    ),

                "semantic_score":
                    round(
                        semantic_score,
                        4
                    ),

                "match_method":
                    method,

                "curriculum_source":
                    curriculum_record.get(
                        "source",
                        ""
                    ),

                "module":
                    curriculum_record.get(
                        "module",
                        ""
                    ),

                "topic":
                    curriculum_record.get(
                        "topic",
                        ""
                    ),

            }


    # ========================================================
    # NO MATCH
    # ========================================================

    if best_match is None:

        return {

            "jd_skill":
                jd_skill,

            "jd_canonical":
                jd_canonical,

            "curriculum_skill":
                "",

            "curriculum_canonical":
                "",

            "score":
                0.0,

            "token_score":
                0.0,

            "semantic_score":
                0.0,

            "match_method":
                "none",

            "curriculum_source":
                "",

            "module":
                "",

            "topic":
                "",

        }


    return best_match


# ============================================================
# 13. CLASSIFY MATCH
# ============================================================

def classify_match(
    score,
    method,
):
    """
    Convert numerical match into:

    covered
    partial
    missing
    """

    if (
        method
        in
        {
            "exact",
            "alias",
            "semantic",
        }
        and
        score >= 0.78
    ):

        return "covered"


    if score >= TOKEN_THRESHOLD:

        return "covered"


    if score >= PARTIAL_THRESHOLD:

        return "partial"


    return "missing"


# ============================================================
# 14. MATCH ALL JD SKILLS
# ============================================================

def match_jd_to_curriculum(
    jd_intelligence,
    curriculum_intelligence,
    use_semantic=True,
):
    """
    Match every JD skill against curriculum skills.
    """

    jd_records = jd_intelligence.get(
        "skill_records",
        []
    )


    curriculum_records = (
        curriculum_intelligence.get(
            "skill_records",
            []
        )
    )


    results = []


    for jd_record in jd_records:

        match = match_one_skill(

            jd_record,

            curriculum_records,

            use_semantic=use_semantic,

        )


        if match is None:

            continue


        # ----------------------------------------------------
        # Priority
        # ----------------------------------------------------

        priority = safe_text(

            jd_record.get(
                "priority"
            ),

            "required",

        )


        # ----------------------------------------------------
        # Category
        # ----------------------------------------------------

        category = safe_text(

            jd_record.get(
                "category"
            ),

            "other",

        )


        # ----------------------------------------------------
        # Classification
        # ----------------------------------------------------

        status = classify_match(

            match[
                "score"
            ],

            match[
                "match_method"
            ],

        )


        results.append({

            **match,

            "priority":
                priority,

            "jd_category":
                category,

            "status":
                status,

        })


    return results


# ============================================================
# 15. COVERAGE SUMMARY
# ============================================================

def calculate_match_summary(
    results,
):
    """
    Calculate overall and priority-specific coverage.
    """

    total = len(
        results
    )


    covered = [

        item

        for item in results

        if item[
            "status"
        ] == "covered"

    ]


    partial = [

        item

        for item in results

        if item[
            "status"
        ] == "partial"

    ]


    missing = [

        item

        for item in results

        if item[
            "status"
        ] == "missing"

    ]


    required = [

        item

        for item in results

        if item[
            "priority"
        ] == "required"

    ]


    preferred = [

        item

        for item in results

        if item[
            "priority"
        ] == "preferred"

    ]


    def percentage(
        numerator,
        denominator,
    ):

        if denominator == 0:

            return 0.0

        return round(

            (
                numerator
                /
                denominator
            )
            * 100,

            2,

        )


    required_covered = sum(

        1

        for item in required

        if item[
            "status"
        ] == "covered"

    )


    required_partial = sum(

        1

        for item in required

        if item[
            "status"
        ] == "partial"

    )


    preferred_covered = sum(

        1

        for item in preferred

        if item[
            "status"
        ] == "covered"

    )


    weighted_covered = (

        len(covered)
        +
        (
            0.5
            *
            len(partial)
        )

    )


    return {

        "total":
            total,

        "covered":
            len(
                covered
            ),

        "partial":
            len(
                partial
            ),

        "missing":
            len(
                missing
            ),

        "coverage_percentage":
            percentage(
                len(covered),
                total,
            ),

        "weighted_coverage_percentage":
            percentage(
                weighted_covered,
                total,
            ),

        "required_total":
            len(
                required
            ),

        "required_covered":
            required_covered,

        "required_partial":
            required_partial,

        "required_missing":
            len(
                required
            )
            -
            required_covered
            -
            required_partial,

        "required_coverage_percentage":
            percentage(
                required_covered,
                len(required),
            ),

        "preferred_total":
            len(
                preferred
            ),

        "preferred_covered":
            preferred_covered,

        "preferred_coverage_percentage":
            percentage(
                preferred_covered,
                len(preferred),
            ),

    }


# ============================================================
# 16. RUN MATCHING
# ============================================================

if (
    jd_available
    and
    curriculum_available
):

    st.divider()

    st.subheader(
        "⚙️ Run Skill Matching"
    )


    semantic_available = (
        load_semantic_model()
        is not None
    )


    if semantic_available:

        st.success(
            """
            🧠 Semantic matching is available.
            """
        )

    else:

        st.info(
            """
            ℹ️ Semantic model is not available.

            Exact, alias and token matching will still work.
            """
        )


    use_semantic = st.checkbox(

        "Use semantic similarity when needed",

        value=semantic_available,

        disabled=not semantic_available,

        key="use_semantic_jd_matching",

    )


    if st.button(

        "🔗 Match JD Skills with Curriculum",

        type="primary",

        use_container_width=True,

        key="run_jd_curriculum_match",

    ):

        with st.spinner(
            "Comparing industry skills against curriculum..."
        ):

            match_results = (
                match_jd_to_curriculum(

                    jd_skill_intelligence,

                    curriculum_skill_intelligence,

                    use_semantic=use_semantic,

                )
            )


            match_summary = (
                calculate_match_summary(
                    match_results
                )
            )


        # ====================================================
        # SAVE RESULTS
        # ====================================================

        st.session_state[
            "skill_match_results"
        ] = match_results


        st.session_state[
            "skill_match_summary"
        ] = match_summary


        # ====================================================
        # SAVE STATUS LISTS
        # ====================================================

        st.session_state[
            "covered_skills"
        ] = [

            item

            for item in match_results

            if item[
                "status"
            ] == "covered"

        ]


        st.session_state[
            "partial_skills"
        ] = [

            item

            for item in match_results

            if item[
                "status"
            ] == "partial"

        ]


        st.session_state[
            "missing_skills"
        ] = [

            item

            for item in match_results

            if item[
                "status"
            ] == "missing"

        ]


        st.success(
            """
            ✅ JD ↔ Curriculum skill matching completed.
            """
        )


# ============================================================
# 17. LOAD MATCH RESULTS
# ============================================================

skill_match_results = (
    st.session_state.get(
        "skill_match_results",
        []
    )
)


skill_match_summary = (
    st.session_state.get(
        "skill_match_summary",
        {}
    )
)


# ============================================================
# 18. DISPLAY SUMMARY
# ============================================================

if skill_match_results:

    st.divider()

    st.subheader(
        "📊 JD ↔ Curriculum Coverage"
    )


    # ========================================================
    # SUMMARY METRICS
    # ========================================================

    col1, col2, col3, col4, col5 = (
        st.columns(5)
    )


    with col1:

        st.metric(
            "JD Skills",
            skill_match_summary.get(
                "total",
                0
            ),
        )


    with col2:

        st.metric(
            "🟢 Covered",
            skill_match_summary.get(
                "covered",
                0
            ),
        )


    with col3:

        st.metric(
            "🟡 Partial",
            skill_match_summary.get(
                "partial",
                0
            ),
        )


    with col4:

        st.metric(
            "🔴 Missing",
            skill_match_summary.get(
                "missing",
                0
            ),
        )


    with col5:

        st.metric(
            "Coverage",
            f"""
            {
                skill_match_summary.get(
                    "coverage_percentage",
                    0
                )
            }%
            """,
        )


    # ========================================================
    # REQUIRED SKILL COVERAGE
    # ========================================================

    st.markdown(
        "### 🎯 Required Skill Coverage"
    )


    required_col1, required_col2 = (
        st.columns(2)
    )


    with required_col1:

        st.metric(

            "Required Skills",

            skill_match_summary.get(
                "required_total",
                0
            ),

        )


    with required_col2:

        st.metric(

            "Required Coverage",

            f"""
            {
                skill_match_summary.get(
                    "required_coverage_percentage",
                    0
                )
            }%
            """,

        )


    # ========================================================
    # COVERAGE BAR
    # ========================================================

    coverage = (
        skill_match_summary.get(
            "coverage_percentage",
            0
        )
        /
        100
    )


    st.progress(
        min(
            max(
                coverage,
                0.0
            ),
            1.0
        )
    )


    # ========================================================
    # MATCH TABLE
    # ========================================================

    st.markdown(
        "### 🔍 Detailed Skill Matching"
    )


    display_rows = []


    for item in skill_match_results:

        status = item.get(
            "status"
        )


        if status == "covered":

            status_label = (
                "🟢 Covered"
            )

        elif status == "partial":

            status_label = (
                "🟡 Partial"
            )

        else:

            status_label = (
                "🔴 Missing"
            )


        display_rows.append({

            "JD Skill":
                item.get(
                    "jd_skill"
                ),

            "Curriculum Match":
                item.get(
                    "curriculum_skill"
                )
                or
                "—",

            "Status":
                status_label,

            "Priority":
                item.get(
                    "priority"
                ),

            "Confidence":
                f"""
                {
                    round(
                        item.get(
                            "score",
                            0
                        )
                        * 100
                    )
                }%
                """,

            "Method":
                item.get(
                    "match_method"
                ),

            "Module":
                item.get(
                    "module"
                )
                or
                "—",

            "Topic":
                item.get(
                    "topic"
                )
                or
                "—",

        })


    match_df = pd.DataFrame(
        display_rows
    )


    st.dataframe(

        match_df,

        use_container_width=True,

        hide_index=True,

    )


    # ========================================================
    # RESULT TABS
    # ========================================================

    result_tabs = st.tabs([

        "🟢 Covered",

        "🟡 Partial",

        "🔴 Missing",

    ])


    # ========================================================
    # COVERED
    # ========================================================

    with result_tabs[0]:

        covered = st.session_state.get(
            "covered_skills",
            []
        )


        if covered:

            for item in covered:

                st.success(

                    f"""
                    **{item["jd_skill"]}**

                    Curriculum:
                    {item["curriculum_skill"]}

                    Module:
                    {item.get("module") or "Not specified"}
                    """
                )

        else:

            st.info(
                "No fully covered skills."
            )


    # ========================================================
    # PARTIAL
    # ========================================================

    with result_tabs[1]:

        partial = st.session_state.get(
            "partial_skills",
            []
        )


        if partial:

            for item in partial:

                st.warning(

                    f"""
                    **{item["jd_skill"]}**

                    Closest curriculum match:
                    {item["curriculum_skill"] or "None"}

                    Confidence:
                    {round(item["score"] * 100)}%

                    Module:
                    {item.get("module") or "Not specified"}
                    """
                )

        else:

            st.info(
                "No partial matches."
            )


    # ========================================================
    # MISSING
    # ========================================================

    with result_tabs[2]:

        missing = st.session_state.get(
            "missing_skills",
            []
        )


        if missing:

            for item in missing:

                st.error(

                    f"""
                    **{item["jd_skill"]}**

                    🔴 Curriculum coverage not identified.

                    Priority:
                    {item.get("priority", "unknown")}

                    Category:
                    {item.get("jd_category", "unknown")}
                    """
                )

        else:

            st.success(
                "🎉 No missing skills identified."
            )


# ============================================================
# 19. MISSING REQUIRED SKILLS
# ============================================================

if skill_match_results:

    missing_required = [

        item

        for item in skill_match_results

        if (

            item.get(
                "status"
            ) == "missing"

            and

            item.get(
                "priority"
            ) == "required"

        )

    ]


    if missing_required:

        st.divider()

        st.error(
            f"""
            🚨 **{len(missing_required)} required industry
            skill(s) are currently missing from the curriculum.**
            """
        )


        missing_required_df = pd.DataFrame([

            {

                "Required Skill":
                    item.get(
                        "jd_skill"
                    ),

                "Category":
                    item.get(
                        "jd_category"
                    ),

                "Closest Curriculum Match":
                    item.get(
                        "curriculum_skill"
                    )
                    or
                    "None",

                "Confidence":
                    f"""
                    {
                        round(
                            item.get(
                                "score",
                                0
                            )
                            * 100
                        )
                    }%
                    """,

            }

            for item in missing_required

        ])


        st.dataframe(

            missing_required_df,

            use_container_width=True,

            hide_index=True,

        )


# ============================================================
# 20. SAVE COMPLETION FLAG
# ============================================================

st.session_state[
    "skill_matching_complete"
] = bool(
    skill_match_results
)


# ============================================================
# END OF CHUNK 6
# ============================================================
# ============================================================
# CHUNK 7/10
# INDUSTRY GAP ANALYSIS
# ============================================================

"""
Purpose
-------
Convert JD ↔ Curriculum matching results into a detailed
Industry Gap Analysis.

Gap Types
---------

1. Skill Gap
2. Technology Gap
3. Tool Gap
4. Framework Gap
5. Programming Language Gap
6. Cloud Gap
7. Database Gap
8. Concept Gap
9. Project Gap
10. Depth Gap

Outputs
-------

st.session_state["industry_gap_analysis"]

st.session_state["industry_gap_summary"]

st.session_state["critical_gaps"]

st.session_state["skill_gaps"]

st.session_state["technology_gaps"]

st.session_state["project_gaps"]

"""

# ============================================================
# 1. SECTION HEADER
# ============================================================

st.divider()

st.header(
    "6️⃣ 🔍 Industry & Curriculum Gap Analysis"
)

st.markdown(
    """
This module analyses where the current curriculum does not
adequately satisfy industry requirements.

The analysis considers:

- Missing industry skills
- Partially covered skills
- Missing technologies
- Missing tools and frameworks
- Missing concepts
- Missing cloud/database technologies
- Missing project exposure
- Required vs preferred skills
- Industry-critical gaps
"""
)


# ============================================================
# 2. LOAD MATCH RESULTS
# ============================================================

skill_match_results = st.session_state.get(
    "skill_match_results",
    []
)

skill_match_summary = st.session_state.get(
    "skill_match_summary",
    {}
)

jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence",
    {}
)

curriculum_skill_intelligence = st.session_state.get(
    "curriculum_skill_intelligence",
    {}
)


# ============================================================
# 3. VALIDATION
# ============================================================

if not skill_match_results:

    st.warning(
        """
        ⚠️ Skill matching results are not available.

        Complete **Chunk 6 — JD ↔ Curriculum Skill Matching**
        before running Gap Analysis.
        """
    )


# ============================================================
# 4. GAP SEVERITY
# ============================================================

GAP_SEVERITY = {

    "critical": 5,

    "high": 4,

    "medium": 3,

    "low": 2,

    "informational": 1,

}


# ============================================================
# 5. GAP CATEGORY MAPPING
# ============================================================

CATEGORY_GAP_MAP = {

    "ai_ml":
        "AI / Machine Learning",

    "data":
        "Data / Analytics",

    "programming_languages":
        "Programming Language",

    "frameworks":
        "Framework",

    "cloud":
        "Cloud",

    "databases":
        "Database",

    "devops":
        "DevOps / MLOps",

    "business_intelligence":
        "Business Intelligence",

    "engineering":
        "Software Engineering",

    "tools":
        "Tool",

    "other":
        "Other",

}


# ============================================================
# 6. DETERMINE GAP CATEGORY
# ============================================================

def determine_gap_category(
    item,
):
    """
    Determine the most appropriate gap category.
    """

    category = safe_text(
        item.get(
            "jd_category"
        ),
        "other",
    )


    return CATEGORY_GAP_MAP.get(
        category,
        category.title()
        if category
        else "Other",
    )


# ============================================================
# 7. DETERMINE GAP TYPE
# ============================================================

def determine_gap_type(
    item,
):
    """
    Convert JD category into a more useful gap type.
    """

    category = safe_text(
        item.get(
            "jd_category"
        ),
        "other",
    )


    mapping = {

        "ai_ml":
            "Concept / AI Skill Gap",

        "data":
            "Data Skill Gap",

        "programming_languages":
            "Programming Language Gap",

        "frameworks":
            "Framework Gap",

        "cloud":
            "Cloud Technology Gap",

        "databases":
            "Database Gap",

        "devops":
            "DevOps / MLOps Gap",

        "business_intelligence":
            "BI Tool Gap",

        "engineering":
            "Engineering Skill Gap",

        "tools":
            "Tool Gap",

    }


    return mapping.get(
        category,
        "General Skill Gap",
    )


# ============================================================
# 8. DETERMINE SEVERITY
# ============================================================

def determine_gap_severity(
    item,
):
    """
    Determine severity based on:

    - required/preferred
    - missing/partial
    - skill category
    """

    status = safe_text(
        item.get(
            "status"
        )
    )


    priority = safe_text(
        item.get(
            "priority"
        ),
        "required",
    )


    category = safe_text(
        item.get(
            "jd_category"
        ),
        "other",
    )


    # --------------------------------------------------------
    # Missing Required
    # --------------------------------------------------------

    if (
        status == "missing"
        and
        priority == "required"
    ):

        if category in {
            "ai_ml",
            "programming_languages",
            "frameworks",
            "cloud",
            "devops",
        }:

            return "critical"


        return "high"


    # --------------------------------------------------------
    # Partial Required
    # --------------------------------------------------------

    if (
        status == "partial"
        and
        priority == "required"
    ):

        return "high"


    # --------------------------------------------------------
    # Missing Preferred
    # --------------------------------------------------------

    if (
        status == "missing"
        and
        priority == "preferred"
    ):

        return "medium"


    # --------------------------------------------------------
    # Partial Preferred
    # --------------------------------------------------------

    if (
        status == "partial"
        and
        priority == "preferred"
    ):

        return "low"


    return "informational"


# ============================================================
# 9. BUILD ONE GAP RECORD
# ============================================================

def build_gap_record(
    item,
):
    """
    Convert a matching result into a gap record.
    """

    status = safe_text(
        item.get(
            "status"
        )
    )


    if status == "covered":

        return None


    severity = determine_gap_severity(
        item
    )


    gap_category = determine_gap_category(
        item
    )


    gap_type = determine_gap_type(
        item
    )


    jd_skill = safe_text(
        item.get(
            "jd_skill"
        )
    )


    curriculum_skill = safe_text(
        item.get(
            "curriculum_skill"
        )
    )


    score = float(
        item.get(
            "score",
            0.0
        )
        or
        0.0
    )


    # --------------------------------------------------------
    # Gap magnitude
    # --------------------------------------------------------

    gap_score = round(
        (
            1.0
            -
            score
        )
        * 100,
        2,
    )


    # --------------------------------------------------------
    # Priority weight
    # --------------------------------------------------------

    priority = safe_text(
        item.get(
            "priority"
        ),
        "required",
    )


    priority_weight = (

        1.0

        if priority == "required"

        else 0.6

    )


    weighted_gap_score = round(

        gap_score
        *
        priority_weight,

        2,

    )


    return {

        "jd_skill":
            jd_skill,

        "curriculum_match":
            curriculum_skill
            or
            "Not identified",

        "status":
            status,

        "priority":
            priority,

        "severity":
            severity,

        "gap_type":
            gap_type,

        "gap_category":
            gap_category,

        "confidence":
            round(
                score * 100,
                2,
            ),

        "gap_score":
            gap_score,

        "weighted_gap_score":
            weighted_gap_score,

        "match_method":
            item.get(
                "match_method",
                "",
            ),

        "module":
            item.get(
                "module",
                "",
            ),

        "topic":
            item.get(
                "topic",
                "",
            ),

        "jd_category":
            item.get(
                "jd_category",
                "other",
            ),

    }


# ============================================================
# 10. BUILD GAP ANALYSIS
# ============================================================

def build_gap_analysis(
    match_results,
):
    """
    Generate all curriculum-industry gaps.
    """

    gaps = []


    for item in match_results:

        gap = build_gap_record(
            item
        )


        if gap:

            gaps.append(
                gap
            )


    return gaps


# ============================================================
# 11. CLASSIFY GAP LISTS
# ============================================================

def classify_gap_lists(
    gaps,
):
    """
    Separate gaps into useful groups.
    """

    result = {

        "critical_gaps": [],

        "high_gaps": [],

        "medium_gaps": [],

        "low_gaps": [],

        "skill_gaps": [],

        "technology_gaps": [],

        "tool_gaps": [],

        "framework_gaps": [],

        "cloud_gaps": [],

        "database_gaps": [],

        "programming_gaps": [],

        "concept_gaps": [],

        "project_gaps": [],

        "other_gaps": [],

    }


    for gap in gaps:

        severity = gap.get(
            "severity"
        )


        if severity == "critical":

            result[
                "critical_gaps"
            ].append(
                gap
            )


        elif severity == "high":

            result[
                "high_gaps"
            ].append(
                gap
            )


        elif severity == "medium":

            result[
                "medium_gaps"
            ].append(
                gap
            )


        elif severity == "low":

            result[
                "low_gaps"
            ].append(
                gap
            )


        # ----------------------------------------------------
        # Gap Type
        # ----------------------------------------------------

        gap_type = safe_text(
            gap.get(
                "gap_type"
            )
        )


        if "Technology" in gap_type:

            result[
                "technology_gaps"
            ].append(
                gap
            )

        elif "Tool" in gap_type:

            result[
                "tool_gaps"
            ].append(
                gap
            )

        elif "Framework" in gap_type:

            result[
                "framework_gaps"
            ].append(
                gap
            )

        elif "Cloud" in gap_type:

            result[
                "cloud_gaps"
            ].append(
                gap
            )

        elif "Database" in gap_type:

            result[
                "database_gaps"
            ].append(
                gap
            )

        elif "Programming" in gap_type:

            result[
                "programming_gaps"
            ].append(
                gap
            )

        elif "Concept" in gap_type:

            result[
                "concept_gaps"
            ].append(
                gap
            )

        else:

            result[
                "skill_gaps"
            ].append(
                gap
            )


    return result


# ============================================================
# 12. GENERATE GAP ANALYSIS
# ============================================================

if skill_match_results:

    if st.button(

        "🔍 Generate Industry Gap Analysis",

        type="primary",

        use_container_width=True,

        key="generate_industry_gap_analysis",

    ):

        with st.spinner(
            "Analysing curriculum-industry gaps..."
        ):

            gaps = build_gap_analysis(
                skill_match_results
            )


            gap_lists = classify_gap_lists(
                gaps
            )


        # ====================================================
        # SAVE MASTER GAP DATA
        # ====================================================

        industry_gap_analysis = {

            "gaps":
                gaps,

            "total_gaps":
                len(
                    gaps
                ),

            **gap_lists,

            "generated_at":
                datetime.now().isoformat(),

        }


        st.session_state[
            "industry_gap_analysis"
        ] = industry_gap_analysis


        # ====================================================
        # BACKWARD COMPATIBILITY
        # ====================================================

        st.session_state[
            "critical_gaps"
        ] = gap_lists[
            "critical_gaps"
        ]


        st.session_state[
            "skill_gaps"
        ] = gap_lists[
            "skill_gaps"
        ]


        st.session_state[
            "technology_gaps"
        ] = gap_lists[
            "technology_gaps"
        ]


        st.session_state[
            "tool_gaps"
        ] = gap_lists[
            "tool_gaps"
        ]


        st.session_state[
            "framework_gaps"
        ] = gap_lists[
            "framework_gaps"
        ]


        st.session_state[
            "project_gaps"
        ] = gap_lists[
            "project_gaps"
        ]


        st.success(
            "✅ Industry Gap Analysis generated."
        )


# ============================================================
# 13. LOAD GAP ANALYSIS
# ============================================================

industry_gap_analysis = (
    st.session_state.get(
        "industry_gap_analysis"
    )
)


# ============================================================
# 14. GAP SUMMARY
# ============================================================

if industry_gap_analysis:

    gaps = industry_gap_analysis.get(
        "gaps",
        []
    )


    st.divider()

    st.subheader(
        "📊 Industry Gap Summary"
    )


    # ========================================================
    # METRICS
    # ========================================================

    col1, col2, col3, col4, col5 = (
        st.columns(5)
    )


    with col1:

        st.metric(
            "Total Gaps",
            len(
                gaps
            ),
        )


    with col2:

        st.metric(
            "🔴 Critical",
            len(
                industry_gap_analysis.get(
                    "critical_gaps",
                    []
                )
            ),
        )


    with col3:

        st.metric(
            "🟠 High",
            len(
                industry_gap_analysis.get(
                    "high_gaps",
                    []
                )
            ),
        )


    with col4:

        st.metric(
            "🟡 Medium",
            len(
                industry_gap_analysis.get(
                    "medium_gaps",
                    []
                )
            ),
        )


    with col5:

        st.metric(
            "🔵 Low",
            len(
                industry_gap_analysis.get(
                    "low_gaps",
                    []
                )
            ),
        )


    # ========================================================
    # GAP TABLE
    # ========================================================

    st.markdown(
        "### 🔍 Detailed Gap Analysis"
    )


    gap_rows = []


    for gap in gaps:

        severity = gap.get(
            "severity"
        )


        if severity == "critical":

            severity_label = "🔴 Critical"

        elif severity == "high":

            severity_label = "🟠 High"

        elif severity == "medium":

            severity_label = "🟡 Medium"

        else:

            severity_label = "🔵 Low"


        gap_rows.append({

            "JD Requirement":
                gap.get(
                    "jd_skill"
                ),

            "Curriculum Coverage":
                gap.get(
                    "curriculum_match"
                ),

            "Status":
                (
                    "🔴 Missing"
                    if gap.get(
                        "status"
                    ) == "missing"
                    else
                    "🟡 Partial"
                ),

            "Priority":
                gap.get(
                    "priority"
                ),

            "Severity":
                severity_label,

            "Gap Type":
                gap.get(
                    "gap_type"
                ),

            "Confidence":
                f"""
                {
                    gap.get(
                        "confidence",
                        0
                    )
                }%
                """,

            "Module":
                gap.get(
                    "module"
                )
                or
                "—",

        })


    gap_df = pd.DataFrame(
        gap_rows
    )


    st.dataframe(

        gap_df,

        use_container_width=True,

        hide_index=True,

    )


# ============================================================
# 15. GAP CATEGORY TABS
# ============================================================

if industry_gap_analysis:

    gap_tabs = st.tabs([

        "🚨 Critical",

        "🧠 Skills",

        "⚙️ Technology",

        "🛠 Tools",

        "🧩 Frameworks",

        "☁️ Cloud",

        "🗄️ Database",

        "💻 Programming",

    ])


    # ========================================================
    # CRITICAL
    # ========================================================

    with gap_tabs[0]:

        values = industry_gap_analysis.get(
            "critical_gaps",
            []
        )


        if values:

            for gap in values:

                st.error(

                    f"""
                    ### 🔴 {gap["jd_skill"]}

                    **Type:** {gap["gap_type"]}

                    **Priority:** {gap["priority"]}

                    **Curriculum Match:**
                    {gap["curriculum_match"]}

                    **Gap Score:**
                    {gap["gap_score"]}%

                    **Recommendation Required:**
                    Curriculum enhancement should be considered.
                    """
                )

        else:

            st.success(
                "No critical gaps identified."
            )


    # ========================================================
    # SKILLS
    # ========================================================

    with gap_tabs[1]:

        values = industry_gap_analysis.get(
            "skill_gaps",
            []
        )


        if values:

            st.dataframe(

                pd.DataFrame(
                    values
                ),

                use_container_width=True,

                hide_index=True,

            )

        else:

            st.info(
                "No general skill gaps identified."
            )


    # ========================================================
    # TECHNOLOGY
    # ========================================================

    with gap_tabs[2]:

        values = industry_gap_analysis.get(
            "technology_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(

                    f"""
                    **{gap["jd_skill"]}**

                    Curriculum:
                    {gap["curriculum_match"]}

                    Severity:
                    {gap["severity"]}
                    """
                )

        else:

            st.info(
                "No technology gaps identified."
            )


    # ========================================================
    # TOOLS
    # ========================================================

    with gap_tabs[3]:

        values = industry_gap_analysis.get(
            "tool_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(
                    f"""
                    🛠 **{gap["jd_skill"]}**

                    Current curriculum:
                    {gap["curriculum_match"]}
                    """
                )

        else:

            st.info(
                "No tool gaps identified."
            )


    # ========================================================
    # FRAMEWORKS
    # ========================================================

    with gap_tabs[4]:

        values = industry_gap_analysis.get(
            "framework_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(
                    f"""
                    🧩 **{gap["jd_skill"]}**

                    Current curriculum:
                    {gap["curriculum_match"]}
                    """
                )

        else:

            st.info(
                "No framework gaps identified."
            )


    # ========================================================
    # CLOUD
    # ========================================================

    with gap_tabs[5]:

        values = industry_gap_analysis.get(
            "cloud_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(
                    f"""
                    ☁️ **{gap["jd_skill"]}**

                    Current curriculum:
                    {gap["curriculum_match"]}
                    """
                )

        else:

            st.info(
                "No cloud gaps identified."
            )


    # ========================================================
    # DATABASE
    # ========================================================

    with gap_tabs[6]:

        values = industry_gap_analysis.get(
            "database_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(
                    f"""
                    🗄️ **{gap["jd_skill"]}**

                    Current curriculum:
                    {gap["curriculum_match"]}
                    """
                )

        else:

            st.info(
                "No database gaps identified."
            )


    # ========================================================
    # PROGRAMMING
    # ========================================================

    with gap_tabs[7]:

        values = industry_gap_analysis.get(
            "programming_gaps",
            []
        )


        if values:

            for gap in values:

                st.warning(
                    f"""
                    💻 **{gap["jd_skill"]}**

                    Current curriculum:
                    {gap["curriculum_match"]}
                    """
                )

        else:

            st.info(
                "No programming-language gaps identified."
            )


# ============================================================
# 16. REQUIRED INDUSTRY GAPS
# ============================================================

if industry_gap_analysis:

    required_gaps = [

        gap

        for gap in industry_gap_analysis.get(
            "gaps",
            []
        )

        if gap.get(
            "priority"
        ) == "required"

    ]


    if required_gaps:

        st.divider()

        st.subheader(
            "🎯 Required Industry Gaps"
        )


        required_df = pd.DataFrame([

            {

                "Required Skill":
                    gap.get(
                        "jd_skill"
                    ),

                "Status":
                    gap.get(
                        "status"
                    ),

                "Severity":
                    gap.get(
                        "severity"
                    ),

                "Gap Type":
                    gap.get(
                        "gap_type"
                    ),

                "Current Curriculum":
                    gap.get(
                        "curriculum_match"
                    ),

                "Gap Score":
                    f"""
                    {
                        gap.get(
                            "gap_score",
                            0
                        )
                    }%
                    """,

            }

            for gap in required_gaps

        ])


        st.dataframe(

            required_df,

            use_container_width=True,

            hide_index=True,

        )


# ============================================================
# 17. INDUSTRY GAP SCORE
# ============================================================

if industry_gap_analysis:

    gaps = industry_gap_analysis.get(
        "gaps",
        []
    )


    if gaps:

        total_weighted_gap = sum(

            gap.get(
                "weighted_gap_score",
                0
            )

            for gap in gaps

        )


        max_weighted_gap = sum(

            100

            if gap.get(
                "priority"
            ) == "required"

            else 60

            for gap in gaps

        )


        if max_weighted_gap > 0:

            industry_gap_score = round(

                (
                    total_weighted_gap
                    /
                    max_weighted_gap
                )
                * 100,

                2,

            )

        else:

            industry_gap_score = 0.0


        curriculum_alignment_score = round(

            100
            -
            industry_gap_score,

            2,

        )


        st.divider()

        st.subheader(
            "📈 Industry Alignment Score"
        )


        score_col1, score_col2 = (
            st.columns(2)
        )


        with score_col1:

            st.metric(

                "Industry Gap Score",

                f"{industry_gap_score}%",

            )


        with score_col2:

            st.metric(

                "Curriculum Alignment",

                f"{curriculum_alignment_score}%",

            )


        st.progress(

            min(
                max(
                    curriculum_alignment_score
                    /
                    100,

                    0.0,

                ),

                1.0,

            )

        )


        # ----------------------------------------------------
        # SAVE SCORES
        # ----------------------------------------------------

        industry_gap_analysis[
            "industry_gap_score"
        ] = industry_gap_score


        industry_gap_analysis[
            "curriculum_alignment_score"
        ] = curriculum_alignment_score


        st.session_state[
            "industry_gap_score"
        ] = industry_gap_score


        st.session_state[
            "curriculum_alignment_score"
        ] = curriculum_alignment_score


# ============================================================
# 18. GAP PRIORITY MATRIX
# ============================================================

if industry_gap_analysis:

    gaps = industry_gap_analysis.get(
        "gaps",
        []
    )


    if gaps:

        with st.expander(
            "📌 Gap Priority Matrix",
            expanded=False,
        ):

            priority_rows = []


            for gap in gaps:

                priority_rows.append({

                    "Skill":
                        gap.get(
                            "jd_skill"
                        ),

                    "Priority":
                        gap.get(
                            "priority"
                        ),

                    "Severity":
                        gap.get(
                            "severity"
                        ),

                    "Gap Score":
                        gap.get(
                            "gap_score"
                        ),

                    "Weighted Score":
                        gap.get(
                            "weighted_gap_score"
                        ),

                })


            priority_df = pd.DataFrame(
                priority_rows
            )


            priority_df = priority_df.sort_values(

                by=[
                    "Weighted Score"
                ],

                ascending=False,

            )


            st.dataframe(

                priority_df,

                use_container_width=True,

                hide_index=True,

            )


# ============================================================
# 19. SAVE COMPLETION STATE
# ============================================================

st.session_state[
    "industry_gap_analysis_complete"
] = bool(
    industry_gap_analysis
)


# ============================================================
# END OF CHUNK 7
# ============================================================
# ============================================================
# CHUNK 8/10
# RAG + LLM INDUSTRY INTELLIGENCE
# ============================================================

"""
Purpose
-------
Use LLM + optional RAG to enrich deterministic curriculum
gap analysis.

Flow
----

Chunk 6
    ↓
JD ↔ Curriculum Matching
    ↓
Chunk 7
    ↓
Deterministic Gap Analysis
    ↓
Chunk 8
    ↓
RAG Context
    +
Groq / Llama
    ↓
Industry Intelligence
    ↓
Recommendations

Output
------

st.session_state["llm_gap_analysis"]

st.session_state["industry_recommendations"]

st.session_state["llm_gap_summary"]
"""


# ============================================================
# 1. SECTION HEADER
# ============================================================

st.divider()

st.header(
    "7️⃣ 🧠 RAG + LLM Industry Intelligence"
)

st.markdown(
    """
The deterministic gap analysis is now enriched using
industry knowledge and LLM reasoning.

The AI analyses each important gap and recommends:

- What should be added
- Why it matters
- Concepts to teach
- Tools to teach
- Technologies to introduce
- Project ideas
- Expected learning depth
- Industry relevance
- Suggested curriculum placement
"""
)


# ============================================================
# 2. LOAD DATA
# ============================================================

industry_gap_analysis = st.session_state.get(
    "industry_gap_analysis"
)

jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence",
    {}
)

curriculum_skill_intelligence = st.session_state.get(
    "curriculum_skill_intelligence",
    {}
)


# ============================================================
# 3. VALIDATION
# ============================================================

if not industry_gap_analysis:

    st.warning(
        """
        ⚠️ Industry Gap Analysis is not available.

        Complete Chunk 7 first.
        """
    )


# ============================================================
# 4. GROQ IMPORT
# ============================================================

def get_groq_client():
    """
    Create Groq client.

    Requires:

        GROQ_API_KEY

    in environment variables or Streamlit secrets.
    """

    try:

        from groq import Groq

    except ImportError:

        return None


    api_key = os.getenv(
        "GROQ_API_KEY"
    )


    # --------------------------------------------------------
    # Streamlit secrets fallback
    # --------------------------------------------------------

    if not api_key:

        try:

            api_key = st.secrets[
                "GROQ_API_KEY"
            ]

        except Exception:

            api_key = None


    if not api_key:

        return None


    try:

        return Groq(
            api_key=api_key
        )

    except Exception:

        return None


# ============================================================
# 5. MODEL CONFIGURATION
# ============================================================

DEFAULT_GROQ_MODEL = (
    "llama-3.3-70b-versatile"
)


def get_groq_model():
    """
    Allow model override through Streamlit secrets.
    """

    try:

        model = st.secrets.get(
            "GROQ_MODEL"
        )

        if model:

            return model

    except Exception:

        pass


    return os.getenv(
        "GROQ_MODEL",
        DEFAULT_GROQ_MODEL,
    )


# ============================================================
# 6. PREPARE CURRICULUM CONTEXT
# ============================================================

def build_curriculum_context(
    curriculum_intelligence,
):
    """
    Create compact curriculum context for LLM.
    """

    if not curriculum_intelligence:

        return {}


    return {

        "modules": curriculum_intelligence.get(
            "modules",
            []
        ),

        "topics": curriculum_intelligence.get(
            "topics",
            []
        ),

        "concepts": curriculum_intelligence.get(
            "concepts",
            []
        ),

        "skills": curriculum_intelligence.get(
            "all_skills",
            []
        ),

        "tools": curriculum_intelligence.get(
            "tools",
            []
        ),

        "technologies": curriculum_intelligence.get(
            "technologies",
            []
        ),

        "frameworks": curriculum_intelligence.get(
            "frameworks",
            []
        ),

        "languages": curriculum_intelligence.get(
            "languages",
            []
        ),

        "cloud": curriculum_intelligence.get(
            "cloud",
            []
        ),

        "databases": curriculum_intelligence.get(
            "databases",
            []
        ),

        "projects": curriculum_intelligence.get(
            "projects",
            []
        ),

    }


# ============================================================
# 7. PREPARE JD CONTEXT
# ============================================================

def build_jd_context(
    jd_intelligence,
):
    """
    Create compact JD context.
    """

    if not jd_intelligence:

        return {}


    return {

        "required_skills":
            jd_intelligence.get(
                "required_skills",
                []
            ),

        "preferred_skills":
            jd_intelligence.get(
                "preferred_skills",
                []
            ),

        "technical_skills":
            jd_intelligence.get(
                "technical_skills",
                []
            ),

        "tools":
            jd_intelligence.get(
                "tools",
                []
            ),

        "technologies":
            jd_intelligence.get(
                "technologies",
                []
            ),

        "frameworks":
            jd_intelligence.get(
                "frameworks",
                []
            ),

        "languages":
            jd_intelligence.get(
                "languages",
                []
            ),

        "cloud":
            jd_intelligence.get(
                "cloud",
                []
            ),

        "databases":
            jd_intelligence.get(
                "databases",
                []
            ),

        "domain_skills":
            jd_intelligence.get(
                "domain_skills",
                []
            ),

        "project_skills":
            jd_intelligence.get(
                "project_skills",
                []
            ),

    }


# ============================================================
# 8. PREPARE GAP CONTEXT
# ============================================================

def build_gap_context(
    gap_analysis,
    max_gaps=30,
):
    """
    Select the most important gaps for LLM analysis.
    """

    if not gap_analysis:

        return []


    gaps = gap_analysis.get(
        "gaps",
        []
    )


    # --------------------------------------------------------
    # Sort by weighted gap score
    # --------------------------------------------------------

    gaps = sorted(

        gaps,

        key=lambda item:
            float(
                item.get(
                    "weighted_gap_score",
                    0
                )
                or
                0
            ),

        reverse=True,

    )


    return gaps[
        :max_gaps
    ]


# ============================================================
# 9. OPTIONAL RAG RETRIEVAL
# ============================================================

def retrieve_rag_context(
    query,
    top_k=5,
):
    """
    Retrieve industry knowledge from the application's
    existing vector store.

    This function intentionally supports multiple possible
    vector-store implementations.

    If no vector store is available, it returns an empty list.

    The deterministic gap analysis continues to work.
    """

    # --------------------------------------------------------
    # Existing vector store from session
    # --------------------------------------------------------

    vectorstore = st.session_state.get(
        "vectorstore"
    )


    if vectorstore is None:

        vectorstore = st.session_state.get(
            "industry_vectorstore"
        )


    if vectorstore is None:

        vectorstore = st.session_state.get(
            "rag_vectorstore"
        )


    if vectorstore is None:

        return []


    try:

        # ----------------------------------------------------
        # LangChain retriever
        # ----------------------------------------------------

        if hasattr(
            vectorstore,
            "as_retriever"
        ):

            retriever = vectorstore.as_retriever(

                search_kwargs={
                    "k": top_k
                }

            )


            documents = retriever.invoke(
                query
            )


            results = []


            for document in documents:

                if hasattr(
                    document,
                    "page_content"
                ):

                    results.append(
                        document.page_content
                    )

                elif isinstance(
                    document,
                    dict,
                ):

                    results.append(
                        safe_text(
                            document.get(
                                "page_content"
                            )
                        )
                    )


            return results


        # ----------------------------------------------------
        # Direct similarity search
        # ----------------------------------------------------

        if hasattr(
            vectorstore,
            "similarity_search"
        ):

            documents = (
                vectorstore.similarity_search(
                    query,
                    k=top_k
                )
            )


            results = []


            for document in documents:

                if hasattr(
                    document,
                    "page_content"
                ):

                    results.append(
                        document.page_content
                    )


            return results


    except Exception:

        return []


    return []


# ============================================================
# 10. BUILD RAG CONTEXT
# ============================================================

def build_rag_context(
    gaps,
):
    """
    Retrieve supporting knowledge for important gaps.
    """

    rag_results = []


    for gap in gaps:

        skill = safe_text(
            gap.get(
                "jd_skill"
            )
        )


        if not skill:

            continue


        query = f"""
        Industry requirements for {skill}.
        Skills, tools, technologies, projects,
        concepts and practical applications required
        for employment.
        """


        documents = retrieve_rag_context(
            query,
            top_k=3,
        )


        if documents:

            rag_results.append({

                "skill":
                    skill,

                "documents":
                    documents,

            })


    return rag_results


# ============================================================
# 11. CONVERT RAG TO TEXT
# ============================================================

def rag_context_to_text(
    rag_results,
):
    """
    Convert retrieved documents into compact LLM context.
    """

    if not rag_results:

        return (
            "No external RAG evidence was retrieved."
        )


    sections = []


    for item in rag_results:

        skill = item.get(
            "skill",
            ""
        )


        documents = item.get(
            "documents",
            []
        )


        sections.append(

            f"SKILL: {skill}\n"
            +
            "\n".join(

                f"- {doc}"

                for doc in documents

            )

        )


    return "\n\n".join(
        sections
    )


# ============================================================
# 12. LLM SYSTEM PROMPT
# ============================================================

INDUSTRY_INTELLIGENCE_SYSTEM_PROMPT = """
You are an expert curriculum and industry skills analyst.

Your job is to analyse gaps between:

1. Job Description requirements
2. Existing academic curriculum
3. Industry expectations

You must NOT blindly recommend technologies.

Only recommend additions that are:

- relevant to the JD
- useful for employability
- appropriate for the subject
- realistically teachable
- connected to the existing curriculum

For every gap determine:

- industry importance
- why it matters
- whether it should be added
- concepts to teach
- tools to teach
- technologies to teach
- project exposure required
- recommended depth
- suggested module
- suggested learning outcome

Be conservative.

If evidence is insufficient, explicitly state:

"Insufficient evidence"

Do not invent facts about a company.

Return valid JSON only.
"""


# ============================================================
# 13. BUILD LLM PROMPT
# ============================================================

def build_industry_prompt(
    gaps,
    curriculum_context,
    jd_context,
    rag_context,
):
    """
    Build structured LLM request.
    """

    prompt = f"""
Analyse the following curriculum against industry requirements.

========================
JOB DESCRIPTION CONTEXT
========================

{json.dumps(
    jd_context,
    indent=2,
    ensure_ascii=False,
)}

========================
CURRENT CURRICULUM
========================

{json.dumps(
    curriculum_context,
    indent=2,
    ensure_ascii=False,
)}

========================
IDENTIFIED GAPS
========================

{json.dumps(
    gaps,
    indent=2,
    ensure_ascii=False,
)}

========================
RAG / INDUSTRY CONTEXT
========================

{rag_context}

========================
REQUIRED OUTPUT
========================

Return JSON with exactly this structure:

{{
    "overall_assessment": "",
    "industry_alignment": "",
    "critical_observations": [],
    "recommendations": [
        {{
            "skill": "",
            "importance": "critical|high|medium|low",
            "reason": "",
            "add_to_curriculum": true,
            "concepts": [],
            "tools": [],
            "technologies": [],
            "frameworks": [],
            "projects": [],
            "recommended_depth": "",
            "recommended_module": "",
            "learning_outcomes": [],
            "industry_relevance": "",
            "implementation_priority": 1
        }}
    ],
    "curriculum_improvements": [],
    "project_recommendations": [],
    "overall_recommendation": ""
}}

Rules:

1. Do not recommend irrelevant technologies.
2. Do not assume a skill is missing if the curriculum contains
   an equivalent concept.
3. Distinguish concept knowledge from practical tool knowledge.
4. Distinguish tool exposure from project-level proficiency.
5. Prioritize required JD skills over preferred skills.
6. Give practical recommendations.
7. Avoid generic statements.
8. Use RAG evidence when available.
9. Return JSON only.
"""


    return prompt


# ============================================================
# 14. CALL GROQ
# ============================================================

def call_groq_industry_analysis(
    prompt,
):
    """
    Send prompt to Groq / Llama.
    """

    client = get_groq_client()


    if client is None:

        return None


    model = get_groq_model()


    try:

        response = client.chat.completions.create(

            model=model,

            messages=[

                {

                    "role":
                        "system",

                    "content":
                        INDUSTRY_INTELLIGENCE_SYSTEM_PROMPT,

                },

                {

                    "role":
                        "user",

                    "content":
                        prompt,

                },

            ],

            temperature=0.2,

            max_tokens=8000,

        )


        content = (
            response
            .choices[0]
            .message
            .content
        )


        return content


    except Exception as exc:

        st.error(
            f"""
            ❌ Groq request failed:

            {exc}
            """
        )


        return None


# ============================================================
# 15. PARSE LLM JSON
# ============================================================

def parse_llm_json(
    response,
):
    """
    Safely parse JSON returned by LLM.
    """

    if not response:

        return None


    text = safe_text(
        response
    )


    # --------------------------------------------------------
    # Remove markdown code fences
    # --------------------------------------------------------

    text = re.sub(

        r"^```json\s*",

        "",

        text,

        flags=re.IGNORECASE,

    )


    text = re.sub(

        r"^```\s*",

        "",

        text,

    )


    text = re.sub(

        r"\s*```$",

        "",

        text,

    )


    text = text.strip()


    # --------------------------------------------------------
    # Direct JSON
    # --------------------------------------------------------

    try:

        return json.loads(
            text
        )

    except Exception:

        pass


    # --------------------------------------------------------
    # Extract JSON object
    # --------------------------------------------------------

    start = text.find(
        "{"
    )

    end = text.rfind(
        "}"
    )


    if (
        start >= 0
        and
        end > start
    ):

        candidate = text[
            start:end + 1
        ]


        try:

            return json.loads(
                candidate
            )

        except Exception:

            return None


    return None


# ============================================================
# 16. VALIDATE LLM RESULT
# ============================================================

def validate_llm_result(
    result,
):
    """
    Ensure expected keys exist.
    """

    if not isinstance(
        result,
        dict,
    ):

        return None


    result.setdefault(
        "overall_assessment",
        "",
    )


    result.setdefault(
        "industry_alignment",
        "",
    )


    result.setdefault(
        "critical_observations",
        [],
    )


    result.setdefault(
        "recommendations",
        [],
    )


    result.setdefault(
        "curriculum_improvements",
        [],
    )


    result.setdefault(
        "project_recommendations",
        [],
    )


    result.setdefault(
        "overall_recommendation",
        "",
    )


    if not isinstance(
        result[
            "recommendations"
        ],
        list,
    ):

        result[
            "recommendations"
        ] = []


    return result


# ============================================================
# 17. NORMALIZE RECOMMENDATIONS
# ============================================================

def normalize_llm_recommendations(
    result,
):
    """
    Clean recommendation objects returned by LLM.
    """

    recommendations = result.get(
        "recommendations",
        []
    )


    normalized = []


    for index, item in enumerate(
        recommendations,
        start=1,
    ):

        if not isinstance(
            item,
            dict,
        ):

            continue


        normalized.append({

            "id":
                index,

            "skill":
                safe_text(
                    item.get(
                        "skill"
                    )
                ),

            "importance":
                safe_text(
                    item.get(
                        "importance"
                    ),
                    "medium",
                ),

            "reason":
                safe_text(
                    item.get(
                        "reason"
                    )
                ),

            "add_to_curriculum":
                bool(
                    item.get(
                        "add_to_curriculum",
                        True
                    )
                ),

            "concepts":
                normalize_list(
                    item.get(
                        "concepts"
                    )
                ),

            "tools":
                normalize_list(
                    item.get(
                        "tools"
                    )
                ),

            "technologies":
                normalize_list(
                    item.get(
                        "technologies"
                    )
                ),

            "frameworks":
                normalize_list(
                    item.get(
                        "frameworks"
                    )
                ),

            "projects":
                normalize_list(
                    item.get(
                        "projects"
                    )
                ),

            "recommended_depth":
                safe_text(
                    item.get(
                        "recommended_depth"
                    )
                ),

            "recommended_module":
                safe_text(
                    item.get(
                        "recommended_module"
                    )
                ),

            "learning_outcomes":
                normalize_list(
                    item.get(
                        "learning_outcomes"
                    )
                ),

            "industry_relevance":
                safe_text(
                    item.get(
                        "industry_relevance"
                    )
                ),

            "implementation_priority":
                int(
                    item.get(
                        "implementation_priority",
                        index
                    )
                    or
                    index
                ),

        })


    return normalized


# ============================================================
# 18. GENERATE AI INDUSTRY INTELLIGENCE
# ============================================================

if industry_gap_analysis:

    gaps_for_llm = build_gap_context(

        industry_gap_analysis,

        max_gaps=30,

    )


    if gaps_for_llm:

        st.divider()

        st.subheader(
            "🧠 AI Industry Analysis"
        )


        # ----------------------------------------------------
        # Configuration
        # ----------------------------------------------------

        col1, col2 = (
            st.columns(2)
        )


        with col1:

            st.write(
                f"""
                **Gaps selected for AI analysis:**
                {len(gaps_for_llm)}
                """
            )


        with col2:

            st.write(
                f"""
                **Groq Model:**
                `{get_groq_model()}`
                """
            )


        rag_enabled = st.checkbox(

            "Use RAG industry context",

            value=True,

            key="industry_rag_enabled",

        )


        # ----------------------------------------------------
        # Generate
        # ----------------------------------------------------

        if st.button(

            "🧠 Analyse Gaps with Llama + RAG",

            type="primary",

            use_container_width=True,

            key="run_llm_industry_analysis",

        ):

            with st.spinner(

                "Analysing industry gaps with Llama..."

            ):

                curriculum_context = (
                    build_curriculum_context(
                        curriculum_skill_intelligence
                    )
                )


                jd_context = (
                    build_jd_context(
                        jd_skill_intelligence
                    )
                )


                if rag_enabled:

                    rag_results = (
                        build_rag_context(
                            gaps_for_llm
                        )
                    )


                    rag_context = (
                        rag_context_to_text(
                            rag_results
                        )
                    )

                else:

                    rag_results = []

                    rag_context = (
                        "RAG disabled."
                    )


                prompt = build_industry_prompt(

                    gaps_for_llm,

                    curriculum_context,

                    jd_context,

                    rag_context,

                )


                raw_response = (
                    call_groq_industry_analysis(
                        prompt
                    )
                )


                parsed_result = (
                    parse_llm_json(
                        raw_response
                    )
                )


                parsed_result = (
                    validate_llm_result(
                        parsed_result
                    )
                )


                if parsed_result:

                    recommendations = (
                        normalize_llm_recommendations(
                            parsed_result
                        )
                    )


                    parsed_result[
                        "recommendations"
                    ] = recommendations


                    parsed_result[
                        "rag_used"
                    ] = bool(
                        rag_results
                    )


                    parsed_result[
                        "generated_at"
                    ] = datetime.now().isoformat()


                    st.session_state[
                        "llm_gap_analysis"
                    ] = parsed_result


                    st.session_state[
                        "industry_recommendations"
                    ] = recommendations


                    st.session_state[
                        "industry_rag_results"
                    ] = rag_results


                    st.success(
                        """
                        ✅ AI Industry Intelligence generated.
                        """
                    )

                else:

                    st.error(
                        """
                        ❌ The LLM did not return valid JSON.

                        Try running the analysis again.
                        """
                    )


# ============================================================
# 19. LOAD AI RESULT
# ============================================================

llm_gap_analysis = (
    st.session_state.get(
        "llm_gap_analysis"
    )
)


# ============================================================
# 20. DISPLAY AI RESULT
# ============================================================

if llm_gap_analysis:

    st.divider()

    st.subheader(
        "🤖 AI Industry Intelligence"
    )


    # ========================================================
    # OVERALL ASSESSMENT
    # ========================================================

    overall_assessment = safe_text(

        llm_gap_analysis.get(
            "overall_assessment"
        )

    )


    if overall_assessment:

        st.info(
            overall_assessment
        )


    # ========================================================
    # INDUSTRY ALIGNMENT
    # ========================================================

    industry_alignment = safe_text(

        llm_gap_analysis.get(
            "industry_alignment"
        )

    )


    if industry_alignment:

        st.markdown(
            "### 🏭 Industry Alignment"
        )


        st.write(
            industry_alignment
        )


    # ========================================================
    # CRITICAL OBSERVATIONS
    # ========================================================

    observations = normalize_list(

        llm_gap_analysis.get(
            "critical_observations"
        )

    )


    if observations:

        st.markdown(
            "### 🚨 Critical Observations"
        )


        for observation in observations:

            st.warning(
                observation
            )


    # ========================================================
    # RECOMMENDATIONS
    # ========================================================

    recommendations = (
        llm_gap_analysis.get(
            "recommendations",
            []
        )
    )


    if recommendations:

        st.markdown(
            "### 🎯 AI Recommendations"
        )


        for recommendation in recommendations:

            skill = safe_text(
                recommendation.get(
                    "skill"
                )
            )


            importance = safe_text(
                recommendation.get(
                    "importance"
                ),
                "medium",
            )


            if importance == "critical":

                icon = "🔴"

            elif importance == "high":

                icon = "🟠"

            elif importance == "medium":

                icon = "🟡"

            else:

                icon = "🔵"


            with st.expander(

                f"{icon} {skill}",

                expanded=False,

            ):

                st.markdown(
                    "**Why is this important?**"
                )


                st.write(
                    recommendation.get(
                        "reason"
                    )
                )


                # --------------------------------------------
                # Concepts
                # --------------------------------------------

                concepts = normalize_list(

                    recommendation.get(
                        "concepts"
                    )

                )


                if concepts:

                    st.markdown(
                        "**📚 Concepts to Teach**"
                    )


                    st.write(
                        ", ".join(
                            concepts
                        )
                    )


                # --------------------------------------------
                # Technologies
                # --------------------------------------------

                technologies = normalize_list(

                    recommendation.get(
                        "technologies"
                    )

                )


                if technologies:

                    st.markdown(
                        "**⚙️ Technologies**"
                    )


                    st.write(
                        ", ".join(
                            technologies
                        )
                    )


                # --------------------------------------------
                # Tools
                # --------------------------------------------

                tools = normalize_list(

                    recommendation.get(
                        "tools"
                    )

                )


                if tools:

                    st.markdown(
                        "**🛠 Tools**"
                    )


                    st.write(
                        ", ".join(
                            tools
                        )
                    )


                # --------------------------------------------
                # Frameworks
                # --------------------------------------------

                frameworks = normalize_list(

                    recommendation.get(
                        "frameworks"
                    )

                )


                if frameworks:

                    st.markdown(
                        "**🧩 Frameworks**"
                    )


                    st.write(
                        ", ".join(
                            frameworks
                        )
                    )


                # --------------------------------------------
                # Projects
                # --------------------------------------------

                projects = normalize_list(

                    recommendation.get(
                        "projects"
                    )

                )


                if projects:

                    st.markdown(
                        "**🚀 Recommended Projects**"
                    )


                    for project in projects:

                        st.write(
                            f"• {project}"
                        )


                # --------------------------------------------
                # Depth
                # --------------------------------------------

                depth = safe_text(

                    recommendation.get(
                        "recommended_depth"
                    )

                )


                if depth:

                    st.markdown(
                        "**📈 Recommended Depth**"
                    )


                    st.write(
                        depth
                    )


                # --------------------------------------------
                # Module
                # --------------------------------------------

                module = safe_text(

                    recommendation.get(
                        "recommended_module"
                    )

                )


                if module:

                    st.markdown(
                        "**📚 Suggested Module**"
                    )


                    st.write(
                        module
                    )


                # --------------------------------------------
                # Learning Outcomes
                # --------------------------------------------

                outcomes = normalize_list(

                    recommendation.get(
                        "learning_outcomes"
                    )

                )


                if outcomes:

                    st.markdown(
                        "**🎯 Learning Outcomes**"
                    )


                    for outcome in outcomes:

                        st.write(
                            f"• {outcome}"
                        )


                # --------------------------------------------
                # Industry Relevance
                # --------------------------------------------

                relevance = safe_text(

                    recommendation.get(
                        "industry_relevance"
                    )

                )


                if relevance:

                    st.markdown(
                        "**🏭 Industry Relevance**"
                    )


                    st.write(
                        relevance
                    )


    # ========================================================
    # CURRICULUM IMPROVEMENTS
    # ========================================================

    improvements = normalize_list(

        llm_gap_analysis.get(
            "curriculum_improvements"
        )

    )


    if improvements:

        st.markdown(
            "### 📚 Curriculum Improvements"
        )


        for improvement in improvements:

            st.write(
                f"• {improvement}"
            )


    # ========================================================
    # PROJECT RECOMMENDATIONS
    # ========================================================

    project_recommendations = normalize_list(

        llm_gap_analysis.get(
            "project_recommendations"
        )

    )


    if project_recommendations:

        st.markdown(
            "### 🚀 Industry Project Recommendations"
        )


        for project in project_recommendations:

            st.success(
                project
            )


    # ========================================================
    # FINAL RECOMMENDATION
    # ========================================================

    final_recommendation = safe_text(

        llm_gap_analysis.get(
            "overall_recommendation"
        )

    )


    if final_recommendation:

        st.divider()

        st.subheader(
            "🎯 Overall AI Recommendation"
        )


        st.success(
            final_recommendation
        )


# ============================================================
# 21. RAG EVIDENCE
# ============================================================

rag_results = st.session_state.get(
    "industry_rag_results",
    []
)


if rag_results:

    with st.expander(
        "🔎 View RAG Industry Evidence",
        expanded=False,
    ):

        for item in rag_results:

            skill = safe_text(
                item.get(
                    "skill"
                )
            )


            st.markdown(
                f"### {skill}"
            )


            for document in item.get(
                "documents",
                []
            ):

                st.write(
                    document
                )


# ============================================================
# 22. DOWNLOAD AI ANALYSIS
# ============================================================

if llm_gap_analysis:

    st.download_button(

        "⬇️ Download AI Industry Analysis",

        data=serialize_json(
            llm_gap_analysis
        ),

        file_name=(
            "ai_industry_gap_analysis.json"
        ),

        mime="application/json",

        key="download_ai_industry_analysis",

    )


# ============================================================
# 23. COMPLETION FLAG
# ============================================================

st.session_state[
    "llm_industry_analysis_complete"
] = bool(
    llm_gap_analysis
)


# ============================================================
# END OF CHUNK 8
# ============================================================
# ============================================================
# CHUNK 9/10
# MULTI-AGENT CURRICULUM ENHANCEMENT
# ============================================================

"""
Purpose
-------
Use multiple specialized AI agents to review the industry
gaps and produce a final curriculum enhancement proposal.

Agents
------

1. Gap Analyst Agent
2. Industry Expert Agent
3. Curriculum Expert Agent
4. Project / Practical Learning Agent
5. Critic Agent
6. Final Enhancement Agent

Architecture
------------

Gap Analysis
     ↓
Gap Analyst
     ↓
Industry Expert
     ↓
Curriculum Expert
     ↓
Project Expert
     ↓
Critic
     ↓
Final Enhancement Agent
     ↓
Validated Enhancement Plan

Output
------

st.session_state["agent_gap_analysis"]

st.session_state["agent_industry_analysis"]

st.session_state["agent_curriculum_analysis"]

st.session_state["agent_project_analysis"]

st.session_state["agent_critic_analysis"]

st.session_state["final_curriculum_enhancement"]

st.session_state["multi_agent_complete"]
"""


# ============================================================
# 1. SECTION HEADER
# ============================================================

st.divider()

st.header(
    "8️⃣ 🤖 Multi-Agent Curriculum Enhancement"
)

st.markdown(
    """
Multiple specialized AI agents independently analyse the
industry-curriculum gap before a final enhancement agent
creates the recommended curriculum changes.

This reduces the risk of relying on a single LLM response.
"""
)


# ============================================================
# 2. LOAD PREVIOUS RESULTS
# ============================================================

industry_gap_analysis = st.session_state.get(
    "industry_gap_analysis"
)

llm_gap_analysis = st.session_state.get(
    "llm_gap_analysis"
)

industry_recommendations = st.session_state.get(
    "industry_recommendations",
    []
)

jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence",
    {}
)

curriculum_skill_intelligence = st.session_state.get(
    "curriculum_skill_intelligence",
    {}
)


# ============================================================
# 3. VALIDATION
# ============================================================

if not industry_gap_analysis:

    st.warning(
        """
        ⚠️ Gap Analysis is not available.

        Complete Chunk 7 first.
        """
    )


# ============================================================
# 4. AGENT MODEL CONFIGURATION
# ============================================================

AGENT_MODEL = get_groq_model()


# ============================================================
# 5. GENERIC GROQ AGENT CALLER
# ============================================================

def run_agent(
    system_prompt,
    user_prompt,
    temperature=0.2,
):
    """
    Generic Llama/Groq agent execution.

    All specialized agents use the same underlying LLM,
    but have different roles and instructions.
    """

    client = get_groq_client()


    if client is None:

        return {

            "success":
                False,

            "error":
                (
                    "GROQ_API_KEY not configured."
                ),

            "result":
                None,

        }


    try:

        response = client.chat.completions.create(

            model=AGENT_MODEL,

            messages=[

                {

                    "role":
                        "system",

                    "content":
                        system_prompt,

                },

                {

                    "role":
                        "user",

                    "content":
                        user_prompt,

                },

            ],

            temperature=temperature,

            max_tokens=7000,

        )


        content = (
            response
            .choices[0]
            .message
            .content
        )


        parsed = parse_llm_json(
            content
        )


        if parsed is None:

            return {

                "success":
                    False,

                "error":
                    "Agent returned invalid JSON.",

                "raw":
                    content,

                "result":
                    None,

            }


        return {

            "success":
                True,

            "error":
                "",

            "raw":
                content,

            "result":
                parsed,

        }


    except Exception as exc:

        return {

            "success":
                False,

            "error":
                str(
                    exc
                ),

            "result":
                None,

        }


# ============================================================
# 6. BUILD SHARED CONTEXT
# ============================================================

def build_agent_context():
    """
    Build compact context shared by all agents.
    """

    gaps = []


    if industry_gap_analysis:

        gaps = industry_gap_analysis.get(
            "gaps",
            []
        )


    # --------------------------------------------------------
    # Limit context
    # --------------------------------------------------------

    gaps = sorted(

        gaps,

        key=lambda x:
            float(
                x.get(
                    "weighted_gap_score",
                    0
                )
                or
                0
            ),

        reverse=True,

    )[:30]


    return {

        "job_requirements":
            build_jd_context(
                jd_skill_intelligence
            ),

        "curriculum":
            build_curriculum_context(
                curriculum_skill_intelligence
            ),

        "deterministic_gaps":
            gaps,

        "llm_recommendations":
            industry_recommendations,

    }


# ============================================================
# 7. GAP ANALYST AGENT
# ============================================================

GAP_AGENT_SYSTEM = """
You are the Gap Analyst Agent in an AI Curriculum Intelligence
platform.

Your responsibility is to independently verify the gap between
a Job Description and an academic curriculum.

Do not redesign the curriculum yet.

Focus on:

- missing skills
- partially covered skills
- false positives
- equivalent concepts
- industry-critical gaps
- required vs preferred skills
- concept gaps
- technology gaps
- practical gaps
- depth gaps

Be conservative.

A skill should be considered covered when the curriculum
contains an equivalent concept even if the exact terminology
differs.

Return JSON only.
"""


def run_gap_agent(
    context,
):
    """
    Execute Gap Analyst Agent.
    """

    prompt = f"""
Analyse this curriculum and JD.

CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

Return:

{{
    "verified_gaps": [
        {{
            "skill": "",
            "status": "missing|partial|covered",
            "priority": "required|preferred",
            "reason": "",
            "evidence_from_curriculum": "",
            "industry_importance": "critical|high|medium|low",
            "gap_type": ""
        }}
    ],
    "false_positive_gaps": [],
    "critical_gaps": [],
    "overall_assessment": ""
}}

Return JSON only.
"""


    return run_agent(
        GAP_AGENT_SYSTEM,
        prompt,
        temperature=0.1,
    )


# ============================================================
# 8. INDUSTRY EXPERT AGENT
# ============================================================

INDUSTRY_AGENT_SYSTEM = """
You are an Industry Skills Expert.

Your responsibility is to determine what employers actually
expect from candidates for the skills represented in the JD.

Analyse:

- current industry expectations
- practical skill requirements
- tools
- technologies
- frameworks
- cloud platforms
- deployment
- production practices
- portfolio expectations
- project expectations
- skill depth

Do not recommend technology simply because it is popular.

Recommendations must be relevant to the supplied JD and
curriculum.

Return JSON only.
"""


def run_industry_agent(
    context,
    gap_agent_result,
):
    """
    Execute Industry Expert Agent.
    """

    prompt = f"""
You are reviewing an existing Gap Analyst result.

CURRICULUM + JD CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

GAP ANALYST RESULT:

{json.dumps(
    gap_agent_result,
    indent=2,
    ensure_ascii=False,
)}

Determine the industry implications.

Return:

{{
    "industry_skill_priorities": [
        {{
            "skill": "",
            "priority": "critical|high|medium|low",
            "reason": "",
            "required_depth": "",
            "practical_expectation": "",
            "recommended_tools": [],
            "recommended_technologies": []
        }}
    ],
    "industry_trends_relevant_to_jd": [],
    "portfolio_expectations": [],
    "industry_summary": ""
}}

Return JSON only.
"""


    return run_agent(
        INDUSTRY_AGENT_SYSTEM,
        prompt,
        temperature=0.2,
    )


# ============================================================
# 9. CURRICULUM EXPERT AGENT
# ============================================================

CURRICULUM_AGENT_SYSTEM = """
You are a senior university curriculum designer and academic
program expert.

Your responsibility is to translate verified industry gaps
into academically appropriate curriculum improvements.

Consider:

- prerequisite knowledge
- module sequencing
- topic depth
- learning outcomes
- theory vs practical balance
- module duration
- assessment feasibility
- curriculum overload
- prerequisite relationships
- existing module duplication

Do not add unnecessary content.

Prefer enhancement of existing modules before creating new
modules.

Return JSON only.
"""


def run_curriculum_agent(
    context,
    gap_agent_result,
    industry_agent_result,
):
    """
    Execute Curriculum Expert Agent.
    """

    prompt = f"""
Analyse the following curriculum and proposed industry gaps.

BASE CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

GAP ANALYST:

{json.dumps(
    gap_agent_result,
    indent=2,
    ensure_ascii=False,
)}

INDUSTRY EXPERT:

{json.dumps(
    industry_agent_result,
    indent=2,
    ensure_ascii=False,
)}

Create an academic curriculum enhancement proposal.

Return:

{{
    "module_enhancements": [
        {{
            "existing_module": "",
            "action": "enhance|modify|merge|replace|new_module",
            "reason": "",
            "topics_to_add": [],
            "concepts_to_add": [],
            "tools_to_add": [],
            "technologies_to_add": [],
            "recommended_depth": "",
            "estimated_hours": 0,
            "prerequisites": [],
            "learning_outcomes": []
        }}
    ],
    "new_modules": [],
    "topics_to_remove_or_reduce": [],
    "topics_to_merge": [],
    "academic_risks": [],
    "curriculum_summary": ""
}}

Return JSON only.
"""


    return run_agent(
        CURRICULUM_AGENT_SYSTEM,
        prompt,
        temperature=0.2,
    )


# ============================================================
# 10. PROJECT EXPERT AGENT
# ============================================================

PROJECT_AGENT_SYSTEM = """
You are an Industry Project and Experiential Learning Expert.

Your responsibility is to determine how curriculum gaps should
be converted into practical student projects.

Projects should:

- demonstrate multiple skills
- reflect industry workflows
- be realistic for students
- create portfolio value
- connect directly to JD requirements
- include tools and technologies
- have measurable outcomes

Avoid toy projects.

Return JSON only.
"""


def run_project_agent(
    context,
    industry_agent_result,
    curriculum_agent_result,
):
    """
    Execute Project Expert Agent.
    """

    prompt = f"""
Design practical project recommendations using this context.

BASE CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

INDUSTRY EXPERT:

{json.dumps(
    industry_agent_result,
    indent=2,
    ensure_ascii=False,
)}

CURRICULUM EXPERT:

{json.dumps(
    curriculum_agent_result,
    indent=2,
    ensure_ascii=False,
)}

Return:

{{
    "projects": [
        {{
            "title": "",
            "problem_statement": "",
            "industry_use_case": "",
            "skills": [],
            "concepts": [],
            "tools": [],
            "technologies": [],
            "difficulty": "beginner|intermediate|advanced",
            "duration": "",
            "deliverables": [],
            "evaluation_criteria": []
        }}
    ],
    "project_roadmap": [],
    "portfolio_strategy": ""
}}

Return JSON only.
"""


    return run_agent(
        PROJECT_AGENT_SYSTEM,
        prompt,
        temperature=0.3,
    )


# ============================================================
# 11. CRITIC AGENT
# ============================================================

CRITIC_AGENT_SYSTEM = """
You are a highly critical Curriculum Quality Assurance Agent.

You review recommendations created by other AI agents.

Your job is to identify:

- hallucinated requirements
- unnecessary technologies
- duplicate topics
- unrealistic scope
- curriculum overload
- weak industry relevance
- inappropriate depth
- missing prerequisites
- weak project recommendations
- academic inconsistencies
- unsupported claims

You must challenge the other agents.

Do not simply approve their recommendations.

Return JSON only.
"""


def run_critic_agent(
    context,
    gap_result,
    industry_result,
    curriculum_result,
    project_result,
):
    """
    Execute Critic Agent.
    """

    prompt = f"""
Review the following multi-agent curriculum analysis.

BASE CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

GAP ANALYST:

{json.dumps(
    gap_result,
    indent=2,
    ensure_ascii=False,
)}

INDUSTRY EXPERT:

{json.dumps(
    industry_result,
    indent=2,
    ensure_ascii=False,
)}

CURRICULUM EXPERT:

{json.dumps(
    curriculum_result,
    indent=2,
    ensure_ascii=False,
)}

PROJECT EXPERT:

{json.dumps(
    project_result,
    indent=2,
    ensure_ascii=False,
)}

Return:

{{
    "approved_items": [],
    "rejected_items": [],
    "items_needing_revision": [
        {{
            "item": "",
            "problem": "",
            "correction": ""
        }}
    ],
    "hallucination_risks": [],
    "scope_risks": [],
    "duplication_risks": [],
    "critical_corrections": [],
    "critic_summary": ""
}}

Be strict.

Return JSON only.
"""


    return run_agent(
        CRITIC_AGENT_SYSTEM,
        prompt,
        temperature=0.1,
    )


# ============================================================
# 12. FINAL ENHANCEMENT AGENT
# ============================================================

FINAL_AGENT_SYSTEM = """
You are the Final Curriculum Enhancement Architect.

You receive analyses from:

- Gap Analyst
- Industry Expert
- Curriculum Expert
- Project Expert
- Critic Agent

You must synthesize them into one final curriculum enhancement
proposal.

The Critic Agent has authority to flag weak recommendations.

Rules:

1. Do not blindly accept every recommendation.
2. Prioritize critical industry gaps.
3. Preserve valid existing curriculum.
4. Enhance existing modules when possible.
5. Avoid curriculum overload.
6. Clearly distinguish:
   - Must Add
   - Should Add
   - Optional
7. Include concepts, tools and projects.
8. Include recommended depth.
9. Include module placement.
10. Include learning outcomes.
11. Include practical implementation.
12. Do not invent unsupported industry requirements.

Return JSON only.
"""


def run_final_agent(
    context,
    gap_result,
    industry_result,
    curriculum_result,
    project_result,
    critic_result,
):
    """
    Execute final synthesis agent.
    """

    prompt = f"""
Create the final curriculum enhancement proposal.

BASE CONTEXT:

{json.dumps(
    context,
    indent=2,
    ensure_ascii=False,
)}

GAP ANALYST:

{json.dumps(
    gap_result,
    indent=2,
    ensure_ascii=False,
)}

INDUSTRY EXPERT:

{json.dumps(
    industry_result,
    indent=2,
    ensure_ascii=False,
)}

CURRICULUM EXPERT:

{json.dumps(
    curriculum_result,
    indent=2,
    ensure_ascii=False,
)}

PROJECT EXPERT:

{json.dumps(
    project_result,
    indent=2,
    ensure_ascii=False,
)}

CRITIC:

{json.dumps(
    critic_result,
    indent=2,
    ensure_ascii=False,
)}

Return:

{{
    "executive_summary": "",

    "industry_alignment": "",

    "must_add": [
        {{
            "item": "",
            "reason": "",
            "module": "",
            "concepts": [],
            "tools": [],
            "technologies": [],
            "projects": [],
            "depth": "",
            "hours": 0,
            "learning_outcomes": []
        }}
    ],

    "should_add": [
        {{
            "item": "",
            "reason": "",
            "module": "",
            "concepts": [],
            "tools": [],
            "technologies": [],
            "projects": [],
            "depth": "",
            "hours": 0,
            "learning_outcomes": []
        }}
    ],

    "optional_additions": [],

    "module_changes": [
        {{
            "module": "",
            "action": "",
            "changes": [],
            "reason": ""
        }}
    ],

    "projects": [
        {{
            "title": "",
            "skills": [],
            "technologies": [],
            "problem_statement": "",
            "deliverables": []
        }}
    ],

    "topics_to_reduce": [],

    "topics_to_remove": [],

    "topics_to_merge": [],

    "estimated_additional_hours": 0,

    "implementation_priority": [],

    "final_recommendation": ""
}}

Return JSON only.
"""


    return run_agent(
        FINAL_AGENT_SYSTEM,
        prompt,
        temperature=0.15,
    )


# ============================================================
# 13. RUN MULTI-AGENT PIPELINE
# ============================================================

if industry_gap_analysis:

    st.divider()

    st.subheader(
        "🤖 Multi-Agent Analysis"
    )


    st.markdown(
        """
The agents execute sequentially. Each agent receives the
output of the previous stage.
"""
    )


    if st.button(

        "🚀 Run Multi-Agent Curriculum Analysis",

        type="primary",

        use_container_width=True,

        key="run_multi_agent_curriculum_analysis",

    ):

        context = build_agent_context()


        # ====================================================
        # AGENT 1 — GAP ANALYST
        # ====================================================

        with st.status(
            "Running Gap Analyst Agent...",
            expanded=True,
        ) as status:

            gap_response = run_gap_agent(
                context
            )


            if gap_response[
                "success"
            ]:

                gap_result = gap_response[
                    "result"
                ]


                st.write(
                    "✅ Gap Analyst completed."
                )

            else:

                gap_result = {

                    "error":
                        gap_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Gap Analyst failed."
                )


            status.update(
                label="Gap Analyst completed",
                state="complete",
            )


        # ====================================================
        # AGENT 2 — INDUSTRY EXPERT
        # ====================================================

        with st.status(
            "Running Industry Expert Agent...",
            expanded=True,
        ) as status:

            industry_response = (
                run_industry_agent(

                    context,

                    gap_result,

                )
            )


            if industry_response[
                "success"
            ]:

                industry_result = (
                    industry_response[
                        "result"
                    ]
                )


                st.write(
                    "✅ Industry Expert completed."
                )

            else:

                industry_result = {

                    "error":
                        industry_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Industry Expert failed."
                )


            status.update(
                label="Industry Expert completed",
                state="complete",
            )


        # ====================================================
        # AGENT 3 — CURRICULUM EXPERT
        # ====================================================

        with st.status(
            "Running Curriculum Expert Agent...",
            expanded=True,
        ) as status:

            curriculum_response = (
                run_curriculum_agent(

                    context,

                    gap_result,

                    industry_result,

                )
            )


            if curriculum_response[
                "success"
            ]:

                curriculum_result = (
                    curriculum_response[
                        "result"
                    ]
                )


                st.write(
                    "✅ Curriculum Expert completed."
                )

            else:

                curriculum_result = {

                    "error":
                        curriculum_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Curriculum Expert failed."
                )


            status.update(
                label="Curriculum Expert completed",
                state="complete",
            )


        # ====================================================
        # AGENT 4 — PROJECT EXPERT
        # ====================================================

        with st.status(
            "Running Project Expert Agent...",
            expanded=True,
        ) as status:

            project_response = (
                run_project_agent(

                    context,

                    industry_result,

                    curriculum_result,

                )
            )


            if project_response[
                "success"
            ]:

                project_result = (
                    project_response[
                        "result"
                    ]
                )


                st.write(
                    "✅ Project Expert completed."
                )

            else:

                project_result = {

                    "error":
                        project_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Project Expert failed."
                )


            status.update(
                label="Project Expert completed",
                state="complete",
            )


        # ====================================================
        # AGENT 5 — CRITIC
        # ====================================================

        with st.status(
            "Running Critic Agent...",
            expanded=True,
        ) as status:

            critic_response = (
                run_critic_agent(

                    context,

                    gap_result,

                    industry_result,

                    curriculum_result,

                    project_result,

                )
            )


            if critic_response[
                "success"
            ]:

                critic_result = (
                    critic_response[
                        "result"
                    ]
                )


                st.write(
                    "✅ Critic completed."
                )

            else:

                critic_result = {

                    "error":
                        critic_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Critic failed."
                )


            status.update(
                label="Critic Agent completed",
                state="complete",
            )


        # ====================================================
        # AGENT 6 — FINAL ENHANCEMENT
        # ====================================================

        with st.status(
            "Running Final Enhancement Agent...",
            expanded=True,
        ) as status:

            final_response = (
                run_final_agent(

                    context,

                    gap_result,

                    industry_result,

                    curriculum_result,

                    project_result,

                    critic_result,

                )
            )


            if final_response[
                "success"
            ]:

                final_result = (
                    final_response[
                        "result"
                    ]
                )


                st.write(
                    "✅ Final Enhancement completed."
                )

            else:

                final_result = {

                    "error":
                        final_response.get(
                            "error"
                        )

                }


                st.write(
                    "⚠️ Final Enhancement failed."
                )


            status.update(
                label="Final Enhancement completed",
                state="complete",
            )


        # ====================================================
        # SAVE AGENT RESULTS
        # ====================================================

        st.session_state[
            "agent_gap_analysis"
        ] = gap_result


        st.session_state[
            "agent_industry_analysis"
        ] = industry_result


        st.session_state[
            "agent_curriculum_analysis"
        ] = curriculum_result


        st.session_state[
            "agent_project_analysis"
        ] = project_result


        st.session_state[
            "agent_critic_analysis"
        ] = critic_result


        st.session_state[
            "final_curriculum_enhancement"
        ] = final_result


        st.session_state[
            "multi_agent_complete"
        ] = True


        st.success(
            """
            🎉 Multi-Agent Curriculum Enhancement completed.
            """
        )


# ============================================================
# 14. LOAD FINAL RESULT
# ============================================================

final_curriculum_enhancement = (
    st.session_state.get(
        "final_curriculum_enhancement"
    )
)


# ============================================================
# 15. DISPLAY FINAL RESULT
# ============================================================

if final_curriculum_enhancement:

    st.divider()

    st.subheader(
        "🎯 Final Curriculum Enhancement"
    )


    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================

    executive_summary = safe_text(

        final_curriculum_enhancement.get(
            "executive_summary"
        )

    )


    if executive_summary:

        st.info(
            executive_summary
        )


    # ========================================================
    # INDUSTRY ALIGNMENT
    # ========================================================

    alignment = safe_text(

        final_curriculum_enhancement.get(
            "industry_alignment"
        )

    )


    if alignment:

        st.markdown(
            "### 🏭 Industry Alignment"
        )


        st.write(
            alignment
        )


    # ========================================================
    # MUST ADD
    # ========================================================

    must_add = (
        final_curriculum_enhancement.get(
            "must_add",
            []
        )
    )


    if must_add:

        st.markdown(
            "### 🔴 Must Add"
        )


        for item in must_add:

            if not isinstance(
                item,
                dict,
            ):

                continue


            title = safe_text(
                item.get(
                    "item"
                )
            )


            with st.expander(
                f"🔴 {title}",
                expanded=True,
            ):

                st.write(
                    item.get(
                        "reason"
                    )
                )


                concepts = normalize_list(
                    item.get(
                        "concepts"
                    )
                )


                if concepts:

                    st.markdown(
                        "**Concepts:**"
                    )

                    st.write(
                        ", ".join(
                            concepts
                        )
                    )


                tools = normalize_list(
                    item.get(
                        "tools"
                    )
                )


                if tools:

                    st.markdown(
                        "**Tools:**"
                    )

                    st.write(
                        ", ".join(
                            tools
                        )
                    )


                technologies = normalize_list(
                    item.get(
                        "technologies"
                    )
                )


                if technologies:

                    st.markdown(
                        "**Technologies:**"
                    )

                    st.write(
                        ", ".join(
                            technologies
                        )
                    )


                projects = normalize_list(
                    item.get(
                        "projects"
                    )
                )


                if projects:

                    st.markdown(
                        "**Projects:**"
                    )

                    for project in projects:

                        st.write(
                            f"• {project}"
                        )


                st.markdown(
                    f"""
                    **Module:** {item.get("module", "")}

                    **Depth:** {item.get("depth", "")}

                    **Estimated Hours:** {item.get("hours", 0)}
                    """
                )


                outcomes = normalize_list(
                    item.get(
                        "learning_outcomes"
                    )
                )


                if outcomes:

                    st.markdown(
                        "**Learning Outcomes:**"
                    )


                    for outcome in outcomes:

                        st.write(
                            f"• {outcome}"
                        )


    # ========================================================
    # SHOULD ADD
    # ========================================================

    should_add = (
        final_curriculum_enhancement.get(
            "should_add",
            []
        )
    )


    if should_add:

        st.markdown(
            "### 🟠 Should Add"
        )


        for item in should_add:

            if not isinstance(
                item,
                dict,
            ):

                continue


            title = safe_text(
                item.get(
                    "item"
                )
            )


            with st.expander(
                f"🟠 {title}",
                expanded=False,
            ):

                st.write(
                    item.get(
                        "reason"
                    )
                )


                concepts = normalize_list(
                    item.get(
                        "concepts"
                    )
                )


                if concepts:

                    st.write(
                        "**Concepts:** "
                        +
                        ", ".join(
                            concepts
                        )
                    )


                tools = normalize_list(
                    item.get(
                        "tools"
                    )
                )


                if tools:

                    st.write(
                        "**Tools:** "
                        +
                        ", ".join(
                            tools
                        )
                    )


                technologies = normalize_list(
                    item.get(
                        "technologies"
                    )
                )


                if technologies:

                    st.write(
                        "**Technologies:** "
                        +
                        ", ".join(
                            technologies
                        )
                    )


                st.write(
                    f"""
                    **Module:** {item.get("module", "")}

                    **Depth:** {item.get("depth", "")}

                    **Hours:** {item.get("hours", 0)}
                    """
                )


    # ========================================================
    # OPTIONAL
    # ========================================================

    optional_additions = normalize_list(

        final_curriculum_enhancement.get(
            "optional_additions"
        )

    )


    if optional_additions:

        st.markdown(
            "### 🔵 Optional Additions"
        )


        for item in optional_additions:

            st.write(
                f"• {item}"
            )


    # ========================================================
    # MODULE CHANGES
    # ========================================================

    module_changes = (
        final_curriculum_enhancement.get(
            "module_changes",
            []
        )
    )


    if module_changes:

        st.markdown(
            "### 📚 Recommended Module Changes"
        )


        module_rows = []


        for item in module_changes:

            if isinstance(
                item,
                dict,
            ):

                module_rows.append({

                    "Module":
                        item.get(
                            "module",
                            ""
                        ),

                    "Action":
                        item.get(
                            "action",
                            ""
                        ),

                    "Changes":
                        "; ".join(
                            normalize_list(
                                item.get(
                                    "changes"
                                )
                            )
                        ),

                    "Reason":
                        item.get(
                            "reason",
                            ""
                        ),

                })


        if module_rows:

            st.dataframe(

                pd.DataFrame(
                    module_rows
                ),

                use_container_width=True,

                hide_index=True,

            )


    # ========================================================
    # PROJECTS
    # ========================================================

    final_projects = (
        final_curriculum_enhancement.get(
            "projects",
            []
        )
    )


    if final_projects:

        st.markdown(
            "### 🚀 Recommended Industry Projects"
        )


        for project in final_projects:

            if not isinstance(
                project,
                dict,
            ):

                continue


            title = safe_text(
                project.get(
                    "title"
                )
            )


            with st.expander(
                f"🚀 {title}",
                expanded=False,
            ):

                st.write(
                    project.get(
                        "problem_statement",
                        ""
                    )
                )


                skills = normalize_list(
                    project.get(
                        "skills"
                    )
                )


                if skills:

                    st.write(
                        "**Skills:** "
                        +
                        ", ".join(
                            skills
                        )
                    )


                technologies = normalize_list(
                    project.get(
                        "technologies"
                    )
                )


                if technologies:

                    st.write(
                        "**Technologies:** "
                        +
                        ", ".join(
                            technologies
                        )
                    )


                deliverables = normalize_list(
                    project.get(
                        "deliverables"
                    )
                )


                if deliverables:

                    st.markdown(
                        "**Deliverables:**"
                    )


                    for deliverable in deliverables:

                        st.write(
                            f"• {deliverable}"
                        )


    # ========================================================
    # REMOVE / REDUCE
    # ========================================================

    topics_to_reduce = normalize_list(

        final_curriculum_enhancement.get(
            "topics_to_reduce"
        )

    )


    topics_to_remove = normalize_list(

        final_curriculum_enhancement.get(
            "topics_to_remove"
        )

    )


    topics_to_merge = normalize_list(

        final_curriculum_enhancement.get(
            "topics_to_merge"
        )

    )


    if (
        topics_to_reduce
        or
        topics_to_remove
        or
        topics_to_merge
    ):

        st.markdown(
            "### 🔄 Curriculum Optimization"
        )


        if topics_to_reduce:

            st.write(
                "**Reduce:** "
                +
                ", ".join(
                    topics_to_reduce
                )
            )


        if topics_to_remove:

            st.write(
                "**Remove:** "
                +
                ", ".join(
                    topics_to_remove
                )
            )


        if topics_to_merge:

            st.write(
                "**Merge:** "
                +
                ", ".join(
                    topics_to_merge
                )
            )


    # ========================================================
    # ADDITIONAL HOURS
    # ========================================================

    additional_hours = (
        final_curriculum_enhancement.get(
            "estimated_additional_hours",
            0
        )
    )


    st.metric(
        "Estimated Additional Hours",
        additional_hours,
    )


    # ========================================================
    # IMPLEMENTATION PRIORITY
    # ========================================================

    implementation_priority = normalize_list(

        final_curriculum_enhancement.get(
            "implementation_priority"
        )

    )


    if implementation_priority:

        st.markdown(
            "### 🗓 Implementation Priority"
        )


        for index, item in enumerate(

            implementation_priority,

            start=1,

        ):

            st.write(
                f"**{index}.** {item}"
            )


    # ========================================================
    # FINAL RECOMMENDATION
    # ========================================================

    final_recommendation = safe_text(

        final_curriculum_enhancement.get(
            "final_recommendation"
        )

    )


    if final_recommendation:

        st.divider()

        st.subheader(
            "🏆 Final Recommendation"
        )


        st.success(
            final_recommendation
        )


# ============================================================
# 16. AGENT AUDIT TRAIL
# ============================================================

with st.expander(
    "🔎 Multi-Agent Audit Trail",
    expanded=False,
):

    agent_results = {

        "Gap Analyst":
            st.session_state.get(
                "agent_gap_analysis"
            ),

        "Industry Expert":
            st.session_state.get(
                "agent_industry_analysis"
            ),

        "Curriculum Expert":
            st.session_state.get(
                "agent_curriculum_analysis"
            ),

        "Project Expert":
            st.session_state.get(
                "agent_project_analysis"
            ),

        "Critic":
            st.session_state.get(
                "agent_critic_analysis"
            ),

        "Final Enhancement":
            st.session_state.get(
                "final_curriculum_enhancement"
            ),

    }


    for agent_name, result in (
        agent_results.items()
    ):

        if result:

            st.markdown(
                f"### 🤖 {agent_name}"
            )


            st.json(
                result
            )


# ============================================================
# 17. DOWNLOAD FINAL ENHANCEMENT
# ============================================================

if final_curriculum_enhancement:

    st.download_button(

        "⬇️ Download Final Curriculum Enhancement",

        data=serialize_json(

            final_curriculum_enhancement

        ),

        file_name=(
            "final_curriculum_enhancement.json"
        ),

        mime="application/json",

        key="download_final_curriculum_enhancement",

    )


# ============================================================
# 18. SAVE COMPLETION FLAG
# ============================================================

st.session_state[
    "multi_agent_complete"
] = bool(
    final_curriculum_enhancement
)


# ============================================================
# END OF CHUNK 9
# ============================================================
# ============================================================
# CHUNK 10/10
# FINAL INDUSTRY INTELLIGENCE DASHBOARD
# ============================================================

"""
Purpose
-------
Final integration layer for Industry & JD Intelligence.

Consumes:

    JD Extraction
    JD Skill Intelligence
    Curriculum Skill Intelligence
    JD ↔ Curriculum Matching
    Industry Gap Analysis
    RAG + LLM Analysis
    Multi-Agent Analysis
    Final Curriculum Enhancement

Produces:

    Industry Intelligence Dashboard
    JD Coverage Report
    Gap Report
    AI Recommendation Report
    Final Enhancement Package

Session State Outputs
---------------------

industry_intelligence_package

industry_report_data

gap_report_data

enhancement_report_data

industry_analysis_complete
"""


# ============================================================
# 1. SECTION HEADER
# ============================================================

st.divider()

st.header(
    "9️⃣ 📊 Industry & JD Intelligence Dashboard"
)

st.markdown(
    """
This dashboard consolidates the complete analysis of:

**Job Description → Skills → Curriculum → Coverage → Gaps
→ Industry Intelligence → Multi-Agent Recommendations**
"""
)


# ============================================================
# 2. LOAD ALL RESULTS
# ============================================================

jd_data = st.session_state.get(
    "jd_data",
    {}
)


jd_skill_intelligence = st.session_state.get(
    "jd_skill_intelligence",
    {}
)


curriculum_skill_intelligence = st.session_state.get(
    "curriculum_skill_intelligence",
    {}
)


skill_match_results = st.session_state.get(
    "skill_match_results",
    []
)


skill_match_summary = st.session_state.get(
    "skill_match_summary",
    {}
)


industry_gap_analysis = st.session_state.get(
    "industry_gap_analysis",
    {}
)


llm_gap_analysis = st.session_state.get(
    "llm_gap_analysis",
    {}
)


industry_recommendations = st.session_state.get(
    "industry_recommendations",
    []
)


final_curriculum_enhancement = st.session_state.get(
    "final_curriculum_enhancement",
    {}
)


# ============================================================
# 3. SAFE HELPERS
# ============================================================

def safe_number(
    value,
    default=0,
):
    """
    Convert value to numeric safely.
    """

    try:

        return float(
            value
        )

    except Exception:

        return default


def clean_report_list(
    values,
):
    """
    Convert arbitrary list data into clean strings.
    """

    if not values:

        return []


    if not isinstance(
        values,
        list,
    ):

        values = [
            values
        ]


    result = []


    for value in values:

        if isinstance(
            value,
            dict,
        ):

            text = (
                value.get(
                    "name"
                )
                or
                value.get(
                    "title"
                )
                or
                value.get(
                    "skill"
                )
                or
                value.get(
                    "item"
                )
                or
                value.get(
                    "description"
                )
            )

        else:

            text = value


        text = safe_text(
            text
        )


        if text:

            result.append(
                text
            )


    return unique_values(
        result
    )


# ============================================================
# 4. EXTRACT BASIC METADATA
# ============================================================

jd_title = safe_text(

    jd_data.get(
        "job_title"
    )
    or
    jd_data.get(
        "title"
    )
    or
    jd_skill_intelligence.get(
        "job_title"
    )
    or
    "Industry Job Description"

)


company_name = safe_text(

    jd_data.get(
        "company"
    )
    or
    jd_data.get(
        "company_name"
    )
    or
    jd_skill_intelligence.get(
        "company"
    )
    or
    "Not specified"

)


location = safe_text(

    jd_data.get(
        "location"
    )
    or
    jd_data.get(
        "job_location"
    )
    or
    "Not specified"

)


# ============================================================
# 5. COVERAGE METRICS
# ============================================================

coverage_percentage = safe_number(

    skill_match_summary.get(
        "coverage_percentage",
        0,
    )

)


weighted_coverage_percentage = safe_number(

    skill_match_summary.get(
        "weighted_coverage_percentage",
        0,
    )

)


required_coverage_percentage = safe_number(

    skill_match_summary.get(
        "required_coverage_percentage",
        0,
    )

)


industry_gap_score = safe_number(

    industry_gap_analysis.get(
        "industry_gap_score",
        st.session_state.get(
            "industry_gap_score",
            0,
        ),
    )

)


curriculum_alignment_score = safe_number(

    industry_gap_analysis.get(
        "curriculum_alignment_score",
        st.session_state.get(
            "curriculum_alignment_score",
            0,
        ),
    )

)


# ============================================================
# 6. GAP COUNTS
# ============================================================

covered_count = len([

    item

    for item in skill_match_results

    if item.get(
        "status"
    ) == "covered"

])


partial_count = len([

    item

    for item in skill_match_results

    if item.get(
        "status"
    ) == "partial"

])


missing_count = len([

    item

    for item in skill_match_results

    if item.get(
        "status"
    ) == "missing"

])


required_count = len([

    item

    for item in skill_match_results

    if item.get(
        "priority"
    ) == "required"

])


required_missing_count = len([

    item

    for item in skill_match_results

    if (

        item.get(
            "priority"
        ) == "required"

        and

        item.get(
            "status"
        ) == "missing"

    )

])


# ============================================================
# 7. MAIN KPI DASHBOARD
# ============================================================

st.subheader(
    "🎯 Executive KPI Summary"
)


kpi1, kpi2, kpi3, kpi4, kpi5, kpi6 = (
    st.columns(6)
)


with kpi1:

    st.metric(

        "JD Skills",

        len(
            skill_match_results
        ),

    )


with kpi2:

    st.metric(

        "🟢 Covered",

        covered_count,

    )


with kpi3:

    st.metric(

        "🟡 Partial",

        partial_count,

    )


with kpi4:

    st.metric(

        "🔴 Missing",

        missing_count,

    )


with kpi5:

    st.metric(

        "Coverage",

        f"{coverage_percentage:.1f}%",

    )


with kpi6:

    st.metric(

        "Required Coverage",

        f"{required_coverage_percentage:.1f}%",

    )


# ============================================================
# 8. ALIGNMENT SCORE
# ============================================================

st.divider()

score_col1, score_col2, score_col3 = (
    st.columns(3)
)


with score_col1:

    st.metric(

        "Industry Gap Score",

        f"{industry_gap_score:.1f}%",

    )


with score_col2:

    st.metric(

        "Curriculum Alignment",

        f"{curriculum_alignment_score:.1f}%",

    )


with score_col3:

    st.metric(

        "Required Missing",

        required_missing_count,

    )


# ============================================================
# 9. COVERAGE VISUALIZATION
# ============================================================

st.subheader(
    "📈 JD Skill Coverage"
)


coverage_data = pd.DataFrame({

    "Status": [

        "Covered",

        "Partial",

        "Missing",

    ],

    "Skills": [

        covered_count,

        partial_count,

        missing_count,

    ],

})


if not coverage_data.empty:

    st.bar_chart(

        coverage_data.set_index(
            "Status"
        )

    )


# ============================================================
# 10. CRITICAL GAPS
# ============================================================

critical_gaps = (
    industry_gap_analysis.get(
        "critical_gaps",
        []
    )
)


if critical_gaps:

    st.divider()

    st.subheader(
        "🚨 Critical Industry Gaps"
    )


    critical_rows = []


    for gap in critical_gaps:

        critical_rows.append({

            "JD Skill":
                gap.get(
                    "jd_skill"
                ),

            "Priority":
                gap.get(
                    "priority"
                ),

            "Gap Type":
                gap.get(
                    "gap_type"
                ),

            "Curriculum Match":
                gap.get(
                    "curriculum_match"
                ),

            "Gap Score":
                f"""
                {
                    gap.get(
                        "gap_score",
                        0
                    )
                }%
                """,

        })


    st.dataframe(

        pd.DataFrame(
            critical_rows
        ),

        use_container_width=True,

        hide_index=True,

    )


# ============================================================
# 11. TOP MISSING SKILLS
# ============================================================

missing_skills = [

    item

    for item in skill_match_results

    if item.get(
        "status"
    ) == "missing"

]


missing_skills = sorted(

    missing_skills,

    key=lambda item:

        (
            0
            if item.get(
                "priority"
            ) == "required"
            else 1
        ),

)


if missing_skills:

    st.divider()

    st.subheader(
        "🔴 Missing Industry Skills"
    )


    missing_rows = []


    for item in missing_skills:

        missing_rows.append({

            "Skill":
                item.get(
                    "jd_skill"
                ),

            "Priority":
                item.get(
                    "priority"
                ),

            "Category":
                item.get(
                    "jd_category"
                ),

            "Confidence":
                f"""
                {
                    round(
                        safe_number(
                            item.get(
                                "score",
                                0
                            )
                        )
                        * 100
                    )
                }%
                """,

        })


    st.dataframe(

        pd.DataFrame(
            missing_rows
        ),

        use_container_width=True,

        hide_index=True,

    )


# ============================================================
# 12. AI RECOMMENDATION SUMMARY
# ============================================================

if industry_recommendations:

    st.divider()

    st.subheader(
        "🧠 AI Industry Recommendations"
    )


    recommendation_rows = []


    for item in industry_recommendations:

        if not isinstance(
            item,
            dict,
        ):

            continue


        recommendation_rows.append({

            "Skill":
                item.get(
                    "skill"
                ),

            "Importance":
                item.get(
                    "importance"
                ),

            "Add":
                (
                    "Yes"
                    if item.get(
                        "add_to_curriculum",
                        True
                    )
                    else
                    "No"
                ),

            "Module":
                item.get(
                    "recommended_module"
                ),

            "Depth":
                item.get(
                    "recommended_depth"
                ),

            "Tools":
                ", ".join(
                    clean_report_list(
                        item.get(
                            "tools"
                        )
                    )
                ),

            "Technologies":
                ", ".join(
                    clean_report_list(
                        item.get(
                            "technologies"
                        )
                    )
                ),

        })


    if recommendation_rows:

        st.dataframe(

            pd.DataFrame(
                recommendation_rows
            ),

            use_container_width=True,

            hide_index=True,

        )


# ============================================================
# 13. FINAL ENHANCEMENT SUMMARY
# ============================================================

if final_curriculum_enhancement:

    st.divider()

    st.subheader(
        "🎯 Final Curriculum Enhancement Summary"
    )


    must_add = (
        final_curriculum_enhancement.get(
            "must_add",
            []
        )
    )


    should_add = (
        final_curriculum_enhancement.get(
            "should_add",
            []
        )
    )


    optional_additions = clean_report_list(

        final_curriculum_enhancement.get(
            "optional_additions"
        )

    )


    module_changes = (
        final_curriculum_enhancement.get(
            "module_changes",
            []
        )
    )


    projects = (
        final_curriculum_enhancement.get(
            "projects",
            []
        )
    )


    summary_col1, summary_col2, summary_col3, summary_col4 = (
        st.columns(4)
    )


    with summary_col1:

        st.metric(
            "Must Add",
            len(
                must_add
            ),
        )


    with summary_col2:

        st.metric(
            "Should Add",
            len(
                should_add
            ),
        )


    with summary_col3:

        st.metric(
            "Module Changes",
            len(
                module_changes
            ),
        )


    with summary_col4:

        st.metric(
            "Projects",
            len(
                projects
            ),
        )


    # --------------------------------------------------------
    # Must Add
    # --------------------------------------------------------

    if must_add:

        st.markdown(
            "### 🔴 Must Add"
        )


        for item in must_add:

            if isinstance(
                item,
                dict,
            ):

                st.error(

                    f"""
                    **{item.get("item", "")}**

                    Module:
                    {item.get("module", "")}

                    Depth:
                    {item.get("depth", "")}

                    Hours:
                    {item.get("hours", 0)}
                    """

                )

            else:

                st.error(
                    str(
                        item
                    )
                )


    # --------------------------------------------------------
    # Should Add
    # --------------------------------------------------------

    if should_add:

        st.markdown(
            "### 🟠 Should Add"
        )


        for item in should_add:

            if isinstance(
                item,
                dict,
            ):

                st.warning(

                    f"""
                    **{item.get("item", "")}**

                    Module:
                    {item.get("module", "")}

                    Depth:
                    {item.get("depth", "")}
                    """

                )

            else:

                st.warning(
                    str(
                        item
                    )
                )


    # --------------------------------------------------------
    # Projects
    # --------------------------------------------------------

    if projects:

        st.markdown(
            "### 🚀 Recommended Projects"
        )


        for project in projects:

            if isinstance(
                project,
                dict,
            ):

                st.success(

                    f"""
                    **{project.get("title", "")}**

                    {project.get(
                        "problem_statement",
                        ""
                    )}
                    """

                )

            else:

                st.success(
                    str(
                        project
                    )
                )


# ============================================================
# 14. FINAL EXECUTIVE RECOMMENDATION
# ============================================================

if final_curriculum_enhancement:

    final_recommendation = safe_text(

        final_curriculum_enhancement.get(
            "final_recommendation"
        )

    )


    if final_recommendation:

        st.divider()

        st.subheader(
            "🏆 Executive Recommendation"
        )


        st.success(
            final_recommendation
        )


# ============================================================
# 15. BUILD MASTER INDUSTRY PACKAGE
# ============================================================

industry_intelligence_package = {

    "metadata": {

        "job_title":
            jd_title,

        "company":
            company_name,

        "location":
            location,

        "generated_at":
            datetime.now().isoformat(),

    },

    # --------------------------------------------------------
    # JD
    # --------------------------------------------------------

    "jd_intelligence":
        jd_skill_intelligence,

    # --------------------------------------------------------
    # Curriculum
    # --------------------------------------------------------

    "curriculum_intelligence":
        curriculum_skill_intelligence,

    # --------------------------------------------------------
    # Matching
    # --------------------------------------------------------

    "skill_matching": {

        "results":
            skill_match_results,

        "summary":
            skill_match_summary,

    },

    # --------------------------------------------------------
    # Gap
    # --------------------------------------------------------

    "gap_analysis":
        industry_gap_analysis,

    # --------------------------------------------------------
    # LLM
    # --------------------------------------------------------

    "llm_industry_analysis":
        llm_gap_analysis,

    # --------------------------------------------------------
    # Agentic
    # --------------------------------------------------------

    "multi_agent_analysis": {

        "gap_agent":
            st.session_state.get(
                "agent_gap_analysis"
            ),

        "industry_agent":
            st.session_state.get(
                "agent_industry_analysis"
            ),

        "curriculum_agent":
            st.session_state.get(
                "agent_curriculum_analysis"
            ),

        "project_agent":
            st.session_state.get(
                "agent_project_analysis"
            ),

        "critic_agent":
            st.session_state.get(
                "agent_critic_analysis"
            ),

    },

    # --------------------------------------------------------
    # Final Enhancement
    # --------------------------------------------------------

    "final_curriculum_enhancement":
        final_curriculum_enhancement,

}


# ============================================================
# 16. SAVE MASTER PACKAGE
# ============================================================

st.session_state[
    "industry_intelligence_package"
] = industry_intelligence_package


# ============================================================
# 17. PREPARE REPORT DATA
# ============================================================

industry_report_data = {

    "job_title":
        jd_title,

    "company":
        company_name,

    "coverage_percentage":
        coverage_percentage,

    "weighted_coverage_percentage":
        weighted_coverage_percentage,

    "required_coverage_percentage":
        required_coverage_percentage,

    "covered_skills":
        covered_count,

    "partial_skills":
        partial_count,

    "missing_skills":
        missing_count,

    "required_skills":
        required_count,

    "required_missing":
        required_missing_count,

    "industry_gap_score":
        industry_gap_score,

    "curriculum_alignment_score":
        curriculum_alignment_score,

    "critical_gaps":
        critical_gaps,

    "skill_matches":
        skill_match_results,

}


gap_report_data = {

    "industry_gap_score":
        industry_gap_score,

    "curriculum_alignment_score":
        curriculum_alignment_score,

    "total_gaps":
        len(
            industry_gap_analysis.get(
                "gaps",
                []
            )
        ),

    "critical_gaps":
        industry_gap_analysis.get(
            "critical_gaps",
            []
        ),

    "high_gaps":
        industry_gap_analysis.get(
            "high_gaps",
            []
        ),

    "medium_gaps":
        industry_gap_analysis.get(
            "medium_gaps",
            []
        ),

    "low_gaps":
        industry_gap_analysis.get(
            "low_gaps",
            []
        ),

    "skill_gaps":
        industry_gap_analysis.get(
            "skill_gaps",
            []
        ),

    "technology_gaps":
        industry_gap_analysis.get(
            "technology_gaps",
            []
        ),

    "tool_gaps":
        industry_gap_analysis.get(
            "tool_gaps",
            []
        ),

    "framework_gaps":
        industry_gap_analysis.get(
            "framework_gaps",
            []
        ),

    "cloud_gaps":
        industry_gap_analysis.get(
            "cloud_gaps",
            []
        ),

    "database_gaps":
        industry_gap_analysis.get(
            "database_gaps",
            []
        ),

}


enhancement_report_data = {

    "ai_recommendations":
        industry_recommendations,

    "final_enhancement":
        final_curriculum_enhancement,

    "must_add":
        final_curriculum_enhancement.get(
            "must_add",
            []
        ),

    "should_add":
        final_curriculum_enhancement.get(
            "should_add",
            []
        ),

    "optional_additions":
        final_curriculum_enhancement.get(
            "optional_additions",
            []
        ),

    "module_changes":
        final_curriculum_enhancement.get(
            "module_changes",
            []
        ),

    "projects":
        final_curriculum_enhancement.get(
            "projects",
            []
        ),

    "topics_to_reduce":
        final_curriculum_enhancement.get(
            "topics_to_reduce",
            []
        ),

    "topics_to_remove":
        final_curriculum_enhancement.get(
            "topics_to_remove",
            []
        ),

    "topics_to_merge":
        final_curriculum_enhancement.get(
            "topics_to_merge",
            []
        ),

}


# ============================================================
# 18. SAVE REPORT DATA
# ============================================================

st.session_state[
    "industry_report_data"
] = industry_report_data


st.session_state[
    "gap_report_data"
] = gap_report_data


st.session_state[
    "enhancement_report_data"
] = enhancement_report_data


# ============================================================
# 19. EXPORT MASTER JSON
# ============================================================

st.divider()

st.subheader(
    "📦 Export Industry Intelligence"
)


st.markdown(
    """
The complete analysis can now be passed to the Reports page
or downloaded as a machine-readable JSON package.
"""
)


st.download_button(

    "⬇️ Download Complete Industry Intelligence JSON",

    data=serialize_json(
        industry_intelligence_package
    ),

    file_name=(
        "industry_jd_intelligence_complete.json"
    ),

    mime="application/json",

    key="download_complete_industry_package",

    use_container_width=True,

)


# ============================================================
# 20. EXPORT GAP JSON
# ============================================================

st.download_button(

    "⬇️ Download Gap Analysis JSON",

    data=serialize_json(
        gap_report_data
    ),

    file_name=(
        "industry_curriculum_gap_analysis.json"
    ),

    mime="application/json",

    key="download_gap_analysis_json",

    use_container_width=True,

)


# ============================================================
# 21. EXPORT ENHANCEMENT JSON
# ============================================================

if final_curriculum_enhancement:

    st.download_button(

        "⬇️ Download Curriculum Enhancement JSON",

        data=serialize_json(
            enhancement_report_data
        ),

        file_name=(
            "curriculum_enhancement_recommendations.json"
        ),

        mime="application/json",

        key="download_enhancement_json",

        use_container_width=True,

    )


# ============================================================
# 22. PIPELINE STATUS
# ============================================================

st.divider()

st.subheader(
    "🔄 Industry Intelligence Pipeline Status"
)


pipeline_status = {

    "JD Extraction":
        bool(
            jd_data
            or
            jd_skill_intelligence
        ),

    "JD Skill Intelligence":
        bool(
            jd_skill_intelligence
        ),

    "Curriculum Skill Intelligence":
        bool(
            curriculum_skill_intelligence
        ),

    "JD ↔ Curriculum Matching":
        bool(
            skill_match_results
        ),

    "Deterministic Gap Analysis":
        bool(
            industry_gap_analysis
        ),

    "RAG + LLM Analysis":
        bool(
            llm_gap_analysis
        ),

    "Multi-Agent Analysis":
        bool(
            st.session_state.get(
                "multi_agent_complete",
                False,
            )
        ),

    "Final Enhancement":
        bool(
            final_curriculum_enhancement
        ),

}


status_rows = []


for stage, completed in (
    pipeline_status.items()
):

    status_rows.append({

        "Stage":
            stage,

        "Status":
            (
                "✅ Complete"
                if completed
                else
                "⏳ Pending"
            ),

    })


st.dataframe(

    pd.DataFrame(
        status_rows
    ),

    use_container_width=True,

    hide_index=True,

)


# ============================================================
# 23. COMPLETION MESSAGE
# ============================================================

all_major_stages_complete = all(

    pipeline_status.values()

)


if all_major_stages_complete:

    st.success(
        """
        🎉 **Industry & JD Intelligence Analysis Complete**

        The system has completed:

        JD Analysis → Skill Matching → Gap Analysis →
        RAG Intelligence → Multi-Agent Review →
        Final Curriculum Enhancement.

        The resulting package is ready for the
        **Gap & Enhancement module** and the
        **Reports module**.
        """
    )

else:

    st.info(
        """
        ℹ️ Industry Intelligence is partially complete.

        Complete the pending stages before generating the
        final curriculum enhancement.
        """
    )


# ============================================================
# 24. SESSION STATE FOR PAGE 04
# ============================================================

st.session_state[
    "industry_analysis_complete"
] = all_major_stages_complete


st.session_state[
    "ready_for_gap_enhancement"
] = bool(
    industry_gap_analysis
)


st.session_state[
    "ready_for_reports"
] = bool(
    industry_intelligence_package
)


# ============================================================
# END OF CHUNK 10
# ============================================================
