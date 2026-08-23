# ============================================================
# pages/02_📚_Curriculum_Intelligence.py
# ============================================================

"""
PragyanAI
AI Curriculum & Syllabus Intelligence Platform

PAGE 02
Curriculum Intelligence

Responsibilities
----------------
1. Compare primary syllabus with benchmark syllabus
2. Module-wise comparison
3. Topic-wise comparison
4. Similarity analysis
5. Difference analysis
6. Concept intelligence
7. RAG / Web / LLM based enrichment
8. Identify missing concepts
9. Identify emerging concepts
10. Identify industry-relevant concepts
11. Recommend tools
12. Recommend project areas
13. Generate downloadable analysis
"""

# ============================================================
# IMPORTS
# ============================================================

import io
import json
import os
from datetime import datetime
import re
from typing import Any


import pandas as pd

import streamlit as st


# ============================================================
# INTERNAL MODULES
# ============================================================

from curriculum.comparator import (
    compare_curricula,
)


from curriculum.concept_intelligence import (
    enrich_topics,
)


from curriculum.extractor import (
    extract_syllabus,
)


from rag.document_loader import (
    process_uploaded_file,
)


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(

    page_title=(
        "Curriculum Intelligence | PragyanAI"
    ),

    page_icon="📚",

    layout="wide",

    initial_sidebar_state="expanded",

)


# ============================================================
# CONSTANTS
# ============================================================

SUPPORTED_FILE_TYPES = [

    "pdf",

    "docx",

    "png",

    "jpg",

    "jpeg",

]


# Maximum number of topics that can be
# sent for deep AI analysis in one run.

MAX_CONCEPT_ANALYSIS = 50


# Maximum raw text displayed in UI.

MAX_TEXT_PREVIEW = 15000


# ============================================================
# SESSION STATE
# ============================================================

"""
Session state allows Page 02 to communicate with:

    Page 01
        ↓
    Page 02
        ↓
    Page 03
        ↓
    Page 04
        ↓
    Page 05

without repeatedly uploading documents.
"""


DEFAULT_SESSION_STATE = {

    # --------------------------------------------------------
    # Benchmark syllabus
    # --------------------------------------------------------

    "benchmark_syllabus": None,


    # --------------------------------------------------------
    # Multiple benchmark curricula
    # Future extension
    # --------------------------------------------------------

    "benchmark_syllabi": [],


    # --------------------------------------------------------
    # Benchmark filename
    # --------------------------------------------------------

    "benchmark_filename": "",


    # --------------------------------------------------------
    # Benchmark extracted text
    # --------------------------------------------------------

    "benchmark_text": "",


    # --------------------------------------------------------
    # Curriculum comparison result
    # --------------------------------------------------------

    "curriculum_comparison": None,


    # --------------------------------------------------------
    # Deep concept intelligence result
    # --------------------------------------------------------

    "concept_enrichment": [],


    # --------------------------------------------------------
    # Concept analysis status
    # --------------------------------------------------------

    "concept_analysis_complete": False,


    # --------------------------------------------------------
    # Timestamp
    # --------------------------------------------------------

    "curriculum_intelligence_timestamp": None,


    # --------------------------------------------------------
    # Complete report
    # --------------------------------------------------------

    "curriculum_intelligence_report": None,

}


# ============================================================
# INITIALIZE SESSION STATE
# ============================================================

for key, default_value in DEFAULT_SESSION_STATE.items():

    if key not in st.session_state:

        st.session_state[key] = default_value


# ============================================================
# CHUNK 2/10
# HELPER FUNCTIONS
# ============================================================


# ============================================================
# 1. SAFE TEXT CONVERSION
# ============================================================

def safe_text(
    value: Any,
    default: str = "Not Available",
) -> str:
    """
    Safely convert any value into displayable text.

    Parameters
    ----------
    value:
        Input value.

    default:
        Value returned when input is empty or None.

    Returns
    -------
    str
    """

    if value is None:

        return default


    if isinstance(
        value,
        str,
    ):

        value = value.strip()

        if not value:

            return default

        return value


    return str(value)


# ============================================================
# 2. GET TOPICS FROM ONE MODULE
# ============================================================

def get_topics_from_module(
    module: dict,
) -> list[str]:
    """
    Extract topic names from a single curriculum module.

    Supports both structures:

    {
        "name": "Machine Learning",
        "topics": [
            {
                "name": "Regression"
            },
            {
                "name": "Classification"
            }
        ]
    }

    and:

    {
        "name": "Machine Learning",
        "topics": [
            "Regression",
            "Classification"
        ]
    }
    """

    topics = []


    if not isinstance(
        module,
        dict,
    ):

        return topics


    module_topics = module.get(
        "topics",
        [],
    )


    if not isinstance(
        module_topics,
        list,
    ):

        return topics


    for topic in module_topics:


        # ----------------------------------------------------
        # Topic is a dictionary
        # ----------------------------------------------------

        if isinstance(
            topic,
            dict,
        ):

            name = topic.get(
                "name",
                "",
            )


        # ----------------------------------------------------
        # Topic is plain string
        # ----------------------------------------------------

        else:

            name = str(
                topic
            )


        name = str(
            name
        ).strip()


        if name:

            topics.append(
                name
            )


    return topics


# ============================================================
# 3. GET ALL CURRICULUM TOPICS
# ============================================================

def get_all_topics(
    curriculum: dict,
) -> list[str]:
    """
    Extract all unique topics from a curriculum.

    Returns alphabetically sorted topic names.
    """

    topics = []


    if not isinstance(
        curriculum,
        dict,
    ):

        return topics


    modules = curriculum.get(
        "modules",
        [],
    )


    if not isinstance(
        modules,
        list,
    ):

        return topics


    for module in modules:

        topics.extend(
            get_topics_from_module(
                module
            )
        )


    # Remove duplicates while
    # preserving normalized names.

    unique_topics = {}

    for topic in topics:

        normalized = topic.strip()

        key = normalized.lower()


        if key not in unique_topics:

            unique_topics[key] = normalized


    return sorted(
        unique_topics.values()
    )


# ============================================================
# 4. GET ALL MODULE NAMES
# ============================================================

def get_all_modules(
    curriculum: dict,
) -> list[str]:
    """
    Extract all module names from a curriculum.
    """

    modules = []


    if not isinstance(
        curriculum,
        dict,
    ):

        return modules


    for module in curriculum.get(
        "modules",
        [],
    ):

        if not isinstance(
            module,
            dict,
        ):

            continue


        name = module.get(
            "name",
            "",
        )


        name = str(
            name
        ).strip()


        if name:

            modules.append(
                name
            )


    return modules


# ============================================================
# 5. EXTRACT DOCUMENT TEXT
# ============================================================

def extract_document_text(
    uploaded_file,
) -> tuple[list[dict], str]:
    """
    Process an uploaded benchmark document and return:

        (
            normalized_pages,
            full_text
        )

    The shared RAG document loader is the primary extraction path.

    IMPORTANT:
    ``rag.document_loader.process_uploaded_file()`` has existed in
    multiple versions of this project. Some versions return:

        {
            "success": True,
            "filename": "...",
            "extension": ".docx",
            "text": "...",
            "character_count": 123,
            "error": None,
        }

    while older versions return:

        [
            {
                "page": 1,
                "text": "...",
                "source": "..."
            },
            ...
        ]

    This function supports BOTH formats.

    For DOCX files, a direct python-docx fallback is also provided.
    This prevents Page 02 from reporting "no readable text" when the
    shared loader has extracted the document successfully but returned
    its result in a different structure.
    """

    if uploaded_file is None:
        raise ValueError(
            "No benchmark document was uploaded."
        )

    filename = str(
        getattr(
            uploaded_file,
            "name",
            "uploaded_document",
        )
    )

    extension = (
        os.path.splitext(
            filename
        )[1]
        .lower()
    )

    # ------------------------------------------------------------
    # Read bytes once.
    #
    # Streamlit UploadedFile supports getvalue(), which does not
    # disturb the current read position.
    # ------------------------------------------------------------

    file_bytes = b""

    try:
        if hasattr(
            uploaded_file,
            "getvalue",
        ):
            file_bytes = uploaded_file.getvalue()

        elif hasattr(
            uploaded_file,
            "read",
        ):
            file_bytes = uploaded_file.read()

    except Exception as exc:
        raise ValueError(
            f"Unable to read benchmark document '{filename}': {exc}"
        ) from exc

    if file_bytes is None:
        file_bytes = b""

    if not isinstance(
        file_bytes,
        (bytes, bytearray),
    ):
        try:
            file_bytes = bytes(
                file_bytes
            )
        except Exception:
            file_bytes = b""

    # ------------------------------------------------------------
    # Primary path: shared RAG document loader
    # ------------------------------------------------------------

    loader_result = None
    loader_error = None

    try:
        # Reset the UploadedFile position because some versions of
        # document_loader call uploaded_file.read().
        if hasattr(
            uploaded_file,
            "seek",
        ):
            try:
                uploaded_file.seek(0)
            except Exception:
                pass

        loader_result = process_uploaded_file(
            uploaded_file
        )

    except Exception as exc:
        loader_error = exc

        try:
            import logging

            logging.getLogger(
                __name__
            ).warning(
                "Shared document loader failed for %s: %s",
                filename,
                exc,
            )
        except Exception:
            pass

    # ------------------------------------------------------------
    # Normalize every supported return shape from document_loader
    # ------------------------------------------------------------

    normalized_pages: list[dict] = []

    # ------------------------------------------------------------
    # Case 1:
    #
    # New document_loader:
    #
    # {
    #     "success": True,
    #     "text": "...",
    #     ...
    # }
    # ------------------------------------------------------------

    if isinstance(
        loader_result,
        dict,
    ):

        success = loader_result.get(
            "success",
            True,
        )

        loader_text = (
            loader_result.get(
                "text",
                "",
            )
            or ""
        )

        if isinstance(
            loader_text,
            (list, tuple),
        ):
            loader_text = "\n".join(
                str(
                    item
                )
                for item in loader_text
                if item is not None
            )

        loader_text = str(
            loader_text
        ).strip()

        if loader_text:

            normalized_pages.append({

                "page":
                    loader_result.get(
                        "page",
                        1,
                    ),

                "text":
                    loader_text,

                "source":
                    loader_result.get(
                        "source",
                        filename,
                    ),

            })

        elif success is False:
            loader_error = ValueError(
                loader_result.get(
                    "error",
                    "Document loader returned success=False.",
                )
            )

    # ------------------------------------------------------------
    # Case 2:
    #
    # Older document_loader:
    #
    # [
    #     {"page": 1, "text": "..."},
    #     {"page": 2, "text": "..."}
    # ]
    # ------------------------------------------------------------

    elif isinstance(
        loader_result,
        (list, tuple),
    ):

        for index, page in enumerate(
            loader_result,
            start=1,
        ):

            # --------------------------------------------
            # Normal page dictionary
            # --------------------------------------------

            if isinstance(
                page,
                dict,
            ):

                page_text = (
                    page.get(
                        "text",
                        "",
                    )
                    or ""
                )

                if isinstance(
                    page_text,
                    (list, tuple),
                ):

                    page_text = "\n".join(
                        str(
                            item
                        )
                        for item in page_text
                        if item is not None
                    )

                page_text = str(
                    page_text
                ).strip()

                if page_text:

                    normalized_pages.append({

                        "page":
                            page.get(
                                "page",
                                index,
                            ),

                        "text":
                            page_text,

                        "source":
                            page.get(
                                "source",
                                filename,
                            ),

                    })

            # --------------------------------------------
            # Some loaders return plain strings
            # --------------------------------------------

            elif isinstance(
                page,
                str,
            ):

                page_text = page.strip()

                if page_text:

                    normalized_pages.append({

                        "page":
                            index,

                        "text":
                            page_text,

                        "source":
                            filename,

                    })

    # ------------------------------------------------------------
    # Case 3:
    #
    # Very old/simple loader returning a plain string.
    # ------------------------------------------------------------

    elif isinstance(
        loader_result,
        str,
    ):

        loader_text = loader_result.strip()

        if loader_text:

            normalized_pages.append({

                "page":
                    1,

                "text":
                    loader_text,

                "source":
                    filename,

            })

    # ------------------------------------------------------------
    # Build text from shared loader.
    # ------------------------------------------------------------

    text_parts = []

    for page in normalized_pages:

        page_number = page.get(
            "page",
            "",
        )

        page_text = (
            page.get(
                "text",
                "",
            )
            or ""
        )

        page_text = str(
            page_text
        ).strip()

        if not page_text:
            continue

        text_parts.append(
            f"===== PAGE {page_number} =====\n"
            f"{page_text}"
        )

    full_text = "\n\n".join(
        text_parts
    ).strip()

    # ------------------------------------------------------------
    # DOCX FALLBACK
    #
    # This is especially important for the current error:
    #
    # "The uploaded document was processed, but no readable text
    #  was found."
    #
    # We directly read paragraphs AND tables from DOCX.
    # ------------------------------------------------------------

    if not full_text and extension == ".docx":

        try:

            from docx import Document

            document = Document(
                io.BytesIO(
                    file_bytes
                )
            )

            docx_parts = []

            # --------------------------------------------
            # Paragraphs
            # --------------------------------------------

            for paragraph in document.paragraphs:

                paragraph_text = str(
                    paragraph.text or ""
                ).strip()

                if paragraph_text:

                    docx_parts.append(
                        paragraph_text
                    )

            # --------------------------------------------
            # Tables
            #
            # Syllabus documents frequently store:
            #
            # Course Code | Course Name | Credits
            # Module | Topics | Hours
            #
            # in tables rather than paragraphs.
            # --------------------------------------------

            for table_index, table in enumerate(
                document.tables,
                start=1,
            ):

                docx_parts.append(
                    f"===== TABLE {table_index} ====="
                )

                for row in table.rows:

                    cells = []

                    for cell in row.cells:

                        cell_text = str(
                            cell.text or ""
                        ).strip()

                        # Normalize internal whitespace.
                        cell_text = re.sub(
                            r"[ \t]+",
                            " ",
                            cell_text,
                        )

                        cells.append(
                            cell_text
                        )

                    if any(
                        cells
                    ):

                        docx_parts.append(
                            " | ".join(
                                cells
                            )
                        )

            full_text = "\n".join(
                docx_parts
            ).strip()

            if full_text:

                # Replace shared-loader pages with a reliable
                # DOCX text page.
                normalized_pages = [

                    {
                        "page": 1,
                        "text": full_text,
                        "source": filename,
                    }

                ]

        except Exception as exc:

            try:

                import logging

                logging.getLogger(
                    __name__
                ).warning(
                    "Direct DOCX fallback failed for %s: %s",
                    filename,
                    exc,
                )

            except Exception:
                pass

    # ------------------------------------------------------------
    # Additional generic text fallback
    #
    # This handles a plain-text document accidentally uploaded
    # with a supported extension, without affecting normal PDF/
    # DOCX processing.
    # ------------------------------------------------------------

    if not full_text and file_bytes:

        try:

            decoded = file_bytes.decode(
                "utf-8",
                errors="ignore",
            ).strip()

            if decoded:

                # Only use this when it actually looks like text.
                printable_count = sum(
                    1
                    for char in decoded
                    if char.isprintable()
                    or char in "\n\r\t"
                )

                printable_ratio = (
                    printable_count
                    / max(
                        len(decoded),
                        1,
                    )
                )

                if printable_ratio >= 0.85:

                    full_text = decoded

                    normalized_pages = [

                        {
                            "page": 1,
                            "text": full_text,
                            "source": filename,
                        }

                    ]

        except Exception:
            pass

    # ------------------------------------------------------------
    # Final validation
    # ------------------------------------------------------------

    if not full_text:

        if loader_error is not None:

            raise ValueError(
                "The uploaded document was processed, "
                "but no readable text was found. "
                f"Document loader error: {loader_error}"
            ) from loader_error

        raise ValueError(
            "The uploaded document was processed, "
            "but no readable text was found."
        )

    # ------------------------------------------------------------
    # Clean excessive blank lines while preserving structure.
    # ------------------------------------------------------------

    full_text = re.sub(
        r"\n{4,}",
        "\n\n",
        full_text,
    ).strip()

    # ------------------------------------------------------------
    # Log extraction information.
    # ------------------------------------------------------------

    try:

        import logging

        logging.getLogger(
            __name__
        ).info(
            "Benchmark document extracted: %s | type=%s | chars=%d | pages=%d",
            filename,
            extension or "unknown",
            len(full_text),
            len(normalized_pages),
        )

    except Exception:
        pass

    return (
        normalized_pages,
        full_text,
    )


# ============================================================
# 6. BUILD MODULE COMPARISON DATAFRAME
# ============================================================

def build_module_dataframe(
    comparison: dict,
) -> pd.DataFrame:
    """
    Convert module comparison results into a Pandas DataFrame.

    Expected comparison structure:

    {
        "module_comparisons": [
            {
                "primary_module": "...",
                "benchmark_module": "...",
                "similarity_pct": 80,
                "similar_concepts": [],
                "primary_only": [],
                "benchmark_only": []
            }
        ]
    }
    """

    rows = []


    if not isinstance(
        comparison,
        dict,
    ):

        return pd.DataFrame()


    module_comparisons = comparison.get(
        "module_comparisons",
        [],
    )


    if not isinstance(
        module_comparisons,
        list,
    ):

        return pd.DataFrame()


    for item in module_comparisons:


        if not isinstance(
            item,
            dict,
        ):

            continue


        similar = item.get(
            "similar_concepts",
            [],
        )


        primary_only = item.get(
            "primary_only",
            [],
        )


        benchmark_only = item.get(
            "benchmark_only",
            [],
        )


        if not isinstance(
            similar,
            list,
        ):

            similar = []


        if not isinstance(
            primary_only,
            list,
        ):

            primary_only = []


        if not isinstance(
            benchmark_only,
            list,
        ):

            benchmark_only = []


        similarity = item.get(
            "similarity_pct",
            0,
        )


        try:

            similarity = float(
                similarity
            )

        except (
            TypeError,
            ValueError,
        ):

            similarity = 0.0


        rows.append({

            "Primary Module":
                safe_text(
                    item.get(
                        "primary_module"
                    ),
                    "Unknown",
                ),

            "Benchmark Module":
                safe_text(
                    item.get(
                        "benchmark_module"
                    ),
                    "No Matching Module",
                ),

            "Similarity %":
                similarity,

            "Similar Concepts":
                len(
                    similar
                ),

            "Primary Only":
                len(
                    primary_only
                ),

            "Benchmark Only":
                len(
                    benchmark_only
                ),

        })


    return pd.DataFrame(
        rows
    )


# ============================================================
# 7. BUILD TOPIC COMPARISON DATAFRAME
# ============================================================

def build_topic_dataframe(
    comparison: dict,
) -> pd.DataFrame:
    """
    Flatten module-level comparison into topic-level rows.

    Each row contains:

        Module
        Benchmark Module
        Topic
        Status
        Similarity %
    """

    rows = []


    if not isinstance(
        comparison,
        dict,
    ):

        return pd.DataFrame()


    module_comparisons = comparison.get(
        "module_comparisons",
        [],
    )


    if not isinstance(
        module_comparisons,
        list,
    ):

        return pd.DataFrame()


    for module in module_comparisons:


        if not isinstance(
            module,
            dict,
        ):

            continue


        primary_module = safe_text(

            module.get(
                "primary_module"
            ),

            "Unknown",

        )


        benchmark_module = safe_text(

            module.get(
                "benchmark_module"
            ),

            "No Matching Module",

        )


        similarity = module.get(
            "similarity_pct",
            0,
        )


        try:

            similarity = float(
                similarity
            )

        except (
            TypeError,
            ValueError,
        ):

            similarity = 0.0


        # ----------------------------------------------------
        # Similar concepts
        # ----------------------------------------------------

        similar_concepts = module.get(
            "similar_concepts",
            [],
        )


        if not isinstance(
            similar_concepts,
            list,
        ):

            similar_concepts = []


        for topic in similar_concepts:


            rows.append({

                "Module":
                    primary_module,

                "Benchmark Module":
                    benchmark_module,

                "Topic":
                    str(
                        topic
                    ),

                "Status":
                    "Similar",

                "Similarity %":
                    similarity,

            })


        # ----------------------------------------------------
        # Primary-only concepts
        # ----------------------------------------------------

        primary_only = module.get(
            "primary_only",
            [],
        )


        if not isinstance(
            primary_only,
            list,
        ):

            primary_only = []


        for topic in primary_only:


            rows.append({

                "Module":
                    primary_module,

                "Benchmark Module":
                    benchmark_module,

                "Topic":
                    str(
                        topic
                    ),

                "Status":
                    "Primary Only",

                "Similarity %":
                    similarity,

            })


        # ----------------------------------------------------
        # Benchmark-only concepts
        # ----------------------------------------------------

        benchmark_only = module.get(
            "benchmark_only",
            [],
        )


        if not isinstance(
            benchmark_only,
            list,
        ):

            benchmark_only = []


        for topic in benchmark_only:


            rows.append({

                "Module":
                    primary_module,

                "Benchmark Module":
                    benchmark_module,

                "Topic":
                    str(
                        topic
                    ),

                "Status":
                    "Benchmark Only",

                "Similarity %":
                    similarity,

            })


    return pd.DataFrame(
        rows
    )


# ============================================================
# 8. BUILD CONCEPT INTELLIGENCE DATAFRAME
# ============================================================

def build_concept_dataframe(
    enrichment: list[dict],
) -> pd.DataFrame:
    """
    Convert deep concept intelligence results
    into a summary DataFrame.
    """

    rows = []


    if not isinstance(
        enrichment,
        list,
    ):

        return pd.DataFrame()


    for item in enrichment:


        if not isinstance(
            item,
            dict,
        ):

            continue


        core = item.get(
            "core_concepts",
            [],
        )


        advanced = item.get(
            "advanced_concepts",
            [],
        )


        industry = item.get(
            "industry_concepts",
            [],
        )


        emerging = item.get(
            "emerging_concepts",
            [],
        )


        missing = item.get(
            "missing_concepts",
            [],
        )


        tools = item.get(
            "tools",
            [],
        )


        projects = item.get(
            "project_areas",
            [],
        )


        # Ensure lists

        if not isinstance(
            core,
            list,
        ):

            core = []


        if not isinstance(
            advanced,
            list,
        ):

            advanced = []


        if not isinstance(
            industry,
            list,
        ):

            industry = []


        if not isinstance(
            emerging,
            list,
        ):

            emerging = []


        if not isinstance(
            missing,
            list,
        ):

            missing = []


        if not isinstance(
            tools,
            list,
        ):

            tools = []


        if not isinstance(
            projects,
            list,
        ):

            projects = []


        rows.append({

            "Topic":
                safe_text(
                    item.get(
                        "topic"
                    ),
                    "Unknown Topic",
                ),

            "Core Concepts":
                len(
                    core
                ),

            "Advanced Concepts":
                len(
                    advanced
                ),

            "Industry Concepts":
                len(
                    industry
                ),

            "Emerging Concepts":
                len(
                    emerging
                ),

            "Potential Gaps":
                len(
                    missing
                ),

            "Tools":
                len(
                    tools
                ),

            "Projects":
                len(
                    projects
                ),

        })


    return pd.DataFrame(
        rows
    )


# ============================================================
# 9. COLLECT UNIQUE CONCEPT VALUES
# ============================================================

def collect_concept_values(
    enrichment: list[dict],
    field: str,
) -> list[str]:
    """
    Collect unique values from a specified field.

    Example:

        collect_concept_values(
            enrichment,
            "industry_concepts"
        )
    """

    values = set()


    if not isinstance(
        enrichment,
        list,
    ):

        return []


    for item in enrichment:


        if not isinstance(
            item,
            dict,
        ):

            continue


        field_values = item.get(
            field,
            [],
        )


        if not isinstance(
            field_values,
            list,
        ):

            continue


        for value in field_values:


            if value is None:

                continue


            # ----------------------------------------------
            # Dictionary value
            # ----------------------------------------------

            if isinstance(
                value,
                dict,
            ):

                # Prefer name/title/concept

                value = (

                    value.get(
                        "name"
                    )

                    or

                    value.get(
                        "title"
                    )

                    or

                    value.get(
                        "concept"
                    )

                    or

                    str(
                        value
                    )

                )


            value = str(
                value
            ).strip()


            if value:

                values.add(
                    value
                )


    return sorted(
        values,
        key=lambda x: x.lower(),
    )


# ============================================================
# 10. BUILD GLOBAL CONCEPT SUMMARY
# ============================================================

def build_concept_summary(
    enrichment: list[dict],
) -> dict:
    """
    Create global concept intelligence summary.
    """

    return {

        "topics_analyzed":
            len(
                enrichment
            ),

        "core_concepts":
            collect_concept_values(
                enrichment,
                "core_concepts",
            ),

        "advanced_concepts":
            collect_concept_values(
                enrichment,
                "advanced_concepts",
            ),

        "industry_concepts":
            collect_concept_values(
                enrichment,
                "industry_concepts",
            ),

        "emerging_concepts":
            collect_concept_values(
                enrichment,
                "emerging_concepts",
            ),

        "missing_concepts":
            collect_concept_values(
                enrichment,
                "missing_concepts",
            ),

        "tools":
            collect_concept_values(
                enrichment,
                "tools",
            ),

        "project_areas":
            collect_concept_values(
                enrichment,
                "project_areas",
            ),

    }


# ============================================================
# 11. BUILD COMPLETE REPORT
# ============================================================

def build_full_report(
    primary: dict,
    benchmark: dict | None,
    comparison: dict,
    enrichment: list[dict],
) -> dict:
    """
    Build the complete Curriculum Intelligence report.

    This object will later be consumed by:

        Page 03
            Industry & JD Intelligence

        Page 04
            Gap & Enhancement

        Page 05
            Reports
    """

    concept_summary = build_concept_summary(
        enrichment
    )


    report = {

        "report_type":
            "Curriculum Intelligence",

        "version":
            "1.0",

        "generated_at":
            datetime.now().isoformat(),

        # ----------------------------------------------------
        # Primary curriculum
        # ----------------------------------------------------

        "primary_curriculum":
            primary,

        # ----------------------------------------------------
        # Benchmark curriculum
        # ----------------------------------------------------

        "benchmark_curriculum":
            benchmark,

        # ----------------------------------------------------
        # Curriculum comparison
        # ----------------------------------------------------

        "comparison":
            comparison,

        # ----------------------------------------------------
        # Concept intelligence
        # ----------------------------------------------------

        "concept_intelligence": {

            "summary":
                concept_summary,

            "topic_analysis":
                enrichment,

        },

    }


    return report


# ============================================================
# 12. DISPLAY LIST
# ============================================================

def display_value_list(
    values,
    empty_message: str = (
        "No information available."
    ),
):
    """
    Display a Python list safely in Streamlit.
    """

    if not values:

        st.info(
            empty_message
        )

        return


    if not isinstance(
        values,
        list,
    ):

        values = [
            values
        ]


    for value in values:


        if isinstance(
            value,
            dict,
        ):

            st.json(
                value
            )


        elif isinstance(
            value,
            list,
        ):

            st.write(
                value
            )


        else:

            st.markdown(
                f"- {value}"
            )


# ============================================================
# 13. DISPLAY TOPIC LIST WITH STATUS
# ============================================================

def display_topic_list(
    values,
    status: str = "normal",
):
    """
    Display topics using different Streamlit
    visual styles.

    status:

        normal
        success
        warning
        error
    """

    if not values:

        st.info(
            "None identified."
        )

        return


    if not isinstance(
        values,
        list,
    ):

        values = [
            values
        ]


    for value in values:


        value = str(
            value
        ).strip()


        if not value:

            continue


        if status == "success":

            st.success(
                value
            )


        elif status == "warning":

            st.warning(
                value
            )


        elif status == "error":

            st.error(
                value
            )


        else:

            st.write(
                f"• {value}"
            )


# ============================================================
# 14. GET COMPARISON SUMMARY
# ============================================================

def get_comparison_summary(
    comparison: dict,
) -> dict:
    """
    Normalize comparison summary values.
    """

    if not isinstance(
        comparison,
        dict,
    ):

        return {

            "similarity_pct": 0.0,

            "modules_compared": 0,

            "similar_topics": 0,

            "different_topics": 0,

        }


    summary = comparison.get(
        "summary",
        {},
    )


    if not isinstance(
        summary,
        dict,
    ):

        summary = {}


    def safe_float(
        value,
    ):

        try:

            return float(
                value or 0
            )

        except (
            TypeError,
            ValueError,
        ):

            return 0.0


    def safe_int(
        value,
    ):

        try:

            return int(
                value or 0
            )

        except (
            TypeError,
            ValueError,
        ):

            return 0


    return {

        "similarity_pct":
            safe_float(
                summary.get(
                    "similarity_pct",
                    0,
                )
            ),

        "modules_compared":
            safe_int(
                summary.get(
                    "modules_compared",
                    0,
                )
            ),

        "similar_topics":
            safe_int(
                summary.get(
                    "similar_topics",
                    0,
                )
            ),

        "different_topics":
            safe_int(
                summary.get(
                    "different_topics",
                    0,
                )
            ),

    }


# ============================================================
# 15. GET BENCHMARK-ONLY TOPICS
# ============================================================

def get_benchmark_only_topics(
    comparison: dict,
) -> list[str]:
    """
    Find topics that exist in benchmark curriculum
    but not in the primary curriculum.
    """

    topics = set()


    if not isinstance(
        comparison,
        dict,
    ):

        return []


    for module in comparison.get(
        "module_comparisons",
        [],
    ):


        if not isinstance(
            module,
            dict,
        ):

            continue


        values = module.get(
            "benchmark_only",
            [],
        )


        if not isinstance(
            values,
            list,
        ):

            continue


        for value in values:

            value = str(
                value
            ).strip()


            if value:

                topics.add(
                    value
                )


    return sorted(
        topics,
        key=lambda x: x.lower(),
    )


# ============================================================
# 16. GET PRIMARY-ONLY TOPICS
# ============================================================

def get_primary_only_topics(
    comparison: dict,
) -> list[str]:
    """
    Find topics present in primary curriculum
    but not benchmark curriculum.
    """

    topics = set()


    if not isinstance(
        comparison,
        dict,
    ):

        return []


    for module in comparison.get(
        "module_comparisons",
        [],
    ):


        if not isinstance(
            module,
            dict,
        ):

            continue


        values = module.get(
            "primary_only",
            [],
        )


        if not isinstance(
            values,
            list,
        ):

            continue


        for value in values:

            value = str(
                value
            ).strip()


            if value:

                topics.add(
                    value
                )


    return sorted(
        topics,
        key=lambda x: x.lower(),
    )


# ============================================================
# 17. GET SIMILAR TOPICS
# ============================================================

def get_similar_topics(
    comparison: dict,
) -> list[str]:
    """
    Find topics common to primary and benchmark curriculum.
    """

    topics = set()


    if not isinstance(
        comparison,
        dict,
    ):

        return []


    for module in comparison.get(
        "module_comparisons",
        [],
    ):


        if not isinstance(
            module,
            dict,
        ):

            continue


        values = module.get(
            "similar_concepts",
            [],
        )


        if not isinstance(
            values,
            list,
        ):

            continue


        for value in values:

            value = str(
                value
            ).strip()


            if value:

                topics.add(
                    value
                )


    return sorted(
        topics,
        key=lambda x: x.lower(),
    )


# ============================================================
# 18. GET ALL COMPARISON TOPICS
# ============================================================

def get_all_comparison_topics(
    comparison: dict,
) -> list[str]:
    """
    Return all topics appearing in comparison.
    """

    topics = set()


    if not isinstance(
        comparison,
        dict,
    ):

        return []


    for module in comparison.get(
        "module_comparisons",
        [],
    ):


        if not isinstance(
            module,
            dict,
        ):

            continue


        for field in [

            "similar_concepts",

            "primary_only",

            "benchmark_only",

        ]:


            values = module.get(
                field,
                [],
            )


            if not isinstance(
                values,
                list,
            ):

                continue


            for value in values:

                value = str(
                    value
                ).strip()


                if value:

                    topics.add(
                        value
                    )


    return sorted(
        topics,
        key=lambda x: x.lower(),
    )


# ============================================================
# 19. JSON SERIALIZATION HELPER
# ============================================================

def serialize_json(
    data: Any,
) -> str:
    """
    Convert Python object to readable JSON.
    """

    try:

        return json.dumps(

            data,

            indent=2,

            ensure_ascii=False,

            default=str,

        )

    except Exception:

        return json.dumps(

            {
                "error":
                    "Unable to serialize data."
            },

            indent=2,

        )


# ============================================================
# 20. BUILD DOWNLOAD DATA
# ============================================================

def build_download_data(
    report: dict,
) -> bytes:
    """
    Convert report to UTF-8 bytes for Streamlit download.
    """

    json_string = serialize_json(
        report
    )

    return json_string.encode(
        "utf-8"
    )

# ============================================================
# CHUNK 3/10
# PAGE HEADER + PRIMARY SYLLABUS
# ============================================================


# ============================================================
# PAGE HEADER
# ============================================================

st.title(
    "📚 Curriculum Intelligence"
)


st.markdown(
    """
## AI-Powered Curriculum / Syllabus Intelligence

Compare your **primary university syllabus** with similar
or benchmark curricula and identify:

- Curriculum similarity
- Module-wise differences
- Topic-wise differences
- Concepts present / missing
- Advanced concepts
- Industry-relevant concepts
- Emerging concepts
- Tools and technologies
- Project opportunities
- Curriculum enhancement opportunities
"""
)


st.divider()


# ============================================================
# PRIMARY SYLLABUS VALIDATION
# ============================================================

primary = st.session_state.get(
    "primary_syllabus"
)


if not primary:

    st.error(
        """
        ❌ **Primary syllabus not found.**

        Please go to:

        **📥 Extract Syllabus**

        Upload the university syllabus and click:

        **🚀 Extract Complete Syllabus**

        Then return to this page.
        """
    )


    st.stop()


# ============================================================
# PRIMARY FILE INFORMATION
# ============================================================

primary_filename = st.session_state.get(
    "primary_filename",
    "",
)


primary_text = st.session_state.get(
    "primary_syllabus_text",
    "",
)


# ============================================================
# PRIMARY CURRICULUM HEADER
# ============================================================

st.subheader(
    "1️⃣ Primary Curriculum"
)


if primary_filename:

    st.caption(
        f"Source Document: **{primary_filename}**"
    )


# ============================================================
# INSTITUTION INFORMATION
# ============================================================

st.markdown(
    "### 🏫 Institution Information"
)


institution_col1, institution_col2, institution_col3, institution_col4 = (
    st.columns(4)
)


with institution_col1:

    st.metric(
        "College",
        safe_text(
            primary.get(
                "college"
            )
        ),
    )


with institution_col2:

    st.metric(
        "University",
        safe_text(
            primary.get(
                "university"
            )
        ),
    )


with institution_col3:

    st.metric(
        "Program",
        safe_text(
            primary.get(
                "program"
            )
        ),
    )


with institution_col4:

    st.metric(
        "Department",
        safe_text(
            primary.get(
                "department"
            )
        ),
    )


# ============================================================
# ACADEMIC INFORMATION
# ============================================================

st.markdown(
    "### 🎓 Academic Information"
)


academic_col1, academic_col2, academic_col3, academic_col4 = (
    st.columns(4)
)


with academic_col1:

    st.metric(
        "Academic Year",
        safe_text(
            primary.get(
                "academic_year"
            )
        ),
    )


with academic_col2:

    st.metric(
        "Regulation",
        safe_text(
            primary.get(
                "regulation"
            )
        ),
    )


with academic_col3:

    st.metric(
        "Semester",
        safe_text(
            primary.get(
                "semester"
            )
        ),
    )


with academic_col4:

    st.metric(
        "Category",
        safe_text(
            primary.get(
                "category"
            )
        ),
    )


# ============================================================
# SUBJECT INFORMATION
# ============================================================

st.divider()


st.markdown(
    "### 📘 Subject Information"
)


subject_col1, subject_col2, subject_col3, subject_col4 = (
    st.columns(4)
)


with subject_col1:

    st.metric(
        "Subject Name",
        safe_text(
            primary.get(
                "subject_name"
            )
        ),
    )


with subject_col2:

    st.metric(
        "Subject Code",
        safe_text(
            primary.get(
                "subject_code"
            )
        ),
    )


with subject_col3:

    st.metric(
        "Credits",
        safe_text(
            primary.get(
                "credits"
            )
        ),
    )


with subject_col4:

    st.metric(
        "Contact Hours",
        safe_text(
            primary.get(
                "contact_hours"
            )
        ),
    )


# ============================================================
# COURSE TYPE / L-T-P
# ============================================================

st.markdown(
    "### 📋 Course Structure"
)


structure_col1, structure_col2, structure_col3, structure_col4 = (
    st.columns(4)
)


with structure_col1:

    st.metric(
        "Course Type",
        safe_text(
            primary.get(
                "course_type"
            )
        ),
    )


with structure_col2:

    st.metric(
        "Lecture Hours",
        safe_text(
            primary.get(
                "lecture_hours"
            )
        ),
    )


with structure_col3:

    st.metric(
        "Tutorial Hours",
        safe_text(
            primary.get(
                "tutorial_hours"
            )
        ),
    )


with structure_col4:

    st.metric(
        "Practical Hours",
        safe_text(
            primary.get(
                "practical_hours"
            )
        ),
    )


# ============================================================
# PRIMARY CURRICULUM COUNTS
# ============================================================

st.divider()


st.markdown(
    "### 📊 Curriculum Structure Summary"
)


primary_modules = primary.get(
    "modules",
    [],
)


primary_topics = get_all_topics(
    primary
)


primary_module_names = get_all_modules(
    primary
)


primary_course_outcomes = primary.get(
    "course_outcomes",
    [],
)


primary_program_outcomes = primary.get(
    "program_outcomes",
    [],
)


primary_psos = primary.get(
    "program_specific_outcomes",
    [],
)


primary_tools = primary.get(
    "tools",
    [],
)


primary_projects = primary.get(
    "projects_case_studies",
    [],
)


# Ensure list structures

if not isinstance(
    primary_modules,
    list,
):

    primary_modules = []


if not isinstance(
    primary_topics,
    list,
):

    primary_topics = []


if not isinstance(
    primary_course_outcomes,
    list,
):

    primary_course_outcomes = []


if not isinstance(
    primary_program_outcomes,
    list,
):

    primary_program_outcomes = []


if not isinstance(
    primary_psos,
    list,
):

    primary_psos = []


if not isinstance(
    primary_tools,
    list,
):

    primary_tools = []


if not isinstance(
    primary_projects,
    list,
):

    primary_projects = []


# ============================================================
# SUMMARY METRICS
# ============================================================

summary_col1, summary_col2, summary_col3, summary_col4 = (
    st.columns(4)
)


with summary_col1:

    st.metric(
        "Modules",
        len(
            primary_modules
        ),
    )


with summary_col2:

    st.metric(
        "Topics",
        len(
            primary_topics
        ),
    )


with summary_col3:

    st.metric(
        "Course Outcomes",
        len(
            primary_course_outcomes
        ),
    )


with summary_col4:

    st.metric(
        "Tools / Technologies",
        len(
            primary_tools
        ),
    )


# ============================================================
# SECONDARY SUMMARY
# ============================================================

secondary_col1, secondary_col2, secondary_col3, secondary_col4 = (
    st.columns(4)
)


with secondary_col1:

    st.metric(
        "Program Outcomes",
        len(
            primary_program_outcomes
        ),
    )


with secondary_col2:

    st.metric(
        "PSOs",
        len(
            primary_psos
        ),
    )


with secondary_col3:

    st.metric(
        "Projects / Case Studies",
        len(
            primary_projects
        ),
    )


with secondary_col4:

    text_length = len(
        primary_text
    )


    st.metric(
        "Extracted Characters",
        f"{text_length:,}",
    )


# ============================================================
# COURSE DESCRIPTION
# ============================================================

st.divider()


st.subheader(
    "📖 Course Description"
)


course_description = primary.get(
    "course_description",
    "",
)


if course_description:

    st.write(
        course_description
    )

else:

    st.info(
        "Course description was not detected "
        "in the syllabus."
    )


# ============================================================
# COURSE OBJECTIVES
# ============================================================

st.divider()


st.subheader(
    "🎯 Course Objectives"
)


course_objectives = primary.get(
    "course_objectives",
    [],
)


if course_objectives:

    for index, objective in enumerate(
        course_objectives,
        start=1,
    ):

        if isinstance(
            objective,
            dict,
        ):

            objective_text = (

                objective.get(
                    "description"
                )

                or

                objective.get(
                    "text"
                )

                or

                objective.get(
                    "name"
                )

                or

                str(
                    objective
                )

            )

        else:

            objective_text = str(
                objective
            )


        st.markdown(
            f"**COBJ-{index}:** "
            f"{objective_text}"
        )

else:

    st.warning(
        "No course objectives were detected."
    )


# ============================================================
# PREREQUISITES
# ============================================================

st.divider()


st.subheader(
    "📌 Prerequisites"
)


prerequisites = primary.get(
    "prerequisites",
    [],
)


if prerequisites:

    display_value_list(
        prerequisites,
        "No prerequisites detected.",
    )

else:

    st.info(
        "No prerequisites were specified "
        "or detected."
    )


# ============================================================
# MODULES & TOPICS
# ============================================================

st.divider()


st.subheader(
    "2️⃣ Curriculum Modules & Topics"
)


if not primary_modules:

    st.error(
        """
        ❌ No modules were detected in the primary syllabus.

        Please verify the extraction result on:

        **📥 Extract Syllabus**
        """
    )

else:

    st.caption(
        f"{len(primary_modules)} modules and "
        f"{len(primary_topics)} unique topics detected."
    )


    # --------------------------------------------------------
    # MODULE LOOP
    # --------------------------------------------------------

    for module_index, module in enumerate(
        primary_modules,
        start=1,
    ):


        if not isinstance(
            module,
            dict,
        ):

            module = {

                "name":
                    str(
                        module
                    ),

                "topics":
                    [],

            }


        module_name = safe_text(

            module.get(
                "name"
            ),

            f"Module {module_index}",

        )


        module_hours = module.get(
            "hours"
        )


        module_code = module.get(
            "code"
        )


        # ----------------------------------------------------
        # Build title
        # ----------------------------------------------------

        module_title = (
            f"Module {module_index}: "
            f"{module_name}"
        )


        if module_code:

            module_title += (
                f" [{module_code}]"
            )


        if module_hours:

            module_title += (
                f" — {module_hours} Hours"
            )


        # ----------------------------------------------------
        # Module Expander
        # ----------------------------------------------------

        with st.expander(
            module_title,
            expanded=False,
        ):


            # ------------------------------------------------
            # Module Description
            # ------------------------------------------------

            module_description = module.get(
                "description",
                "",
            )


            if module_description:

                st.markdown(
                    "#### 📝 Module Description"
                )


                st.write(
                    module_description
                )


            # ------------------------------------------------
            # Module Hours
            # ------------------------------------------------

            hours_col1, hours_col2, hours_col3 = (
                st.columns(3)
            )


            with hours_col1:

                st.metric(
                    "Module",
                    module_index,
                )


            with hours_col2:

                st.metric(
                    "Hours",
                    safe_text(
                        module.get(
                            "hours"
                        ),
                        "N/A",
                    ),
                )


            with hours_col3:

                module_topics = (
                    get_topics_from_module(
                        module
                    )
                )


                st.metric(
                    "Topics",
                    len(
                        module_topics
                    ),
                )


            st.divider()


            # ------------------------------------------------
            # Topics
            # ------------------------------------------------

            st.markdown(
                "#### 📚 Topics"
            )


            if not module_topics:

                st.warning(
                    "No topics detected for this module."
                )

            else:

                for topic_index, topic in enumerate(
                    module_topics,
                    start=1,
                ):

                    st.markdown(
                        f"**{module_index}."
                        f"{topic_index}** "
                        f"{topic}"
                    )


            # ------------------------------------------------
            # Module Concepts
            # ------------------------------------------------

            module_concepts = module.get(
                "concepts",
                [],
            )


            if module_concepts:

                st.markdown(
                    "#### 🧠 Concepts"
                )


                display_value_list(
                    module_concepts
                )


            # ------------------------------------------------
            # Module Tools
            # ------------------------------------------------

            module_tools = module.get(
                "tools",
                [],
            )


            if module_tools:

                st.markdown(
                    "#### 🛠 Tools / Technologies"
                )


                display_topic_list(
                    module_tools,
                    status="success",
                )


            # ------------------------------------------------
            # Module Projects
            # ------------------------------------------------

            module_projects = module.get(
                "projects",
                [],
            )


            if module_projects:

                st.markdown(
                    "#### 🚀 Projects / Applications"
                )


                display_value_list(
                    module_projects
                )


# ============================================================
# COURSE OUTCOMES
# ============================================================

st.divider()


st.subheader(
    "3️⃣ Course Outcomes — CO"
)


if not primary_course_outcomes:

    st.warning(
        "No Course Outcomes were detected."
    )

else:

    for index, outcome in enumerate(
        primary_course_outcomes,
        start=1,
    ):


        # ----------------------------------------------------
        # Support dictionary or string
        # ----------------------------------------------------

        if isinstance(
            outcome,
            dict,
        ):

            co_code = (

                outcome.get(
                    "code"
                )

                or

                f"CO{index}"

            )


            outcome_text = (

                outcome.get(
                    "description"
                )

                or

                outcome.get(
                    "text"
                )

                or

                outcome.get(
                    "outcome"
                )

                or

                str(
                    outcome
                )

            )

        else:

            co_code = f"CO{index}"

            outcome_text = str(
                outcome
            )


        st.markdown(
            f"**{co_code}:** {outcome_text}"
        )


# ============================================================
# PROGRAM OUTCOMES
# ============================================================

st.divider()


st.subheader(
    "4️⃣ Program Outcomes — PO"
)


if primary_program_outcomes:

    display_value_list(
        primary_program_outcomes
    )

else:

    st.info(
        "Program Outcomes were not detected "
        "in this syllabus."
    )


# ============================================================
# PROGRAM SPECIFIC OUTCOMES
# ============================================================

st.divider()


st.subheader(
    "5️⃣ Program Specific Outcomes — PSO"
)


if primary_psos:

    display_value_list(
        primary_psos
    )

else:

    st.info(
        "Program Specific Outcomes were not detected."
    )


# ============================================================
# CO-PO MAPPING
# ============================================================

st.divider()


st.subheader(
    "6️⃣ CO-PO / CO-PSO Mapping"
)


co_po_mapping = primary.get(
    "co_po_mapping"
)


if co_po_mapping:

    if isinstance(
        co_po_mapping,
        dict,
    ):

        st.json(
            co_po_mapping
        )

    elif isinstance(
        co_po_mapping,
        list,
    ):

        mapping_df = pd.DataFrame(
            co_po_mapping
        )


        if not mapping_df.empty:

            st.dataframe(
                mapping_df,
                use_container_width=True,
                hide_index=True,
            )

        else:

            st.json(
                co_po_mapping
            )

    else:

        st.write(
            co_po_mapping
        )

else:

    st.info(
        "CO-PO / CO-PSO mapping was not detected."
    )


# ============================================================
# TEACHING & LEARNING METHODS
# ============================================================

st.divider()


st.subheader(
    "7️⃣ Teaching & Learning Methods"
)


teaching_methods = primary.get(
    "teaching_learning_methods",
    [],
)


if teaching_methods:

    display_value_list(
        teaching_methods
    )

else:

    st.info(
        "Teaching and learning methods "
        "were not detected."
    )


# ============================================================
# PRACTICAL / LAB COMPONENTS
# ============================================================

st.divider()


st.subheader(
    "8️⃣ Practical / Laboratory Components"
)


practical_components = primary.get(
    "practical_components",
    [],
)


if practical_components:

    display_value_list(
        practical_components
    )

else:

    st.info(
        "No practical or laboratory components "
        "were detected."
    )


# ============================================================
# ASSESSMENT PATTERN
# ============================================================

st.divider()


st.subheader(
    "9️⃣ Assessment Pattern"
)


assessment_pattern = primary.get(
    "assessment_pattern"
)


if assessment_pattern:

    if isinstance(
        assessment_pattern,
        dict,
    ):

        assessment_df = pd.DataFrame(
            [
                assessment_pattern
            ]
        )


        st.dataframe(
            assessment_df,
            use_container_width=True,
            hide_index=True,
        )


    elif isinstance(
        assessment_pattern,
        list,
    ):

        display_value_list(
            assessment_pattern
        )


    else:

        st.write(
            assessment_pattern
        )

else:

    st.info(
        "Assessment pattern was not detected."
    )


# ============================================================
# TEXTBOOKS
# ============================================================

st.divider()


st.subheader(
    "🔟 Textbooks"
)


textbooks = primary.get(
    "textbooks",
    [],
)


if textbooks:

    display_value_list(
        textbooks
    )

else:

    st.info(
        "No textbooks were detected."
    )


# ============================================================
# REFERENCE BOOKS
# ============================================================

st.divider()


st.subheader(
    "📚 Reference Books"
)


reference_books = primary.get(
    "reference_books",
    [],
)


if reference_books:

    display_value_list(
        reference_books
    )

else:

    st.info(
        "No reference books were detected."
    )


# ============================================================
# ONLINE RESOURCES
# ============================================================

st.divider()


st.subheader(
    "🌐 Online Resources"
)


online_resources = primary.get(
    "online_resources",
    [],
)


if online_resources:

    display_value_list(
        online_resources
    )

else:

    st.info(
        "No online resources were detected."
    )


# ============================================================
# TOOLS & TECHNOLOGIES
# ============================================================

st.divider()


st.subheader(
    "🛠 Tools & Technologies"
)


if primary_tools:

    tool_columns = st.columns(
        min(
            max(
                len(
                    primary_tools
                ),
                1,
            ),
            4,
        )
    )


    for index, tool in enumerate(
        primary_tools
    ):

        column = tool_columns[
            index
            % len(
                tool_columns
            )
        ]


        with column:

            st.success(
                str(
                    tool
                )
            )

else:

    st.info(
        "No tools or technologies were detected."
    )


# ============================================================
# DATASETS
# ============================================================

st.divider()


st.subheader(
    "🗃️ Datasets"
)


datasets = primary.get(
    "datasets",
    [],
)


if datasets:

    display_value_list(
        datasets
    )

else:

    st.info(
        "No datasets were detected."
    )


# ============================================================
# PROJECTS & CASE STUDIES
# ============================================================

st.divider()


st.subheader(
    "🚀 Projects & Case Studies"
)


if primary_projects:

    display_value_list(
        primary_projects
    )

else:

    st.info(
        "No projects or case studies were detected."
    )


# ============================================================
# RAW EXTRACTED SYLLABUS TEXT
# ============================================================

st.divider()


st.subheader(
    "📄 Raw Extracted Syllabus Text"
)


with st.expander(
    "View extracted document text",
    expanded=False,
):

    if primary_text:

        if len(
            primary_text
        ) > MAX_TEXT_PREVIEW:

            st.text(
                primary_text[
                    :MAX_TEXT_PREVIEW
                ]
            )


            st.caption(
                f"Showing first "
                f"{MAX_TEXT_PREVIEW:,} characters "
                f"of {len(primary_text):,}."
            )

        else:

            st.text(
                primary_text
            )

    else:

        st.info(
            "Raw extracted syllabus text is not available."
        )


# ============================================================
# PRIMARY CURRICULUM READY STATUS
# ============================================================

st.divider()


st.success(
    """
    ✅ **Primary curriculum loaded successfully.**

    The curriculum is ready for:

    **Benchmark Curriculum → Module Comparison
    → Topic Comparison → Concept Intelligence**
    """
)


# ============================================================
# DEBUG / STRUCTURED DATA
# ============================================================

with st.expander(
    "🔧 Developer: View Primary Curriculum JSON",
    expanded=False,
):

    st.json(
        primary
    )


# ============================================================
# CHUNK 4/10
# BENCHMARK SYLLABUS UPLOAD & EXTRACTION
# ============================================================


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.subheader(
    "🔎 2️⃣ Upload Similar / Benchmark Curriculum"
)

st.markdown(
    """
Upload another curriculum / syllabus that should be compared
with the primary syllabus.

### Recommended benchmark sources

- Same subject from another university
- Same subject from another college
- New regulation vs old regulation
- Autonomous university curriculum
- Industry-oriented curriculum
- International university curriculum
- Professional certification curriculum
- MOOC / online curriculum
"""
)


# ============================================================
# BENCHMARK FILE UPLOAD
# ============================================================

benchmark_file = st.file_uploader(

    "📤 Upload Benchmark / Similar Syllabus",

    type=SUPPORTED_FILE_TYPES,

    key="curriculum_benchmark_file",

    help=(
        "Supported formats: PDF, DOCX, PNG, JPG and JPEG. "
        "Scanned documents can be processed using OCR."
    ),

)


# ============================================================
# NO FILE SELECTED
# ============================================================

if benchmark_file is None:

    st.info(
        """
        👆 Upload a benchmark syllabus to begin
        curriculum comparison.
        """
    )


# ============================================================
# FILE SELECTED
# ============================================================

else:

    # --------------------------------------------------------
    # FILE INFORMATION
    # --------------------------------------------------------

    st.success(
        f"Selected benchmark file: "
        f"**{benchmark_file.name}**"
    )


    file_name = benchmark_file.name


    file_suffix = (
        file_name
        .split(".")[-1]
        .upper()
        if "." in file_name
        else "UNKNOWN"
    )


    try:

        file_size_mb = (
            len(
                benchmark_file.getvalue()
            )
            / (
                1024 * 1024
            )
        )

    except Exception:

        file_size_mb = 0


    file_col1, file_col2, file_col3 = (
        st.columns(3)
    )


    with file_col1:

        st.metric(
            "File Type",
            file_suffix,
        )


    with file_col2:

        st.metric(
            "File Size",
            f"{file_size_mb:.2f} MB",
        )


    with file_col3:

        st.metric(
            "Supported",
            "Yes",
        )


    # --------------------------------------------------------
    # EXTRACTION INFORMATION
    # --------------------------------------------------------

    st.caption(
        """
        The document will be processed in the following
        sequence:

        Document → PDF/DOCX/OCR → Raw Text →
        LLM Structured Curriculum → Comparison
        """
    )


    # ========================================================
    # EXTRACTION BUTTON
    # ========================================================

    extract_benchmark = st.button(

        "🚀 Extract & Compare Benchmark Curriculum",

        type="primary",

        use_container_width=True,

        key="extract_benchmark_curriculum_button",

    )


    # ========================================================
    # PROCESS BENCHMARK
    # ========================================================

    if extract_benchmark:

        progress = st.progress(
            0
        )

        status = st.empty()


        try:

            # =================================================
            # STEP 1
            # DOCUMENT READING
            # =================================================

            status.info(
                "📄 Step 1/5 — Reading benchmark document..."
            )


            progress.progress(
                10
            )


            benchmark_pages, benchmark_text = (
                extract_document_text(
                    benchmark_file
                )
            )


            if not benchmark_text.strip():

                raise ValueError(
                    "No readable text was extracted "
                    "from the benchmark document."
                )


            # =================================================
            # STEP 2
            # DOCUMENT STATISTICS
            # =================================================

            status.info(
                "📑 Step 2/5 — Processing extracted pages..."
            )


            progress.progress(
                25
            )


            page_count = len(
                benchmark_pages
            )


            character_count = len(
                benchmark_text
            )


            word_count = len(
                benchmark_text.split()
            )


            # =================================================
            # STEP 3
            # LLM CURRICULUM EXTRACTION
            # =================================================

            status.info(
                "🧠 Step 3/5 — AI is extracting "
                "the benchmark curriculum structure..."
            )


            progress.progress(
                45
            )


            benchmark = extract_syllabus(
                benchmark_text
            )


            # =================================================
            # VALIDATE LLM RESULT
            # =================================================

            if not isinstance(
                benchmark,
                dict,
            ):

                raise ValueError(
                    "The syllabus extraction engine "
                    "returned an invalid curriculum object."
                )


            benchmark_modules = benchmark.get(
                "modules",
                [],
            )


            if not isinstance(
                benchmark_modules,
                list,
            ):

                benchmark_modules = []


            benchmark_topics = get_all_topics(
                benchmark
            )


            # =================================================
            # STEP 4
            # CURRICULUM COMPARISON
            # =================================================

            status.info(
                "🔍 Step 4/5 — Comparing primary and "
                "benchmark curriculum..."
            )


            progress.progress(
                70
            )


            comparison = compare_curricula(

                primary,

                benchmark,

            )


            # =================================================
            # VALIDATE COMPARISON
            # =================================================

            if not isinstance(
                comparison,
                dict,
            ):

                raise ValueError(
                    "Curriculum comparator returned "
                    "an invalid result."
                )


            # =================================================
            # STEP 5
            # SAVE RESULTS
            # =================================================

            status.info(
                "💾 Step 5/5 — Saving benchmark "
                "curriculum and comparison..."
            )


            progress.progress(
                90
            )


            # -------------------------------------------------
            # Store benchmark syllabus
            # -------------------------------------------------

            st.session_state[
                "benchmark_syllabus"
            ] = benchmark


            # -------------------------------------------------
            # Store filename
            # -------------------------------------------------

            st.session_state[
                "benchmark_filename"
            ] = benchmark_file.name


            # -------------------------------------------------
            # Store raw text
            # -------------------------------------------------

            st.session_state[
                "benchmark_text"
            ] = benchmark_text


            # -------------------------------------------------
            # Store pages + complete source information
            # -------------------------------------------------

            st.session_state[
                "benchmark_syllabi"
            ] = [

                {

                    "filename":
                        benchmark_file.name,

                    "syllabus":
                        benchmark,

                    "text":
                        benchmark_text,

                    "pages":
                        benchmark_pages,

                    "page_count":
                        page_count,

                    "character_count":
                        character_count,

                    "word_count":
                        word_count,

                }

            ]


            # -------------------------------------------------
            # Store comparison
            # -------------------------------------------------

            st.session_state[
                "curriculum_comparison"
            ] = comparison


            # -------------------------------------------------
            # Reset concept intelligence
            #
            # Because benchmark curriculum has changed,
            # previous concept analysis is no longer valid.
            # -------------------------------------------------

            st.session_state[
                "concept_enrichment"
            ] = []


            st.session_state[
                "concept_analysis_complete"
            ] = False


            # -------------------------------------------------
            # Reset report
            # -------------------------------------------------

            st.session_state[
                "curriculum_intelligence_report"
            ] = None


            # -------------------------------------------------
            # Timestamp
            # -------------------------------------------------

            st.session_state[
                "curriculum_intelligence_timestamp"
            ] = datetime.now().isoformat()


            progress.progress(
                100
            )


            status.success(
                """
                ✅ Benchmark syllabus extracted and
                curriculum comparison completed successfully.
                """
            )


            # =================================================
            # EXTRACTION SUMMARY
            # =================================================

            st.divider()

            st.subheader(
                "📊 Benchmark Extraction Summary"
            )


            summary_col1, summary_col2, summary_col3, summary_col4 = (
                st.columns(4)
            )


            with summary_col1:

                st.metric(
                    "Pages",
                    page_count,
                )


            with summary_col2:

                st.metric(
                    "Words",
                    f"{word_count:,}",
                )


            with summary_col3:

                st.metric(
                    "Modules",
                    len(
                        benchmark_modules
                    ),
                )


            with summary_col4:

                st.metric(
                    "Topics",
                    len(
                        benchmark_topics
                    ),
                )


            # =================================================
            # BENCHMARK IDENTIFICATION
            # =================================================

            st.divider()

            st.subheader(
                "🏫 Extracted Benchmark Curriculum"
            )


            benchmark_col1, benchmark_col2, benchmark_col3, benchmark_col4 = (
                st.columns(4)
            )


            with benchmark_col1:

                st.metric(
                    "College",
                    safe_text(
                        benchmark.get(
                            "college"
                        )
                    ),
                )


            with benchmark_col2:

                st.metric(
                    "University",
                    safe_text(
                        benchmark.get(
                            "university"
                        )
                    ),
                )


            with benchmark_col3:

                st.metric(
                    "Subject",
                    safe_text(
                        benchmark.get(
                            "subject_name"
                        )
                    ),
                )


            with benchmark_col4:

                st.metric(
                    "Subject Code",
                    safe_text(
                        benchmark.get(
                            "subject_code"
                        )
                    ),
                )


            # =================================================
            # BENCHMARK ACADEMIC INFORMATION
            # =================================================

            st.markdown(
                "### 🎓 Academic Information"
            )


            academic_col1, academic_col2, academic_col3, academic_col4 = (
                st.columns(4)
            )


            with academic_col1:

                st.metric(
                    "Program",
                    safe_text(
                        benchmark.get(
                            "program"
                        )
                    ),
                )


            with academic_col2:

                st.metric(
                    "Department",
                    safe_text(
                        benchmark.get(
                            "department"
                        )
                    ),
                )


            with academic_col3:

                st.metric(
                    "Semester",
                    safe_text(
                        benchmark.get(
                            "semester"
                        )
                    ),
                )


            with academic_col4:

                st.metric(
                    "Regulation",
                    safe_text(
                        benchmark.get(
                            "regulation"
                        )
                    ),
                )


            # =================================================
            # BENCHMARK MODULE PREVIEW
            # =================================================

            st.markdown(
                "### 📚 Benchmark Modules"
            )


            if benchmark_modules:

                for module_index, module in enumerate(
                    benchmark_modules,
                    start=1,
                ):


                    if not isinstance(
                        module,
                        dict,
                    ):

                        module = {

                            "name":
                                str(
                                    module
                                ),

                            "topics":
                                [],

                        }


                    module_name = safe_text(

                        module.get(
                            "name"
                        ),

                        f"Module {module_index}",

                    )


                    module_hours = module.get(
                        "hours"
                    )


                    module_topics = (
                        get_topics_from_module(
                            module
                        )
                    )


                    if module_hours:

                        module_title = (

                            f"Module "
                            f"{module_index}: "
                            f"{module_name} "
                            f"— {module_hours} Hours"

                        )

                    else:

                        module_title = (

                            f"Module "
                            f"{module_index}: "
                            f"{module_name}"

                        )


                    with st.expander(
                        module_title,
                        expanded=False,
                    ):


                        st.write(
                            f"**Topics:** "
                            f"{len(module_topics)}"
                        )


                        if module_topics:

                            for topic_index, topic in enumerate(
                                module_topics,
                                start=1,
                            ):

                                st.markdown(

                                    f"**{module_index}."
                                    f"{topic_index}** "
                                    f"{topic}"

                                )

                        else:

                            st.info(
                                "No topics detected "
                                "for this module."
                            )


            else:

                st.warning(
                    "No modules were detected "
                    "in the benchmark curriculum."
                )


            # =================================================
            # EXTRACTION QUALITY WARNING
            # =================================================

            extraction_warnings = []


            if not benchmark.get(
                "subject_name"
            ):

                extraction_warnings.append(
                    "Subject name was not confidently extracted."
                )


            if not benchmark.get(
                "subject_code"
            ):

                extraction_warnings.append(
                    "Subject code was not confidently extracted."
                )


            if not benchmark.get(
                "university"
            ):

                extraction_warnings.append(
                    "University was not confidently extracted."
                )


            if not benchmark_modules:

                extraction_warnings.append(
                    "No modules were detected."
                )


            if extraction_warnings:

                st.divider()

                st.warning(
                    """
                    ⚠️ **Benchmark extraction requires review.**

                    The AI extraction completed, but the following
                    information may need manual verification:
                    """
                )


                for warning in extraction_warnings:

                    st.write(
                        f"• {warning}"
                    )


            else:

                st.success(
                    "✅ Benchmark curriculum structure "
                    "appears complete."
                )


        # =====================================================
        # ERROR HANDLING
        # =====================================================

        except Exception as exc:

            progress.empty()


            status.error(
                "❌ Benchmark curriculum processing failed."
            )


            st.error(
                str(
                    exc
                )
            )


            with st.expander(
                "🔧 Technical Error Details",
                expanded=False,
            ):

                st.exception(
                    exc
                )


# ============================================================
# LOAD BENCHMARK FROM SESSION STATE
# ============================================================

benchmark = st.session_state.get(
    "benchmark_syllabus"
)


# ============================================================
# EXISTING BENCHMARK STATUS
# ============================================================

if benchmark:

    st.divider()

    st.subheader(
        "✅ Benchmark Curriculum Loaded"
    )


    stored_filename = st.session_state.get(
        "benchmark_filename",
        "Unknown",
    )


    benchmark_topics = get_all_topics(
        benchmark
    )


    benchmark_modules = benchmark.get(
        "modules",
        [],
    )


    if not isinstance(
        benchmark_modules,
        list,
    ):

        benchmark_modules = []


    loaded_col1, loaded_col2, loaded_col3, loaded_col4 = (
        st.columns(4)
    )


    with loaded_col1:

        st.metric(
            "Benchmark Modules",
            len(
                benchmark_modules
            ),
        )


    with loaded_col2:

        st.metric(
            "Benchmark Topics",
            len(
                benchmark_topics
            ),
        )


    with loaded_col3:

        st.metric(
            "Benchmark COs",
            len(
                benchmark.get(
                    "course_outcomes",
                    [],
                )
            ),
        )


    with loaded_col4:

        st.metric(
            "Benchmark Tools",
            len(
                benchmark.get(
                    "tools",
                    [],
                )
            ),
        )


    st.caption(
        f"Loaded from: **{stored_filename}**"
    )


# ============================================================
# BENCHMARK RAW TEXT
# ============================================================

if benchmark:

    with st.expander(
        "📄 View Raw Benchmark Text",
        expanded=False,
    ):

        stored_benchmark_text = (
            st.session_state.get(
                "benchmark_text",
                "",
            )
        )


        if stored_benchmark_text:

            if len(
                stored_benchmark_text
            ) > MAX_TEXT_PREVIEW:

                st.text(
                    stored_benchmark_text[
                        :MAX_TEXT_PREVIEW
                    ]
                )


                st.caption(
                    f"Preview limited to "
                    f"{MAX_TEXT_PREVIEW:,} characters."
                )

            else:

                st.text(
                    stored_benchmark_text
                )

        else:

            st.info(
                "No benchmark raw text available."
            )


# ============================================================
# BENCHMARK JSON PREVIEW
# ============================================================

if benchmark:

    with st.expander(
        "🔧 Developer: View Benchmark Curriculum JSON",
        expanded=False,
    ):

        st.json(
            benchmark
        )

# ============================================================
# CHUNK 5/10
# OVERALL + MODULE-WISE CURRICULUM COMPARISON
# ============================================================


# ============================================================
# LOAD COMPARISON RESULT
# ============================================================

comparison = st.session_state.get(
    "curriculum_comparison"
)


# ============================================================
# VALIDATE COMPARISON
# ============================================================

if not comparison:

    st.warning(
        """
        ⚠️ Curriculum comparison is not available yet.

        Upload and process a benchmark syllabus above.
        """
    )

else:

    # ========================================================
    # SECTION HEADER
    # ========================================================

    st.divider()

    st.subheader(
        "3️⃣ Overall Curriculum Comparison"
    )


    st.markdown(
        """
        The comparison engine evaluates the primary curriculum
        against the benchmark curriculum across:

        **Modules → Topics → Concepts → Similarity → Differences**
        """
    )


    # ========================================================
    # NORMALIZED SUMMARY
    # ========================================================

    comparison_summary = get_comparison_summary(
        comparison
    )


    similarity_pct = comparison_summary[
        "similarity_pct"
    ]


    modules_compared = comparison_summary[
        "modules_compared"
    ]


    similar_topics_count = comparison_summary[
        "similar_topics"
    ]


    different_topics_count = comparison_summary[
        "different_topics"
    ]


    # ========================================================
    # OVERALL METRICS
    # ========================================================

    metric_col1, metric_col2, metric_col3, metric_col4 = (
        st.columns(4)
    )


    with metric_col1:

        st.metric(
            "Overall Similarity",
            f"{similarity_pct:.1f}%",
        )


    with metric_col2:

        st.metric(
            "Modules Compared",
            modules_compared,
        )


    with metric_col3:

        st.metric(
            "Similar Concepts",
            similar_topics_count,
        )


    with metric_col4:

        st.metric(
            "Different Concepts",
            different_topics_count,
        )


    # ========================================================
    # SIMILARITY PROGRESS BAR
    # ========================================================

    st.markdown(
        "### 📈 Curriculum Similarity"
    )


    similarity_for_progress = max(
        0.0,
        min(
            similarity_pct / 100,
            1.0,
        ),
    )


    st.progress(
        similarity_for_progress
    )


    # ========================================================
    # SIMILARITY INTERPRETATION
    # ========================================================

    if similarity_pct >= 85:

        st.success(
            """
            🟢 **Very High Similarity**

            The primary and benchmark curricula have strong
            structural and conceptual alignment.
            """
        )


    elif similarity_pct >= 70:

        st.success(
            """
            🟢 **High Similarity**

            The curricula are substantially aligned, with some
            opportunities for additional topics or modernization.
            """
        )


    elif similarity_pct >= 50:

        st.warning(
            """
            🟡 **Moderate Similarity**

            The curricula share important concepts but have
            meaningful differences that require deeper analysis.
            """
        )


    elif similarity_pct >= 30:

        st.warning(
            """
            🟠 **Low Similarity**

            Significant differences exist between the curricula.
            Benchmark topics should be investigated for relevance.
            """
        )


    else:

        st.error(
            """
            🔴 **Very Low Similarity**

            The benchmark curriculum is substantially different
            from the primary curriculum.
            """
        )


    # ========================================================
    # SUBJECT COMPARISON
    # ========================================================

    st.divider()

    st.subheader(
        "📘 Subject-level Comparison"
    )


    subject_col1, subject_col2 = (
        st.columns(2)
    )


    with subject_col1:

        st.markdown(
            "### 🟦 Primary Curriculum"
        )


        st.info(
            safe_text(
                comparison.get(
                    "primary_subject"
                )
                or primary.get(
                    "subject_name"
                ),
                "Primary Subject",
            )
        )


        st.caption(
            f"University: "
            f"{safe_text(primary.get('university'))}"
        )


    with subject_col2:

        st.markdown(
            "### 🟩 Benchmark Curriculum"
        )


        st.info(
            safe_text(
                comparison.get(
                    "benchmark_subject"
                )
                or benchmark.get(
                    "subject_name"
                ),
                "Benchmark Subject",
            )
        )


        st.caption(
            f"University: "
            f"{safe_text(benchmark.get('university'))}"
        )


    # ========================================================
    # STRUCTURAL COMPARISON
    # ========================================================

    st.divider()

    st.subheader(
        "🏗️ Curriculum Structure Comparison"
    )


    primary_modules = primary.get(
        "modules",
        [],
    )


    benchmark_modules = benchmark.get(
        "modules",
        [],
    )


    if not isinstance(
        primary_modules,
        list,
    ):

        primary_modules = []


    if not isinstance(
        benchmark_modules,
        list,
    ):

        benchmark_modules = []


    primary_topic_count = len(
        get_all_topics(
            primary
        )
    )


    benchmark_topic_count = len(
        get_all_topics(
            benchmark
        )
    )


    structure_df = pd.DataFrame({

        "Metric": [

            "Modules",

            "Topics",

            "Course Outcomes",

            "Program Outcomes",

            "PSOs",

            "Tools / Technologies",

            "Projects / Case Studies",

        ],

        "Primary Curriculum": [

            len(
                primary_modules
            ),

            primary_topic_count,

            len(
                primary.get(
                    "course_outcomes",
                    [],
                )
            ),

            len(
                primary.get(
                    "program_outcomes",
                    [],
                )
            ),

            len(
                primary.get(
                    "program_specific_outcomes",
                    [],
                )
            ),

            len(
                primary.get(
                    "tools",
                    [],
                )
            ),

            len(
                primary.get(
                    "projects_case_studies",
                    [],
                )
            ),

        ],

        "Benchmark Curriculum": [

            len(
                benchmark_modules
            ),

            benchmark_topic_count,

            len(
                benchmark.get(
                    "course_outcomes",
                    [],
                )
            ),

            len(
                benchmark.get(
                    "program_outcomes",
                    [],
                )
            ),

            len(
                benchmark.get(
                    "program_specific_outcomes",
                    [],
                )
            ),

            len(
                benchmark.get(
                    "tools",
                    [],
                )
            ),

            len(
                benchmark.get(
                    "projects_case_studies",
                    [],
                )
            ),

        ],

    })


    st.dataframe(

        structure_df,

        use_container_width=True,

        hide_index=True,

    )


    # ========================================================
    # STRUCTURAL DIFFERENCE
    # ========================================================

    st.markdown(
        "### 🔍 Structural Difference"
    )


    structure_difference_rows = []


    for _, row in structure_df.iterrows():

        primary_value = row[
            "Primary Curriculum"
        ]


        benchmark_value = row[
            "Benchmark Curriculum"
        ]


        difference = (
            benchmark_value
            - primary_value
        )


        if difference > 0:

            interpretation = (
                f"Benchmark has "
                f"{difference} more"
            )


        elif difference < 0:

            interpretation = (
                f"Primary has "
                f"{abs(difference)} more"
            )


        else:

            interpretation = (
                "Same"
            )


        structure_difference_rows.append({

            "Metric":
                row["Metric"],

            "Difference":
                difference,

            "Interpretation":
                interpretation,

        })


    structure_difference_df = pd.DataFrame(
        structure_difference_rows
    )


    st.dataframe(

        structure_difference_df,

        use_container_width=True,

        hide_index=True,

    )


    # ========================================================
    # MODULE-WISE COMPARISON
    # ========================================================

    st.divider()

    st.subheader(
        "4️⃣ Module-wise Comparison"
    )


    st.markdown(
        """
        Each primary module is compared with the most relevant
        benchmark module.

        The analysis identifies:

        🟢 Similar concepts  
        🔵 Concepts only in primary syllabus  
        🔴 Concepts only in benchmark syllabus
        """
    )


    module_comparisons = comparison.get(
        "module_comparisons",
        [],
    )


    if not isinstance(
        module_comparisons,
        list,
    ):

        module_comparisons = []


    # ========================================================
    # NO MODULE DATA
    # ========================================================

    if not module_comparisons:

        st.warning(
            """
            No module-level comparison was returned by the
            comparison engine.

            Check `curriculum/comparator.py`.
            """
        )


    else:

        # ====================================================
        # MODULE SUMMARY TABLE
        # ====================================================

        module_df = build_module_dataframe(
            comparison
        )


        if not module_df.empty:

            st.markdown(
                "### 📊 Module Comparison Summary"
            )


            st.dataframe(

                module_df,

                use_container_width=True,

                hide_index=True,

            )


            # ------------------------------------------------
            # MODULE CSV
            # ------------------------------------------------

            module_csv = module_df.to_csv(
                index=False
            )


            st.download_button(

                label=(
                    "⬇️ Download Module Comparison CSV"
                ),

                data=module_csv,

                file_name=(
                    "module_wise_curriculum_comparison.csv"
                ),

                mime="text/csv",

                key="download_module_comparison_csv",

            )


        # ====================================================
        # MODULE DETAILS
        # ====================================================

        st.markdown(
            "### 🔬 Detailed Module Analysis"
        )


        for module_index, module in enumerate(
            module_comparisons,
            start=1,
        ):


            if not isinstance(
                module,
                dict,
            ):

                continue


            primary_module_name = safe_text(

                module.get(
                    "primary_module"
                ),

                f"Module {module_index}",

            )


            benchmark_module_name = safe_text(

                module.get(
                    "benchmark_module"
                ),

                "No Matching Benchmark Module",

            )


            try:

                module_similarity = float(

                    module.get(
                        "similarity_pct",
                        0,
                    )

                    or 0

                )

            except (
                TypeError,
                ValueError,
            ):

                module_similarity = 0.0


            # ------------------------------------------------
            # Module Expander
            # ------------------------------------------------

            with st.expander(

                (
                    f"Module {module_index}: "
                    f"{primary_module_name} "
                    f" ↔ "
                    f"{benchmark_module_name} "
                    f" | "
                    f"{module_similarity:.1f}%"
                ),

                expanded=False,

            ):


                # ============================================
                # MODULE HEADER
                # ============================================

                module_header_col1, module_header_col2 = (
                    st.columns(2)
                )


                with module_header_col1:

                    st.markdown(
                        "### 🟦 Primary Module"
                    )


                    st.info(
                        primary_module_name
                    )


                with module_header_col2:

                    st.markdown(
                        "### 🟩 Benchmark Module"
                    )


                    st.info(
                        benchmark_module_name
                    )


                # ============================================
                # MODULE SIMILARITY
                # ============================================

                st.markdown(
                    "### 📈 Module Similarity"
                )


                module_progress = max(

                    0.0,

                    min(
                        module_similarity / 100,
                        1.0,
                    ),

                )


                st.progress(
                    module_progress
                )


                if module_similarity >= 80:

                    st.success(
                        f"High module similarity: "
                        f"{module_similarity:.1f}%"
                    )


                elif module_similarity >= 60:

                    st.warning(
                        f"Moderate module similarity: "
                        f"{module_similarity:.1f}%"
                    )


                else:

                    st.error(
                        f"Low module similarity: "
                        f"{module_similarity:.1f}%"
                    )


                # ============================================
                # MODULE CONCEPT COUNTS
                # ============================================

                similar_concepts = module.get(
                    "similar_concepts",
                    [],
                )


                primary_only = module.get(
                    "primary_only",
                    [],
                )


                benchmark_only = module.get(
                    "benchmark_only",
                    [],
                )


                if not isinstance(
                    similar_concepts,
                    list,
                ):

                    similar_concepts = []


                if not isinstance(
                    primary_only,
                    list,
                ):

                    primary_only = []


                if not isinstance(
                    benchmark_only,
                    list,
                ):

                    benchmark_only = []


                count_col1, count_col2, count_col3 = (
                    st.columns(3)
                )


                with count_col1:

                    st.metric(
                        "Similar",
                        len(
                            similar_concepts
                        ),
                    )


                with count_col2:

                    st.metric(
                        "Primary Only",
                        len(
                            primary_only
                        ),
                    )


                with count_col3:

                    st.metric(
                        "Benchmark Only",
                        len(
                            benchmark_only
                        ),
                    )


                st.divider()


                # ============================================
                # THREE-WAY CONCEPT COMPARISON
                # ============================================

                similar_col, primary_col, benchmark_col = (
                    st.columns(3)
                )


                # ------------------------------------------------
                # SIMILAR CONCEPTS
                # ------------------------------------------------

                with similar_col:

                    st.markdown(
                        "### 🟢 Similar Concepts"
                    )


                    if similar_concepts:

                        for concept in (
                            similar_concepts
                        ):

                            st.success(
                                str(
                                    concept
                                )
                            )

                    else:

                        st.info(
                            "No strongly similar concepts."
                        )


                # ------------------------------------------------
                # PRIMARY ONLY
                # ------------------------------------------------

                with primary_col:

                    st.markdown(
                        "### 🔵 Primary Only"
                    )


                    if primary_only:

                        for concept in (
                            primary_only
                        ):

                            st.write(
                                f"• {concept}"
                            )

                    else:

                        st.info(
                            "No primary-only concepts."
                        )


                # ------------------------------------------------
                # BENCHMARK ONLY
                # ------------------------------------------------

                with benchmark_col:

                    st.markdown(
                        "### 🔴 Benchmark Only"
                    )


                    if benchmark_only:

                        for concept in (
                            benchmark_only
                        ):

                            st.error(
                                str(
                                    concept
                                )
                            )

                    else:

                        st.info(
                            "No benchmark-only concepts."
                        )


                # ============================================
                # MODULE GAP SIGNAL
                # ============================================

                st.divider()


                st.markdown(
                    "### 🎯 Module Gap Signal"
                )


                benchmark_gap_count = len(
                    benchmark_only
                )


                primary_unique_count = len(
                    primary_only
                )


                if benchmark_gap_count == 0:

                    st.success(
                        """
                        No benchmark-only concepts were detected
                        for this module.
                        """
                    )


                elif benchmark_gap_count <= 3:

                    st.warning(
                        f"""
                        Minor potential gap:
                        **{benchmark_gap_count}**
                        benchmark concepts are not present in
                        the primary curriculum.
                        """
                    )


                elif benchmark_gap_count <= 7:

                    st.warning(
                        f"""
                        Moderate potential gap:
                        **{benchmark_gap_count}**
                        benchmark concepts are not present in
                        the primary curriculum.
                        """
                    )


                else:

                    st.error(
                        f"""
                        Significant potential gap:
                        **{benchmark_gap_count}**
                        benchmark concepts are not present in
                        the primary curriculum.
                        """
                    )


                if primary_unique_count > 0:

                    st.caption(
                        f"The primary curriculum also contains "
                        f"{primary_unique_count} concept(s) that "
                        f"were not found in the benchmark."
                    )


# ============================================================
# END OF CHUNK 5
# ============================================================
# ============================================================
# CHUNK 6/10
# TOPIC-LEVEL COMPARISON & DIFFERENCE INTELLIGENCE
# ============================================================


# ============================================================
# LOAD CURRENT COMPARISON
# ============================================================

comparison = st.session_state.get(
    "curriculum_comparison"
)


# ============================================================
# VALIDATE COMPARISON
# ============================================================

if not comparison:

    st.warning(
        """
        ⚠️ Topic-level comparison is not available.

        Please upload and process a benchmark syllabus first.
        """
    )

else:

    # ========================================================
    # SECTION HEADER
    # ========================================================

    st.divider()

    st.subheader(
        "5️⃣ Topic-Level Comparison"
    )

    st.markdown(
        """
        This analysis goes one level deeper than module
        comparison.

        Every topic is classified as:

        🟢 **Similar** — present in both curricula

        🔵 **Primary Only** — present in the primary syllabus

        🔴 **Benchmark Only** — present in the benchmark syllabus

        These differences become the foundation for the
        **Curriculum Gap & Enhancement** module.
        """
    )


    # ========================================================
    # BUILD TOPIC DATAFRAME
    # ========================================================

    topic_df = build_topic_dataframe(
        comparison
    )


    # ========================================================
    # NO TOPIC DATA
    # ========================================================

    if topic_df.empty:

        st.warning(
            """
            No topic-level comparison data was returned.

            Check the output of:

            `curriculum/comparator.py`
            """
        )

    else:

        # ====================================================
        # TOPIC COUNTS
        # ====================================================

        total_topics = len(
            topic_df
        )


        similar_df = topic_df[
            topic_df[
                "Status"
            ]
            == "Similar"
        ]


        primary_only_df = topic_df[
            topic_df[
                "Status"
            ]
            == "Primary Only"
        ]


        benchmark_only_df = topic_df[
            topic_df[
                "Status"
            ]
            == "Benchmark Only"
        ]


        # ====================================================
        # METRICS
        # ====================================================

        topic_metric_col1, topic_metric_col2, topic_metric_col3, topic_metric_col4 = (
            st.columns(4)
        )


        with topic_metric_col1:

            st.metric(
                "Total Compared Topics",
                total_topics,
            )


        with topic_metric_col2:

            st.metric(
                "🟢 Similar",
                len(
                    similar_df
                ),
            )


        with topic_metric_col3:

            st.metric(
                "🔵 Primary Only",
                len(
                    primary_only_df
                ),
            )


        with topic_metric_col4:

            st.metric(
                "🔴 Benchmark Only",
                len(
                    benchmark_only_df
                ),
            )


        # ====================================================
        # TOPIC COVERAGE
        # ====================================================

        st.markdown(
            "### 📊 Topic Coverage"
        )


        if total_topics > 0:

            similar_percentage = (
                len(
                    similar_df
                )
                / total_topics
                * 100
            )


            primary_only_percentage = (
                len(
                    primary_only_df
                )
                / total_topics
                * 100
            )


            benchmark_only_percentage = (
                len(
                    benchmark_only_df
                )
                / total_topics
                * 100
            )

        else:

            similar_percentage = 0

            primary_only_percentage = 0

            benchmark_only_percentage = 0


        coverage_col1, coverage_col2, coverage_col3 = (
            st.columns(3)
        )


        with coverage_col1:

            st.metric(
                "Similar %",
                f"{similar_percentage:.1f}%",
            )


        with coverage_col2:

            st.metric(
                "Primary Unique %",
                f"{primary_only_percentage:.1f}%",
            )


        with coverage_col3:

            st.metric(
                "Benchmark Unique %",
                f"{benchmark_only_percentage:.1f}%",
            )


        # ====================================================
        # VISUAL DISTRIBUTION
        # ====================================================

        distribution_df = pd.DataFrame({

            "Status": [

                "Similar",

                "Primary Only",

                "Benchmark Only",

            ],

            "Topics": [

                len(
                    similar_df
                ),

                len(
                    primary_only_df
                ),

                len(
                    benchmark_only_df
                ),

            ],

        })


        st.bar_chart(

            distribution_df.set_index(
                "Status"
            )

        )


        # ====================================================
        # COMPLETE TOPIC TABLE
        # ====================================================

        st.divider()

        st.markdown(
            "### 📋 Complete Topic Comparison"
        )


        st.dataframe(

            topic_df,

            use_container_width=True,

            hide_index=True,

        )


        # ====================================================
        # TOPIC FILTER
        # ====================================================

        st.divider()

        st.markdown(
            "### 🔎 Filter Topic Differences"
        )


        filter_col1, filter_col2 = (
            st.columns(2)
        )


        with filter_col1:

            topic_status_filter = st.selectbox(

                "Topic Status",

                options=[

                    "All",

                    "Similar",

                    "Primary Only",

                    "Benchmark Only",

                ],

                key="topic_status_filter",

            )


        with filter_col2:

            module_filter_options = [

                "All Modules"

            ]


            if "Module" in topic_df.columns:

                module_filter_options.extend(

                    sorted(

                        topic_df[
                            "Module"
                        ]
                        .dropna()
                        .astype(str)
                        .unique()
                        .tolist()

                    )

                )


            selected_module = st.selectbox(

                "Module",

                options=module_filter_options,

                key="topic_module_filter",

            )


        # ====================================================
        # APPLY FILTER
        # ====================================================

        filtered_topic_df = topic_df.copy()


        if topic_status_filter != "All":

            filtered_topic_df = (
                filtered_topic_df[
                    filtered_topic_df[
                        "Status"
                    ]
                    == topic_status_filter
                ]
            )


        if selected_module != "All Modules":

            filtered_topic_df = (
                filtered_topic_df[
                    filtered_topic_df[
                        "Module"
                    ]
                    == selected_module
                ]
            )


        st.dataframe(

            filtered_topic_df,

            use_container_width=True,

            hide_index=True,

        )


        # ====================================================
        # DOWNLOAD FILTERED TOPICS
        # ====================================================

        filtered_topic_csv = (
            filtered_topic_df.to_csv(
                index=False
            )
        )


        st.download_button(

            label=(
                "⬇️ Download Filtered Topic Comparison"
            ),

            data=filtered_topic_csv,

            file_name=(
                "filtered_topic_comparison.csv"
            ),

            mime="text/csv",

            key="download_filtered_topic_comparison",

        )


        # ====================================================
        # SIMILAR TOPICS
        # ====================================================

        st.divider()

        st.subheader(
            "🟢 Similar Topics / Concepts"
        )


        if similar_df.empty:

            st.info(
                "No similar topics were detected."
            )

        else:

            similar_topics = sorted(

                set(

                    similar_df[
                        "Topic"
                    ]
                    .astype(str)
                    .str.strip()

                ),

                key=lambda x: x.lower(),

            )


            for index, topic in enumerate(
                similar_topics,
                start=1,
            ):

                st.success(
                    f"{index}. {topic}"
                )


        # ====================================================
        # PRIMARY-ONLY TOPICS
        # ====================================================

        st.divider()

        st.subheader(
            "🔵 Topics Present Only in Primary Curriculum"
        )


        primary_only_topics = (
            get_primary_only_topics(
                comparison
            )
        )


        if primary_only_topics:

            for index, topic in enumerate(
                primary_only_topics,
                start=1,
            ):

                st.write(
                    f"**{index}.** {topic}"
                )

        else:

            st.info(
                "No primary-only topics were detected."
            )


        # ====================================================
        # BENCHMARK-ONLY TOPICS
        # ====================================================

        st.divider()

        st.subheader(
            "🔴 Benchmark Topics Missing from Primary Curriculum"
        )


        benchmark_only_topics = (
            get_benchmark_only_topics(
                comparison
            )
        )


        if benchmark_only_topics:

            st.markdown(
                """
                These topics exist in the benchmark curriculum
                but were not detected in the primary curriculum.

                **They are potential curriculum gaps, not automatic
                recommendations.**

                They will be evaluated further using RAG + LLM
                intelligence in the next stage.
                """
            )


            for index, topic in enumerate(
                benchmark_only_topics,
                start=1,
            ):

                st.error(
                    f"{index}. {topic}"
                )

        else:

            st.success(
                """
                No benchmark-only topics were detected.
                """
            )


        # ====================================================
        # TOPIC GAP PRIORITY
        # ====================================================

        st.divider()

        st.subheader(
            "🎯 Potential Curriculum Gap Priority"
        )


        if benchmark_only_topics:

            gap_priority_rows = []


            for topic in benchmark_only_topics:

                # --------------------------------------------
                # Determine number of modules containing topic
                # --------------------------------------------

                occurrences = 0


                for module in comparison.get(
                    "module_comparisons",
                    [],
                ):


                    if not isinstance(
                        module,
                        dict,
                    ):

                        continue


                    benchmark_only = module.get(
                        "benchmark_only",
                        [],
                    )


                    if not isinstance(
                        benchmark_only,
                        list,
                    ):

                        continue


                    for benchmark_topic in benchmark_only:

                        if str(
                            benchmark_topic
                        ).strip().lower() == topic.lower():

                            occurrences += 1


                # --------------------------------------------
                # Initial priority
                #
                # This is NOT final AI priority.
                # Page 04 will make final recommendation.
                # --------------------------------------------

                if occurrences >= 3:

                    priority = "High"


                elif occurrences == 2:

                    priority = "Medium"


                else:

                    priority = "Review"


                gap_priority_rows.append({

                    "Topic":
                        topic,

                    "Benchmark Occurrences":
                        occurrences,

                    "Initial Priority":
                        priority,

                })


            gap_priority_df = pd.DataFrame(
                gap_priority_rows
            )


            st.dataframe(

                gap_priority_df,

                use_container_width=True,

                hide_index=True,

            )


            st.caption(
                """
                ⚠️ This is only an initial structural signal.
                Final gap priority will consider industry demand,
                prerequisite dependencies, curriculum level,
                learning outcomes and expert AI review.
                """
            )

        else:

            st.success(
                "No initial benchmark-only gaps identified."
            )


        # ====================================================
        # MODULE COVERAGE MATRIX
        # ====================================================

        st.divider()

        st.subheader(
            "🗺️ Module Coverage Matrix"
        )


        module_rows = []


        for module in comparison.get(
            "module_comparisons",
            [],
        ):


            if not isinstance(
                module,
                dict,
            ):

                continue


            module_rows.append({

                "Primary Module":
                    safe_text(
                        module.get(
                            "primary_module"
                        ),
                        "Unknown",
                    ),

                "Benchmark Module":
                    safe_text(
                        module.get(
                            "benchmark_module"
                        ),
                        "Not Matched",
                    ),

                "Similarity %":
                    module.get(
                        "similarity_pct",
                        0,
                    ),

                "Similar":
                    len(
                        module.get(
                            "similar_concepts",
                            [],
                        )
                    ),

                "Primary Only":
                    len(
                        module.get(
                            "primary_only",
                            [],
                        )
                    ),

                "Benchmark Only":
                    len(
                        module.get(
                            "benchmark_only",
                            [],
                        )
                    ),

            })


        module_matrix_df = pd.DataFrame(
            module_rows
        )


        if module_matrix_df.empty:

            st.info(
                "No module coverage matrix available."
            )

        else:

            st.dataframe(

                module_matrix_df,

                use_container_width=True,

                hide_index=True,

            )


        # ====================================================
        # LOW-COVERAGE MODULES
        # ====================================================

        if not module_matrix_df.empty:

            st.divider()

            st.subheader(
                "⚠️ Modules Requiring Attention"
            )


            low_coverage_df = (
                module_matrix_df[
                    module_matrix_df[
                        "Similarity %"
                    ]
                    < 60
                ]
            )


            if low_coverage_df.empty:

                st.success(
                    """
                    No module has similarity below 60%.
                    """
                )

            else:

                st.warning(
                    f"""
                    {len(low_coverage_df)}
                    module(s) have similarity below 60%.
                    """
                )


                st.dataframe(

                    low_coverage_df,

                    use_container_width=True,

                    hide_index=True,

                )


        # ====================================================
        # TOPIC DOWNLOAD
        # ====================================================

        st.divider()

        st.subheader(
            "⬇️ Export Topic Analysis"
        )


        complete_topic_csv = (
            topic_df.to_csv(
                index=False
            )
        )


        st.download_button(

            label=(
                "⬇️ Download Complete Topic Comparison CSV"
            ),

            data=complete_topic_csv,

            file_name=(
                "complete_topic_comparison.csv"
            ),

            mime="text/csv",

            key="download_complete_topic_comparison",

        )


        # ====================================================
        # TOPIC ANALYSIS JSON
        # ====================================================

        topic_analysis_json = {

            "total_topics":
                total_topics,

            "similar_topics":
                (
                    similar_df[
                        "Topic"
                    ]
                    .astype(str)
                    .tolist()
                ),

            "primary_only_topics":
                (
                    primary_only_df[
                        "Topic"
                    ]
                    .astype(str)
                    .tolist()
                ),

            "benchmark_only_topics":
                (
                    benchmark_only_df[
                        "Topic"
                    ]
                    .astype(str)
                    .tolist()
                ),

            "topic_comparison":
                topic_df.to_dict(
                    orient="records"
                ),

        }


        st.download_button(

            label=(
                "⬇️ Download Topic Analysis JSON"
            ),

            data=serialize_json(
                topic_analysis_json
            ),

            file_name=(
                "topic_level_curriculum_analysis.json"
            ),

            mime="application/json",

            key="download_topic_analysis_json",

        )


# ============================================================
# PREPARE TOPICS FOR DEEP INTELLIGENCE
# ============================================================

st.divider()

st.subheader(
    "🧠 Topics Available for Deep AI Analysis"
)


# ============================================================
# GET TOPICS FROM COMPARISON
# ============================================================

comparison_topics = (
    get_all_comparison_topics(
        comparison
    )
)


# ============================================================
# FALLBACK TO PRIMARY CURRICULUM
# ============================================================

if not comparison_topics:

    comparison_topics = (
        get_all_topics(
            primary
        )
    )


# ============================================================
# TOPIC SUMMARY
# ============================================================

topic_count_col1, topic_count_col2 = (
    st.columns(2)
)


with topic_count_col1:

    st.metric(
        "Unique Topics Available",
        len(
            comparison_topics
        ),
    )


with topic_count_col2:

    st.metric(
        "Maximum AI Analysis",
        MAX_CONCEPT_ANALYSIS,
    )


# ============================================================
# TOPIC PREVIEW
# ============================================================

if comparison_topics:

    with st.expander(
        "📚 View All Topics",
        expanded=False,
    ):

        for index, topic in enumerate(
            comparison_topics,
            start=1,
        ):

            st.write(
                f"{index}. {topic}"
            )

else:

    st.warning(
        "No topics are currently available for deep analysis."
    )


# ============================================================
# SAVE TOPIC LIST TO SESSION
# ============================================================

st.session_state[
    "curriculum_analysis_topics"
] = comparison_topics


# ============================================================
# END OF CHUNK 6
# ============================================================
# ============================================================
# CHUNK 7/10
# DEEP CONCEPT INTELLIGENCE
# ============================================================

"""
Deep Concept Intelligence

For every selected curriculum topic, identify:

1. Core concepts
2. Advanced concepts
3. Industry-relevant concepts
4. Emerging concepts
5. Potentially missing concepts
6. Tools / technologies
7. Project areas
8. Practical skills
9. Prerequisites
10. Recommended learning direction

This layer is intentionally separated from the basic
curriculum comparison.

Comparison tells us:

    "What is different?"

Concept Intelligence asks:

    "What should we know about the difference?"

The resulting data is later consumed by:

    Page 03 → Industry / JD Intelligence
    Page 04 → Gap & Enhancement
    Page 05 → Reports
"""


# ============================================================
# LOAD TOPICS
# ============================================================

analysis_topics = st.session_state.get(
    "curriculum_analysis_topics",
    [],
)


# ============================================================
# FALLBACK
# ============================================================

if not analysis_topics:

    analysis_topics = get_all_topics(
        primary
    )


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.subheader(
    "6️⃣ 🧠 Deep Concept Intelligence"
)


st.markdown(
    """
Basic syllabus comparison tells us **what topics differ**.

Deep Concept Intelligence investigates **why those differences
matter and whether the curriculum should be enhanced**.

The AI analysis can consider:

- Existing syllabus concepts
- Benchmark syllabus concepts
- Industry relevance
- Emerging technologies
- Tools and frameworks
- Practical applications
- Project opportunities
- Prerequisite knowledge
- Potential curriculum gaps
"""
)


# ============================================================
# NO TOPICS
# ============================================================

if not analysis_topics:

    st.warning(
        """
        No topics are available for deep analysis.

        Please complete:

        **Primary Syllabus → Benchmark Syllabus →
        Curriculum Comparison**
        """
    )

else:

    # ========================================================
    # TOPIC COUNT
    # ========================================================

    st.info(
        f"""
        **{len(analysis_topics)} unique topics** are available
        for deep AI analysis.
        """
    )


    # ========================================================
    # ANALYSIS MODE
    # ========================================================

    st.markdown(
        "### ⚙️ Analysis Configuration"
    )


    config_col1, config_col2 = (
        st.columns(2)
    )


    with config_col1:

        analysis_limit = st.number_input(

            "Maximum topics to analyze",

            min_value=1,

            max_value=MAX_CONCEPT_ANALYSIS,

            value=min(
                len(
                    analysis_topics
                ),
                MAX_CONCEPT_ANALYSIS,
            ),

            step=1,

            help=(
                "Limit the number of topics sent to "
                "the AI analysis engine."
            ),

            key="deep_analysis_limit",

        )


    with config_col2:

        analysis_scope = st.selectbox(

            "Analysis Scope",

            options=[

                "Comprehensive",

                "Industry Focused",

                "Academic Focused",

                "Emerging Technology Focused",

            ],

            index=0,

            key="deep_analysis_scope",

        )


    # ========================================================
    # TOPIC SELECTION
    # ========================================================

    st.markdown(
        "### 📚 Select Topics"
    )


    selected_topics = st.multiselect(

        "Topics for deep intelligence",

        options=analysis_topics,

        default=analysis_topics[
            :min(
                len(
                    analysis_topics
                ),
                int(
                    analysis_limit
                ),
            )
        ],

        max_selections=int(
            analysis_limit
        ),

        key="selected_deep_analysis_topics",

    )


    # ========================================================
    # SELECTION SUMMARY
    # ========================================================

    selection_col1, selection_col2 = (
        st.columns(2)
    )


    with selection_col1:

        st.metric(
            "Selected Topics",
            len(
                selected_topics
            ),
        )


    with selection_col2:

        st.metric(
            "Remaining Topics",
            max(
                0,
                len(
                    analysis_topics
                )
                - len(
                    selected_topics
                ),
            ),
        )


    # ========================================================
    # ANALYSIS INSTRUCTIONS
    # ========================================================

    with st.expander(
        "🔍 What will the AI analyse?",
        expanded=False,
    ):

        st.markdown(
            """
            ### For each selected topic, the AI will attempt
            to identify:

            **Core**
            - Fundamental concepts
            - Important terminology
            - Essential theory

            **Advanced**
            - Advanced algorithms
            - Advanced architectures
            - Modern techniques
            - Higher-level concepts

            **Industry**
            - Industry skills
            - Real-world usage
            - Production technologies
            - Job-relevant knowledge

            **Emerging**
            - Recent technologies
            - Modern frameworks
            - New approaches
            - Future-relevant skills

            **Gap**
            - Concepts potentially missing
            - Concepts that need expansion
            - Practical components missing

            **Tools**
            - Programming languages
            - Frameworks
            - Platforms
            - Libraries
            - Cloud / deployment tools

            **Projects**
            - Practical project ideas
            - Case studies
            - Industry applications
            """
        )


    # ========================================================
    # RUN ANALYSIS
    # ========================================================

    run_deep_analysis = st.button(

        "🚀 Run Deep Concept Intelligence",

        type="primary",

        use_container_width=True,

        key="run_deep_concept_intelligence",

    )


    # ========================================================
    # EXECUTE ANALYSIS
    # ========================================================

    if run_deep_analysis:

        if not selected_topics:

            st.error(
                "Please select at least one topic."
            )

        else:

            analysis_progress = st.progress(
                0
            )


            analysis_status = st.empty()


            try:

                # ============================================
                # BUILD CONTEXT
                # ============================================

                analysis_status.info(
                    "📚 Preparing curriculum context..."
                )


                analysis_progress.progress(
                    10
                )


                primary_context = serialize_json(
                    primary
                )


                benchmark_context = serialize_json(
                    benchmark
                )


                # ============================================
                # ADD SCOPE INSTRUCTION
                # ============================================

                scope_instructions = {

                    "Comprehensive":
                        """
                        Analyse academic fundamentals,
                        advanced concepts, industry relevance,
                        emerging technologies, tools,
                        projects and curriculum gaps.
                        """,

                    "Industry Focused":
                        """
                        Prioritize job skills, production
                        technologies, tools, frameworks,
                        projects, deployment and industry
                        relevance.
                        """,

                    "Academic Focused":
                        """
                        Prioritize theoretical foundations,
                        academic depth, prerequisites,
                        algorithms, mathematical concepts
                        and conceptual progression.
                        """,

                    "Emerging Technology Focused":
                        """
                        Prioritize modern and emerging
                        technologies, current tools,
                        advanced concepts and future skills.
                        """,

                }


                selected_scope_instruction = (
                    scope_instructions.get(
                        analysis_scope,
                        scope_instructions[
                            "Comprehensive"
                        ],
                    )
                )


                # ============================================
                # PREPARE TOPIC BATCH
                # ============================================

                selected_topics = selected_topics[
                    :int(
                        analysis_limit
                    )
                ]


                # ============================================
                # AI ANALYSIS
                # ============================================

                analysis_status.info(
                    f"""
                    🧠 Analysing
                    {len(selected_topics)}
                    topics using the Concept Intelligence engine...
                    """
                )


                analysis_progress.progress(
                    25
                )


                # ------------------------------------------------
                # enrich_topics() is responsible for:
                #
                # RAG retrieval
                # LLM reasoning
                # Web enrichment, if configured
                # Structured output
                # ------------------------------------------------

                enrichment = enrich_topics(

                    selected_topics,

                    primary_curriculum=primary,

                    benchmark_curriculum=benchmark,

                    context={

                        "analysis_scope":
                            analysis_scope,

                        "scope_instruction":
                            selected_scope_instruction,

                        "primary_context":
                            primary_context,

                        "benchmark_context":
                            benchmark_context,

                    },

                )


                # ============================================
                # VALIDATE RESULT
                # ============================================

                if not isinstance(
                    enrichment,
                    list,
                ):

                    raise ValueError(
                        """
                        Concept Intelligence returned an
                        invalid result.

                        Expected a list of topic analyses.
                        """
                    )


                # ============================================
                # NORMALIZE RESULTS
                # ============================================

                normalized_enrichment = []


                for result in enrichment:

                    if not isinstance(
                        result,
                        dict,
                    ):

                        continue


                    # ----------------------------------------
                    # Ensure topic
                    # ----------------------------------------

                    topic = (

                        result.get(
                            "topic"
                        )

                        or

                        result.get(
                            "name"
                        )

                        or

                        "Unknown Topic"

                    )


                    # ----------------------------------------
                    # Normalize list fields
                    # ----------------------------------------

                    list_fields = [

                        "core_concepts",

                        "advanced_concepts",

                        "industry_concepts",

                        "emerging_concepts",

                        "missing_concepts",

                        "tools",

                        "project_areas",

                        "prerequisites",

                        "practical_skills",

                        "recommendations",

                    ]


                    normalized_result = dict(
                        result
                    )


                    normalized_result[
                        "topic"
                    ] = str(
                        topic
                    ).strip()


                    for field in list_fields:

                        value = normalized_result.get(
                            field,
                            [],
                        )


                        if value is None:

                            value = []


                        if not isinstance(
                            value,
                            list,
                        ):

                            value = [
                                value
                            ]


                        normalized_result[
                            field
                        ] = value


                    normalized_enrichment.append(
                        normalized_result
                    )


                # ============================================
                # SAVE RESULTS
                # ============================================

                st.session_state[
                    "concept_enrichment"
                ] = normalized_enrichment


                st.session_state[
                    "concept_analysis_complete"
                ] = True


                # ============================================
                # UPDATE COMPLETE REPORT
                # ============================================

                st.session_state[
                    "curriculum_intelligence_report"
                ] = build_full_report(

                    primary,

                    benchmark,

                    comparison,

                    normalized_enrichment,

                )


                # ============================================
                # FINISH
                # ============================================

                analysis_progress.progress(
                    100
                )


                analysis_status.success(
                    """
                    ✅ Deep Concept Intelligence completed.
                    """
                )


            except Exception as exc:

                analysis_progress.empty()


                analysis_status.error(
                    "❌ Deep Concept Intelligence failed."
                )


                st.error(
                    str(
                        exc
                    )
                )


                with st.expander(
                    "🔧 Technical Error Details",
                    expanded=False,
                ):

                    st.exception(
                        exc
                    )


# ============================================================
# LOAD STORED ENRICHMENT
# ============================================================

enrichment = st.session_state.get(
    "concept_enrichment",
    [],
)


# ============================================================
# DISPLAY RESULTS
# ============================================================

if enrichment:

    st.divider()

    st.subheader(
        "7️⃣ 🧠 Deep Concept Intelligence Results"
    )


    st.success(
        f"""
        AI analysis completed for
        **{len(enrichment)} topics**.
        """
    )


    # ========================================================
    # GLOBAL SUMMARY
    # ========================================================

    concept_summary = build_concept_summary(
        enrichment
    )


    summary_col1, summary_col2, summary_col3, summary_col4 = (
        st.columns(4)
    )


    with summary_col1:

        st.metric(
            "Topics Analysed",
            concept_summary[
                "topics_analyzed"
            ],
        )


    with summary_col2:

        st.metric(
            "Industry Concepts",
            len(
                concept_summary[
                    "industry_concepts"
                ]
            ),
        )


    with summary_col3:

        st.metric(
            "Emerging Concepts",
            len(
                concept_summary[
                    "emerging_concepts"
                ]
            ),
        )


    with summary_col4:

        st.metric(
            "Potential Gaps",
            len(
                concept_summary[
                    "missing_concepts"
                ]
            ),
        )


    # ========================================================
    # GLOBAL CONCEPT CATEGORIES
    # ========================================================

    st.divider()

    st.subheader(
        "🌐 Global Concept Intelligence"
    )


    concept_category_tabs = st.tabs(

        [

            "Core",

            "Advanced",

            "Industry",

            "Emerging",

            "Potential Gaps",

            "Tools",

            "Projects",

        ]

    )


    # ========================================================
    # CORE
    # ========================================================

    with concept_category_tabs[0]:

        display_topic_list(

            concept_summary[
                "core_concepts"
            ]

        )


    # ========================================================
    # ADVANCED
    # ========================================================

    with concept_category_tabs[1]:

        display_topic_list(

            concept_summary[
                "advanced_concepts"
            ]

        )


    # ========================================================
    # INDUSTRY
    # ========================================================

    with concept_category_tabs[2]:

        display_topic_list(

            concept_summary[
                "industry_concepts"
            ],

            status="success",

        )


    # ========================================================
    # EMERGING
    # ========================================================

    with concept_category_tabs[3]:

        display_topic_list(

            concept_summary[
                "emerging_concepts"
            ],

            status="success",

        )


    # ========================================================
    # POTENTIAL GAPS
    # ========================================================

    with concept_category_tabs[4]:

        display_topic_list(

            concept_summary[
                "missing_concepts"
            ],

            status="error",

        )


    # ========================================================
    # TOOLS
    # ========================================================

    with concept_category_tabs[5]:

        display_topic_list(

            concept_summary[
                "tools"
            ],

            status="success",

        )


    # ========================================================
    # PROJECTS
    # ========================================================

    with concept_category_tabs[6]:

        display_topic_list(

            concept_summary[
                "project_areas"
            ]

        )


    # ========================================================
    # TOPIC SUMMARY TABLE
    # ========================================================

    st.divider()

    st.subheader(
        "📊 Topic Intelligence Summary"
    )


    concept_df = build_concept_dataframe(
        enrichment
    )


    if not concept_df.empty:

        st.dataframe(

            concept_df,

            use_container_width=True,

            hide_index=True,

        )


    # ========================================================
    # TOPIC-BY-TOPIC DETAILS
    # ========================================================

    st.divider()

    st.subheader(
        "🔬 Topic-by-Topic Deep Analysis"
    )


    for index, item in enumerate(
        enrichment,
        start=1,
    ):


        if not isinstance(
            item,
            dict,
        ):

            continue


        topic_name = safe_text(

            item.get(
                "topic"
            ),

            f"Topic {index}",

        )


        # ----------------------------------------------------
        # Get similarity information if available
        # ----------------------------------------------------

        topic_similarity = item.get(
            "similarity_pct"
        )


        if topic_similarity is not None:

            try:

                topic_similarity = float(
                    topic_similarity
                )

            except (
                TypeError,
                ValueError,
            ):

                topic_similarity = None


        # ----------------------------------------------------
        # Expander title
        # ----------------------------------------------------

        if topic_similarity is not None:

            topic_title = (

                f"{index}. "
                f"{topic_name} "
                f"| Similarity: "
                f"{topic_similarity:.1f}%"

            )

        else:

            topic_title = (

                f"{index}. "
                f"{topic_name}"

            )


        with st.expander(
            topic_title,
            expanded=False,
        ):


            # ================================================
            # CORE CONCEPTS
            # ================================================

            st.markdown(
                "### 🧱 Core Concepts"
            )


            display_topic_list(

                item.get(
                    "core_concepts",
                    [],
                )

            )


            # ================================================
            # ADVANCED CONCEPTS
            # ================================================

            st.markdown(
                "### 🚀 Advanced Concepts"
            )


            display_topic_list(

                item.get(
                    "advanced_concepts",
                    [],
                ),

            )


            # ================================================
            # INDUSTRY CONCEPTS
            # ================================================

            st.markdown(
                "### 🏭 Industry-Relevant Concepts"
            )


            display_topic_list(

                item.get(
                    "industry_concepts",
                    [],
                ),

                status="success",

            )


            # ================================================
            # EMERGING CONCEPTS
            # ================================================

            st.markdown(
                "### 🔥 Emerging Concepts"
            )


            display_topic_list(

                item.get(
                    "emerging_concepts",
                    [],
                ),

                status="success",

            )


            # ================================================
            # POTENTIAL MISSING CONCEPTS
            # ================================================

            st.markdown(
                "### ⚠️ Potential Missing Concepts"
            )


            display_topic_list(

                item.get(
                    "missing_concepts",
                    [],
                ),

                status="error",

            )


            # ================================================
            # TOOLS
            # ================================================

            st.markdown(
                "### 🛠 Recommended Tools / Technologies"
            )


            display_topic_list(

                item.get(
                    "tools",
                    [],
                ),

                status="success",

            )


            # ================================================
            # PROJECTS
            # ================================================

            st.markdown(
                "### 🚀 Recommended Project Areas"
            )


            display_topic_list(

                item.get(
                    "project_areas",
                    [],
                )

            )


            # ================================================
            # PREREQUISITES
            # ================================================

            st.markdown(
                "### 📌 Prerequisites"
            )


            display_topic_list(

                item.get(
                    "prerequisites",
                    [],
                )

            )


            # ================================================
            # PRACTICAL SKILLS
            # ================================================

            st.markdown(
                "### 💻 Practical Skills"
            )


            display_topic_list(

                item.get(
                    "practical_skills",
                    [],
                )

            )


            # ================================================
            # RECOMMENDATIONS
            # ================================================

            st.markdown(
                "### 💡 AI Recommendations"
            )


            recommendations = item.get(
                "recommendations",
                [],
            )


            if recommendations:

                display_value_list(
                    recommendations
                )

            else:

                st.info(
                    "No specific recommendations returned."
                )


            # ================================================
            # SOURCES
            # ================================================

            sources = item.get(
                "sources",
                [],
            )


            if sources:

                st.markdown(
                    "### 🔗 Intelligence Sources"
                )


                display_value_list(
                    sources
                )


            # ================================================
            # RAW AI RESULT
            # ================================================

            with st.expander(
                "🔧 View Raw AI Topic Analysis",
                expanded=False,
            ):

                st.json(
                    item
                )


# ============================================================
# NO RESULTS
# ============================================================

elif st.session_state.get(
    "concept_analysis_complete",
    False,
):

    st.warning(
        """
        Deep Concept Intelligence was executed, but no
        structured topic results were returned.
        """
    )


# ============================================================
# PREPARE GAP INPUT
# ============================================================

"""
The following session-state objects are intentionally prepared
for Page 04.

Page 04 will combine:

    Curriculum Comparison
            +
    Topic Intelligence
            +
    Industry / JD Intelligence
            +
    Agent / Critic Agent
            ↓
    Final Gap & Enhancement
"""


st.session_state[
    "curriculum_gap_candidates"
] = {

    "benchmark_only_topics":
        get_benchmark_only_topics(
            comparison
        )
        if comparison
        else [],

    "primary_only_topics":
        get_primary_only_topics(
            comparison
        )
        if comparison
        else [],

    "similar_topics":
        get_similar_topics(
            comparison
        )
        if comparison
        else [],

    "concept_intelligence":
        enrichment,

}


# ============================================================
# DOWNLOAD DEEP INTELLIGENCE JSON
# ============================================================

if enrichment:

    st.divider()

    st.subheader(
        "⬇️ Export Deep Concept Intelligence"
    )


    enrichment_json = {

        "generated_at":
            datetime.now().isoformat(),

        "topic_count":
            len(
                enrichment
            ),

        "analysis_scope":
            st.session_state.get(
                "deep_analysis_scope",
                "Comprehensive",
            ),

        "summary":
            build_concept_summary(
                enrichment
            ),

        "topic_analysis":
            enrichment,

    }


    st.download_button(

        label=(
            "⬇️ Download Deep Concept Intelligence JSON"
        ),

        data=serialize_json(
            enrichment_json
        ),

        file_name=(
            "deep_concept_intelligence.json"
        ),

        mime="application/json",

        key="download_deep_concept_intelligence",

    )


# ============================================================
# END OF CHUNK 7
# ============================================================
# ============================================================
# CHUNK 8/10
# CONCEPT RESULTS + GAP SIGNALS + ANALYSIS PAYLOAD
# ============================================================

"""
This chunk performs post-processing on the results generated
by Chunk 7.

Responsibilities
----------------
1. Display concept-level metrics
2. Identify curriculum gap candidates
3. Identify industry-oriented concepts
4. Identify emerging technologies
5. Identify recommended tools
6. Identify project opportunities
7. Build a topic intelligence matrix
8. Build a curriculum gap matrix
9. Build the payload for Page 03 / Page 04
10. Save all structured results in session state
"""


# ============================================================
# LOAD DATA
# ============================================================

comparison = st.session_state.get(
    "curriculum_comparison"
)


enrichment = st.session_state.get(
    "concept_enrichment",
    [],
)


analysis_topics = st.session_state.get(
    "curriculum_analysis_topics",
    [],
)


# ============================================================
# SECTION HEADER
# ============================================================

st.divider()

st.subheader(
    "8️⃣ 🎯 Curriculum Intelligence & Gap Signals"
)


st.markdown(
    """
The system now combines:

**Syllabus Comparison + Topic Intelligence + AI Concept Analysis**

to generate preliminary signals for curriculum enhancement.

> These are **AI-generated signals**, not final curriculum
> recommendations. Final recommendations are generated in
> the Gap & Enhancement module after Industry/JD analysis
> and expert/critic review.
"""
)


# ============================================================
# VALIDATION
# ============================================================

if not enrichment:

    st.info(
        """
        🧠 Deep Concept Intelligence has not produced results yet.

        Run **Deep Concept Intelligence** above to generate
        concept-level curriculum intelligence.
        """
    )

else:

    # ========================================================
    # GLOBAL COUNTS
    # ========================================================

    summary = build_concept_summary(
        enrichment
    )


    core_concepts = summary.get(
        "core_concepts",
        [],
    )


    advanced_concepts = summary.get(
        "advanced_concepts",
        [],
    )


    industry_concepts = summary.get(
        "industry_concepts",
        [],
    )


    emerging_concepts = summary.get(
        "emerging_concepts",
        [],
    )


    missing_concepts = summary.get(
        "missing_concepts",
        [],
    )


    tools = summary.get(
        "tools",
        [],
    )


    project_areas = summary.get(
        "project_areas",
        [],
    )


    # ========================================================
    # INTELLIGENCE METRICS
    # ========================================================

    st.markdown(
        "### 📊 Intelligence Overview"
    )


    intelligence_col1, intelligence_col2, intelligence_col3, intelligence_col4 = (
        st.columns(4)
    )


    with intelligence_col1:

        st.metric(
            "Core Concepts",
            len(
                core_concepts
            ),
        )


    with intelligence_col2:

        st.metric(
            "Advanced Concepts",
            len(
                advanced_concepts
            ),
        )


    with intelligence_col3:

        st.metric(
            "Industry Concepts",
            len(
                industry_concepts
            ),
        )


    with intelligence_col4:

        st.metric(
            "Emerging Concepts",
            len(
                emerging_concepts
            ),
        )


    intelligence_col5, intelligence_col6, intelligence_col7, intelligence_col8 = (
        st.columns(4)
    )


    with intelligence_col5:

        st.metric(
            "Potential Gaps",
            len(
                missing_concepts
            ),
        )


    with intelligence_col6:

        st.metric(
            "Tools",
            len(
                tools
            ),
        )


    with intelligence_col7:

        st.metric(
            "Project Areas",
            len(
                project_areas
            ),
        )


    with intelligence_col8:

        st.metric(
            "Topics Analysed",
            len(
                enrichment
            ),
        )


    # ========================================================
    # CURRICULUM GAP CANDIDATES
    # ========================================================

    st.divider()

    st.subheader(
        "⚠️ Preliminary Curriculum Gap Candidates"
    )


    benchmark_only_topics = []

    primary_only_topics = []

    similar_topics = []


    if comparison:

        benchmark_only_topics = (
            get_benchmark_only_topics(
                comparison
            )
        )


        primary_only_topics = (
            get_primary_only_topics(
                comparison
            )
        )


        similar_topics = (
            get_similar_topics(
                comparison
            )
        )


    # ========================================================
    # COMBINE STRUCTURAL + CONCEPT GAPS
    # ========================================================

    structural_gaps = set(
        benchmark_only_topics
    )


    AI_gaps = set(
        missing_concepts
    )


    combined_gap_candidates = (
        structural_gaps
        | AI_gaps
    )


    combined_gap_candidates = sorted(

        combined_gap_candidates,

        key=lambda value: value.lower(),

    )


    # ========================================================
    # GAP METRICS
    # ========================================================

    gap_col1, gap_col2, gap_col3 = (
        st.columns(3)
    )


    with gap_col1:

        st.metric(
            "Benchmark-only Gaps",
            len(
                structural_gaps
            ),
        )


    with gap_col2:

        st.metric(
            "AI-identified Gaps",
            len(
                AI_gaps
            ),
        )


    with gap_col3:

        st.metric(
            "Combined Candidates",
            len(
                combined_gap_candidates
            ),
        )


    # ========================================================
    # GAP TABLE
    # ========================================================

    if combined_gap_candidates:

        gap_rows = []


        for gap in combined_gap_candidates:

            structural = (
                gap in structural_gaps
            )


            ai_identified = (
                gap in AI_gaps
            )


            if structural and ai_identified:

                source = (
                    "Benchmark + AI"
                )

                initial_priority = (
                    "High"
                )


            elif structural:

                source = (
                    "Benchmark"
                )

                initial_priority = (
                    "Review"
                )


            else:

                source = (
                    "AI Concept Intelligence"
                )

                initial_priority = (
                    "Review"
                )


            gap_rows.append({

                "Potential Gap":
                    gap,

                "Benchmark Evidence":
                    "Yes"
                    if structural
                    else "No",

                "AI Evidence":
                    "Yes"
                    if ai_identified
                    else "No",

                "Source":
                    source,

                "Initial Priority":
                    initial_priority,

            })


        gap_df = pd.DataFrame(
            gap_rows
        )


        st.dataframe(

            gap_df,

            use_container_width=True,

            hide_index=True,

        )


    else:

        st.success(
            """
            No preliminary curriculum gaps were identified
            from the current analysis.
            """
        )


    # ========================================================
    # GAP WARNING
    # ========================================================

    st.caption(
        """
        ⚠️ A topic is considered a **candidate gap** only when
        structural comparison or AI concept intelligence identifies
        it. Page 04 will validate whether it is genuinely relevant
        before recommending curriculum changes.
        """
    )


    # ========================================================
    # INDUSTRY INTELLIGENCE
    # ========================================================

    st.divider()

    st.subheader(
        "🏭 Industry-Relevant Concept Signals"
    )


    if industry_concepts:

        industry_df = pd.DataFrame({

            "Industry Concept":
                industry_concepts,

        })


        st.dataframe(

            industry_df,

            use_container_width=True,

            hide_index=True,

        )


    else:

        st.info(
            """
            No industry-specific concepts were identified
            by the current analysis.
            """
        )


    # ========================================================
    # EMERGING TECHNOLOGIES
    # ========================================================

    st.divider()

    st.subheader(
        "🔥 Emerging Technology Signals"
    )


    if emerging_concepts:

        for index, concept in enumerate(
            emerging_concepts,
            start=1,
        ):

            st.success(
                f"{index}. {concept}"
            )

    else:

        st.info(
            "No emerging concepts were identified."
        )


    # ========================================================
    # RECOMMENDED TOOLS
    # ========================================================

    st.divider()

    st.subheader(
        "🛠 Recommended Tools & Technologies"
    )


    if tools:

        tool_rows = []


        for tool in tools:

            tool_rows.append({

                "Tool / Technology":
                    tool,

                "Status":
                    "AI Recommended",

            })


        tool_df = pd.DataFrame(
            tool_rows
        )


        st.dataframe(

            tool_df,

            use_container_width=True,

            hide_index=True,

        )


    else:

        st.info(
            "No additional tools were identified."
        )


    # ========================================================
    # PROJECT OPPORTUNITIES
    # ========================================================

    st.divider()

    st.subheader(
        "🚀 Recommended Project Areas"
    )


    if project_areas:

        for index, project in enumerate(
            project_areas,
            start=1,
        ):

            st.markdown(
                f"**{index}.** {project}"
            )

    else:

        st.info(
            "No project areas were identified."
        )


    # ========================================================
    # TOPIC INTELLIGENCE MATRIX
    # ========================================================

    st.divider()

    st.subheader(
        "🧠 Topic Intelligence Matrix"
    )


    topic_matrix_rows = []


    for item in enrichment:

        if not isinstance(
            item,
            dict,
        ):

            continue


        topic = safe_text(

            item.get(
                "topic"
            ),

            "Unknown",

        )


        core = item.get(
            "core_concepts",
            [],
        )


        advanced = item.get(
            "advanced_concepts",
            [],
        )


        industry = item.get(
            "industry_concepts",
            [],
        )


        emerging = item.get(
            "emerging_concepts",
            [],
        )


        missing = item.get(
            "missing_concepts",
            [],
        )


        tools_for_topic = item.get(
            "tools",
            [],
        )


        projects_for_topic = item.get(
            "project_areas",
            [],
        )


        # --------------------------------------------
        # Normalize list fields
        # --------------------------------------------

        fields = [

            "core",

            "advanced",

            "industry",

            "emerging",

            "missing",

            "tools_for_topic",

            "projects_for_topic",

        ]


        local_values = {

            "core":
                core,

            "advanced":
                advanced,

            "industry":
                industry,

            "emerging":
                emerging,

            "missing":
                missing,

            "tools_for_topic":
                tools_for_topic,

            "projects_for_topic":
                projects_for_topic,

        }


        for field in fields:

            if not isinstance(
                local_values[field],
                list,
            ):

                local_values[field] = [
                    local_values[field]
                ]


        topic_matrix_rows.append({

            "Topic":
                topic,

            "Core":
                len(
                    local_values["core"]
                ),

            "Advanced":
                len(
                    local_values["advanced"]
                ),

            "Industry":
                len(
                    local_values["industry"]
                ),

            "Emerging":
                len(
                    local_values["emerging"]
                ),

            "Potential Gaps":
                len(
                    local_values["missing"]
                ),

            "Tools":
                len(
                    local_values["tools_for_topic"]
                ),

            "Projects":
                len(
                    local_values["projects_for_topic"]
                ),

        })


    topic_matrix_df = pd.DataFrame(
        topic_matrix_rows
    )


    if topic_matrix_df.empty:

        st.info(
            "Topic intelligence matrix is empty."
        )

    else:

        st.dataframe(

            topic_matrix_df,

            use_container_width=True,

            hide_index=True,

        )


    # ========================================================
    # HIGH PRIORITY TOPICS
    # ========================================================

    if not topic_matrix_df.empty:

        st.divider()

        st.subheader(
            "🚨 Topics Requiring Deeper Review"
        )


        priority_rows = []


        for _, row in topic_matrix_df.iterrows():

            priority_score = 0


            # --------------------------------------------
            # Gap signal
            # --------------------------------------------

            if row[
                "Potential Gaps"
            ] > 0:

                priority_score += 3


            # --------------------------------------------
            # Industry signal
            # --------------------------------------------

            if row[
                "Industry"
            ] > 0:

                priority_score += 2


            # --------------------------------------------
            # Emerging signal
            # --------------------------------------------

            if row[
                "Emerging"
            ] > 0:

                priority_score += 2


            # --------------------------------------------
            # Tools signal
            # --------------------------------------------

            if row[
                "Tools"
            ] > 0:

                priority_score += 1


            # --------------------------------------------
            # Project signal
            # --------------------------------------------

            if row[
                "Projects"
            ] > 0:

                priority_score += 1


            # --------------------------------------------
            # Priority classification
            # --------------------------------------------

            if priority_score >= 7:

                priority = "High"


            elif priority_score >= 4:

                priority = "Medium"


            else:

                priority = "Low"


            priority_rows.append({

                "Topic":
                    row["Topic"],

                "Priority Score":
                    priority_score,

                "Priority":
                    priority,

                "Potential Gaps":
                    row["Potential Gaps"],

                "Industry Concepts":
                    row["Industry"],

                "Emerging Concepts":
                    row["Emerging"],

                "Tools":
                    row["Tools"],

                "Projects":
                    row["Projects"],

            })


        priority_df = pd.DataFrame(
            priority_rows
        )


        priority_df = (
            priority_df
            .sort_values(
                "Priority Score",
                ascending=False,
            )
            .reset_index(
                drop=True
            )
        )


        st.dataframe(

            priority_df,

            use_container_width=True,

            hide_index=True,

        )


        st.caption(
            """
            The priority score is an internal screening signal.
            It is not the final curriculum recommendation.
            """
        )


    # ========================================================
    # CURRICULUM ENHANCEMENT INPUTS
    # ========================================================

    st.divider()

    st.subheader(
        "🧩 Curriculum Enhancement Inputs"
    )


    enhancement_inputs = {

        "structural_gaps":
            benchmark_only_topics,

        "ai_identified_gaps":
            missing_concepts,

        "combined_gap_candidates":
            combined_gap_candidates,

        "similar_topics":
            similar_topics,

        "primary_only_topics":
            primary_only_topics,

        "industry_concepts":
            industry_concepts,

        "emerging_concepts":
            emerging_concepts,

        "recommended_tools":
            tools,

        "recommended_project_areas":
            project_areas,

    }


    # ========================================================
    # SAVE ENHANCEMENT INPUTS
    # ========================================================

    st.session_state[
        "curriculum_enhancement_inputs"
    ] = enhancement_inputs


    # ========================================================
    # DISPLAY ENHANCEMENT INPUT SUMMARY
    # ========================================================

    enhancement_summary_df = pd.DataFrame({

        "Intelligence Category": [

            "Structural Gaps",

            "AI-identified Gaps",

            "Combined Gap Candidates",

            "Similar Topics",

            "Primary-only Topics",

            "Industry Concepts",

            "Emerging Concepts",

            "Recommended Tools",

            "Project Areas",

        ],

        "Count": [

            len(
                benchmark_only_topics
            ),

            len(
                missing_concepts
            ),

            len(
                combined_gap_candidates
            ),

            len(
                similar_topics
            ),

            len(
                primary_only_topics
            ),

            len(
                industry_concepts
            ),

            len(
                emerging_concepts
            ),

            len(
                tools
            ),

            len(
                project_areas
            ),

        ],

    })


    st.dataframe(

        enhancement_summary_df,

        use_container_width=True,

        hide_index=True,

    )


    # ========================================================
    # COMPLETE INTELLIGENCE PAYLOAD
    # ========================================================

    curriculum_intelligence_payload = {

        "generated_at":
            datetime.now().isoformat(),

        "primary_curriculum":
            primary,

        "benchmark_curriculum":
            benchmark,

        "comparison":
            comparison,

        "topic_analysis":
            topic_matrix_rows,

        "concept_intelligence":
            enrichment,

        "gap_candidates":
            {

                "benchmark_only":
                    benchmark_only_topics,

                "ai_identified":
                    missing_concepts,

                "combined":
                    combined_gap_candidates,

            },

        "industry_signals":
            industry_concepts,

        "emerging_signals":
            emerging_concepts,

        "recommended_tools":
            tools,

        "project_areas":
            project_areas,

    }


    # ========================================================
    # SAVE COMPLETE PAYLOAD
    # ========================================================

    st.session_state[
        "curriculum_intelligence_payload"
    ] = curriculum_intelligence_payload


    # ========================================================
    # UPDATE MASTER REPORT
    # ========================================================

    st.session_state[
        "curriculum_intelligence_report"
    ] = curriculum_intelligence_payload


    # ========================================================
    # EXPORT INTELLIGENCE PAYLOAD
    # ========================================================

    st.divider()

    st.subheader(
        "⬇️ Export Curriculum Intelligence"
    )


    payload_json = serialize_json(
        curriculum_intelligence_payload
    )


    st.download_button(

        label=(
            "⬇️ Download Complete Curriculum Intelligence JSON"
        ),

        data=payload_json,

        file_name=(
            "curriculum_intelligence_complete.json"
        ),

        mime="application/json",

        key=(
            "download_complete_curriculum_intelligence"
        ),

    )


    # ========================================================
    # CSV EXPORT — TOPIC MATRIX
    # ========================================================

    if not topic_matrix_df.empty:

        topic_matrix_csv = (
            topic_matrix_df.to_csv(
                index=False
            )
        )


        st.download_button(

            label=(
                "⬇️ Download Topic Intelligence CSV"
            ),

            data=topic_matrix_csv,

            file_name=(
                "topic_intelligence_matrix.csv"
            ),

            mime="text/csv",

            key=(
                "download_topic_intelligence_matrix"
            ),

        )


    # ========================================================
    # CSV EXPORT — GAP MATRIX
    # ========================================================

    if "gap_df" in locals() and not gap_df.empty:

        gap_csv = gap_df.to_csv(
            index=False
        )


        st.download_button(

            label=(
                "⬇️ Download Curriculum Gap Candidates CSV"
            ),

            data=gap_csv,

            file_name=(
                "curriculum_gap_candidates.csv"
            ),

            mime="text/csv",

            key=(
                "download_curriculum_gap_candidates"
            ),

        )


    # ========================================================
    # DATA HANDOFF STATUS
    # ========================================================

    st.divider()

    st.success(
        """
        ✅ **Curriculum Intelligence data prepared successfully.**

        The structured output is now ready for:

        **Industry & JD Intelligence → Gap Analysis →
        Agent Review → Curriculum Enhancement**
        """
    )


# ============================================================
# SESSION STATE HANDOFF
# ============================================================

st.session_state[
    "curriculum_intelligence_ready"
] = bool(
    comparison
    and enrichment
)


# ============================================================
# END OF CHUNK 8
# ============================================================
# ============================================================
# CHUNK 9/10
# REPORT GENERATION + DOWNLOADS + PAGE HANDOFF
# ============================================================

"""
Responsibilities
----------------
1. Build final Curriculum Intelligence report
2. Generate executive summary
3. Generate module comparison report
4. Generate topic comparison report
5. Generate gap candidate report
6. Generate concept intelligence report
7. Generate Markdown report
8. Generate JSON report
9. Generate CSV exports
10. Save report package to session state
11. Prepare clean handoff for Page 05
"""


# ============================================================
# LOAD CURRENT DATA
# ============================================================

primary = st.session_state.get(
    "primary_syllabus",
    {}
)

benchmark = st.session_state.get(
    "benchmark_syllabus"
)

comparison = st.session_state.get(
    "curriculum_comparison"
)

enrichment = st.session_state.get(
    "concept_enrichment",
    []
)

curriculum_payload = st.session_state.get(
    "curriculum_intelligence_payload"
)


# ============================================================
# REPORT SECTION
# ============================================================

st.divider()

st.subheader(
    "9️⃣ 📊 Curriculum Intelligence Report"
)


st.markdown(
    """
The report consolidates the complete analysis performed on
this page.

### Report sections

1. Primary syllabus profile
2. Benchmark syllabus profile
3. Curriculum structure comparison
4. Module-wise comparison
5. Topic-wise comparison
6. Similar concepts
7. Primary-only concepts
8. Benchmark-only concepts
9. Deep concept intelligence
10. Potential curriculum gaps
11. Industry concept signals
12. Emerging technologies
13. Recommended tools
14. Project opportunities
15. Preliminary enhancement inputs
"""
)


# ============================================================
# CHECK REPORT READINESS
# ============================================================

report_ready = bool(

    primary

    and benchmark

    and comparison

)


if not report_ready:

    st.warning(
        """
        ⚠️ Complete the following before generating the
        final Curriculum Intelligence report:

        **Primary Syllabus + Benchmark Syllabus +
        Curriculum Comparison**
        """
    )

else:

    # ========================================================
    # BUILD CORE DATA
    # ========================================================

    comparison_summary = get_comparison_summary(
        comparison
    )


    module_df = build_module_dataframe(
        comparison
    )


    topic_df = build_topic_dataframe(
        comparison
    )


    concept_df = build_concept_dataframe(
        enrichment
    )


    # ========================================================
    # GAP DATA
    # ========================================================

    benchmark_only_topics = (
        get_benchmark_only_topics(
            comparison
        )
    )


    primary_only_topics = (
        get_primary_only_topics(
            comparison
        )
    )


    similar_topics = (
        get_similar_topics(
            comparison
        )
    )


    concept_summary = build_concept_summary(
        enrichment
    )


    combined_gaps = sorted(

        set(
            benchmark_only_topics
        )
        |
        set(
            concept_summary.get(
                "missing_concepts",
                [],
            )
        ),

        key=lambda value: value.lower(),

    )


    # ========================================================
    # REPORT OBJECT
    # ========================================================

    report = {

        "report_metadata": {

            "report_type":
                "Curriculum Intelligence Report",

            "version":
                "1.0",

            "generated_at":
                datetime.now().isoformat(),

            "platform":
                "PragyanAI AI Curriculum Platform",

        },


        # ----------------------------------------------------
        # PRIMARY PROFILE
        # ----------------------------------------------------

        "primary_curriculum_profile": {

            "college":
                primary.get(
                    "college"
                ),

            "university":
                primary.get(
                    "university"
                ),

            "program":
                primary.get(
                    "program"
                ),

            "department":
                primary.get(
                    "department"
                ),

            "academic_year":
                primary.get(
                    "academic_year"
                ),

            "regulation":
                primary.get(
                    "regulation"
                ),

            "semester":
                primary.get(
                    "semester"
                ),

            "subject_name":
                primary.get(
                    "subject_name"
                ),

            "subject_code":
                primary.get(
                    "subject_code"
                ),

            "credits":
                primary.get(
                    "credits"
                ),

            "contact_hours":
                primary.get(
                    "contact_hours"
                ),

            "module_count":
                len(
                    primary.get(
                        "modules",
                        [],
                    )
                ),

            "topic_count":
                len(
                    get_all_topics(
                        primary
                    )
                ),

        },


        # ----------------------------------------------------
        # BENCHMARK PROFILE
        # ----------------------------------------------------

        "benchmark_curriculum_profile": {

            "college":
                benchmark.get(
                    "college"
                ),

            "university":
                benchmark.get(
                    "university"
                ),

            "program":
                benchmark.get(
                    "program"
                ),

            "department":
                benchmark.get(
                    "department"
                ),

            "academic_year":
                benchmark.get(
                    "academic_year"
                ),

            "regulation":
                benchmark.get(
                    "regulation"
                ),

            "semester":
                benchmark.get(
                    "semester"
                ),

            "subject_name":
                benchmark.get(
                    "subject_name"
                ),

            "subject_code":
                benchmark.get(
                    "subject_code"
                ),

            "credits":
                benchmark.get(
                    "credits"
                ),

            "module_count":
                len(
                    benchmark.get(
                        "modules",
                        [],
                    )
                ),

            "topic_count":
                len(
                    get_all_topics(
                        benchmark
                    )
                ),

        },


        # ----------------------------------------------------
        # OVERALL COMPARISON
        # ----------------------------------------------------

        "overall_comparison": {

            "similarity_pct":
                comparison_summary[
                    "similarity_pct"
                ],

            "modules_compared":
                comparison_summary[
                    "modules_compared"
                ],

            "similar_topics":
                comparison_summary[
                    "similar_topics"
                ],

            "different_topics":
                comparison_summary[
                    "different_topics"
                ],

        },


        # ----------------------------------------------------
        # MODULE COMPARISON
        # ----------------------------------------------------

        "module_comparison":
            comparison.get(
                "module_comparisons",
                [],
            ),


        # ----------------------------------------------------
        # TOPIC COMPARISON
        # ----------------------------------------------------

        "topic_comparison":
            (
                topic_df.to_dict(
                    orient="records"
                )
                if not topic_df.empty
                else []
            ),


        # ----------------------------------------------------
        # TOPIC CATEGORIES
        # ----------------------------------------------------

        "topic_categories": {

            "similar":
                similar_topics,

            "primary_only":
                primary_only_topics,

            "benchmark_only":
                benchmark_only_topics,

        },


        # ----------------------------------------------------
        # DEEP CONCEPT INTELLIGENCE
        # ----------------------------------------------------

        "concept_intelligence": {

            "summary":
                concept_summary,

            "topic_analysis":
                enrichment,

        },


        # ----------------------------------------------------
        # GAP CANDIDATES
        # ----------------------------------------------------

        "gap_candidates": {

            "benchmark_gaps":
                benchmark_only_topics,

            "ai_identified_gaps":
                concept_summary.get(
                    "missing_concepts",
                    [],
                ),

            "combined_candidates":
                combined_gaps,

        },


        # ----------------------------------------------------
        # INDUSTRY SIGNALS
        # ----------------------------------------------------

        "industry_signals":
            concept_summary.get(
                "industry_concepts",
                [],
            ),


        # ----------------------------------------------------
        # EMERGING TECHNOLOGY SIGNALS
        # ----------------------------------------------------

        "emerging_signals":
            concept_summary.get(
                "emerging_concepts",
                [],
            ),


        # ----------------------------------------------------
        # TOOLS
        # ----------------------------------------------------

        "recommended_tools":
            concept_summary.get(
                "tools",
                [],
            ),


        # ----------------------------------------------------
        # PROJECTS
        # ----------------------------------------------------

        "recommended_project_areas":
            concept_summary.get(
                "project_areas",
                [],
            ),

    }


    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================

    similarity = comparison_summary[
        "similarity_pct"
    ]


    if similarity >= 85:

        alignment_level = (
            "Very High"
        )

    elif similarity >= 70:

        alignment_level = (
            "High"
        )

    elif similarity >= 50:

        alignment_level = (
            "Moderate"
        )

    elif similarity >= 30:

        alignment_level = (
            "Low"
        )

    else:

        alignment_level = (
            "Very Low"
        )


    executive_summary = {

        "primary_subject":
            safe_text(
                primary.get(
                    "subject_name"
                )
            ),

        "benchmark_subject":
            safe_text(
                benchmark.get(
                    "subject_name"
                )
            ),

        "overall_similarity_pct":
            similarity,

        "alignment_level":
            alignment_level,

        "benchmark_topics_missing":
            len(
                benchmark_only_topics
            ),

        "ai_potential_gaps":
            len(
                concept_summary.get(
                    "missing_concepts",
                    [],
                )
            ),

        "industry_concepts":
            len(
                concept_summary.get(
                    "industry_concepts",
                    [],
                )
            ),

        "emerging_concepts":
            len(
                concept_summary.get(
                    "emerging_concepts",
                    [],
                )
            ),

        "recommended_tools":
            len(
                concept_summary.get(
                    "tools",
                    [],
                )
            ),

        "project_areas":
            len(
                concept_summary.get(
                    "project_areas",
                    [],
                )
            ),

    }


    report[
        "executive_summary"
    ] = executive_summary


    # ========================================================
    # SAVE REPORT
    # ========================================================

    st.session_state[
        "curriculum_intelligence_report"
    ] = report


    # ========================================================
    # BUILD REPORT MARKDOWN
    # ========================================================

    markdown_parts = []


    markdown_parts.append(
        "# Curriculum Intelligence Report"
    )


    markdown_parts.append(
        ""
    )


    markdown_parts.append(
        f"Generated: "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # EXECUTIVE SUMMARY
    # ========================================================

    markdown_parts.append(
        "## 1. Executive Summary"
    )


    markdown_parts.append(
        ""
    )


    markdown_parts.append(

        f"- Primary Subject: "
        f"{safe_text(primary.get('subject_name'))}"

    )


    markdown_parts.append(

        f"- Benchmark Subject: "
        f"{safe_text(benchmark.get('subject_name'))}"

    )


    markdown_parts.append(

        f"- Overall Similarity: "
        f"{similarity:.1f}%"

    )


    markdown_parts.append(

        f"- Alignment Level: "
        f"{alignment_level}"

    )


    markdown_parts.append(

        f"- Benchmark-only Topics: "
        f"{len(benchmark_only_topics)}"

    )


    markdown_parts.append(

        f"- AI Potential Gaps: "
        f"{len(concept_summary.get('missing_concepts', []))}"

    )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # PRIMARY CURRICULUM
    # ========================================================

    markdown_parts.append(
        "## 2. Primary Curriculum"
    )


    markdown_parts.append(
        ""
    )


    markdown_parts.append(

        f"- College: "
        f"{safe_text(primary.get('college'))}"

    )


    markdown_parts.append(

        f"- University: "
        f"{safe_text(primary.get('university'))}"

    )


    markdown_parts.append(

        f"- Program: "
        f"{safe_text(primary.get('program'))}"

    )


    markdown_parts.append(

        f"- Subject: "
        f"{safe_text(primary.get('subject_name'))}"

    )


    markdown_parts.append(

        f"- Subject Code: "
        f"{safe_text(primary.get('subject_code'))}"

    )


    markdown_parts.append(

        f"- Modules: "
        f"{len(primary.get('modules', []))}"

    )


    markdown_parts.append(

        f"- Topics: "
        f"{len(get_all_topics(primary))}"

    )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # BENCHMARK CURRICULUM
    # ========================================================

    markdown_parts.append(
        "## 3. Benchmark Curriculum"
    )


    markdown_parts.append(
        ""
    )


    markdown_parts.append(

        f"- College: "
        f"{safe_text(benchmark.get('college'))}"

    )


    markdown_parts.append(

        f"- University: "
        f"{safe_text(benchmark.get('university'))}"

    )


    markdown_parts.append(

        f"- Program: "
        f"{safe_text(benchmark.get('program'))}"

    )


    markdown_parts.append(

        f"- Subject: "
        f"{safe_text(benchmark.get('subject_name'))}"

    )


    markdown_parts.append(

        f"- Subject Code: "
        f"{safe_text(benchmark.get('subject_code'))}"

    )


    markdown_parts.append(

        f"- Modules: "
        f"{len(benchmark.get('modules', []))}"

    )


    markdown_parts.append(

        f"- Topics: "
        f"{len(get_all_topics(benchmark))}"

    )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # MODULE COMPARISON
    # ========================================================

    markdown_parts.append(
        "## 4. Module-wise Comparison"
    )


    markdown_parts.append(
        ""
    )


    for index, module in enumerate(

        comparison.get(
            "module_comparisons",
            [],
        ),

        start=1,

    ):


        if not isinstance(
            module,
            dict,
        ):

            continue


        markdown_parts.append(

            f"### Module {index}: "
            f"{safe_text(module.get('primary_module'))}"

        )


        markdown_parts.append(

            f"- Benchmark Module: "
            f"{safe_text(module.get('benchmark_module'))}"

        )


        try:

            module_similarity = float(
                module.get(
                    "similarity_pct",
                    0,
                )
                or 0
            )

        except (
            TypeError,
            ValueError,
        ):

            module_similarity = 0


        markdown_parts.append(

            f"- Similarity: "
            f"{module_similarity:.1f}%"

        )


        similar = module.get(
            "similar_concepts",
            [],
        )


        primary_only = module.get(
            "primary_only",
            [],
        )


        benchmark_only = module.get(
            "benchmark_only",
            [],
        )


        if similar:

            markdown_parts.append(
                "- Similar Concepts:"
            )


            for item in similar:

                markdown_parts.append(
                    f"  - {item}"
                )


        if primary_only:

            markdown_parts.append(
                "- Primary-only Concepts:"
            )


            for item in primary_only:

                markdown_parts.append(
                    f"  - {item}"
                )


        if benchmark_only:

            markdown_parts.append(
                "- Benchmark-only Concepts:"
            )


            for item in benchmark_only:

                markdown_parts.append(
                    f"  - {item}"
                )


        markdown_parts.append(
            ""
        )


    # ========================================================
    # GAP SECTION
    # ========================================================

    markdown_parts.append(
        "## 5. Potential Curriculum Gaps"
    )


    markdown_parts.append(
        ""
    )


    if combined_gaps:

        for gap in combined_gaps:

            markdown_parts.append(
                f"- {gap}"
            )

    else:

        markdown_parts.append(
            "No potential gaps identified."
        )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # INDUSTRY SECTION
    # ========================================================

    markdown_parts.append(
        "## 6. Industry-Relevant Concepts"
    )


    markdown_parts.append(
        ""
    )


    industry_values = concept_summary.get(
        "industry_concepts",
        [],
    )


    if industry_values:

        for item in industry_values:

            markdown_parts.append(
                f"- {item}"
            )

    else:

        markdown_parts.append(
            "No industry concepts identified."
        )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # EMERGING SECTION
    # ========================================================

    markdown_parts.append(
        "## 7. Emerging Concepts"
    )


    markdown_parts.append(
        ""
    )


    emerging_values = concept_summary.get(
        "emerging_concepts",
        [],
    )


    if emerging_values:

        for item in emerging_values:

            markdown_parts.append(
                f"- {item}"
            )

    else:

        markdown_parts.append(
            "No emerging concepts identified."
        )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # TOOLS SECTION
    # ========================================================

    markdown_parts.append(
        "## 8. Recommended Tools & Technologies"
    )


    markdown_parts.append(
        ""
    )


    tool_values = concept_summary.get(
        "tools",
        [],
    )


    if tool_values:

        for item in tool_values:

            markdown_parts.append(
                f"- {item}"
            )

    else:

        markdown_parts.append(
            "No additional tools identified."
        )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # PROJECT SECTION
    # ========================================================

    markdown_parts.append(
        "## 9. Recommended Project Areas"
    )


    markdown_parts.append(
        ""
    )


    project_values = concept_summary.get(
        "project_areas",
        [],
    )


    if project_values:

        for item in project_values:

            markdown_parts.append(
                f"- {item}"
            )

    else:

        markdown_parts.append(
            "No project areas identified."
        )


    markdown_parts.append(
        ""
    )


    # ========================================================
    # RECOMMENDATION NOTE
    # ========================================================

    markdown_parts.append(
        "## 10. Recommendation Status"
    )


    markdown_parts.append(
        ""
    )


    markdown_parts.append(
        """
This report contains AI-generated curriculum intelligence
and preliminary gap signals.

Final curriculum enhancement decisions should consider:

- Industry / Job Description analysis
- Faculty / subject expert review
- Student level
- Course credits and contact hours
- Prerequisites
- Accreditation requirements
- Program outcomes
- Assessment structure
- Available infrastructure
- Teaching capacity
"""
    )


    markdown_report = "\n".join(
        markdown_parts
    )


    # ========================================================
    # SAVE MARKDOWN
    # ========================================================

    st.session_state[
        "curriculum_intelligence_markdown"
    ] = markdown_report


    # ========================================================
    # REPORT PREVIEW
    # ========================================================

    st.divider()

    st.subheader(
        "📄 Report Preview"
    )


    with st.expander(
        "View Executive Summary",
        expanded=True,
    ):

        st.markdown(
            f"""
### Curriculum Alignment

**Primary:** {safe_text(primary.get("subject_name"))}

**Benchmark:** {safe_text(benchmark.get("subject_name"))}

**Overall Similarity:** {similarity:.1f}%

**Alignment:** {alignment_level}

### Potential Gaps

{len(combined_gaps)}

### Industry Concepts

{len(industry_values)}

### Emerging Concepts

{len(emerging_values)}

### Recommended Tools

{len(tool_values)}

### Project Areas

{len(project_values)}
"""
        )


    # ========================================================
    # REPORT DOWNLOADS
    # ========================================================

    st.divider()

    st.subheader(
        "⬇️ Download Curriculum Intelligence Reports"
    )


    report_col1, report_col2 = (
        st.columns(2)
    )


    # ========================================================
    # JSON REPORT
    # ========================================================

    with report_col1:

        report_json = serialize_json(
            report
        )


        st.download_button(

            label=(
                "⬇️ Download Complete JSON Report"
            ),

            data=report_json,

            file_name=(
                "curriculum_intelligence_report.json"
            ),

            mime="application/json",

            use_container_width=True,

            key=(
                "download_final_curriculum_json"
            ),

        )


    # ========================================================
    # MARKDOWN REPORT
    # ========================================================

    with report_col2:

        st.download_button(

            label=(
                "⬇️ Download Markdown Report"
            ),

            data=markdown_report,

            file_name=(
                "curriculum_intelligence_report.md"
            ),

            mime="text/markdown",

            use_container_width=True,

            key=(
                "download_final_curriculum_markdown"
            ),

        )


    # ========================================================
    # MODULE CSV
    # ========================================================

    if not module_df.empty:

        st.download_button(

            label=(
                "⬇️ Download Module Comparison CSV"
            ),

            data=module_df.to_csv(
                index=False
            ),

            file_name=(
                "curriculum_module_comparison.csv"
            ),

            mime="text/csv",

            key=(
                "download_final_module_csv"
            ),

        )


    # ========================================================
    # TOPIC CSV
    # ========================================================

    if not topic_df.empty:

        st.download_button(

            label=(
                "⬇️ Download Topic Comparison CSV"
            ),

            data=topic_df.to_csv(
                index=False
            ),

            file_name=(
                "curriculum_topic_comparison.csv"
            ),

            mime="text/csv",

            key=(
                "download_final_topic_csv"
            ),

        )


    # ========================================================
    # CONCEPT CSV
    # ========================================================

    if not concept_df.empty:

        st.download_button(

            label=(
                "⬇️ Download Concept Intelligence CSV"
            ),

            data=concept_df.to_csv(
                index=False
            ),

            file_name=(
                "curriculum_concept_intelligence.csv"
            ),

            mime="text/csv",

            key=(
                "download_final_concept_csv"
            ),

        )


    # ========================================================
    # GAP CSV
    # ========================================================

    if combined_gaps:

        gap_report_df = pd.DataFrame({

            "Potential Curriculum Gap":
                combined_gaps,

        })


        st.download_button(

            label=(
                "⬇️ Download Gap Candidates CSV"
            ),

            data=gap_report_df.to_csv(
                index=False
            ),

            file_name=(
                "curriculum_gap_candidates.csv"
            ),

            mime="text/csv",

            key=(
                "download_final_gap_csv"
            ),

        )


    # ========================================================
    # FINAL SESSION STATE PACKAGE
    # ========================================================

    st.session_state[
        "curriculum_report_package"
    ] = {

        "report":
            report,

        "markdown":
            markdown_report,

        "module_dataframe":
            (
                module_df.to_dict(
                    orient="records"
                )
                if not module_df.empty
                else []
            ),

        "topic_dataframe":
            (
                topic_df.to_dict(
                    orient="records"
                )
                if not topic_df.empty
                else []
            ),

        "concept_dataframe":
            (
                concept_df.to_dict(
                    orient="records"
                )
                if not concept_df.empty
                else []
            ),

        "gap_candidates":
            combined_gaps,

    }


    # ========================================================
    # PAGE 05 HANDOFF
    # ========================================================

    st.session_state[
        "reports_page_data"
    ] = {

        "curriculum_intelligence":
            report,

        "curriculum_markdown":
            markdown_report,

        "module_comparison":
            (
                module_df.to_dict(
                    orient="records"
                )
                if not module_df.empty
                else []
            ),

        "topic_comparison":
            (
                topic_df.to_dict(
                    orient="records"
                )
                if not topic_df.empty
                else []
            ),

        "concept_intelligence":
            (
                concept_df.to_dict(
                    orient="records"
                )
                if not concept_df.empty
                else []
            ),

        "gap_candidates":
            combined_gaps,

    }


    # ========================================================
    # COMPLETION STATUS
    # ========================================================

    st.divider()

    st.success(
        """
        ✅ **Curriculum Intelligence Report Generated**

        Page 02 is now complete.

        The output is ready for:

        **💼 Industry & JD Intelligence**

        and subsequently:

        **🔍 Gap & Enhancement**
        """
    )


# ============================================================
# END OF CHUNK 9
# ============================================================
# ============================================================
# CHUNK 10/10
# FINAL VALIDATION + MASTER EXPORT + HANDOFF
# ============================================================

st.divider()
st.header("✅ Curriculum Intelligence — Final Status")


# ============================================================
# 1. LOAD SESSION STATE
# ============================================================

primary = st.session_state.get("primary_syllabus", {})
benchmark = st.session_state.get("benchmark_syllabus")
comparison = st.session_state.get("curriculum_comparison")
enrichment = st.session_state.get("concept_enrichment", [])
report = st.session_state.get("curriculum_intelligence_report")


# ============================================================
# 2. PIPELINE STATUS
# ============================================================

primary_ready = bool(primary)
benchmark_ready = bool(benchmark)
comparison_ready = bool(comparison)
concept_ready = bool(enrichment)
report_ready = bool(report)

pipeline_complete = all([
    primary_ready,
    benchmark_ready,
    comparison_ready,
    concept_ready,
    report_ready,
])


status_df = pd.DataFrame([
    {
        "Component": "Primary Syllabus",
        "Status": "Ready" if primary_ready else "Pending",
    },
    {
        "Component": "Benchmark Curriculum",
        "Status": "Ready" if benchmark_ready else "Pending",
    },
    {
        "Component": "Curriculum Comparison",
        "Status": "Ready" if comparison_ready else "Pending",
    },
    {
        "Component": "Deep Concept Intelligence",
        "Status": "Ready" if concept_ready else "Pending",
    },
    {
        "Component": "Curriculum Report",
        "Status": "Ready" if report_ready else "Pending",
    },
])


st.dataframe(
    status_df,
    use_container_width=True,
    hide_index=True,
)


# ============================================================
# 3. CURRICULUM SUMMARY
# ============================================================

if primary_ready:

    primary_topics = get_all_topics(primary)

    st.subheader("📘 Primary Curriculum")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric(
            "Modules",
            len(primary.get("modules", [])),
        )

    with col2:
        st.metric(
            "Topics",
            len(primary_topics),
        )

    with col3:
        st.metric(
            "COs",
            len(primary.get("course_outcomes", [])),
        )

    with col4:
        st.metric(
            "POs",
            len(primary.get("program_outcomes", [])),
        )


# ============================================================
# 4. BENCHMARK SUMMARY
# ============================================================

if benchmark_ready:

    benchmark_topics = get_all_topics(benchmark)

    st.subheader("📗 Benchmark Curriculum")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric(
            "Modules",
            len(benchmark.get("modules", [])),
        )

    with col2:
        st.metric(
            "Topics",
            len(benchmark_topics),
        )

    with col3:
        st.metric(
            "Tools",
            len(benchmark.get("tools", [])),
        )


# ============================================================
# 5. COMPARISON SUMMARY
# ============================================================

comparison_summary = {}

if comparison_ready:

    comparison_summary = get_comparison_summary(
        comparison
    )

    similarity = float(
        comparison_summary.get(
            "similarity_pct",
            0,
        ) or 0
    )

    st.subheader("📊 Curriculum Comparison")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric(
            "Similarity",
            f"{similarity:.1f}%",
        )

    with col2:
        st.metric(
            "Modules Compared",
            comparison_summary.get(
                "modules_compared",
                0,
            ),
        )

    with col3:
        st.metric(
            "Similar Topics",
            comparison_summary.get(
                "similar_topics",
                0,
            ),
        )

    with col4:
        st.metric(
            "Different Topics",
            comparison_summary.get(
                "different_topics",
                0,
            ),
        )

    st.progress(
        max(
            0,
            min(
                similarity / 100,
                1,
            ),
        )
    )


# ============================================================
# 6. GAP SUMMARY
# ============================================================

structural_gaps = []

ai_gaps = []

if comparison_ready:

    structural_gaps = get_benchmark_only_topics(
        comparison
    )

if concept_ready:

    concept_summary = build_concept_summary(
        enrichment
    )

    ai_gaps = concept_summary.get(
        "missing_concepts",
        []
    )


combined_gaps = sorted(
    set(structural_gaps) | set(ai_gaps),
    key=str.lower,
)


st.subheader("⚠️ Curriculum Gap Signals")

col1, col2, col3 = st.columns(3)

with col1:
    st.metric(
        "Benchmark Gaps",
        len(structural_gaps),
    )

with col2:
    st.metric(
        "AI Gaps",
        len(ai_gaps),
    )

with col3:
    st.metric(
        "Combined Candidates",
        len(combined_gaps),
    )


if combined_gaps:

    with st.expander(
        "View Gap Candidates"
    ):

        for index, gap in enumerate(
            combined_gaps,
            start=1,
        ):

            st.write(
                f"{index}. {gap}"
            )


# ============================================================
# 7. CONCEPT INTELLIGENCE SUMMARY
# ============================================================

industry_concepts = []
emerging_concepts = []
recommended_tools = []
project_areas = []

if concept_ready:

    industry_concepts = concept_summary.get(
        "industry_concepts",
        []
    )

    emerging_concepts = concept_summary.get(
        "emerging_concepts",
        []
    )

    recommended_tools = concept_summary.get(
        "tools",
        []
    )

    project_areas = concept_summary.get(
        "project_areas",
        []
    )


    st.subheader(
        "🧠 Concept Intelligence"
    )

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric(
            "Industry Concepts",
            len(industry_concepts),
        )

    with col2:
        st.metric(
            "Emerging Concepts",
            len(emerging_concepts),
        )

    with col3:
        st.metric(
            "Tools",
            len(recommended_tools),
        )

    with col4:
        st.metric(
            "Projects",
            len(project_areas),
        )


# ============================================================
# 8. MASTER PACKAGE
# ============================================================

master_package = {
    "metadata": {
        "platform": "PragyanAI AI Curriculum Platform",
        "module": "Curriculum Intelligence",
        "page": "02_Curriculum_Intelligence",
        "version": "1.0",
        "generated_at": datetime.now().isoformat(),
    },

    "primary_syllabus": primary,

    "benchmark_syllabus": benchmark,

    "curriculum_comparison": comparison,

    "concept_intelligence": enrichment,

    "gap_candidates": {
        "structural": structural_gaps,
        "ai_identified": ai_gaps,
        "combined": combined_gaps,
    },

    "industry_signals": industry_concepts,

    "emerging_signals": emerging_concepts,

    "recommended_tools": recommended_tools,

    "project_areas": project_areas,

    "report": report,

    "summary": {
        "primary_subject": primary.get(
            "subject_name"
        ),

        "benchmark_subject": (
            benchmark.get("subject_name")
            if benchmark
            else None
        ),

        "similarity_pct": (
            similarity
            if comparison_ready
            else None
        ),

        "primary_modules": len(
            primary.get("modules", [])
        ),

        "primary_topics": len(
            get_all_topics(primary)
        ),

        "benchmark_modules": (
            len(
                benchmark.get("modules", [])
            )
            if benchmark
            else 0
        ),

        "benchmark_topics": (
            len(
                get_all_topics(benchmark)
            )
            if benchmark
            else 0
        ),

        "concept_topics": len(
            enrichment
        ),

        "gap_candidates": len(
            combined_gaps
        ),

        "industry_concepts": len(
            industry_concepts
        ),

        "emerging_concepts": len(
            emerging_concepts
        ),

        "recommended_tools": len(
            recommended_tools
        ),

        "project_areas": len(
            project_areas
        ),
    },
}


# ============================================================
# 9. SAVE MASTER PACKAGE
# ============================================================

st.session_state[
    "curriculum_master_package"
] = master_package


st.session_state[
    "curriculum_intelligence_status"
] = {
    "primary_ready": primary_ready,
    "benchmark_ready": benchmark_ready,
    "comparison_ready": comparison_ready,
    "concept_ready": concept_ready,
    "report_ready": report_ready,
    "complete": pipeline_complete,
}


# ============================================================
# 10. MASTER JSON DOWNLOAD
# ============================================================

master_json = serialize_json(
    master_package
)


st.subheader(
    "📦 Export Curriculum Intelligence"
)


st.download_button(
    label="⬇️ Download Master Intelligence Package",
    data=master_json,
    file_name="curriculum_intelligence_master.json",
    mime="application/json",
    use_container_width=True,
)


# ============================================================
# 11. DOWNSTREAM HANDOFF
# ============================================================

page03_ready = (
    primary_ready
    and comparison_ready
)

page04_ready = (
    primary_ready
    and comparison_ready
    and concept_ready
)

page05_ready = report_ready


st.divider()

st.subheader(
    "🔗 Module Handoff"
)


col1, col2, col3 = st.columns(3)


with col1:

    st.markdown(
        "### 💼 Page 03"
    )

    st.caption(
        "Industry & JD Intelligence"
    )

    if page03_ready:
        st.success("READY")
    else:
        st.warning("PENDING")


with col2:

    st.markdown(
        "### 🔍 Page 04"
    )

    st.caption(
        "Gap & Curriculum Enhancement"
    )

    if page04_ready:
        st.success("READY")
    else:
        st.warning("PENDING")


with col3:

    st.markdown(
        "### 📊 Page 05"
    )

    st.caption(
        "Reports"
    )

    if page05_ready:
        st.success("READY")
    else:
        st.warning("PENDING")


# ============================================================
# 12. HANDOFF OBJECTS
# ============================================================

st.session_state[
    "industry_jd_input"
] = {
    "primary_syllabus": primary,
    "curriculum_comparison": comparison,
    "curriculum_intelligence": master_package,
}


st.session_state[
    "gap_enhancement_input"
] = {
    "primary_syllabus": primary,
    "benchmark_syllabus": benchmark,
    "curriculum_comparison": comparison,
    "concept_intelligence": enrichment,
    "gap_candidates": combined_gaps,
    "industry_signals": industry_concepts,
    "emerging_signals": emerging_concepts,
    "recommended_tools": recommended_tools,
    "project_areas": project_areas,
}


st.session_state[
    "reports_input"
] = master_package


# ============================================================
# 13. DEVELOPER DIAGNOSTICS
# ============================================================

with st.expander(
    "🛠 Developer Diagnostics"
):

    st.json({
        "primary_ready": primary_ready,
        "benchmark_ready": benchmark_ready,
        "comparison_ready": comparison_ready,
        "concept_ready": concept_ready,
        "report_ready": report_ready,
        "page03_ready": page03_ready,
        "page04_ready": page04_ready,
        "page05_ready": page05_ready,
        "pipeline_complete": pipeline_complete,
    })


# ============================================================
# 14. FINAL STATUS
# ============================================================

if pipeline_complete:

    st.success(
        """
        🎉 **Curriculum Intelligence Complete**

        Page 02 has successfully produced the structured
        curriculum intelligence package.

        Next:

        **Page 03 → Industry & JD Intelligence**

        **Page 04 → Gap & Curriculum Enhancement**

        **Page 05 → Reports**
        """
    )

else:

    st.info(
        """
        Curriculum Intelligence is not yet complete.

        Complete the pending stages above.
        """
    )


# ============================================================
# 15. FINAL PAGE FLAG
# ============================================================

st.session_state[
    "curriculum_intelligence_page_complete"
] = pipeline_complete


# ============================================================
# END OF CHUNK 10
# ============================================================
