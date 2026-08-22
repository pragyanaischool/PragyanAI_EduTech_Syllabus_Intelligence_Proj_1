# ============================================================
# rag/docx_loader.py
# CHUNK 1/10
#
# DOCX DOCUMENT LOADER
#
# Purpose:
#   Load Microsoft Word DOCX documents for the RAG /
#   Curriculum Intelligence pipeline.
#
# Features:
#   - Paragraph extraction
#   - Heading detection
#   - Table extraction
#   - Header / footer extraction
#   - Hyperlink extraction
#   - Document metadata
#   - Page-break awareness
#   - Text cleaning
#   - RAG-ready document objects
#   - LangChain compatibility
#   - Batch loading
#
# Dependencies:
#   pip install python-docx
#
# Optional:
#   pip install langchain-core
#
# ============================================================

from __future__ import annotations

import logging
import re

from dataclasses import dataclass, field

from pathlib import Path

from typing import (
    Any,
    Dict,
    Iterable,
    List,
    Optional,
    Sequence,
    Union,
)


# ============================================================
# LOGGING
# ============================================================

logger = logging.getLogger(__name__)


# ============================================================
# VERSION
# ============================================================

DOCX_LOADER_VERSION = "1.0.0"


# ============================================================
# CONSTANTS
# ============================================================

DEFAULT_MIN_TEXT_LENGTH = 10

SUPPORTED_DOCX_EXTENSIONS = {
    ".docx",
}


# ============================================================
# OPTIONAL PYTHON-DOCX IMPORT
# ============================================================

try:
    from docx import Document as PythonDocxDocument
except ImportError:
    PythonDocxDocument = None


# ============================================================
# OPTIONAL LANGCHAIN IMPORT
# ============================================================

try:
    from langchain_core.documents import Document as LangChainDocument
except ImportError:
    LangChainDocument = None


# ============================================================
# DOCUMENT BLOCK
# ============================================================

@dataclass
class DOCXBlock:

    block_type: str = "paragraph"

    text: str = ""

    index: int = 0

    section: str = ""

    heading_level: int = 0

    style: str = ""

    metadata: Dict[str, Any] = field(
        default_factory=dict
    )


# ============================================================
# DOCX DOCUMENT
# ============================================================

@dataclass
class DOCXDocument:

    text: str = ""

    metadata: Dict[str, Any] = field(
        default_factory=dict
    )

    source: str = ""

    filename: str = ""

    file_path: str = ""

    block_index: int = 0

    block_type: str = "paragraph"

    heading_level: int = 0

    section: str = ""

    extraction_method: str = "python-docx"

    error: Optional[str] = None


# ============================================================
# DOCX LOAD RESULT
# ============================================================

@dataclass
class DOCXLoadResult:

    documents: List[DOCXDocument] = field(
        default_factory=list
    )

    blocks: List[DOCXBlock] = field(
        default_factory=list
    )

    metadata: Dict[str, Any] = field(
        default_factory=dict
    )

    paragraph_count: int = 0

    table_count: int = 0

    table_row_count: int = 0

    table_cell_count: int = 0

    heading_count: int = 0

    hyperlink_count: int = 0

    success: bool = True

    error: Optional[str] = None

    warnings: List[str] = field(
        default_factory=list
    )


# ============================================================
# END CHUNK 1
# ============================================================
# ============================================================
# CHUNK 2/10
#
# BASIC UTILITIES
# ============================================================


# ============================================================
# CHECK PYTHON-DOCX
# ============================================================

def is_python_docx_available() -> bool:

    return PythonDocxDocument is not None


# ============================================================
# REQUIRE PYTHON-DOCX
# ============================================================

def require_python_docx() -> None:

    if PythonDocxDocument is None:

        raise ImportError(
            "python-docx is required for DOCX loading. "
            "Install it using: pip install python-docx"
        )


# ============================================================
# CLEAN TEXT
# ============================================================

def clean_text(
    text: Any,
) -> str:

    if text is None:
        return ""

    text = str(text)

    # Normalize line endings.
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Remove null characters.
    text = text.replace("\x00", "")

    # Normalize tabs.
    text = text.replace("\t", " ")

    # Normalize spaces.
    text = re.sub(
        r"[ ]{2,}",
        " ",
        text,
    )

    # Normalize newlines.
    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text,
    )

    lines = []

    for line in text.split("\n"):

        line = line.strip()

        lines.append(line)

    # Remove repeated blank lines.
    output = []

    previous_blank = False

    for line in lines:

        if not line:

            if previous_blank:
                continue

            previous_blank = True

            output.append("")

        else:

            previous_blank = False

            output.append(line)

    return "\n".join(output).strip()


# ============================================================
# NORMALIZE FILENAME
# ============================================================

def normalize_filename(
    filename: Any,
) -> str:

    if filename is None:
        return ""

    return Path(
        str(filename)
    ).name


# ============================================================
# NORMALIZE SOURCE
# ============================================================

def normalize_source(
    source: Any,
) -> str:

    if source is None:
        return ""

    return str(source).strip()


# ============================================================
# NORMALIZE STYLE NAME
# ============================================================

def normalize_style_name(
    style: Any,
) -> str:

    if style is None:
        return ""

    try:
        name = style.name
    except Exception:
        name = str(style)

    return str(name or "").strip()


# ============================================================
# END CHUNK 2
# ============================================================
# ============================================================
# CHUNK 3/10
#
# PARAGRAPH + HEADING EXTRACTION
# ============================================================


# ============================================================
# GET HEADING LEVEL
# ============================================================

def get_heading_level(
    paragraph: Any,
) -> int:

    try:

        style_name = normalize_style_name(
            paragraph.style
        )

    except Exception:

        return 0

    # Standard Word heading styles.
    match = re.search(
        r"(?:Heading|Title)\s*(\d+)?",
        style_name,
        flags=re.IGNORECASE,
    )

    if not match:
        return 0

    number = match.group(1)

    if number is None:
        return 1

    try:
        return int(number)
    except ValueError:
        return 1


# ============================================================
# IS HEADING
# ============================================================

def is_heading(
    paragraph: Any,
) -> bool:

    return get_heading_level(
        paragraph
    ) > 0


# ============================================================
# EXTRACT PARAGRAPH
# ============================================================

def extract_paragraph(
    paragraph: Any,
    index: int,
    current_section: str = "",
) -> DOCXBlock:

    text = clean_text(
        paragraph.text
    )

    style_name = normalize_style_name(
        getattr(
            paragraph,
            "style",
            None,
        )
    )

    heading_level = get_heading_level(
        paragraph
    )

    if heading_level > 0:

        block_type = "heading"

        section = text or current_section

    else:

        block_type = "paragraph"

        section = current_section

    metadata = {
        "style": style_name,
        "index": index,
        "heading_level": heading_level,
    }

    # --------------------------------------------------------
    # Paragraph alignment
    # --------------------------------------------------------

    try:

        alignment = paragraph.alignment

        if alignment is not None:

            metadata["alignment"] = str(
                alignment
            )

    except Exception:
        pass

    return DOCXBlock(

        block_type=block_type,

        text=text,

        index=index,

        section=section,

        heading_level=heading_level,

        style=style_name,

        metadata=metadata,

    )


# ============================================================
# EXTRACT PARAGRAPHS
# ============================================================

def extract_paragraphs(
    document: Any,
) -> List[DOCXBlock]:

    blocks = []

    current_section = ""

    for index, paragraph in enumerate(
        document.paragraphs
    ):

        block = extract_paragraph(

            paragraph=paragraph,

            index=index,

            current_section=current_section,

        )

        if block.heading_level > 0:

            current_section = block.text

            block.section = current_section

        blocks.append(block)

    return blocks


# ============================================================
# FILTER TEXT PARAGRAPHS
# ============================================================

def get_text_paragraphs(
    document: Any,
    skip_empty: bool = True,
) -> List[DOCXBlock]:

    blocks = extract_paragraphs(
        document
    )

    if not skip_empty:
        return blocks

    return [

        block

        for block in blocks

        if block.text.strip()

    ]


# ============================================================
# END CHUNK 3
# ============================================================
# ============================================================
# CHUNK 4/10
#
# TABLE EXTRACTION
# ============================================================


# ============================================================
# EXTRACT TABLE CELL
# ============================================================

def extract_table_cell(
    cell: Any,
) -> str:

    if cell is None:
        return ""

    try:

        paragraphs = cell.paragraphs

        texts = [

            clean_text(
                paragraph.text
            )

            for paragraph
            in paragraphs

        ]

        texts = [
            text
            for text in texts
            if text
        ]

        if texts:
            return "\n".join(texts)

    except Exception:
        pass

    try:
        return clean_text(
            cell.text
        )
    except Exception:
        return ""


# ============================================================
# EXTRACT TABLE
# ============================================================

def extract_table(
    table: Any,
    table_index: int,
    start_index: int,
) -> List[DOCXBlock]:

    blocks = []

    rows = getattr(
        table,
        "rows",
        [],
    )

    for row_index, row in enumerate(rows):

        cells = getattr(
            row,
            "cells",
            [],
        )

        values = [

            extract_table_cell(
                cell
            )

            for cell
            in cells

        ]

        values = [
            value
            for value in values
            if value
        ]

        if not values:
            continue

        row_text = " | ".join(
            values
        )

        block = DOCXBlock(

            block_type="table_row",

            text=row_text,

            index=start_index + row_index,

            section=f"Table {table_index + 1}",

            heading_level=0,

            style="",

            metadata={

                "table_index":
                    table_index,

                "row_index":
                    row_index,

                "cell_count":
                    len(cells),

                "cells":
                    values,

            },

        )

        blocks.append(block)

    return blocks


# ============================================================
# EXTRACT ALL TABLES
# ============================================================

def extract_tables(
    document: Any,
    start_index: int = 0,
) -> List[DOCXBlock]:

    blocks = []

    tables = getattr(
        document,
        "tables",
        [],
    )

    current_index = start_index

    for table_index, table in enumerate(
        tables
    ):

        table_blocks = extract_table(

            table=table,

            table_index=table_index,

            start_index=current_index,

        )

        blocks.extend(
            table_blocks
        )

        current_index += len(
            table_blocks
        )

    return blocks


# ============================================================
# TABLE AS TEXT
# ============================================================

def table_to_text(
    table: Any,
) -> str:

    rows = []

    for row in getattr(
        table,
        "rows",
        [],
    ):

        cells = [

            extract_table_cell(
                cell
            )

            for cell
            in getattr(
                row,
                "cells",
                [],
            )

        ]

        cells = [
            cell
            for cell in cells
            if cell
        ]

        if cells:

            rows.append(
                " | ".join(cells)
            )

    return "\n".join(
        rows
    )


# ============================================================
# END CHUNK 4
# ============================================================
# ============================================================
# CHUNK 5/10
#
# HEADERS / FOOTERS / HYPERLINKS
# ============================================================


# ============================================================
# EXTRACT HEADER FOOTER
# ============================================================

def extract_headers_footers(
    document: Any,
) -> Dict[str, List[str]]:

    headers = []

    footers = []

    sections = getattr(
        document,
        "sections",
        [],
    )

    for section_index, section in enumerate(
        sections
    ):

        # ----------------------------------------------------
        # Header
        # ----------------------------------------------------

        try:

            for paragraph in section.header.paragraphs:

                text = clean_text(
                    paragraph.text
                )

                if text:

                    headers.append(text)

        except Exception as exc:

            logger.debug(
                "Header extraction failed: %s",
                exc,
            )

        # ----------------------------------------------------
        # Footer
        # ----------------------------------------------------

        try:

            for paragraph in section.footer.paragraphs:

                text = clean_text(
                    paragraph.text
                )

                if text:

                    footers.append(text)

        except Exception as exc:

            logger.debug(
                "Footer extraction failed: %s",
                exc,
            )

    return {

        "headers":
            list(
                dict.fromkeys(headers)
            ),

        "footers":
            list(
                dict.fromkeys(footers)
            ),

    }


# ============================================================
# EXTRACT HYPERLINKS
# ============================================================

def extract_hyperlinks(
    document: Any,
) -> List[Dict[str, str]]:

    hyperlinks = []

    # python-docx does not expose hyperlinks as a simple
    # first-class API, so we inspect the underlying XML.

    try:

        relationships = document.part.rels

        relationship_map = {

            rel_id:
                rel.target_ref

            for rel_id, rel
            in relationships.items()

            if "hyperlink"
            in str(
                rel.reltype
            ).lower()

        }

    except Exception:

        relationship_map = {}

    # --------------------------------------------------------
    # Paragraph hyperlinks
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        try:

            for hyperlink in paragraph._p.xpath(
                ".//w:hyperlink"
            ):

                rel_id = hyperlink.get(

                    "{http://schemas.openxmlformats.org/"
                    "officeDocument/2006/relationships}id"

                )

                url = relationship_map.get(
                    rel_id,
                    "",
                )

                text_nodes = hyperlink.xpath(
                    ".//w:t"
                )

                text = "".join(

                    node.text or ""

                    for node
                    in text_nodes

                )

                text = clean_text(
                    text
                )

                if text or url:

                    hyperlinks.append({

                        "text":
                            text,

                        "url":
                            url,

                    })

        except Exception:
            continue

    # --------------------------------------------------------
    # Deduplicate
    # --------------------------------------------------------

    unique = []

    seen = set()

    for item in hyperlinks:

        key = (

            item.get("text", ""),

            item.get("url", ""),

        )

        if key in seen:
            continue

        seen.add(key)

        unique.append(item)

    return unique


# ============================================================
# END CHUNK 5
# ============================================================
# ============================================================
# CHUNK 6/10
#
# CORE DOCX LOADING
# ============================================================


# ============================================================
# VALIDATE DOCX PATH
# ============================================================

def validate_docx_path(
    file_path: Union[
        str,
        Path,
    ],
) -> Path:

    path = Path(
        file_path
    )

    if not path.exists():

        raise FileNotFoundError(
            f"DOCX file does not exist: {path}"
        )

    if not path.is_file():

        raise ValueError(
            f"Path is not a file: {path}"
        )

    if path.suffix.lower() not in (
        SUPPORTED_DOCX_EXTENSIONS
    ):

        raise ValueError(

            f"Unsupported file type: {path.suffix}. "
            "Expected .docx"

        )

    return path


# ============================================================
# OPEN DOCX
# ============================================================

def open_docx(
    file_path: Union[
        str,
        Path,
    ],
) -> Any:

    require_python_docx()

    path = validate_docx_path(
        file_path
    )

    try:

        return PythonDocxDocument(
            str(path)
        )

    except Exception as exc:

        logger.exception(
            "Unable to open DOCX: %s",
            path,
        )

        raise RuntimeError(
            f"Unable to open DOCX: {path}"
        ) from exc


# ============================================================
# EXTRACT CORE DOCUMENT METADATA
# ============================================================

def extract_core_metadata(
    document: Any,
) -> Dict[str, Any]:

    metadata = {}

    try:

        core = document.core_properties

        fields = [

            "author",
            "category",
            "comments",
            "content_status",
            "created",
            "identifier",
            "keywords",
            "language",
            "last_modified_by",
            "last_printed",
            "modified",
            "revision",
            "subject",
            "title",
            "version",

        ]

        for field_name in fields:

            try:

                value = getattr(
                    core,
                    field_name,
                    None,
                )

                if value is not None:

                    value = str(
                        value
                    ).strip()

                    if value:

                        metadata[
                            field_name
                        ] = value

            except Exception:
                continue

    except Exception as exc:

        logger.debug(
            "Core metadata extraction failed: %s",
            exc,
        )

    return metadata


# ============================================================
# BUILD DOCUMENT OBJECTS
# ============================================================

def blocks_to_documents(
    blocks: Sequence[DOCXBlock],
    source: str = "",
    filename: str = "",
    file_path: str = "",
    base_metadata: Optional[
        Mapping[str, Any]
    ] = None,
) -> List[DOCXDocument]:

    documents = []

    base = dict(
        base_metadata
        or
        {}
    )

    for block in blocks:

        if not block.text.strip():
            continue

        metadata = dict(base)

        metadata.update(
            block.metadata
        )

        metadata.update({

            "block_index":
                block.index,

            "block_type":
                block.block_type,

            "section":
                block.section,

            "heading_level":
                block.heading_level,

            "style":
                block.style,

        })

        documents.append(

            DOCXDocument(

                text=block.text,

                metadata=metadata,

                source=source,

                filename=filename,

                file_path=file_path,

                block_index=block.index,

                block_type=block.block_type,

                heading_level=block.heading_level,

                section=block.section,

                extraction_method="python-docx",

            )

        )

    return documents


# ============================================================
# LOAD DOCX
# ============================================================

def load_docx(
    file_path: Union[
        str,
        Path,
    ],
    include_tables: bool = True,
    include_headers_footers: bool = False,
    include_empty_blocks: bool = False,
) -> DOCXLoadResult:

    path = validate_docx_path(
        file_path
    )

    try:

        document = open_docx(
            path
        )

        core_metadata = extract_core_metadata(
            document
        )

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        paragraph_blocks = extract_paragraphs(
            document
        )

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        table_blocks = []

        if include_tables:

            table_blocks = extract_tables(

                document,

                start_index=len(
                    paragraph_blocks
                ),

            )

        blocks = (

            paragraph_blocks
            +
            table_blocks

        )

        # ----------------------------------------------------
        # Optional headers / footers
        # ----------------------------------------------------

        header_footer = {

            "headers": [],
            "footers": [],

        }

        if include_headers_footers:

            header_footer = extract_headers_footers(
                document
            )

        # ----------------------------------------------------
        # Hyperlinks
        # ----------------------------------------------------

        hyperlinks = extract_hyperlinks(
            document
        )

        # ----------------------------------------------------
        # Base metadata
        # ----------------------------------------------------

        metadata = dict(
            core_metadata
        )

        metadata.update({

            "filename":
                path.name,

            "file_path":
                str(
                    path.resolve()
                ),

            "paragraph_count":
                len(
                    document.paragraphs
                ),

            "table_count":
                len(
                    document.tables
                ),

            "hyperlink_count":
                len(
                    hyperlinks
                ),

            "headers":
                header_footer[
                    "headers"
                ],

            "footers":
                header_footer[
                    "footers"
                ],

            "hyperlinks":
                hyperlinks,

        })

        # ----------------------------------------------------
        # Convert blocks
        # ----------------------------------------------------

        documents = blocks_to_documents(

            blocks=blocks,

            source=str(
                path
            ),

            filename=path.name,

            file_path=str(
                path.resolve()
            ),

            base_metadata=metadata,

        )

        if include_empty_blocks:

            # Preserve empty blocks if explicitly requested.
            documents = []

            for block in blocks:

                metadata_copy = dict(
                    metadata
                )

                metadata_copy.update(
                    block.metadata
                )

                documents.append(

                    DOCXDocument(

                        text=block.text,

                        metadata=metadata_copy,

                        source=str(path),

                        filename=path.name,

                        file_path=str(
                            path.resolve()
                        ),

                        block_index=block.index,

                        block_type=block.block_type,

                        heading_level=block.heading_level,

                        section=block.section,

                        extraction_method="python-docx",

                    )

                )

        # ----------------------------------------------------
        # Statistics
        # ----------------------------------------------------

        table_row_count = sum(

            len(
                table.rows
            )

            for table
            in document.tables

        )

        table_cell_count = sum(

            len(
                row.cells
            )

            for table
            in document.tables

            for row
            in table.rows

        )

        heading_count = sum(

            1

            for block
            in paragraph_blocks

            if block.heading_level > 0

        )

        warnings = []

        if not documents:

            warnings.append(

                "No extractable text was found in the DOCX."

            )

        if len(
            document.tables
        ) > 0 and not include_tables:

            warnings.append(

                "DOCX contains tables but table extraction "
                "was disabled."

            )

        return DOCXLoadResult(

            documents=documents,

            blocks=blocks,

            metadata=metadata,

            paragraph_count=len(
                document.paragraphs
            ),

            table_count=len(
                document.tables
            ),

            table_row_count=table_row_count,

            table_cell_count=table_cell_count,

            heading_count=heading_count,

            hyperlink_count=len(
                hyperlinks
            ),

            success=True,

            warnings=warnings,

        )

    except Exception as exc:

        logger.exception(
            "DOCX loading failed: %s",
            path,
        )

        return DOCXLoadResult(

            documents=[],

            blocks=[],

            metadata={

                "filename":
                    path.name,

                "file_path":
                    str(path),

            },

            success=False,

            error=str(
                exc
            ),

            warnings=[],

        )


# ============================================================
# STRICT LOADING
# ============================================================

def load_docx_strict(
    file_path: Union[
        str,
        Path,
    ],
) -> DOCXLoadResult:

    result = load_docx(
        file_path
    )

    if not result.success:

        raise RuntimeError(

            result.error
            or
            "Unknown DOCX loading error."

        )

    return result


# ============================================================
# END CHUNK 6
# ============================================================
# ============================================================
# CHUNK 7/10
#
# BYTES / STREAM / TEXT CONVERSION
# ============================================================


# ============================================================
# OPEN DOCX BYTES
# ============================================================

def open_docx_bytes(
    data: bytes,
) -> Any:

    require_python_docx()

    if not data:

        raise ValueError(
            "DOCX data is empty."
        )

    try:

        from io import BytesIO

        return PythonDocxDocument(
            BytesIO(data)
        )

    except Exception as exc:

        logger.exception(
            "Unable to open DOCX bytes."
        )

        raise RuntimeError(
            "Unable to open DOCX from bytes."
        ) from exc


# ============================================================
# LOAD DOCX BYTES
# ============================================================

def load_docx_bytes(
    data: bytes,
    filename: str = "document.docx",
    source: str = "memory",
    include_tables: bool = True,
    include_headers_footers: bool = False,
) -> DOCXLoadResult:

    try:

        document = open_docx_bytes(
            data
        )

        core_metadata = extract_core_metadata(
            document
        )

        paragraph_blocks = extract_paragraphs(
            document
        )

        table_blocks = []

        if include_tables:

            table_blocks = extract_tables(

                document,

                start_index=len(
                    paragraph_blocks
                ),

            )

        blocks = (
            paragraph_blocks
            +
            table_blocks
        )

        header_footer = {

            "headers": [],
            "footers": [],

        }

        if include_headers_footers:

            header_footer = extract_headers_footers(
                document
            )

        hyperlinks = extract_hyperlinks(
            document
        )

        metadata = dict(
            core_metadata
        )

        metadata.update({

            "filename":
                normalize_filename(
                    filename
                ),

            "source":
                source,

            "paragraph_count":
                len(
                    document.paragraphs
                ),

            "table_count":
                len(
                    document.tables
                ),

            "hyperlinks":
                hyperlinks,

            "headers":
                header_footer[
                    "headers"
                ],

            "footers":
                header_footer[
                    "footers"
                ],

        })

        documents = blocks_to_documents(

            blocks=blocks,

            source=source,

            filename=normalize_filename(
                filename
            ),

            file_path="",

            base_metadata=metadata,

        )

        table_row_count = sum(

            len(
                table.rows
            )

            for table
            in document.tables

        )

        table_cell_count = sum(

            len(
                row.cells
            )

            for table
            in document.tables

            for row
            in table.rows

        )

        heading_count = sum(

            1

            for block
            in paragraph_blocks

            if block.heading_level > 0

        )

        return DOCXLoadResult(

            documents=documents,

            blocks=blocks,

            metadata=metadata,

            paragraph_count=len(
                document.paragraphs
            ),

            table_count=len(
                document.tables
            ),

            table_row_count=table_row_count,

            table_cell_count=table_cell_count,

            heading_count=heading_count,

            hyperlink_count=len(
                hyperlinks
            ),

            success=True,

            warnings=[],

        )

    except Exception as exc:

        logger.exception(
            "DOCX byte loading failed."
        )

        return DOCXLoadResult(

            success=False,

            error=str(
                exc
            ),

        )


# ============================================================
# LOAD DOCX STREAM
# ============================================================

def load_docx_stream(
    stream: Any,
    filename: str = "document.docx",
    source: str = "stream",
) -> DOCXLoadResult:

    if stream is None:

        raise ValueError(
            "DOCX stream cannot be None."
        )

    if not hasattr(
        stream,
        "read",
    ):

        raise TypeError(

            "Expected a file-like object with read()."

        )

    data = stream.read()

    if not isinstance(
        data,
        bytes,
    ):

        raise TypeError(
            "DOCX stream must return bytes."
        )

    return load_docx_bytes(

        data=data,

        filename=filename,

        source=source,

    )


# ============================================================
# DOCUMENTS → TEXT
# ============================================================

def get_all_text(
    result: DOCXLoadResult,
    separator: str = "\n\n",
    include_headings: bool = True,
    include_tables: bool = True,
) -> str:

    texts = []

    for document in result.documents:

        text = clean_text(
            document.text
        )

        if not text:
            continue

        if (
            document.block_type == "heading"
            and
            not include_headings
        ):
            continue

        if (
            document.block_type == "table_row"
            and
            not include_tables
        ):
            continue

        texts.append(text)

    return separator.join(
        texts
    ).strip()


# ============================================================
# END CHUNK 7
# ============================================================
# ============================================================
# CHUNK 8/10
#
# RAG / LANGCHAIN CONVERSION
# ============================================================


# ============================================================
# DOCUMENT → DICT
# ============================================================

def document_to_dict(
    document: DOCXDocument,
) -> Dict[str, Any]:

    metadata = dict(
        document.metadata
    )

    metadata.update({

        "source":
            document.source,

        "filename":
            document.filename,

        "file_path":
            document.file_path,

        "block_index":
            document.block_index,

        "block_type":
            document.block_type,

        "heading_level":
            document.heading_level,

        "section":
            document.section,

        "extraction_method":
            document.extraction_method,

    })

    return {

        "page_content":
            document.text,

        "metadata":
            metadata,

    }


# ============================================================
# DOCUMENT → LANGCHAIN
# ============================================================

def document_to_langchain(
    document: DOCXDocument,
) -> Any:

    if LangChainDocument is None:

        raise ImportError(

            "langchain-core is required for LangChain "
            "Document conversion. Install using: "
            "pip install langchain-core"

        )

    metadata = dict(
        document.metadata
    )

    metadata.update({

        "source":
            document.source,

        "filename":
            document.filename,

        "file_path":
            document.file_path,

        "block_index":
            document.block_index,

        "block_type":
            document.block_type,

        "heading_level":
            document.heading_level,

        "section":
            document.section,

        "extraction_method":
            document.extraction_method,

    })

    return LangChainDocument(

        page_content=document.text,

        metadata=metadata,

    )


# ============================================================
# RESULT → LANGCHAIN DOCUMENTS
# ============================================================

def to_langchain_documents(
    result: DOCXLoadResult,
    skip_empty: bool = True,
) -> List[Any]:

    documents = []

    for document in result.documents:

        if (
            skip_empty
            and
            not clean_text(
                document.text
            )
        ):
            continue

        documents.append(

            document_to_langchain(
                document
            )

        )

    return documents


# ============================================================
# RESULT → DICTS
# ============================================================

def to_document_dicts(
    result: DOCXLoadResult,
    skip_empty: bool = True,
) -> List[Dict[str, Any]]:

    output = []

    for document in result.documents:

        if (
            skip_empty
            and
            not clean_text(
                document.text
            )
        ):
            continue

        output.append(
            document_to_dict(
                document
            )
        )

    return output


# ============================================================
# GROUP BY SECTION
# ============================================================

def group_by_section(
    result: DOCXLoadResult,
) -> Dict[str, List[DOCXDocument]]:

    groups = {}

    for document in result.documents:

        section = (
            document.section.strip()
            or
            "General"
        )

        groups.setdefault(
            section,
            []
        ).append(
            document
        )

    return groups


# ============================================================
# GET HEADINGS
# ============================================================

def get_headings(
    result: DOCXLoadResult,
) -> List[DOCXDocument]:

    return [

        document

        for document
        in result.documents

        if (
            document.block_type == "heading"
            and
            document.heading_level > 0
        )

    ]


# ============================================================
# GET TABLE DOCUMENTS
# ============================================================

def get_table_documents(
    result: DOCXLoadResult,
) -> List[DOCXDocument]:

    return [

        document

        for document
        in result.documents

        if document.block_type == "table_row"

    ]


# ============================================================
# END CHUNK 8
# ============================================================
# ============================================================
# CHUNK 9/10
#
# BATCH PROCESSING + STATISTICS
# ============================================================


# ============================================================
# DOCUMENT STATISTICS
# ============================================================

def document_statistics(
    result: DOCXLoadResult,
) -> Dict[str, Any]:

    total_characters = 0

    total_words = 0

    heading_count = 0

    table_row_documents = 0

    for document in result.documents:

        text = clean_text(
            document.text
        )

        total_characters += len(
            text
        )

        total_words += len(
            text.split()
        )

        if document.block_type == "heading":

            heading_count += 1

        if document.block_type == "table_row":

            table_row_documents += 1

    return {

        "paragraph_count":
            result.paragraph_count,

        "table_count":
            result.table_count,

        "table_row_count":
            result.table_row_count,

        "table_cell_count":
            result.table_cell_count,

        "heading_count":
            heading_count,

        "hyperlink_count":
            result.hyperlink_count,

        "document_count":
            len(
                result.documents
            ),

        "total_characters":
            total_characters,

        "total_words":
            total_words,

        "table_row_documents":
            table_row_documents,

        "average_words_per_document": (

            round(

                total_words
                /
                max(
                    len(
                        result.documents
                    ),
                    1,
                ),

                2,

            )

        ),

    }


# ============================================================
# LOAD DOCX DIRECTORY
# ============================================================

def load_docx_directory(
    directory: Union[
        str,
        Path,
    ],
    recursive: bool = True,
) -> List[DOCXLoadResult]:

    path = Path(
        directory
    )

    if not path.exists():

        raise FileNotFoundError(
            f"Directory does not exist: {path}"
        )

    if not path.is_dir():

        raise ValueError(
            f"Not a directory: {path}"
        )

    if recursive:

        files = path.rglob(
            "*.docx"
        )

    else:

        files = path.glob(
            "*.docx"
        )

    results = []

    for file_path in sorted(
        files
    ):

        try:

            results.append(
                load_docx(
                    file_path
                )
            )

        except Exception as exc:

            logger.exception(
                "Failed loading DOCX: %s",
                file_path,
            )

            results.append(

                DOCXLoadResult(

                    success=False,

                    error=str(
                        exc
                    ),

                )

            )

    return results


# ============================================================
# BATCH LOAD
# ============================================================

def batch_load_docx(
    file_paths: Iterable[
        Union[
            str,
            Path,
        ]
    ],
) -> List[DOCXLoadResult]:

    results = []

    for file_path in file_paths:

        try:

            results.append(
                load_docx(
                    file_path
                )
            )

        except Exception as exc:

            logger.exception(
                "Batch DOCX loading failed: %s",
                file_path,
            )

            results.append(

                DOCXLoadResult(

                    success=False,

                    error=str(
                        exc
                    ),

                )

            )

    return results


# ============================================================
# COMBINE RESULTS
# ============================================================

def combine_results(
    results: Sequence[
        DOCXLoadResult
    ],
) -> DOCXLoadResult:

    documents = []

    blocks = []

    metadata = {

        "source_count":
            len(
                results
            ),

    }

    warnings = []

    errors = []

    success = True

    paragraph_count = 0

    table_count = 0

    table_row_count = 0

    table_cell_count = 0

    heading_count = 0

    hyperlink_count = 0

    for result in results:

        documents.extend(
            result.documents
        )

        blocks.extend(
            result.blocks
        )

        warnings.extend(
            result.warnings
        )

        paragraph_count += (
            result.paragraph_count
        )

        table_count += (
            result.table_count
        )

        table_row_count += (
            result.table_row_count
        )

        table_cell_count += (
            result.table_cell_count
        )

        heading_count += (
            result.heading_count
        )

        hyperlink_count += (
            result.hyperlink_count
        )

        if not result.success:

            success = False

            if result.error:

                errors.append(
                    result.error
                )

    return DOCXLoadResult(

        documents=documents,

        blocks=blocks,

        metadata=metadata,

        paragraph_count=paragraph_count,

        table_count=table_count,

        table_row_count=table_row_count,

        table_cell_count=table_cell_count,

        heading_count=heading_count,

        hyperlink_count=hyperlink_count,

        success=success,

        error=(

            "; ".join(
                errors
            )

            if errors
            else None

        ),

        warnings=warnings,

    )


# ============================================================
# VALIDATE RESULT
# ============================================================

def validate_load_result(
    result: DOCXLoadResult,
) -> Dict[str, Any]:

    errors = []

    warnings = []

    if not isinstance(
        result,
        DOCXLoadResult,
    ):

        return {

            "valid": False,

            "errors": [
                "Invalid DOCXLoadResult."
            ],

            "warnings": [],

        }

    if not result.success:

        errors.append(

            result.error
            or
            "DOCX loading failed."

        )

    if result.paragraph_count < 0:

        errors.append(
            "Paragraph count cannot be negative."
        )

    if result.table_count < 0:

        errors.append(
            "Table count cannot be negative."
        )

    if result.heading_count < 0:

        errors.append(
            "Heading count cannot be negative."
        )

    if not result.documents:

        warnings.append(
            "No DOCX document blocks were extracted."
        )

    warnings.extend(
        result.warnings
    )

    return {

        "valid":
            len(errors) == 0,

        "errors":
            list(
                dict.fromkeys(
                    errors
                )
            ),

        "warnings":
            list(
                dict.fromkeys(
                    warnings
                )
            ),

    }


# ============================================================
# END CHUNK 9
# ============================================================
# ============================================================
# CHUNK 10/10
#
# GENERIC API + SUMMARY + EXPORTS + SELF TEST
# ============================================================


# ============================================================
# LOADER SUMMARY
# ============================================================

def loader_summary(
    result: DOCXLoadResult,
) -> Dict[str, Any]:

    statistics = document_statistics(
        result
    )

    return {

        "loader_version":
            DOCX_LOADER_VERSION,

        "success":
            result.success,

        "paragraph_count":
            result.paragraph_count,

        "table_count":
            result.table_count,

        "table_row_count":
            result.table_row_count,

        "table_cell_count":
            result.table_cell_count,

        "heading_count":
            result.heading_count,

        "hyperlink_count":
            result.hyperlink_count,

        "document_count":
            len(
                result.documents
            ),

        "total_words":
            statistics[
                "total_words"
            ],

        "total_characters":
            statistics[
                "total_characters"
            ],

        "warnings":
            result.warnings,

        "error":
            result.error,

    }


# ============================================================
# GENERIC LOAD FUNCTION
# ============================================================

def load(
    source: Union[
        str,
        Path,
        bytes,
        Any,
    ],
    filename: str = "document.docx",
) -> DOCXLoadResult:

    # --------------------------------------------------------
    # File path
    # --------------------------------------------------------

    if isinstance(
        source,
        (
            str,
            Path,
        ),
    ):

        return load_docx(
            source
        )

    # --------------------------------------------------------
    # Bytes
    # --------------------------------------------------------

    if isinstance(
        source,
        bytes,
    ):

        return load_docx_bytes(

            source,

            filename=filename,

        )

    # --------------------------------------------------------
    # File-like object
    # --------------------------------------------------------

    if hasattr(
        source,
        "read",
    ):

        return load_docx_stream(

            source,

            filename=filename,

        )

    raise TypeError(

        "Unsupported DOCX source. Expected a path, "
        "bytes, or file-like object."

    )


# ============================================================
# RAG CONTENT EXTRACTION
# ============================================================

def get_rag_documents(
    source: Union[
        str,
        Path,
        bytes,
        Any,
    ],
    filename: str = "document.docx",
) -> List[Dict[str, Any]]:

    result = load(
        source,
        filename=filename,
    )

    if not result.success:

        raise RuntimeError(

            result.error
            or
            "Unable to load DOCX."

        )

    return to_document_dicts(
        result
    )


# ============================================================
# CAPABILITIES
# ============================================================

DOCX_LOADER_CAPABILITIES = [

    "docx_file_loading",

    "docx_bytes_loading",

    "docx_stream_loading",

    "paragraph_extraction",

    "heading_detection",

    "table_extraction",

    "header_extraction",

    "footer_extraction",

    "hyperlink_extraction",

    "core_metadata_extraction",

    "section_detection",

    "rag_document_conversion",

    "langchain_document_conversion",

    "batch_docx_loading",

    "directory_docx_loading",

]


# ============================================================
# PUBLIC EXPORTS
# ============================================================

__all__ = [

    # Version
    "DOCX_LOADER_VERSION",

    # Models
    "DOCXBlock",

    "DOCXDocument",

    "DOCXLoadResult",

    # Availability
    "is_python_docx_available",

    "require_python_docx",

    # Utilities
    "clean_text",

    "normalize_filename",

    "normalize_source",

    "normalize_style_name",

    # Headings / paragraphs
    "get_heading_level",

    "is_heading",

    "extract_paragraph",

    "extract_paragraphs",

    "get_text_paragraphs",

    # Tables
    "extract_table_cell",

    "extract_table",

    "extract_tables",

    "table_to_text",

    # Headers / footers / hyperlinks
    "extract_headers_footers",

    "extract_hyperlinks",

    # Core loading
    "validate_docx_path",

    "open_docx",

    "extract_core_metadata",

    "blocks_to_documents",

    "load_docx",

    "load_docx_strict",

    # Bytes / streams
    "open_docx_bytes",

    "load_docx_bytes",

    "load_docx_stream",

    # Text
    "get_all_text",

    # RAG
    "document_to_dict",

    "document_to_langchain",

    "to_langchain_documents",

    "to_document_dicts",

    "group_by_section",

    "get_headings",

    "get_table_documents",

    "get_rag_documents",

    # Statistics
    "document_statistics",

    "loader_summary",

    # Batch
    "load_docx_directory",

    "batch_load_docx",

    "combine_results",

    # Validation
    "validate_load_result",

    # Capabilities
    "DOCX_LOADER_CAPABILITIES",

]


# ============================================================
# SELF TEST
# ============================================================

if __name__ == "__main__":

    import sys

    print(
        "\n"
        "============================================"
    )

    print(
        "DOCX LOADER SELF TEST"
    )

    print(
        "============================================"
    )

    # --------------------------------------------------------
    # Dependency
    # --------------------------------------------------------

    print(
        "\npython-docx available:"
    )

    print(
        is_python_docx_available()
    )

    # --------------------------------------------------------
    # Text cleaning
    # --------------------------------------------------------

    sample_text = """

        This is     sample DOCX text.


        It contains      extra spaces.


        And multiple blank lines.

    """

    print(
        "\nCleaned Text:"
    )

    print(
        clean_text(
            sample_text
        )
    )

    # --------------------------------------------------------
    # Model test
    # --------------------------------------------------------

    sample_block = DOCXBlock(

        block_type="heading",

        text="Generative AI",

        index=0,

        section="Generative AI",

        heading_level=1,

        style="Heading 1",

        metadata={

            "test":
                True,

        },

    )

    sample_document = DOCXDocument(

        text="Generative AI is an important skill.",

        metadata={

            "section":
                "Generative AI",

        },

        source="demo",

        filename="demo.docx",

        file_path="",

        block_index=1,

        block_type="paragraph",

        heading_level=0,

        section="Generative AI",

    )

    sample_result = DOCXLoadResult(

        documents=[
            sample_document
        ],

        blocks=[
            sample_block
        ],

        metadata={

            "title":
                "Demo Curriculum",

        },

        paragraph_count=1,

        table_count=0,

        table_row_count=0,

        table_cell_count=0,

        heading_count=1,

        hyperlink_count=0,

        success=True,

    )

    # --------------------------------------------------------
    # Statistics
    # --------------------------------------------------------

    print(
        "\nStatistics:"
    )

    print(
        document_statistics(
            sample_result
        )
    )

    # --------------------------------------------------------
    # Summary
    # --------------------------------------------------------

    print(
        "\nSummary:"
    )

    print(
        loader_summary(
            sample_result
        )
    )

    # --------------------------------------------------------
    # Validation
    # --------------------------------------------------------

    print(
        "\nValidation:"
    )

    print(
        validate_load_result(
            sample_result
        )
    )

    # --------------------------------------------------------
    # Dict conversion
    # --------------------------------------------------------

    print(
        "\nDocument Dictionary:"
    )

    print(
        document_to_dict(
            sample_document
        )
    )

    # --------------------------------------------------------
    # Actual file test
    # --------------------------------------------------------

    if len(sys.argv) > 1:

        docx_path = sys.argv[1]

        print(
            "\nLoading:"
        )

        print(
            docx_path
        )

        result = load_docx(
            docx_path
        )

        print(
            "\nLoader Summary:"
        )

        print(
            loader_summary(
                result
            )
        )

        print(
            "\nHeadings:"
        )

        for heading in get_headings(
            result
        ):

            print(
                f"- "
                f"{heading.heading_level}: "
                f"{heading.text}"
            )

        print(
            "\nExtracted Text Preview:"
        )

        print(

            get_all_text(
                result
            )[:5000]

        )

        print(
            "\nSections:"
        )

        for section, documents in group_by_section(
            result
        ).items():

            print(

                f"- {section}: "
                f"{len(documents)} blocks"

            )

    print(
        "\n============================================"
    )

    print(
        "DOCX LOADER TEST COMPLETE"
    )

    print(
        "============================================"
    )


# ============================================================
# END OF rag/docx_loader.py
# ============================================================
