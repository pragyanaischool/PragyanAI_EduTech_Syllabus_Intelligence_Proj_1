# ============================================================
# rag/document_loader.py
# ============================================================
#
# PragyanAI Curriculum Intelligence Platform
#
# Unified document loading layer.
#
# Supported:
#   PDF
#   DOCX
#   TXT
#   MD
#   CSV
#   JSON
#   PNG
#   JPG
#   JPEG
#   WEBP
#
# Main API:
#
#   process_uploaded_file(uploaded_file)
#
# ============================================================

from __future__ import annotations

import io
import json
import logging

from pathlib import Path
from typing import Any, Dict, Optional


logger = logging.getLogger(
    "pragyanai.rag.document_loader"
)


# ============================================================
# SUPPORTED EXTENSIONS
# ============================================================

SUPPORTED_EXTENSIONS = {

    ".pdf",

    ".docx",

    ".txt",

    ".md",

    ".csv",

    ".json",

    ".png",

    ".jpg",

    ".jpeg",

    ".webp",

}


# ============================================================
# END CHUNK 1
# ============================================================
# ============================================================
# CHUNK 2/6
#
# BASIC DOCUMENT LOADERS
# ============================================================


def load_text_file(
    file_bytes: bytes,
) -> str:

    """
    Load TXT / Markdown / CSV files.
    """

    if not file_bytes:

        return ""

    return file_bytes.decode(

        "utf-8",

        errors="ignore",

    )


# ============================================================
# JSON LOADER
# ============================================================

def load_json_file(
    file_bytes: bytes,
) -> str:

    """
    Load JSON and convert it into readable text.
    """

    if not file_bytes:

        return ""

    text = file_bytes.decode(

        "utf-8",

        errors="ignore",

    )

    try:

        data = json.loads(
            text
        )

        return json.dumps(

            data,

            indent=2,

            ensure_ascii=False,

        )

    except json.JSONDecodeError:

        return text


# ============================================================
# DOCX LOADER
# ============================================================

def load_docx_file(
    file_bytes: bytes,
) -> str:

    """
    Extract text from DOCX.
    """

    try:

        from docx import Document

        document = Document(

            io.BytesIO(
                file_bytes
            )

        )

        sections = []

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                sections.append(
                    text
                )

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = [

                    cell.text.strip()

                    for cell
                    in row.cells

                ]

                cells = [

                    cell

                    for cell in cells

                    if cell

                ]

                if cells:

                    sections.append(

                        " | ".join(
                            cells
                        )

                    )

        return "\n".join(
            sections
        )

    except Exception as exc:

        logger.exception(
            "DOCX extraction failed"
        )

        raise RuntimeError(

            f"Unable to extract DOCX: {exc}"

        ) from exc


# ============================================================
# PDF LOADER
# ============================================================

def load_pdf_file(
    file_bytes: bytes,
) -> str:

    """
    Extract text from PDF using PyMuPDF.
    """

    try:

        import fitz

        document = fitz.open(

            stream=file_bytes,

            filetype="pdf",

        )

        pages = []

        for page_number, page in enumerate(

            document,

            start=1,

        ):

            text = page.get_text(
                "text"
            )

            if text.strip():

                pages.append(

                    f"[Page {page_number}]\n"
                    +
                    text.strip()

                )

        document.close()

        return "\n\n".join(
            pages
        )

    except Exception as exc:

        logger.exception(
            "PDF extraction failed"
        )

        raise RuntimeError(

            f"Unable to extract PDF: {exc}"

        ) from exc


# ============================================================
# END CHUNK 2
# ============================================================
# ============================================================
# CHUNK 3/6
#
# IMAGE / OCR
# ============================================================


def load_image_file(
    file_bytes: bytes,
) -> str:

    """
    Extract text from image using Tesseract OCR.

    OCR is optional.

    If Tesseract is not installed, a clear
    RuntimeError is returned instead of causing
    the entire application to crash.
    """

    try:

        from PIL import Image

        import pytesseract

        image = Image.open(

            io.BytesIO(
                file_bytes
            )

        )

        text = pytesseract.image_to_string(
            image
        )

        return text.strip()

    except ImportError as exc:

        raise RuntimeError(

            "OCR dependencies are not installed. "
            "Install Pillow and pytesseract."

        ) from exc

    except Exception as exc:

        message = str(
            exc
        )

        # ----------------------------------------------------
        # Tesseract executable missing
        # ----------------------------------------------------

        if (
            "tesseract" in
            message.lower()
        ):

            raise RuntimeError(

                "Tesseract OCR engine is not "
                "available on this deployment. "
                "PDF and DOCX extraction will still "
                "work, but image OCR requires "
                "Tesseract."

            ) from exc

        raise RuntimeError(

            f"Image OCR failed: {exc}"

        ) from exc


# ============================================================
# OPTIONAL OCR FALLBACK
# ============================================================

def try_ocr(
    file_bytes: bytes,
) -> str:

    """
    Try OCR without crashing the application.
    """

    try:

        return load_image_file(
            file_bytes
        )

    except Exception as exc:

        logger.warning(

            "OCR unavailable: %s",

            exc,

        )

        return ""


# ============================================================
# END CHUNK 3
# ============================================================
# ============================================================
# CHUNK 4/6
#
# UNIFIED FILE PROCESSOR
# ============================================================


def process_file_bytes(
    file_bytes: bytes,
    filename: str,
) -> str:

    """
    Process a document based on its extension.

    Parameters
    ----------
    file_bytes:
        Raw file bytes.

    filename:
        Original filename.

    Returns
    -------
    str:
        Extracted document text.
    """

    if not file_bytes:

        return ""

    extension = (

        Path(
            filename
        )
        .suffix
        .lower()

    )

    logger.info(

        "Processing file: %s (%s)",

        filename,

        extension,

    )

    # --------------------------------------------------------
    # Validate
    # --------------------------------------------------------

    if extension not in SUPPORTED_EXTENSIONS:

        raise ValueError(

            f"Unsupported file type: "
            f"{extension or 'unknown'}"

        )

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    if extension == ".pdf":

        return load_pdf_file(
            file_bytes
        )

    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    if extension == ".docx":

        return load_docx_file(
            file_bytes
        )

    # --------------------------------------------------------
    # TXT / MD / CSV
    # --------------------------------------------------------

    if extension in {

        ".txt",

        ".md",

        ".csv",

    }:

        return load_text_file(
            file_bytes
        )

    # --------------------------------------------------------
    # JSON
    # --------------------------------------------------------

    if extension == ".json":

        return load_json_file(
            file_bytes
        )

    # --------------------------------------------------------
    # Images
    # --------------------------------------------------------

    if extension in {

        ".png",

        ".jpg",

        ".jpeg",

        ".webp",

    }:

        return load_image_file(
            file_bytes
        )

    return ""


# ============================================================
# STREAMLIT UPLOADED FILE
# ============================================================

def process_uploaded_file(
    uploaded_file: Any,
) -> Dict[str, Any]:

    """
    Process a Streamlit UploadedFile.

    Returns:

    {
        "success": True,
        "filename": "...",
        "extension": ".pdf",
        "text": "...",
        "character_count": 12345,
        "error": None
    }
    """

    if uploaded_file is None:

        return {

            "success":
                False,

            "filename":
                None,

            "extension":
                None,

            "text":
                "",

            "character_count":
                0,

            "error":
                "No file supplied.",

        }

    filename = getattr(

        uploaded_file,

        "name",

        "uploaded_file",

    )

    try:

        file_bytes = uploaded_file.read()

    except Exception as exc:

        logger.exception(
            "Unable to read uploaded file"
        )

        return {

            "success":
                False,

            "filename":
                filename,

            "extension":
                Path(
                    filename
                ).suffix.lower(),

            "text":
                "",

            "character_count":
                0,

            "error":
                str(exc),

        }

    try:

        text = process_file_bytes(

            file_bytes=file_bytes,

            filename=filename,

        )

        text = text.strip()

        return {

            "success":
                bool(text),

            "filename":
                filename,

            "extension":
                Path(
                    filename
                ).suffix.lower(),

            "text":
                text,

            "character_count":
                len(text),

            "error":
                None
                if text
                else
                "No text could be extracted.",

        }

    except Exception as exc:

        logger.exception(

            "File processing failed: %s",

            filename,

        )

        return {

            "success":
                False,

            "filename":
                filename,

            "extension":
                Path(
                    filename
                ).suffix.lower(),

            "text":
                "",

            "character_count":
                0,

            "error":
                str(exc),

        }


# ============================================================
# END CHUNK 4
# ============================================================
# ============================================================
# CHUNK 5/6
#
# CURRICULUM DOCUMENT PROCESSING
# ============================================================


def process_curriculum_file(
    uploaded_file: Any,
) -> Dict[str, Any]:

    """
    Specialized wrapper for syllabus/curriculum documents.

    The document is first converted to text.

    The text can then be passed to:

        curriculum.extractor.extract_syllabus()

    or:

        curriculum.extractor.extract_curriculum_from_text()
    """

    result = process_uploaded_file(
        uploaded_file
    )

    if not result[
        "success"
    ]:

        return result

    text = result[
        "text"
    ]

    # --------------------------------------------------------
    # Basic cleanup
    # --------------------------------------------------------

    text = normalize_document_text(
        text
    )

    result[
        "text"
    ] = text

    result[
        "character_count"
    ] = len(
        text
    )

    return result


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize_document_text(
    text: str,
) -> str:

    """
    Normalize extracted document text while
    preserving meaningful structure.
    """

    if not text:

        return ""

    # Normalize Windows line endings

    text = text.replace(
        "\r\n",
        "\n",
    )

    text = text.replace(
        "\r",
        "\n",
    )

    # Remove excessive spaces

    text = "\n".join(

        line.rstrip()

        for line
        in text.splitlines()

    )

    # Remove excessive blank lines

    while "\n\n\n" in text:

        text = text.replace(

            "\n\n\n",

            "\n\n",

        )

    return text.strip()


# ============================================================
# SAVE EXTRACTED TEXT
# ============================================================

def save_extracted_text(
    text: str,
    output_path: str,
) -> str:

    """
    Save extracted text to disk.

    Mainly useful for debugging/local development.
    """

    path = Path(
        output_path
    )

    path.parent.mkdir(

        parents=True,

        exist_ok=True,

    )

    path.write_text(

        text,

        encoding="utf-8",

    )

    return str(
        path
    )


# ============================================================
# END CHUNK 5
# ============================================================
# ============================================================
# CHUNK 6/6
#
# COMPATIBILITY + EXPORTS
# ============================================================


# ============================================================
# COMPATIBILITY ALIAS
# ============================================================

process_document = (
    process_file_bytes
)


# ============================================================
# SUPPORTED FILE CHECK
# ============================================================

def is_supported_file(
    filename: str,
) -> bool:

    if not filename:

        return False

    extension = (

        Path(
            filename
        )
        .suffix
        .lower()

    )

    return (
        extension
        in
        SUPPORTED_EXTENSIONS
    )


# ============================================================
# GET FILE TYPE
# ============================================================

def get_file_type(
    filename: str,
) -> str:

    extension = (

        Path(
            filename
        )
        .suffix
        .lower()

    )

    mapping = {

        ".pdf":
            "PDF",

        ".docx":
            "DOCX",

        ".txt":
            "Text",

        ".md":
            "Markdown",

        ".csv":
            "CSV",

        ".json":
            "JSON",

        ".png":
            "Image",

        ".jpg":
            "Image",

        ".jpeg":
            "Image",

        ".webp":
            "Image",

    }

    return mapping.get(

        extension,

        "Unknown",

    )


# ============================================================
# PUBLIC API
# ============================================================

__all__ = [

    "SUPPORTED_EXTENSIONS",

    "load_text_file",

    "load_json_file",

    "load_docx_file",

    "load_pdf_file",

    "load_image_file",

    "try_ocr",

    "process_file_bytes",

    "process_uploaded_file",

    "process_curriculum_file",

    "process_document",

    "normalize_document_text",

    "save_extracted_text",

    "is_supported_file",

    "get_file_type",

]


# ============================================================
# END OF rag/document_loader.py
# ============================================================
